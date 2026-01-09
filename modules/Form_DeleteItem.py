"""
Form_DeleteItem_Government.py - Delete Item from AMES List
Professional government interface for deleting items from approved AMES lists
With color coding: Green = Added, Red = Deleted
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from hijri_converter import Hijri, Gregorian
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageTk
import os
from helpers.unified_database import UnifiedDatabase
from helpers.resource_path import get_logo_path, get_template_path

try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FormDeleteItem:
    """Government-styled Delete Item Form with color coding"""
    
    def __init__(self, root, parent_window=None):
        self.root = root
        self.parent_window = parent_window
        self.root.title("Sistem Pemadaman Item - Permohonan Pemadaman Item AMES")
        self.root.geometry("1600x1000")
        self.root.configure(bg='#F5F5F5')
        
        # Initialize database
        self.db = UnifiedDatabase()
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': 'white',
            'text_dark': '#1a1a1a',
            'text_light': 'white',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_danger': '#C62828',
            'button_info': '#1976D2',
            'button_hover': '#004d99',
            'item_added': '#4CAF50',      # Green for added
            'item_deleted': '#F44336'      # Red for deleted
        }
        
        # Center window
        self.center_window()
        
        # Make window appear on top
        self.root.attributes('-topmost', True)
        self.root.after_idle(lambda: self.root.attributes('-topmost', False))
        self.root.lift()
        self.root.focus_force()
        
        # Main container
        self.main_container = tk.Frame(root, bg=self.colors['bg_main'])
        self.main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create header
        self.create_header()
        
        # Create scrollable content
        self.create_scrollable_content()
        
        # Auto-update tarikh islam
        self.update_tarikh_islam()
        
        # Handle window close
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
    
    def create_header(self):
        """Create government header"""
        header_frame = tk.Frame(self.main_container, bg=self.colors['primary'], height=100)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Logo section
        logo_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        logo_frame.pack(side=tk.LEFT, padx=20)
        
        try:
            logo_path = get_logo_path()
            if os.path.exists(logo_path):
                logo_image = Image.open(logo_path)
                
                # Convert to RGBA if not already
                if logo_image.mode != 'RGBA':
                    logo_image = logo_image.convert('RGBA')
                
                logo_image = logo_image.resize((60, 60), Image.Resampling.LANCZOS)
                
                # Create a background image with the same color as the frame
                background_color = self.colors['primary']  # '#003366'
                bg_rgb = tuple(int(background_color[i:i+2], 16) for i in (1, 3, 5))
                
                # Create a new RGBA image with the background color
                background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
                
                # Composite the logo onto the background using alpha compositing
                if logo_image.mode == 'RGBA':
                    logo_image = Image.alpha_composite(background, logo_image)
                    # Convert back to RGB for PhotoImage (Tkinter doesn't support RGBA well)
                    logo_image = logo_image.convert('RGB')
                
                logo_photo = ImageTk.PhotoImage(logo_image)
                logo_label = tk.Label(logo_frame, image=logo_photo, bg=self.colors['primary'])
                logo_label.image = logo_photo
                logo_label.pack()
        except Exception as e:
            pass
        
        # Title section
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.pack(side=tk.LEFT, expand=True)
        
        title = tk.Label(title_frame,
                        text="AMES - PEMADAMAN / PENAMBAHAN ITEM",
                        font=('Arial', 18, 'bold'),
                        bg=self.colors['primary'],
                        fg='white')
        title.pack(pady=(25, 5))
        
        subtitle = tk.Label(title_frame,
                           text="Permohonan Pemadaman atau Penambahan Item dari Senarai AMES",
                           font=('Arial', 18, 'italic'),
                           bg=self.colors['primary'],
                           fg='#E0E0E0')
        subtitle.pack()
        
        # Back button
        btn_back = tk.Button(header_frame,
                            text="← KEMBALI",
                            font=('Arial', 18, 'bold'),
                            bg=self.colors['secondary'],
                            fg='white',
                            relief=tk.FLAT,
                            cursor='hand2',
                            width=15,
                            height=2,
                            command=self.on_back_click)
        btn_back.pack(side=tk.RIGHT, padx=20)
        
        # Separator
        separator = tk.Frame(self.main_container, bg=self.colors['accent'], height=2)
        separator.pack(fill=tk.X)
    
    def create_scrollable_content(self):
        """Create scrollable canvas for content"""
        scroll_container = tk.Frame(self.main_container, bg=self.colors['bg_main'])
        scroll_container.pack(fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(scroll_container, bg=self.colors['bg_main'], highlightthickness=0)
        scrollbar_y = ttk.Scrollbar(scroll_container, orient='vertical', command=self.canvas.yview)
        
        self.scrollable_frame = tk.Frame(self.canvas, bg=self.colors['bg_main'])
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        
        self.canvas.configure(yscrollcommand=scrollbar_y.set)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.scrollable_frame.bind('<Configure>', self.on_frame_configure)
        self.canvas.bind('<Configure>', self.on_canvas_configure)
        self.canvas.bind_all('<MouseWheel>', self.on_mousewheel)
        
        # Create form sections
        self.create_form_section()
        self.create_legend_section()
        self.create_table_section()
        self.create_button_section()
    
    def on_frame_configure(self, event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    def on_canvas_configure(self, event):
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas_frame, width=canvas_width)
    
    def on_mousewheel(self, event):
        if event.num == 5 or event.delta < 0:
            self.canvas.yview_scroll(1, 'units')
        elif event.num == 4 or event.delta > 0:
            self.canvas.yview_scroll(-1, 'units')
    
    def create_form_section(self):
        """Create form input section"""
        section_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        section_container.pack(fill=tk.X, padx=30, pady=20)
        
        card = tk.Frame(section_container, bg='white')
        card.pack(fill=tk.X)
        
        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, padx=25, pady=20)
        
        # Section title
        section_title = tk.Label(content,
                                text="Maklumat Permohonan",
                                font=('Arial', 18, 'bold'),
                                bg='white',
                                fg=self.colors['primary'])
        section_title.grid(row=0, column=0, columnspan=4, pady=(0, 15), sticky='w')
        
        row = 1
        
        def create_label(text, row_num, col_num):
            lbl = tk.Label(content, text=text, font=('Arial', 11, 'bold'),
                          bg='white', fg=self.colors['text_dark'])
            lbl.grid(row=row_num, column=col_num, sticky='e', padx=(5, 10), pady=6)
            return lbl
        
        # Row 1: Rujukan & Alamat
        create_label("RUJUKAN KAMI:", row, 0)
        rujukan_frame = tk.Frame(content, bg='white')
        rujukan_frame.grid(row=row, column=1, sticky='w', padx=5, pady=6)
        tk.Label(rujukan_frame, text="KE.JB(90)650/14/AMES/", bg='white', font=('Arial', 11)).pack(side=tk.LEFT)
        self.entry_rujukan = tk.Entry(rujukan_frame, width=22, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_rujukan.pack(side=tk.LEFT)
        
        create_label("ALAMAT:", row, 2)
        self.text_alamat = tk.Text(content, width=45, height=4, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.text_alamat.grid(row=row, column=3, rowspan=3, sticky='w', padx=5, pady=6)
        row += 1
        
        # Row 2: Nama Syarikat
        create_label("NAMA SYARIKAT:", row, 0)
        self.entry_nama = tk.Entry(content, width=42, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_nama.grid(row=row, column=1, sticky='w', padx=5, pady=6)
        row += 1
        
        # Row 3: Tarikh
        create_label("TARIKH:", row, 0)
        self.entry_tarikh = tk.Entry(content, width=42, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_tarikh.insert(0, datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.grid(row=row, column=1, sticky='w', padx=5, pady=6)
        self.entry_tarikh.bind('<KeyRelease>', lambda e: self.update_tarikh_islam())
        row += 1
        
        # Row 4: Tarikh Islam & Nama Pegawai
        create_label("TARIKH ISLAM:", row, 0)
        self.entry_islam = tk.Entry(content, width=42, font=('Arial', 11), state='readonly', relief=tk.SOLID, bd=1)
        self.entry_islam.grid(row=row, column=1, sticky='w', padx=5, pady=6)
        
        create_label("NAMA PEGAWAI:", row, 2)
        self.combo_pegawai = ttk.Combobox(content, width=42, font=('Arial', 11))
        self.combo_pegawai['values'] = ["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                        "KHAIRONE DZARWANY BINTI MOHD KAIRI"]
        self.combo_pegawai.current(0)
        self.combo_pegawai.grid(row=row, column=3, sticky='w', padx=5, pady=6)
        row += 1
        
        # Divider
        divider = tk.Frame(content, bg=self.colors['border'], height=1)
        divider.grid(row=row, column=0, columnspan=4, sticky='ew', pady=10)
        row += 1
        
        # Row 5: No. Kelulusan AMES & Search Button
        create_label("NO. KELULUSAN AMES:", row, 0)
        kelulusan_frame = tk.Frame(content, bg='white')
        kelulusan_frame.grid(row=row, column=1, columnspan=3, sticky='w', padx=5, pady=6)
        
        self.entry_kelulusan = tk.Entry(kelulusan_frame, width=30, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_kelulusan.pack(side=tk.LEFT, padx=(0, 10))
        
        tk.Button(kelulusan_frame, text="🔍 Cari AMES", font=('Arial', 11, 'bold'),
                 bg=self.colors['button_info'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=15,
                 command=self.search_ames).pack(side=tk.LEFT, padx=5)
        
        tk.Button(kelulusan_frame, text="📋 Pilih dari Senarai", font=('Arial', 11, 'bold'),
                 bg=self.colors['button_secondary'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=18,
                 command=self.show_ames_list).pack(side=tk.LEFT, padx=5)
        
        row += 1
        
        # Kategori (read-only, filled by search)
        create_label("KATEGORI:", row, 0)
        self.entry_kategori = tk.Entry(content, width=42, font=('Arial', 11), 
                                      state='readonly', relief=tk.SOLID, bd=1)
        self.entry_kategori.grid(row=row, column=1, sticky='w', padx=5, pady=6)
    
    def create_legend_section(self):
        """Create color legend"""
        legend_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        legend_container.pack(fill=tk.X, padx=30, pady=(0, 10))
        
        card = tk.Frame(legend_container, bg='white')
        card.pack(fill=tk.X)
        
        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.X, padx=25, pady=15)
        
        tk.Label(content, text="LEGENDA WARNA:",
                font=('Arial', 18, 'bold'), bg='white',
                fg=self.colors['primary']).pack(side=tk.LEFT, padx=(0, 20))
        
        # Gray legend
        gray_frame = tk.Frame(content, bg='#F5F5F5', width=20, height=20)
        gray_frame.pack(side=tk.LEFT, padx=5)
        tk.Label(content, text="= Item Asal (dari AMES)",
                font=('Arial', 11), bg='white').pack(side=tk.LEFT, padx=(0, 20))
        
        # Green legend
        green_frame = tk.Frame(content, bg=self.colors['item_added'], width=20, height=20)
        green_frame.pack(side=tk.LEFT, padx=5)
        tk.Label(content, text="= Item Ditambah (Baru)",
                font=('Arial', 11), bg='white').pack(side=tk.LEFT, padx=(0, 20))
        
        # Red legend
        red_frame = tk.Frame(content, bg=self.colors['item_deleted'], width=20, height=20)
        red_frame.pack(side=tk.LEFT, padx=5)
        tk.Label(content, text="= Item Dipadam",
                font=('Arial', 11), bg='white').pack(side=tk.LEFT)
    
    def create_table_section(self):
        """Create items table section"""
        section_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        section_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 20))
        
        card = tk.Frame(section_container, bg='white')
        card.pack(fill=tk.BOTH, expand=True)
        
        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, expand=True, padx=25, pady=20)
        
        tk.Label(content, text="SENARAI ITEM",
                font=('Arial', 11, 'bold'), bg='white', 
                fg=self.colors['primary']).pack(anchor='w', pady=(0, 10))
        
        # Button controls
        btn_frame = tk.Frame(content, bg='white')
        btn_frame.pack(anchor='e', pady=(0, 5))
        
        tk.Button(btn_frame, text="+ Tambah Item (HIJAU)", font=('Arial', 11, 'bold'),
                 bg=self.colors['item_added'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=20,
                 command=self.add_item).pack(side=tk.LEFT, padx=3)
        
        tk.Button(btn_frame, text="- Padam Item (MERAH)", font=('Arial', 11, 'bold'),
                 bg=self.colors['item_deleted'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=20,
                 command=self.delete_item).pack(side=tk.LEFT, padx=3)
        
        tk.Button(btn_frame, text="✎ Edit", font=('Arial', 11, 'bold'),
                 bg=self.colors['button_primary'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=12,
                 command=self.edit_item).pack(side=tk.LEFT, padx=3)
        
        tk.Button(btn_frame, text="✕ Buang dari Senarai", font=('Arial', 11, 'bold'),
                 bg=self.colors['button_secondary'], fg='white',
                 relief=tk.FLAT, cursor='hand2', width=18,
                 command=self.remove_from_list).pack(side=tk.LEFT, padx=3)
        
        # Treeview
        tree_frame = tk.Frame(content, bg='white')
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        scroll_y = ttk.Scrollbar(tree_frame, orient='vertical')
        scroll_x = ttk.Scrollbar(tree_frame, orient='horizontal')
        
        # Configure style for colored rows
        style = ttk.Style()
        style.configure("Treeview", rowheight=25)
        
        self.tree = ttk.Treeview(tree_frame,
                                columns=('BIL', 'STATUS', 'KOD_TARIF', 'DESKRIPSI', 'TARIKH'),
                                show='headings',
                                yscrollcommand=scroll_y.set,
                                xscrollcommand=scroll_x.set,
                                height=10)
        
        scroll_y.config(command=self.tree.yview)
        scroll_x.config(command=self.tree.xview)
        
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        # Column headers
        self.tree.heading('BIL', text='BIL')
        self.tree.heading('STATUS', text='STATUS')
        self.tree.heading('KOD_TARIF', text='KOD TARIF')
        self.tree.heading('DESKRIPSI', text='DESKRIPSI')
        self.tree.heading('TARIKH', text='TARIKH KUATKUASA')
        
        self.tree.column('BIL', width=60, anchor='center')
        self.tree.column('STATUS', width=120, anchor='center')
        self.tree.column('KOD_TARIF', width=150, anchor='w')
        self.tree.column('DESKRIPSI', width=400, anchor='w')
        self.tree.column('TARIKH', width=150, anchor='center')
        
        # Configure tags for colors
        self.tree.tag_configure('added', background='#C8E6C9')     # Light green
        self.tree.tag_configure('deleted', background='#FFCDD2')   # Light red
        self.tree.tag_configure('original', background='#F5F5F5')  # Light gray (original items)
    
    def create_button_section(self):
        """Create bottom buttons"""
        button_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        button_container.pack(fill=tk.X, padx=30, pady=(0, 25))
        
        button_frame = tk.Frame(button_container, bg=self.colors['bg_main'])
        button_frame.pack()
        
        btn_preview = tk.Button(button_frame, text="👁 PRATONTON",
                               font=('Arial', 18, 'bold'),
                               bg=self.colors['button_secondary'], fg='white',
                               relief=tk.FLAT, cursor='hand2',
                               width=18, height=2,
                               command=self.on_preview_click)
        btn_preview.pack(side=tk.LEFT, padx=10)
        
        btn_save = tk.Button(button_frame, text="💾 SIMPAN DOKUMEN",
                            font=('Arial', 18, 'bold'),
                            bg=self.colors['button_success'], fg='white',
                            relief=tk.FLAT, cursor='hand2',
                            width=22, height=2,
                            command=self.on_save_click)
        btn_save.pack(side=tk.LEFT, padx=10)
    
    def update_tarikh_islam(self):
        """Convert Gregorian to Hijri"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    
                    self.entry_islam.config(state='normal')
                    self.entry_islam.delete(0, tk.END)
                    self.entry_islam.insert(0, hijri_text)
                    self.entry_islam.config(state='readonly')
        except:
            pass
    
    def format_tarikh_malay(self):
        """Format date in Malay"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                
                bulan_melayu = ["Januari", "Februari", "Mac", "April", "Mei", "Jun",
                               "Julai", "Ogos", "September", "Oktober", "November", "Disember"]
                
                return f"{day:02d} {bulan_melayu[month-1]} {year}"
        except:
            return self.entry_tarikh.get()
    
    def search_ames(self):
        """Search for AMES application by number"""
        ames_no = self.entry_kelulusan.get().strip()
        
        if not ames_no:
            messagebox.showwarning("Amaran", "Sila masukkan No. Kelulusan AMES")
            return
        
        try:
            # Search in database
            results = self.db.search_applications(ames_no, form_type='ames')
            
            if not results:
                messagebox.showinfo("Tiada Rekod", 
                    f"Tiada rekod AMES dijumpai untuk:\n{ames_no}")
                return
            
            # If multiple results, show selection dialog
            if len(results) > 1:
                self.show_search_results(results)
            else:
                # Load the single result
                self.load_ames_data(results[0]['id'])
                
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat mencari AMES: {str(e)}")
    
    def show_ames_list(self):
        """Show list of all AMES applications"""
        AMESListDialog(self.root, self.db, self.load_ames_data)
    
    def show_search_results(self, results):
        """Show search results selection dialog"""
        SearchResultsDialog(self.root, results, self.load_ames_data)
    
    def load_ames_data(self, application_id):
        """Load AMES data from database"""
        try:
            # Get full application details
            app = self.db.get_application_by_id(application_id)
            
            if not app:
                messagebox.showerror("Ralat", "Rekod tidak dijumpai")
                return
            
            # Fill form fields
            if app.get('rujukan_kami'):
                # Extract just the number part after the last /
                parts = app['rujukan_kami'].split('/')
                if len(parts) > 0:
                    self.entry_rujukan.delete(0, tk.END)
                    self.entry_rujukan.insert(0, parts[-1])
            
            self.entry_nama.delete(0, tk.END)
            self.entry_nama.insert(0, app.get('nama_syarikat', ''))
            
            self.text_alamat.delete("1.0", tk.END)
            self.text_alamat.insert("1.0", app.get('alamat', ''))
            
            # Set pegawai if exists
            pegawai = app.get('nama_pegawai', '')
            if pegawai in self.combo_pegawai['values']:
                self.combo_pegawai.set(pegawai)
            
            # AMES specific data
            ames_details = app.get('ames_details', {})
            
            if ames_details:
                self.entry_kelulusan.delete(0, tk.END)
                self.entry_kelulusan.insert(0, ames_details.get('no_kelulusan', ''))
                
                kategori = ames_details.get('kategori', '')
                self.entry_kategori.config(state='normal')
                self.entry_kategori.delete(0, tk.END)
                self.entry_kategori.insert(0, kategori)
                self.entry_kategori.config(state='readonly')
            
            # Load existing items into table (as reference, in white)
            items = app.get('items', [])
            
            # Clear existing table
            for item in self.tree.get_children():
                self.tree.delete(item)
            
            # Add items (original items shown without color)
            for idx, item_data in enumerate(items, start=1):
                self.tree.insert('', 'end', values=(
                    idx,
                    "ASAL",  # Original item
                    item_data.get('kod_tarif', ''),
                    item_data.get('deskripsi', ''),
                    item_data.get('tarikh_kuatkuasa', '')
                ))
            
            messagebox.showinfo("Berjaya", 
                f"Data AMES berjaya dimuatkan!\n\n"
                f"No. Kelulusan: {ames_details.get('no_kelulusan', '')}\n"
                f"Syarikat: {app.get('nama_syarikat', '')}\n"
                f"Jumlah Item Asal: {len(items)}\n\n"
                f"Anda boleh menambah (HIJAU) atau memadamkan (MERAH) item.")
            
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat memuatkan data: {str(e)}")
    
    def add_item(self):
        """Add new item (GREEN)"""
        dialog = ItemDialog(self.root, "Tambah Item (HIJAU - Baru)")
        if dialog.result:
            kod, desk, tarikh = dialog.result
            bil = len(self.tree.get_children()) + 1
            self.tree.insert('', 'end', values=(bil, "TAMBAH", kod, desk, tarikh), tags=('added',))
    
    def delete_item(self):
        """Mark item for deletion (RED)"""
        dialog = ItemDialog(self.root, "Padam Item (MERAH - Dipadam)")
        if dialog.result:
            kod, desk, tarikh = dialog.result
            bil = len(self.tree.get_children()) + 1
            self.tree.insert('', 'end', values=(bil, "PADAM", kod, desk, tarikh), tags=('deleted',))
    
    def edit_item(self):
        """Edit selected item"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih item untuk diedit")
            return
        
        item = selected[0]
        values = self.tree.item(item)['values']
        tags = self.tree.item(item)['tags']
        
        dialog = ItemDialog(self.root, "Edit Item", 
                          initial_values=(values[2], values[3], values[4]))
        if dialog.result:
            kod, desk, tarikh = dialog.result
            self.tree.item(item, values=(values[0], values[1], kod, desk, tarikh), tags=tags)
    
    def remove_from_list(self):
        """Remove item from list completely"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih item untuk dibuang dari senarai")
            return
        
        if messagebox.askyesno("Pengesahan", 
                              "Adakah anda pasti ingin membuang item ini dari senarai?\n\n"
                              "Item tidak akan dipaparkan dalam dokumen."):
            for item in selected:
                self.tree.delete(item)
            
            # Renumber
            for idx, item in enumerate(self.tree.get_children(), start=1):
                values = list(self.tree.item(item)['values'])
                values[0] = idx
                tags = self.tree.item(item)['tags']
                self.tree.item(item, values=values, tags=tags)
    
    def generate_document(self):
        """Generate Word document"""
        template_file = "delete_item_ames.docx"
        
        # Try to use embedded template storage first
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(template_file)
            if doc is None:
                # Fallback to file system
                template_path = get_template_path(template_file)
                if not os.path.exists(template_path):
                    messagebox.showwarning("Amaran", 
                        f"Templat tidak dijumpai: {template_file}\n\nMenggunakan dokumen kosong...")
                    doc = Document()
                    doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
                else:
                    # Check if file is .doc (old format) - python-docx doesn't support it
                    if template_file.lower().endswith('.doc') and not template_file.lower().endswith('.docx'):
                        messagebox.showerror("Ralat", 
                            f"Templat {template_file} menggunakan format lama (.doc).\n\n"
                            f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                        return None
                    try:
                        doc = Document(template_path)
                    except Exception as e:
                        if template_file.lower().endswith('.doc') and not template_file.lower().endswith('.docx'):
                            messagebox.showerror("Ralat", 
                                f"Templat {template_file} menggunakan format lama (.doc) yang tidak disokong.\n\n"
                                f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                            return None
                        else:
                            messagebox.showwarning("Amaran",
                                f"Ralat memuat templat: {str(e)}\n\nMenggunakan dokumen kosong...")
                            doc = Document()
                            doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
        except ImportError:
            # Fallback to file system if template_storage not available
            template_path = get_template_path(template_file)
            if not os.path.exists(template_path):
                messagebox.showwarning("Amaran", 
                    f"Templat tidak dijumpai: {template_path}\n\nMenggunakan dokumen kosong...")
                doc = Document()
                doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
            else:
                # Check if file is .doc (old format) - python-docx doesn't support it
                if template_file.lower().endswith('.doc') and not template_file.lower().endswith('.docx'):
                    messagebox.showerror("Ralat", 
                        f"Templat {template_file} menggunakan format lama (.doc).\n\n"
                        f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                    return None
                try:
                    doc = Document(template_path)
                except Exception as e:
                    if template_file.lower().endswith('.doc') and not template_file.lower().endswith('.docx'):
                        messagebox.showerror("Ralat", 
                            f"Templat {template_file} menggunakan format lama (.doc) yang tidak disokong.\n\n"
                            f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                        return None
                    else:
                        messagebox.showwarning("Amaran",
                            f"Ralat memuat templat: {str(e)}\n\nMenggunakan dokumen kosong...")
                        doc = Document()
                        doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
        
        # Prepare replacements
        replacements = {
            '<<RUJUKAN_KAMI>>': f"KE.JB(90)650/14/AMES/{self.entry_rujukan.get()}",
            '<<NAMA_SYARIKAT>>': self.entry_nama.get().upper(),
            '<<ALAMAT>>': self.text_alamat.get("1.0", tk.END).strip(),
            '<<TARIKH>>': self.entry_tarikh.get(),
            '<<TARIKH_MALAY>>': self.format_tarikh_malay(),
            '<<TARIKH_ISLAM>>': self.entry_islam.get(),
            '<<NO_KELULUSAN>>': self.entry_kelulusan.get(),
            '<<KATEGORI>>': self.entry_kategori.get() if self.entry_kategori.get() else "N/A",
            '<<NAMA_PEGAWAI>>': self.combo_pegawai.get().upper(),
        }
        
        # Replace placeholders
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value) if value else "")
        
        # Add items table with colors
        self.add_items_table_to_document(doc)
        
        return doc
    
    def add_items_table_to_document(self, doc):
        """Add items table to document with color coding"""
        doc.add_page_break()
        doc.add_heading('LAMPIRAN - SENARAI ITEM', level=1)
        
        # Count added and deleted items
        added_items = []
        deleted_items = []
        
        for item in self.tree.get_children():
            values = self.tree.item(item)['values']
            status = values[1]
            if status == "TAMBAH":
                added_items.append(values)
            elif status == "PADAM":
                deleted_items.append(values)
        
        # Added items section
        if added_items:
            doc.add_heading('ITEM YANG DITAMBAH', level=2)
            para = doc.add_paragraph()
            run = para.add_run('(Ditandakan dengan warna hijau)')
            run.font.color.rgb = RGBColor(76, 175, 80)  # Green
            run.bold = True
            
            table = doc.add_table(rows=1 + len(added_items), cols=4)
            # Try to set table style, fallback if style doesn't exist
            try:
                table.style = 'Light Grid Accent 1'
            except (KeyError, ValueError):
                try:
                    table.style = 'Table Grid'
                except (KeyError, ValueError):
                    try:
                        table.style = 'Light List Accent 1'
                    except (KeyError, ValueError):
                        pass
            
            # Headers
            headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows with green background
            for idx, item_values in enumerate(added_items, start=1):
                row = table.rows[idx]
                # Safely get values with defaults
                row.cells[0].text = str(idx)
                row.cells[1].text = str(item_values[2]) if len(item_values) > 2 else ''
                row.cells[2].text = str(item_values[3]) if len(item_values) > 3 else ''
                row.cells[3].text = str(item_values[4]) if len(item_values) > 4 else ''
                
                # Center align BIL column
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Apply green shading and set font
                for cell in row.cells:
                    self.set_cell_background(cell, "C8E6C9")
                    # Set font for all paragraphs in cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
        
        # Deleted items section
        if deleted_items:
            doc.add_paragraph()
            doc.add_heading('ITEM YANG DIPADAM', level=2)
            para = doc.add_paragraph()
            run = para.add_run('(Ditandakan dengan warna merah)')
            run.font.color.rgb = RGBColor(244, 67, 54)  # Red
            run.bold = True
            
            table = doc.add_table(rows=1 + len(deleted_items), cols=4)
            # Try to set table style, fallback if style doesn't exist
            try:
                table.style = 'Light Grid Accent 1'
            except (KeyError, ValueError):
                try:
                    table.style = 'Table Grid'
                except (KeyError, ValueError):
                    try:
                        table.style = 'Light List Accent 1'
                    except (KeyError, ValueError):
                        pass
            
            # Headers
            headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()  # Clear any existing content
                run = para.add_run(header)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows with red background
            for idx, item_values in enumerate(deleted_items, start=1):
                row = table.rows[idx]
                # Safely get values with defaults
                row.cells[0].text = str(idx)
                row.cells[1].text = str(item_values[2]) if len(item_values) > 2 else ''
                row.cells[2].text = str(item_values[3]) if len(item_values) > 3 else ''
                row.cells[3].text = str(item_values[4]) if len(item_values) > 4 else ''
                
                # Center align BIL column
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Apply red shading and set font
                for cell in row.cells:
                    self.set_cell_background(cell, "FFCDD2")
                    # Set font for all paragraphs in cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
    
    def set_cell_background(self, cell, color_hex):
        """Set cell background color"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            pass
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview_delete_item.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':
                os.startfile(temp_path)
            elif os.name == 'posix':
                import subprocess
                subprocess.call(['open', temp_path])
            
            messagebox.showinfo("Pratonton", "Dokumen dibuka untuk pratonton")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF"""
        if len(self.tree.get_children()) == 0:
            messagebox.showwarning("Amaran", "Sila tambah sekurang-kurangnya satu item")
            return
        
        doc = self.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.get().replace('/', '_').replace('\\', '_')
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("Fail PDF", "*.pdf"), ("Semua fail", "*.*")],
            initialfile=f"DeleteItem_AMES_{safe_name}.pdf"
        )
        
        if not filename:
            return
        
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            temp_docx = "temp_delete_item_conversion.docx"
            doc.save(temp_docx)
            
            if PDF_AVAILABLE:
                convert(temp_docx, filename)
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
            else:
                filename = filename.replace('.pdf', '.docx')
                doc.save(filename)
                messagebox.showinfo("Maklumat", 
                    "Library docx2pdf tidak dijumpai. Dokumen disimpan sebagai Word (.docx)")
            
            # Save to database
            self.save_to_database(filename)
            
            messagebox.showinfo("Berjaya", 
                f"Dokumen berjaya disimpan:\n{filename}\n\nRekod telah disimpan ke pangkalan data.")
            
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menyimpan dokumen: {str(e)}")
            if os.path.exists("temp_delete_item_conversion.docx"):
                os.remove("temp_delete_item_conversion.docx")
    
    def save_to_database(self, document_path):
        """Save application to unified database"""
        try:
            # Basic application data
            application_data = {
                'category': 'AMES',
                'sub_option': 'Delete Item',
                'rujukan_kami': f"KE.JB(90)650/14/AMES/{self.entry_rujukan.get()}",
                'rujukan_tuan': None,
                'nama_syarikat': self.entry_nama.get(),
                'alamat': self.text_alamat.get("1.0", tk.END).strip(),
                'tarikh': self.entry_tarikh.get(),
                'tarikh_islam': self.entry_islam.get(),
                'nama_pegawai': self.combo_pegawai.get(),
                'status': 'DILULUSKAN',
                'document_path': document_path
            }
            
            # AMES-specific details
            specific_details = {
                'no_kelulusan': self.entry_kelulusan.get(),
                'kategori': self.entry_kategori.get(),
                'tarikh_mula': None,
                'tarikh_tamat': None,
                'tempoh_kelulusan': None,
                'items': []
            }
            
            # Get all items
            for item in self.tree.get_children():
                values = self.tree.item(item)['values']
                specific_details['items'].append({
                    'item_type': values[1].lower(),  # 'tambah' or 'padam'
                    'bil': values[0],
                    'kod_tarif': values[2],
                    'deskripsi': values[3],
                    'nisbah': None,
                    'tarikh_kuatkuasa': values[4]
                })
            
            # Save to database
            app_id = self.db.save_application('ames', application_data, specific_details)
            return app_id
            
        except Exception as e:
            messagebox.showerror("Ralat Database", f"Gagal menyimpan ke database: {str(e)}")
            return None
    
    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = 1600
        height = 1000
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def on_back_click(self):
        """Go back to main menu"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.root.destroy()
    
    def on_close(self):
        """Handle window close"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.root.destroy()


class AMESListDialog:
    """Dialog to show all AMES applications"""
    
    def __init__(self, parent, database, callback):
        self.db = database
        self.callback = callback
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Senarai AMES - Pilih Permohonan")
        self.dialog.geometry("1200x600")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.configure(bg='#F5F5F5')
        
        # Center dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - 600
        y = (self.dialog.winfo_screenheight() // 2) - 300
        self.dialog.geometry(f'1200x600+{x}+{y}')
        # Make window appear on top
        self.dialog.attributes('-topmost', True)
        self.dialog.lift()
        self.dialog.focus_force()
        
        self.create_content()
        self.load_ames_list()
    
    def create_content(self):
        """Create dialog content"""
        # Header
        header = tk.Frame(self.dialog, bg='#003366')
        header.pack(fill=tk.X)
        
        tk.Label(header, text="SENARAI PERMOHONAN AMES",
                font=('Arial', 13, 'bold'), bg='#003366', fg='white').pack(pady=15)
        
        # Search frame
        search_frame = tk.Frame(self.dialog, bg='white')
        search_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        tk.Label(search_frame, text="Cari:", font=('Arial', 18, 'bold'),
                bg='white').pack(side=tk.LEFT, padx=(10, 5))
        
        self.search_entry = tk.Entry(search_frame, width=40, font=('Arial', 18))
        self.search_entry.pack(side=tk.LEFT, padx=5)
        self.search_entry.bind('<KeyRelease>', lambda e: self.filter_list())
        
        tk.Button(search_frame, text="🔍", font=('Arial', 18),
                 bg='#1976D2', fg='white', relief=tk.FLAT,
                 cursor='hand2', width=5,
                 command=self.filter_list).pack(side=tk.LEFT, padx=5)
        
        # Table frame
        table_frame = tk.Frame(self.dialog, bg='white')
        table_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))
        
        scroll_y = ttk.Scrollbar(table_frame, orient='vertical')
        scroll_x = ttk.Scrollbar(table_frame, orient='horizontal')
        
        self.tree = ttk.Treeview(table_frame,
                                columns=('ID', 'NO_KELULUSAN', 'SYARIKAT', 
                                        'KATEGORI', 'TARIKH', 'ITEMS'),
                                show='headings',
                                yscrollcommand=scroll_y.set,
                                xscrollcommand=scroll_x.set)
        
        scroll_y.config(command=self.tree.yview)
        scroll_x.config(command=self.tree.xview)
        
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        # Columns
        self.tree.heading('ID', text='ID')
        self.tree.heading('NO_KELULUSAN', text='NO. KELULUSAN')
        self.tree.heading('SYARIKAT', text='NAMA SYARIKAT')
        self.tree.heading('KATEGORI', text='KATEGORI')
        self.tree.heading('TARIKH', text='TARIKH')
        self.tree.heading('ITEMS', text='JML ITEM')
        
        self.tree.column('ID', width=50, anchor='center')
        self.tree.column('NO_KELULUSAN', width=200, anchor='w')
        self.tree.column('SYARIKAT', width=300, anchor='w')
        self.tree.column('KATEGORI', width=100, anchor='center')
        self.tree.column('TARIKH', width=100, anchor='center')
        self.tree.column('ITEMS', width=80, anchor='center')
        
        # Double-click to select
        self.tree.bind('<Double-Button-1>', lambda e: self.on_select())
        
        # Buttons
        btn_frame = tk.Frame(self.dialog, bg='#F5F5F5')
        btn_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        tk.Button(btn_frame, text="Pilih", font=('Arial', 18, 'bold'),
                 bg='#2E7D32', fg='white', relief=tk.FLAT,
                 cursor='hand2', width=15, height=2,
                 command=self.on_select).pack(side=tk.LEFT, padx=5)
        
        tk.Button(btn_frame, text="Batal", font=('Arial', 18, 'bold'),
                 bg='#666666', fg='white', relief=tk.FLAT,
                 cursor='hand2', width=15, height=2,
                 command=self.dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def load_ames_list(self):
        """Load all AMES applications"""
        try:
            # Get all AMES applications
            apps = self.db.get_all_applications(form_type='ames', limit=500)
            
            self.all_apps = apps  # Store for filtering
            
            for app in apps:
                # Get AMES details
                full_app = self.db.get_application_by_id(app['id'])
                ames_details = full_app.get('ames_details', {}) if full_app else {}
                items = full_app.get('items', []) if full_app else []
                
                self.tree.insert('', 'end', values=(
                    app['id'],
                    ames_details.get('no_kelulusan', '-'),
                    app['nama_syarikat'],
                    ames_details.get('kategori', '-'),
                    app.get('tarikh', '-'),
                    len(items)
                ))
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat memuatkan senarai: {str(e)}")
    
    def filter_list(self):
        """Filter list based on search"""
        search_text = self.search_entry.get().strip().lower()
        
        # Clear tree
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Reload with filter
        if not search_text:
            self.load_ames_list()
            return
        
        for app in self.all_apps:
            # Search in multiple fields
            if (search_text in app.get('nama_syarikat', '').lower() or
                search_text in app.get('rujukan_kami', '').lower() or
                search_text in str(app.get('id', '')).lower()):
                
                full_app = self.db.get_application_by_id(app['id'])
                ames_details = full_app.get('ames_details', {}) if full_app else {}
                items = full_app.get('items', []) if full_app else []
                
                self.tree.insert('', 'end', values=(
                    app['id'],
                    ames_details.get('no_kelulusan', '-'),
                    app['nama_syarikat'],
                    ames_details.get('kategori', '-'),
                    app.get('tarikh', '-'),
                    len(items)
                ))
    
    def on_select(self):
        """Handle selection"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih permohonan AMES")
            return
        
        item = selected[0]
        values = self.tree.item(item)['values']
        app_id = values[0]
        
        # Call callback with app_id
        self.callback(app_id)
        self.dialog.destroy()


class SearchResultsDialog:
    """Dialog to show search results"""
    
    def __init__(self, parent, results, callback):
        self.results = results
        self.callback = callback
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Keputusan Carian - Pilih Permohonan")
        self.dialog.geometry("1000x400")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.configure(bg='#F5F5F5')
        
        # Center dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - 500
        y = (self.dialog.winfo_screenheight() // 2) - 200
        self.dialog.geometry(f'1000x400+{x}+{y}')
        # Make window appear on top
        self.dialog.attributes('-topmost', True)
        self.dialog.lift()
        self.dialog.focus_force()
        
        self.create_content()
    
    def create_content(self):
        """Create dialog content"""
        # Header
        header = tk.Frame(self.dialog, bg='#003366')
        header.pack(fill=tk.X)
        
        tk.Label(header, text=f"KEPUTUSAN CARIAN ({len(self.results)} rekod)",
                font=('Arial', 18, 'bold'), bg='#003366', fg='white').pack(pady=15)
        
        # Table
        table_frame = tk.Frame(self.dialog, bg='white')
        table_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(20, 10))
        
        scroll = ttk.Scrollbar(table_frame, orient='vertical')
        
        self.tree = ttk.Treeview(table_frame,
                                columns=('ID', 'RUJUKAN', 'SYARIKAT', 'TARIKH'),
                                show='headings',
                                yscrollcommand=scroll.set)
        
        scroll.config(command=self.tree.yview)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        # Columns
        self.tree.heading('ID', text='ID')
        self.tree.heading('RUJUKAN', text='RUJUKAN')
        self.tree.heading('SYARIKAT', text='NAMA SYARIKAT')
        self.tree.heading('TARIKH', text='TARIKH')
        
        self.tree.column('ID', width=50, anchor='center')
        self.tree.column('RUJUKAN', width=250, anchor='w')
        self.tree.column('SYARIKAT', width=400, anchor='w')
        self.tree.column('TARIKH', width=100, anchor='center')
        
        # Load results
        for result in self.results:
            self.tree.insert('', 'end', values=(
                result['id'],
                result.get('rujukan_kami', '-'),
                result.get('nama_syarikat', '-'),
                result.get('tarikh', '-')
            ))
        
        # Double-click to select
        self.tree.bind('<Double-Button-1>', lambda e: self.on_select())
        
        # Buttons
        btn_frame = tk.Frame(self.dialog, bg='#F5F5F5')
        btn_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        tk.Button(btn_frame, text="Pilih", font=('Arial', 18, 'bold'),
                 bg='#2E7D32', fg='white', relief=tk.FLAT,
                 cursor='hand2', width=15, height=2,
                 command=self.on_select).pack(side=tk.LEFT, padx=5)
        
        tk.Button(btn_frame, text="Batal", font=('Arial', 18, 'bold'),
                 bg='#666666', fg='white', relief=tk.FLAT,
                 cursor='hand2', width=15, height=2,
                 command=self.dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def on_select(self):
        """Handle selection"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih permohonan")
            return
        
        item = selected[0]
        values = self.tree.item(item)['values']
        app_id = values[0]
        
        self.callback(app_id)
        self.dialog.destroy()


class ItemDialog:
    """Dialog for adding/editing items"""
    
    def __init__(self, parent, title, initial_values=None):
        self.result = None
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        self.dialog.geometry("600x300")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.configure(bg='white')
        
        # Center dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - 300
        y = (self.dialog.winfo_screenheight() // 2) - 150
        self.dialog.geometry(f'600x300+{x}+{y}')
        # Make window appear on top
        self.dialog.attributes('-topmost', True)
        self.dialog.lift()
        self.dialog.focus_force()
        
        colors = {'primary': '#003366', 'success': '#2E7D32', 'border': '#CCCCCC'}
        
        # Title
        title_label = tk.Label(self.dialog, text=title,
                              font=('Arial', 18, 'bold'),
                              bg='white', fg=colors['primary'])
        title_label.pack(pady=(20, 20))
        
        # Form frame
        form_frame = tk.Frame(self.dialog, bg='white')
        form_frame.pack(padx=40, pady=10)
        
        # Kod Tarif
        tk.Label(form_frame, text="Kod Tarif:", font=('Arial', 18, 'bold'),
                bg='white').grid(row=0, column=0, padx=10, pady=12, sticky='e')
        self.entry_kod = tk.Entry(form_frame, width=45, font=('Arial', 18), relief=tk.SOLID, bd=1)
        self.entry_kod.grid(row=0, column=1, padx=10, pady=12)
        
        # Deskripsi
        tk.Label(form_frame, text="Deskripsi:", font=('Arial', 18, 'bold'),
                bg='white').grid(row=1, column=0, padx=10, pady=12, sticky='e')
        self.entry_desk = tk.Entry(form_frame, width=45, font=('Arial', 18), relief=tk.SOLID, bd=1)
        self.entry_desk.grid(row=1, column=1, padx=10, pady=12)
        
        # Tarikh Kuatkuasa
        tk.Label(form_frame, text="Tarikh Kuatkuasa:", font=('Arial', 18, 'bold'),
                bg='white').grid(row=2, column=0, padx=10, pady=12, sticky='e')
        self.entry_tarikh = tk.Entry(form_frame, width=45, font=('Arial', 18), relief=tk.SOLID, bd=1)
        self.entry_tarikh.insert(0, datetime.now().strftime("%d %B %Y").upper())
        self.entry_tarikh.grid(row=2, column=1, padx=10, pady=12)
        
        # Pre-fill if editing
        if initial_values:
            self.entry_kod.insert(0, initial_values[0])
            self.entry_desk.delete(0, tk.END)
            self.entry_desk.insert(0, initial_values[1])
            self.entry_tarikh.delete(0, tk.END)
            self.entry_tarikh.insert(0, initial_values[2])
        
        # Buttons
        btn_frame = tk.Frame(self.dialog, bg='white')
        btn_frame.pack(pady=25)
        
        tk.Button(btn_frame, text="Simpan", width=15, font=('Arial', 18, 'bold'),
                 bg=colors['success'], fg='white', relief=tk.FLAT, cursor='hand2',
                 command=self.on_ok).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Batal", width=15, font=('Arial', 18),
                 bg='#666666', fg='white', relief=tk.FLAT, cursor='hand2',
                 command=self.on_cancel).pack(side=tk.LEFT, padx=8)
        
        # Focus on first field
        self.entry_kod.focus_set()
        
        # Wait for dialog
        self.dialog.wait_window()
    
    def on_ok(self):
        """Save and close"""
        kod = self.entry_kod.get().strip()
        desk = self.entry_desk.get().strip()
        tarikh = self.entry_tarikh.get().strip()
        
        if not kod or not desk:
            messagebox.showwarning("Amaran", "Sila isi Kod Tarif dan Deskripsi")
            return
        
        self.result = (kod, desk, tarikh)
        self.dialog.destroy()
    
    def on_cancel(self):
        """Cancel and close"""
        self.dialog.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = FormDeleteItem(root)
    root.mainloop()