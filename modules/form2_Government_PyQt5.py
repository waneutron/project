"""
form2_Government_PyQt5.py - Form2 (PyQt5)
Optimized OOP structure with lazy loading to reduce app load
"""
import os
import sys
from datetime import datetime, timedelta

# PyQt5 imports
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QKeySequence
from PyQt5.QtWidgets import (QAbstractItemView, QButtonGroup, QCheckBox, QComboBox,
                                QDesktopWidget, QDialog, QFileDialog, QFrame, QGridLayout,
                                QHBoxLayout, QLabel, QLineEdit, QMessageBox, QPushButton,
                                QRadioButton, QScrollArea, QTableWidget, QTableWidgetItem,
                                QTextEdit, QVBoxLayout, QWidget)

# ============================================
# CUSTOM WIDGETS
# ============================================

class SafeLineEdit(QLineEdit):
    """QLineEdit that doesn't close dialog on Enter/Return key"""
    def keyPressEvent(self, event):
        # Allow Tab for moving to next field
        if event.key() == Qt.Key_Tab:
            self.focusNextChild()
            return
        # Allow Shift+Tab for moving to previous field
        elif event.key() == Qt.Key_Backtab:
            self.focusPreviousChild()
            return
        # Ignore Return/Enter - don't let it close the dialog
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            event.ignore()
            return
        # Handle everything else normally
        super().keyPressEvent(event)

# Lazy imports - only import when needed
_hijri_converter = None
_docx = None
_docx2pdf = None

def _lazy_import_hijri():
    """Lazy import hijri_converter"""
    global _hijri_converter
    if _hijri_converter is None:
        from hijri_converter import Hijri, Gregorian
        _hijri_converter = (Hijri, Gregorian)
    return _hijri_converter

def _lazy_import_docx():
    """Lazy import docx"""
    global _docx
    if _docx is None:
        from docx import Document
        _docx = Document
    return _docx

def _lazy_import_docx2pdf():
    """Lazy import docx2pdf"""
    global _docx2pdf
    if _docx2pdf is None:
        try:
            from docx2pdf import convert
            _docx2pdf = convert
        except ImportError:
            _docx2pdf = False
    return _docx2pdf

# Helpers
from helpers.resource_path import get_template_path
from helpers.unified_database import UnifiedDatabase

try:
    from helpers.docx_helper import replace_in_document
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value) if value else "")
        return {}


# ============================================
# HELPER CLASSES (Loaded only when needed)
# ============================================

class DateConverter:
    """Handles date conversions - Lazy loaded"""
    __slots__ = ('parent',)
    
    def __init__(self, parent_form):
        self.parent = parent_form
    
    def update_tarikh_islam(self):
        """Convert Gregorian to Hijri date"""
        try:
            Hijri, Gregorian = _lazy_import_hijri()
            tarikh_str = self.parent.entry_tarikh.text()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    self.parent.entry_islam.setText(hijri_text)
        except:
            pass
    
    def format_tarikh_malay(self):
        """Format date in Malay: 03 November 2025"""
        try:
            tarikh_str = self.parent.entry_tarikh.text()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                bulan_melayu = {
                    1: "Januari", 2: "Februari", 3: "Mac", 4: "April",
                    5: "Mei", 6: "Jun", 7: "Julai", 8: "Ogos",
                    9: "September", 10: "Oktober", 11: "November", 12: "Disember"
                }
                return f"{day:02d} {bulan_melayu[month]} {year}"
        except:
            return self.parent.entry_tarikh.text()
    
    def format_date_malay(self, date_str):
        """Format date string (DD/MM/YYYY) to Malay format"""
        try:
            if '/' in date_str:
                day, month, year = map(int, date_str.split('/'))
                return f"{day:02d}/{month:02d}/{year}"
        except:
            return date_str
    
    def calculate_tempoh_pemusnahan(self):
        """Calculate the period between start and end dates"""
        try:
            tarikh_mula = self.parent.entry_pemusnahan_mula.text()
            tarikh_tamat = self.parent.entry_pemusnahan_tamat.text()
            
            if '/' in tarikh_mula and '/' in tarikh_tamat:
                day1, month1, year1 = map(int, tarikh_mula.split('/'))
                day2, month2, year2 = map(int, tarikh_tamat.split('/'))
                
                date1 = datetime(year1, month1, day1)
                date2 = datetime(year2, month2, day2)
                delta = (date2 - date1).days
                
                tempoh_text = self.number_to_malay_words(delta)
                tempoh_formatted = f"{tempoh_text} ({delta}) hari"
                self.parent.entry_tempoh.setText(tempoh_formatted)
        except:
            pass
    
    def number_to_malay_words(self, num):
        """Convert number to Malay words"""
        ones = ["", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "lapan", "sembilan"]
        teens = ["sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas", "lima belas",
                "enam belas", "tujuh belas", "lapan belas", "sembilan belas"]
        tens = ["", "", "dua puluh", "tiga puluh", "empat puluh", "lima puluh",
               "enam puluh", "tujuh puluh", "lapan puluh", "sembilan puluh"]
        
        if num == 0:
            return "kosong"
        elif num < 10:
            return ones[num]
        elif num < 20:
            return teens[num - 10]
        elif num < 100:
            ten = num // 10
            one = num % 10
            if one == 0:
                return tens[ten]
            else:
                return f"{tens[ten]} {ones[one]}"
        elif num < 1000:
            hundred = num // 100
            remainder = num % 100
            if hundred == 1:
                result = "seratus"
            else:
                result = f"{ones[hundred]} ratus"
            if remainder > 0:
                result += f" {self.number_to_malay_words(remainder)}"
            return result
        else:
            return str(num)


class FormFieldManager:
    """Manages form field visibility and state"""
    __slots__ = ('parent',)
    
    def __init__(self, parent_form):
        self.parent = parent_form
    
    def update_field_visibility(self):
        """Show/hide fields based on requirements"""
        # Hide all optional fields first
        widgets_to_hide = [
            self.parent.label_jenis, self.parent.combo_jenis,
            self.parent.label_pengecualian, self.parent.combo_pengecualian,
            self.parent.label_amount, self.parent.entry_amount,
            self.parent.label_sst_adam, self.parent.sst_adam_frame
        ]
        
        for widget in widgets_to_hide:
            if widget:
                widget.hide()
        
        # Show SST/ADAM radio buttons ONLY for penjualan
        if self.parent.sub_option == "penjualan":
            if self.parent.label_sst_adam:
                self.parent.label_sst_adam.show()
            if self.parent.sst_adam_frame:
                self.parent.sst_adam_frame.show()
        
        # Show required fields
        if self.parent.requires_pengecualian:
            if self.parent.label_jenis:
                self.parent.label_jenis.show()
            if self.parent.combo_jenis:
                self.parent.combo_jenis.show()
            if self.parent.label_pengecualian:
                self.parent.label_pengecualian.show()
            if self.parent.combo_pengecualian:
                self.parent.combo_pengecualian.show()
        
        # For tidak_lulus: ALWAYS show amount field (as reason field)
        if self.parent.sub_option == "tidak_lulus":
            if self.parent.label_amount:
                self.parent.label_amount.show()
            if self.parent.entry_amount:
                self.parent.entry_amount.show()
        # For other categories: show amount field only if required
        elif self.parent.requires_amount:
            if self.parent.label_amount:
                self.parent.label_amount.show()
            if self.parent.entry_amount:
                self.parent.entry_amount.show()
        
        # Update pemusnahan visibility
        self.update_pemusnahan_visibility()
    
    def update_pemusnahan_visibility(self):
        """Show/hide pemusnahan section based on process selection"""
        if not hasattr(self.parent, 'combo_process'):
            return
        
        process = self.parent.combo_process.currentText().lower()
        
        pemusnahan_widgets = [
            self.parent.pemusnahan_divider,
            self.parent.check_pemusnahan,
            self.parent.pemusnahan_section_title,
            self.parent.label_pemusnahan_mula,
            self.parent.tempoh_frame,
            self.parent.label_tempoh,
            self.parent.entry_tempoh
        ]
        
        # IMPORTANT: tidak_lulus category should NEVER show pemusnahan details
        if self.parent.sub_option == "tidak_lulus":
            # Hide all pemusnahan widgets including checkbox
            for widget in pemusnahan_widgets:
                if widget:
                    widget.hide()
        elif process == "pemusnahan":
            # Show checkbox and divider, but actual fields depend on checkbox state
            if self.parent.pemusnahan_divider:
                self.parent.pemusnahan_divider.show()
            if self.parent.check_pemusnahan:
                self.parent.check_pemusnahan.show()
                # Check the checkbox by default when process is pemusnahan
                self.parent.check_pemusnahan.setChecked(True)
            # Update field visibility based on checkbox
            self.update_pemusnahan_fields_visibility()
        else:
            # Hide all pemusnahan widgets
            for widget in pemusnahan_widgets:
                if widget:
                    widget.hide()
            # Uncheck checkbox when process is not pemusnahan
            if self.parent.check_pemusnahan:
                self.parent.check_pemusnahan.setChecked(False)
    
    def update_pemusnahan_fields_visibility(self):
        """Show/hide pemusnahan fields based on checkbox state"""
        if not hasattr(self.parent, 'check_pemusnahan'):
            return
        
        is_checked = self.parent.check_pemusnahan.isChecked()
        
        pemusnahan_fields = [
            self.parent.pemusnahan_section_title,
            self.parent.label_pemusnahan_mula,
            self.parent.tempoh_frame,
            self.parent.label_tempoh,
            self.parent.entry_tempoh
        ]
        
        for widget in pemusnahan_fields:
            if widget:
                if is_checked:
                    widget.show()
                else:
                    widget.hide()


class ValidationManager:
    """Handles field validation"""
    __slots__ = ('parent', 'field_states')
    
    def __init__(self, parent_form):
        self.parent = parent_form
        self.field_states = {}
    
    def validate_field_realtime(self, field_name, value, field_widget):
        """Real-time validation with visual feedback"""
        is_valid = True
        
        if field_name == "rujukan" and value:
            if len(value) < 2:
                is_valid = False
        elif field_name == "nama_syarikat" and value:
            if len(value) < 3:
                is_valid = False
        elif field_name == "tarikh" and value:
            try:
                datetime.strptime(value, "%d/%m/%Y")
            except:
                is_valid = False
        elif field_name == "amount" and value:
            try:
                float(value.replace(',', ''))
            except:
                is_valid = False
        
        # Visual feedback
        if is_valid:
            field_widget.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
            self.field_states[field_name] = True
        else:
            field_widget.setStyleSheet("background-color: #FFEBEE; border: 1px solid #F44336; font-size: 14px;")
            self.field_states[field_name] = False
        
        return is_valid
    
    def validate_before_save(self):
        """Validate all required fields before saving"""
        errors = []
        
        if not self.parent.entry_nama.text().strip():
            errors.append("• Nama Syarikat")
            self.parent.entry_nama.setStyleSheet("background-color: #FFEBEE; border: 1px solid #F44336; font-size: 14px;")
        else:
            self.parent.entry_nama.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        
        if not self.parent.entry_rujukan.text().strip():
            errors.append("• Rujukan")
            self.parent.entry_rujukan.setStyleSheet("background-color: #FFEBEE; border: 1px solid #F44336; font-size: 14px;")
        else:
            self.parent.entry_rujukan.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        
        if not self.parent.entry_tarikh.text().strip():
            errors.append("• Tarikh")
            self.parent.entry_tarikh.setStyleSheet("background-color: #FFEBEE; border: 1px solid #F44336; font-size: 14px;")
        else:
            self.parent.entry_tarikh.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        
        if not self.parent.entry_alamat1.text().strip():
            errors.append("• Alamat")
            self.parent.entry_alamat1.setStyleSheet("background-color: #FFEBEE; border: 1px solid #F44336; font-size: 14px;")
        else:
            self.parent.entry_alamat1.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
            # Also set alamat2 and alamat3 to gray
            if hasattr(self.parent, 'entry_alamat2'):
                self.parent.entry_alamat2.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
            if hasattr(self.parent, 'entry_alamat3'):
                self.parent.entry_alamat3.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        
        return errors


class DocumentGenerator:
    """Handles document generation - Lazy loaded"""
    __slots__ = ('parent', '_template_cache')
    
    def __init__(self, parent_form):
        self.parent = parent_form
        self._template_cache = None
    
    def generate_document(self):
        """Generate Word document with replacements"""
        Document = _lazy_import_docx()
        
        # Try to use embedded template storage first
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(self.parent.template_file)
            if doc is None:
                # Fallback to file system
                template_path = get_template_path(self.parent.template_file)
                if not os.path.exists(template_path):
                    QMessageBox.critical(self.parent, "Ralat", 
                        f"Templat tidak dijumpai: {self.parent.template_file}")
                    return None
                doc = Document(template_path)
        except ImportError:
            # Fallback to file system
            template_path = get_template_path(self.parent.template_file)
            if not os.path.exists(template_path):
                QMessageBox.critical(self.parent, "Ralat", 
                    f"Templat tidak dijumpai: {template_path}")
                return None
            doc = Document(template_path)
        except Exception as e:
            QMessageBox.critical(self.parent, "Ralat", f"Ralat memuat templat: {str(e)}")
            return None
        
        try:
            # Use PlaceholderBuilder for standardized placeholders
            from helpers.placeholder_registry import PlaceholderBuilder
            
            # Combine the 3 alamat entries
            alamat_lines = []
            if self.parent.entry_alamat1.text().strip():
                alamat_lines.append(self.parent.entry_alamat1.text().strip())
            if self.parent.entry_alamat2.text().strip():
                alamat_lines.append(self.parent.entry_alamat2.text().strip())
            if self.parent.entry_alamat3.text().strip():
                alamat_lines.append(self.parent.entry_alamat3.text().strip())
            
            # Build replacements using PlaceholderBuilder
            builder = PlaceholderBuilder(self.parent)
            # Form2 uses RUJUKAN (not RUJUKAN_KAMI)
            rujukan_kami = self.parent.entry_rujukan.text()
            builder.add_rujukan(rujukan_kami, "KE.JB(90)650/05-02/")
            if self.parent.check_rujukan.isChecked():
                builder.add_standard('RUJUKAN_TUAN', self.parent.entry_rujukan_syarikat.text())
            builder.add_nama_syarikat(self.parent.entry_nama.text())
            builder.add_alamat(alamat_lines)
            builder.add_tarikh(self.parent.entry_tarikh.text())
            builder.add_tarikh_malay(self.parent.date_converter.format_tarikh_malay())
            builder.add_tarikh_islam(self.parent.entry_islam.text())
            builder.add_nama_pegawai(self.parent.combo_pegawai.currentText())
            builder.add_standard('AMOUNT', self.parent.entry_amount.text())
            
            # Add SST/ADAM type for penjualan
            if self.parent.sub_option == "penjualan":
                builder.add_standard('JENIS_TEMPLATE', self.parent.var_sst_adam.upper())
                builder.add_standard('STATUS', self.parent.var_sst_adam.upper())
            
            # Process-specific replacements
            if self.parent.combo_process.currentText():
                builder.add_standard('PROSES', self.parent.combo_process.currentText())
                
                # Pemusnahan-specific replacements
                if self.parent.combo_process.currentText().lower() == "pemusnahan":
                    builder.add_standard('TARIKH_MULA', self.parent.date_converter.format_date_malay(
                        self.parent.entry_pemusnahan_mula.text()))
                    builder.add_standard('TARIKH_TAMAT', self.parent.date_converter.format_date_malay(
                        self.parent.entry_pemusnahan_tamat.text()))
                    builder.add_standard('TEMPOH', self.parent.entry_tempoh.text())
            
            # Jenis Barang
            if self.parent.combo_jenis.currentText():
                builder.add_standard('JENIS_BARANG', self.parent.combo_jenis.currentText())
            
            # Pengecualian
            if self.parent.combo_pengecualian.currentText():
                pengecualian = self.parent.combo_pengecualian.currentText()
                builder.add_standard('PENGECUALIAN', pengecualian)
                
                # Tajuk surat
                jenis_barang = self.parent.combo_jenis.currentText() or ""
                proses = self.parent.combo_process.currentText() or ""
                
                tajuk_surat = f"pelupusan secara {proses} {jenis_barang} yang diperolehi dengan pengecualian di bawah {pengecualian}"
                builder.add_standard('TAJUK_SURAT', tajuk_surat.upper())
                builder.add_standard('TAJUK_SURAT2', tajuk_surat.lower())
            
            # Rujukan info
            if self.parent.check_rujukan.isChecked():
                builder.add_standard('RUJUKAN_INFO', f"rujukan {self.parent.entry_rujukan_syarikat.text()} ")
            else:
                builder.add_standard('RUJUKAN_INFO', "")
            
            # Get final replacements dict
            replacements = builder.build()
            
            # Replace in document
            replace_in_document(doc, replacements)
            
            # Make TAJUK_SURAT bold
            tajuk_uppercase = replacements.get('<<TAJUK_SURAT>>')
            if tajuk_uppercase:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if tajuk_uppercase in run.text:
                            run.bold = True
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if tajuk_uppercase in run.text:
                                        run.bold = True
            
            return doc
        except Exception as e:
            QMessageBox.critical(self.parent, "Ralat", f"Ralat menjana dokumen: {str(e)}")
            return None


class HelpDialog(QDialog):
    """Help dialog - Separate class for modularity"""
    __slots__ = ('parent_form',)
    
    def __init__(self, parent_form):
        super().__init__(parent_form)
        self.parent_form = parent_form
        self.setWindowTitle("📚 Panduan Pengisian Borang")
        self.setModal(True)
        self.setGeometry(100, 100, 700, 600)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Header
        header = QFrame()
        header.setStyleSheet("background-color: #003366;")
        header.setFixedHeight(60)
        header_layout = QVBoxLayout(header)
        title = QLabel("📚 Panduan Pengisian Borang Form 2")
        title.setStyleSheet("color: white; font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(title)
        layout.addWidget(header)
        btn_configure = QPushButton("⚙️ Template Settings")
        btn_configure.clicked.connect(self.open_template_scanner)

        def open_template_scanner(self):
            """Open template scanner dialog"""
            from helpers.template_completeness_dialog import TemplateScannerDialog
            dialog = TemplateScannerDialog(self, self.placeholder_mapper)
            dialog.exec_()
                
        # Scrollable content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(30, 20, 30, 20)
        
        help_sections = [
            ("📝 RUJUKAN", "Format: KE.JB(90)650/05-02/XXXX\nMasukkan nombor rujukan unik anda."),
            ("🏢 NAMA SYARIKAT", "Masukkan nama penuh syarikat dengan ejaan yang betul."),
            ("📅 TARIKH", "Format: DD/MM/YYYY (contoh: 17/12/2025)\nTarikh Islam akan dikira automatik."),
            ("📍 ALAMAT", "Masukkan alamat lengkap dalam 3 baris:\n- Baris 1: Nombor & nama jalan\n- Baris 2: Kawasan/Taman\n- Baris 3: Poskod & Bandar"),
            ("👤 NAMA PEGAWAI", "Pilih nama pegawai yang bertanggungjawab."),
            ("💰 AMAUN", "Masukkan jumlah dalam RM (contoh: 1234.56)\nHanya angka dan titik perpuluhan."),
            ("✅ PROSES & JENIS", "Pilih proses dan jenis barang yang berkaitan."),
        ]
        
        for title_text, content_text in help_sections:
            section_frame = QFrame()
            section_layout = QVBoxLayout(section_frame)
            
            title_label = QLabel(title_text)
            title_label.setStyleSheet("color: #003366; font-size: 15px; font-weight: bold;")
            section_layout.addWidget(title_label)
            
            content_label = QLabel(content_text)
            content_label.setStyleSheet("color: #333333; font-size: 12px;")
            content_label.setWordWrap(True)
            section_layout.addWidget(content_label)
            
            divider = QFrame()
            divider.setFrameShape(QFrame.HLine)
            divider.setStyleSheet("color: #EEEEEE;")
            section_layout.addWidget(divider)
            
            scroll_layout.addWidget(section_frame)
        
        scroll_area.setWidget(scroll_widget)
        layout.addWidget(scroll_area)
        
        # Close button
        btn_close = QPushButton("Tutup")
        btn_close.setStyleSheet("""
            QPushButton {
                background-color: #003366;
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 5px;
            }
        """)
        btn_close.clicked.connect(self.accept)
        layout.addWidget(btn_close)
        
        self.center_dialog()
    
    def center_dialog(self):
        """Center dialog on screen"""
        self.setGeometry(
            (QDesktopWidget().screenGeometry().width() - self.width()) // 2,
            (QDesktopWidget().screenGeometry().height() - self.height()) // 2,
            self.width(), self.height()
        )


# ============================================
# MAIN FORM CLASS (Optimized with lazy loading)
# ============================================

class Form2(QDialog):
    """Main Form2 - Optimized OOP structure with lazy loading"""
    
    def __init__(self, parent_window):
        super().__init__()
        self.parent_window = parent_window
        self.setWindowTitle("Sistem Pengurusan Dokumen - Pengisian Data")
        self.setGeometry(100, 100, 1600, 1000)
        self.setMinimumSize(1200, 800)  # Prevent window from shrinking too much
        self.category = None
        self.sub_option = None
        self.template_file = None
        self.requires_amount = False
        self.requires_pengecualian = False
        self.template_selected = False
        self.table_vehicles_data = []  # Store Butiran 5D vehicle table data
        from helpers.placeholder_mapper import PlaceholderMapper
        self.placeholder_mapper = PlaceholderMapper()
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': '#FFFFFF',
            'text_dark': '#1a1a1a',
            'text_light': '#FFFFFF',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_hover': '#004d99',
            'success': '#4CAF50',
            'warning': '#FF9800',
            'error': '#F44336',
            'info': '#2196F3'
        }
        
        # Create UI programmatically
        self.create_ui()
        
        # Apply global styles ONCE after UI creation (batch apply)
        self.apply_global_styles()
        
        self.db = UnifiedDatabase()
        
        # Initialize helper classes (lightweight, no heavy imports)
        self.date_converter = DateConverter(self)
        self.field_manager = FormFieldManager(self)
        self.validation_manager = ValidationManager(self)
        self.document_generator = DocumentGenerator(self)
        
        # SST/ADAM variable
        self.var_sst_adam = "sst"
        
        # Center window
        self.center_window()
        
        # Show template selector first
        self.show_template_selector()
    
    def apply_global_styles(self):
        """Apply global styles once to reduce redraws"""
        style_sheet = """
            QFrame {
                border: none;
            }
            QLineEdit, QComboBox, QTextEdit {
                background-color: #F5F5F5;
                border: 1px solid #CCCCCC;
                font-size: 14px;
                padding: 5px;
            }
            QLineEdit:focus, QComboBox:focus, QTextEdit:focus {
                border: 2px solid #006699;
                background-color: #FFFFFF;
            }
            QPushButton {
                font-weight: bold;
                border-radius: 5px;
                padding: 8px;
            }
            QLabel {
                background-color: transparent;
            }
        """
        self.setStyleSheet(style_sheet)
    
    def create_ui(self):
        """Create the main UI"""
        
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Header
        self.create_header(main_layout)
        
        # Scrollable content
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet(f"background-color: {self.colors['bg_main']}; border: none;")
        
        self.scroll_widget = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_widget)
        self.scroll_layout.setContentsMargins(40, 30, 40, 30)
        self.scroll_layout.setSpacing(20)
        
        self.scroll_area.setWidget(self.scroll_widget)
        main_layout.addWidget(self.scroll_area, stretch=1)
    
    def create_header(self, parent_layout):
        """Create government header"""
        header_frame = QFrame()
        header_frame.setFixedHeight(80)
        header_frame.setStyleSheet(f"background-color: {self.colors['primary']};")
        
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(15, 10, 15, 10)
        
        # Back button
        btn_back = QPushButton("← KEMBALI")
        btn_back.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['secondary']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: {self.colors['button_hover']};
            }}
        """)
        btn_back.clicked.connect(self.on_back_click)
        header_layout.addWidget(btn_back)
        
        # Title
        title_widget = QWidget()
        title_layout = QVBoxLayout(title_widget)
        title_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("Borang Pengisian Data Dokumen")
        title.setStyleSheet(f"background-color: {self.colors['primary']}; color: white; font-size: 20px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(title)
        
        self.subtitle = QLabel("Pilih template untuk mula")
        self.subtitle.setStyleSheet(f"background-color: {self.colors['primary']}; color: #E0E0E0; font-size: 14px;")
        self.subtitle.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(self.subtitle)
        
        header_layout.addWidget(title_widget, stretch=1)
        
        # Help button
        btn_help = QPushButton("❓ HELP")
        btn_help.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['accent']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
        """)
        btn_help.clicked.connect(self.show_help_dialog)
        header_layout.addWidget(btn_help)
        
        parent_layout.addWidget(header_frame)
        
        # Separator
        separator = QFrame()
        separator.setFixedHeight(2)
        separator.setStyleSheet(f"background-color: {self.colors['accent']};")
        parent_layout.addWidget(separator)
    
    def show_template_selector(self):
        """Show template selection UI"""
        # Ensure scroll_layout exists
        if not hasattr(self, 'scroll_layout'):
            if hasattr(self, 'scroll_area') and hasattr(self, 'scroll_widget'):
                self.scroll_widget = self.scroll_area.widget()
                if self.scroll_widget:
                    self.scroll_layout = self.scroll_widget.layout()
        
        if not hasattr(self, 'scroll_layout'):
            print("[Form2] Warning: scroll_layout not found in show_template_selector")
            return
        
        # Clear existing content
        while self.scroll_layout.count():
            child = self.scroll_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        
        # Template selector card
        selector_card = QFrame()
        selector_card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 10px;
            }
        """)
        selector_layout = QVBoxLayout(selector_card)
        selector_layout.setContentsMargins(40, 30, 40, 30)
        selector_layout.setSpacing(20)
        
        # Title
        title = QLabel("Pilih Template Dokumen")
        title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 20px; font-weight: bold;")
        selector_layout.addWidget(title)
        
        # Description
        desc = QLabel("Sila pilih template yang ingin anda gunakan:")
        desc.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px;")
        selector_layout.addWidget(desc)
        
        # Template combo
        template_label = QLabel("Template:")
        template_label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 15px; font-weight: bold;")
        selector_layout.addWidget(template_label)
        
        self.template_combo = QComboBox()
        self.template_combo.setStyleSheet("""
            QComboBox {
                font-size: 15px;
                padding: 8px;
                border: 1px solid #CCCCCC;
                border-radius: 5px;
                background-color: white;
            }
        """)
        self.load_available_templates()
        selector_layout.addWidget(self.template_combo)
        
        # Continue button
        btn_continue = QPushButton("Teruskan")
        btn_continue.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 12px 30px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #1B5E20;
            }}
        """)
        btn_continue.clicked.connect(self.on_template_selected)
        selector_layout.addWidget(btn_continue)
        
        self.scroll_layout.addWidget(selector_card)
        self.scroll_layout.addStretch()
    
    def load_available_templates(self):
        """Load available templates"""
        templates = [
            ("pelupusan_pemusnahan.docx", "Pelupusan - Pemusnahan"),
            ("pelupusan_penjualan.docx", "Pelupusan - Penjualan"),
            ("pelupusan_skrap.docx", "Pelupusan - Skrap"),
            ("pelupusan_tidak_lulus.docx", "Pelupusan - Tidak Lulus"),
            ("batal_sijil.doc", "Batal Sijil"),
            ("surat kelulusan butiran 5D (Lulus).docx", "Butiran 5D - Lulus"),
            ("surat kelulusan butiran 5D (tidak lulus).docx", "Butiran 5D - Tidak Lulus"),
        ]
        
        # Also try to get from template storage
        try:
            from helpers.template_storage import get_template_storage
            storage = get_template_storage()
            if hasattr(storage, 'templates'):
                for template_name in storage.templates.keys():
                    # Add if not already in list
                    if not any(t[0] == template_name for t in templates):
                        display_name = template_name.replace('.docx', '').replace('.doc', '').replace('_', ' ').title()
                        templates.append((template_name, display_name))
        except:
            pass
        
        # Add to combo
        for template_file, display_name in templates:
            self.template_combo.addItem(display_name, template_file)
    
    def on_template_selected(self):
        """Handle template selection with auto-check"""
        if self.template_combo.currentIndex() < 0:
            QMessageBox.warning(self, "Amaran", "Sila pilih template terlebih dahulu.")
            return
        
        self.template_file = self.template_combo.currentData()
        if not self.template_file:
            QMessageBox.warning(self, "Amaran", "Template tidak valid.")
            return
        
        # Check if template is configured
        if not self.placeholder_mapper.is_template_configured(self.template_file):
            from helpers.resource_path import get_template_path
            from helpers.template_completeness_dialog import PlaceholderMappingDialog
            
            template_path = get_template_path(self.template_file)
            
            reply = QMessageBox.question(self, "Configure Template",
                f"Template '{self.template_file}' needs one-time configuration.\n\n"
                f"Configure placeholder mappings now?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                dialog = PlaceholderMappingDialog(
                    self, self.template_file, template_path, self.placeholder_mapper
                )
                if not dialog.exec_():
                    return  # User cancelled
            else:
                # Mark as configured (empty mapping)
                self.placeholder_mapper.set_template_mapping(self.template_file, {})
        
        # Continue normal flow
        self.determine_template_properties()
        self.template_selected = True
        self.create_form_content()
        
        QTimer.singleShot(100, lambda: (
            self.date_converter.update_tarikh_islam(),
            self.load_options(),
            self.field_manager.update_field_visibility()
        ))
    
    def determine_template_properties(self):
        """Determine category, sub_option, and requirements from template file"""
        template_lower = self.template_file.lower()
        
        # Determine category and sub_option
        if "pelupusan" in template_lower:
            self.category = "Pelupusan"
            if "pemusnahan" in template_lower:
                self.sub_option = "pemusnahan"
            elif "penjualan" in template_lower:
                self.sub_option = "penjualan"
            elif "skrap" in template_lower:
                self.sub_option = "skrap"
            elif "tidak_lulus" in template_lower or "tidak lulus" in template_lower:
                self.sub_option = "tidak_lulus"
        elif "batal" in template_lower:
            self.category = "Lain-lain"
            self.sub_option = "batal_sijil"
        elif "butiran" in template_lower or "5d" in template_lower:
            self.category = "Lain-lain"
            if "lulus" in template_lower and "tidak" not in template_lower:
                self.sub_option = "butiran_5d_lulus"
            else:
                self.sub_option = "butiran_5d_tidak_lulus"
        else:
            self.category = "Lain-lain"
            self.sub_option = "lain_lain"
        
        # Determine requirements
        self.requires_amount = (self.category == "Pelupusan" and self.sub_option == "penjualan")
        self.requires_pengecualian = (self.category == "Pelupusan")
        
        # Update header subtitle
        if hasattr(self, 'subtitle'):
            self.subtitle.setText(f"Kategori: {self.category} | Sub-Kategori: {self.sub_option}")
    
    def create_form_content(self):
        """Create form content after template selection"""
        # Clear existing content
        while self.scroll_layout.count():
            child = self.scroll_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        
        # Create form
        self.create_form(self.scroll_layout)
    
    def create_form(self, parent_layout):
        """Create form fields"""
        # Card frame
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(30, 30, 30, 30)
        card_layout.setSpacing(15)
        
        # Section title
        section_title = QLabel("Maklumat Dokumen")
        section_title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 17px; font-weight: bold;")
        card_layout.addWidget(section_title)
        
        # Form grid
        form_grid = QGridLayout()
        form_grid.setSpacing(0)  # No spacing for tightest layout
        form_grid.setColumnMinimumWidth(0, 130)  # Label column width
        form_grid.setColumnMinimumWidth(2, 85)  # Right label column width (very minimal for ALAMAT)
        form_grid.setHorizontalSpacing(2)  # Minimal horizontal spacing (same as NAMA PEGAWAI)
        form_grid.setVerticalSpacing(8)  # Vertical spacing between rows
        form_grid.setColumnStretch(2, 0)  # Don't stretch column 2, keep it minimal
        
        row = 0
        
        # Helper function for consistent labels
        def create_label(text):
            label = QLabel(text)
            label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px; font-weight: bold; padding: 0px; margin: 0px;")
            label.setContentsMargins(0, 0, 0, 0)  # Remove any margins
            return label
        
        # Process
        form_grid.addWidget(create_label("PROSES:"), row, 0)
        self.combo_process = QComboBox()
        self.combo_process.setEditable(False)
        self.combo_process.currentTextChanged.connect(self.on_process_change)
        form_grid.addWidget(self.combo_process, row, 1)
        row += 1
        
        # SST/ADAM Radio Buttons (only for penjualan)
        self.label_sst_adam = create_label("JENIS:")
        form_grid.addWidget(self.label_sst_adam, row, 0)
        
        self.sst_adam_frame = QWidget()
        sst_adam_layout = QHBoxLayout(self.sst_adam_frame)
        sst_adam_layout.setContentsMargins(0, 0, 0, 0)
        
        self.radio_sst = QRadioButton("SST-ADM")
        self.radio_sst.setChecked(True)
        self.radio_sst.toggled.connect(lambda: self.on_sst_adam_change())
        sst_adam_layout.addWidget(self.radio_sst)
        
        self.radio_adam = QRadioButton("AMES-03")
        self.radio_adam.toggled.connect(lambda: self.on_sst_adam_change())
        sst_adam_layout.addWidget(self.radio_adam)
        
        form_grid.addWidget(self.sst_adam_frame, row, 1, 1, 3)
        row += 1
        
        # Rujukan
        form_grid.addWidget(create_label("RUJUKAN:"), row, 0)
        rujukan_widget = QWidget()
        rujukan_layout = QHBoxLayout(rujukan_widget)
        rujukan_layout.setContentsMargins(0, 0, 0, 0)
        rujukan_layout.setSpacing(5)
        prefix_label = QLabel("KE.JB(90)650/05-02/")
        prefix_label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px;")
        rujukan_layout.addWidget(prefix_label)
        self.entry_rujukan = SafeLineEdit()
        self.entry_rujukan.setMinimumWidth(200)
        self.entry_rujukan.textChanged.connect(
            lambda: self.validation_manager.validate_field_realtime('rujukan', 
                self.entry_rujukan.text(), self.entry_rujukan))
        rujukan_layout.addWidget(self.entry_rujukan)
        form_grid.addWidget(rujukan_widget, row, 1)
        
        # ALAMAT label - use exact same settings as NAMA PEGAWAI
        alamat_label = create_label("ALAMAT:")
        alamat_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)  # Same alignment as NAMA PEGAWAI
        alamat_label.setMaximumWidth(360)  # Same width as NAMA PEGAWAI
        form_grid.addWidget(alamat_label, row, 2)
        self.entry_alamat1 = SafeLineEdit()
        self.entry_alamat1.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat1, row, 3)
        row += 1
        
        # Nama Syarikat
        form_grid.addWidget(create_label("NAMA SYARIKAT:"), row, 0)
        self.entry_nama = SafeLineEdit()
        self.entry_nama.textChanged.connect(
            lambda: self.validation_manager.validate_field_realtime('nama_syarikat',
                self.entry_nama.text(), self.entry_nama))
        form_grid.addWidget(self.entry_nama, row, 1)
        
        self.entry_alamat2 = SafeLineEdit()
        self.entry_alamat2.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat2, row, 3)
        row += 1
        
        # Tarikh
        form_grid.addWidget(create_label("TARIKH:"), row, 0)
        self.entry_tarikh = SafeLineEdit()
        self.entry_tarikh.setText(datetime.now().strftime("%d/%m/%Y"))
        # Use lambda to avoid weak reference issue with __slots__
        self.entry_tarikh.textChanged.connect(lambda: self.date_converter.update_tarikh_islam())
        form_grid.addWidget(self.entry_tarikh, row, 1)
        
        self.entry_alamat3 = SafeLineEdit()
        self.entry_alamat3.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat3, row, 3)
        row += 1
        
        # Tarikh Islam & Nama Pegawai
        form_grid.addWidget(create_label("TARIKH ISLAM:"), row, 0)
        self.entry_islam = SafeLineEdit()
        self.entry_islam.setReadOnly(True)
        self.entry_islam.setStyleSheet("background-color: #F5F5F5;")
        form_grid.addWidget(self.entry_islam, row, 1)
        
        pegawai_label = create_label("NAMA PEGAWAI:")
        pegawai_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)  # Right align (same as ALAMAT)
        pegawai_label.setMaximumWidth(360)  # Same width as ALAMAT label
        form_grid.addWidget(pegawai_label, row, 2)
        self.combo_pegawai = QComboBox()
        self.combo_pegawai.setEditable(True)  # Allow typing
        self.combo_pegawai.addItems(["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                    "KHAIRONE DZARWANY BINTI MOHD KAIRI"])
        form_grid.addWidget(self.combo_pegawai, row, 3)
        row += 1
        
        # Divider
        divider = QFrame()
        divider.setFrameShape(QFrame.HLine)
        divider.setStyleSheet(f"background-color: {self.colors['border']}; max-height: 1px;")
        divider.setFixedHeight(1)
        form_grid.addWidget(divider, row, 0, 1, 4)
        form_grid.setRowMinimumHeight(row, 15)  # Add spacing after divider
        row += 1
        
        # Jenis Barang
        self.label_jenis = create_label("JENIS BARANG:")
        form_grid.addWidget(self.label_jenis, row, 0)
        self.combo_jenis = QComboBox()
        self.combo_jenis.setEditable(False)
        form_grid.addWidget(self.combo_jenis, row, 1)
        row += 1
        
        # Pengecualian
        self.label_pengecualian = create_label("PENGECUALIAN:")
        form_grid.addWidget(self.label_pengecualian, row, 0)
        self.combo_pengecualian = QComboBox()
        self.combo_pengecualian.setEditable(False)
        form_grid.addWidget(self.combo_pengecualian, row, 1, 1, 3)
        row += 1
        
        # Amount / Reason field
        if self.sub_option == "tidak_lulus":
            self.label_amount = create_label("SEBAB TIDAK DILULUSKAN:")
        else:
            self.label_amount = create_label("AMAUN (RM):")
        form_grid.addWidget(self.label_amount, row, 0)
        self.entry_amount = SafeLineEdit()
        if self.sub_option != "tidak_lulus":
            self.entry_amount.textChanged.connect(
                lambda: self.validation_manager.validate_field_realtime('amount',
                    self.entry_amount.text(), self.entry_amount))
        form_grid.addWidget(self.entry_amount, row, 1)
        row += 1
        
        # Rujukan Syarikat checkbox
        form_grid.addWidget(QLabel(""), row, 0)  # Spacer
        self.check_rujukan = QCheckBox("Rujukan Syarikat")
        self.check_rujukan.toggled.connect(self.toggle_rujukan_syarikat)
        form_grid.addWidget(self.check_rujukan, row, 1)
        
        self.entry_rujukan_syarikat = SafeLineEdit()
        self.entry_rujukan_syarikat.setEnabled(False)
        form_grid.addWidget(self.entry_rujukan_syarikat, row, 2)
        row += 1
        
        # ===== PEMUSNAHAN SECTION =====
        form_grid.setRowMinimumHeight(row, 20)  # Add spacing before section
        self.pemusnahan_divider = QFrame()
        self.pemusnahan_divider.setFrameShape(QFrame.HLine)
        self.pemusnahan_divider.setStyleSheet(f"background-color: {self.colors['border']}; max-height: 1px;")
        self.pemusnahan_divider.setFixedHeight(1)
        form_grid.addWidget(self.pemusnahan_divider, row, 0, 1, 4)
        form_grid.setRowMinimumHeight(row, 15)  # Add spacing after divider
        row += 1
        
        # Checkbox to enable/disable pemusnahan section
        self.check_pemusnahan = QCheckBox("Tunjukkan Maklumat Pemusnahan")
        self.check_pemusnahan.setStyleSheet(f"""
            QCheckBox {{
                color: {self.colors['text_dark']};
                font-size: 15px;
                font-weight: bold;
            }}
            QCheckBox::indicator {{
                width: 20px;
                height: 20px;
            }}
        """)
        self.check_pemusnahan.stateChanged.connect(self.on_pemusnahan_checkbox_changed)
        form_grid.addWidget(self.check_pemusnahan, row, 0, 1, 4)
        form_grid.setRowMinimumHeight(row, 25)  # Add spacing after checkbox
        row += 1
        
        self.pemusnahan_section_title = QLabel("MAKLUMAT PEMUSNAHAN")
        self.pemusnahan_section_title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 17px; font-weight: bold;")
        form_grid.addWidget(self.pemusnahan_section_title, row, 0, 1, 4)
        form_grid.setRowMinimumHeight(row, 25)  # Add spacing after title
        row += 1
        
        # Tempoh Pemusnahan
        self.label_pemusnahan_mula = create_label("TEMPOH PEMUSNAHAN:")
        form_grid.addWidget(self.label_pemusnahan_mula, row, 0)
        
        self.tempoh_frame = QWidget()
        tempoh_layout = QHBoxLayout(self.tempoh_frame)
        tempoh_layout.setContentsMargins(0, 0, 0, 0)
        tempoh_layout.setSpacing(8)
        
        mulai_label = QLabel("Mulai:")
        mulai_label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px;")
        tempoh_layout.addWidget(mulai_label)
        self.entry_pemusnahan_mula = SafeLineEdit()
        today = datetime.now()
        self.entry_pemusnahan_mula.setText(today.strftime("%d/%m/%Y"))
        # Use lambda to avoid weak reference issue with __slots__
        self.entry_pemusnahan_mula.textChanged.connect(lambda: self.date_converter.calculate_tempoh_pemusnahan())
        tempoh_layout.addWidget(self.entry_pemusnahan_mula)
        
        hingga_label = QLabel("hingga:")
        hingga_label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px;")
        tempoh_layout.addWidget(hingga_label)
        self.entry_pemusnahan_tamat = SafeLineEdit()
        end_date = today + timedelta(days=30)
        self.entry_pemusnahan_tamat.setText(end_date.strftime("%d/%m/%Y"))
        # Use lambda to avoid weak reference issue with __slots__
        self.entry_pemusnahan_tamat.textChanged.connect(lambda: self.date_converter.calculate_tempoh_pemusnahan())
        tempoh_layout.addWidget(self.entry_pemusnahan_tamat)
        
        form_grid.addWidget(self.tempoh_frame, row, 1, 1, 3)
        row += 1
        
        # Tempoh (Period)
        self.label_tempoh = create_label("TEMPOH:")
        form_grid.addWidget(self.label_tempoh, row, 0)
        self.entry_tempoh = SafeLineEdit()
        self.entry_tempoh.setReadOnly(True)
        self.entry_tempoh.setStyleSheet("background-color: #F5F5F5;")
        form_grid.addWidget(self.entry_tempoh, row, 1)
        # Calculate initial tempoh
        self.date_converter.calculate_tempoh_pemusnahan()
        row += 1
        
        card_layout.addLayout(form_grid)
        
        # ===== BUTIRAN 5D TABLE SECTION =====
        if self.sub_option and ("butiran_5d" in self.sub_option or "5d" in self.sub_option.lower()):
            self.create_butiran5d_table_section(card_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.setAlignment(Qt.AlignCenter)
        button_layout.setSpacing(10)
        
        btn_preview = QPushButton("👁 PRATONTON")
        btn_preview.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_secondary']};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
        """)
        btn_preview.clicked.connect(self.on_preview_click)
        button_layout.addWidget(btn_preview)
        
        btn_save = QPushButton("💾 SIMPAN DOKUMEN")
        btn_save.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
        """)
        btn_save.clicked.connect(self.on_save_click)
        button_layout.addWidget(btn_save)
        
        card_layout.addLayout(button_layout)
        parent_layout.addWidget(card)
    
    
    def load_options(self):
        """Load dropdown options based on category"""
        # Process options
        process_options = {
            "pemusnahan": ["pemusnahan"],
            "penjualan": ["penjualan", "pemberian"],
            "skrap": ["pemusnahan"],
            "tidak_lulus": ["pemusnahan", "penjualan", "pemberian"]
        }
        
        if self.sub_option in process_options:
            self.combo_process.clear()
            self.combo_process.addItems(process_options[self.sub_option])
            if self.combo_process.count() > 0:
                self.combo_process.setCurrentIndex(0)
        
        # Jenis Barang
        self.combo_jenis.clear()
        self.combo_jenis.addItems(["barang siap", "bahan mentah", "komponen", "bahan pembungkus",
                                   "bahan pembungkusan", "skrap", "sisa dan hampas", "apa-apa barang bercukai"])
        if self.combo_jenis.count() > 0:
            self.combo_jenis.setCurrentIndex(0)
        
        # Pengecualian options
        self.combo_pengecualian.clear()
        self.combo_pengecualian.addItems([
            "perintah cukai jualan (orang yang dikecualikan daripada pembayaran cukai) 2018",
            "skim pengeksport utama diluluskan (ames)",
            "kemudahan pemotongan cukai jualan",
            "subseksyen 35(3) akta cukai jualan 2018"
        ])
        if self.combo_pengecualian.count() > 0:
            self.combo_pengecualian.setCurrentIndex(0)
    
    def on_pemusnahan_checkbox_changed(self, state):
        """Handle pemusnahan checkbox state change"""
        self.field_manager.update_pemusnahan_fields_visibility()
    
    def on_process_change(self):
        """Handle process selection change"""
        self.field_manager.update_pemusnahan_visibility()
    
    def on_sst_adam_change(self):
        """Handle SST/ADAM radio button change"""
        if self.radio_sst.isChecked():
            self.var_sst_adam = "sst"
        else:
            self.var_sst_adam = "adam"
    
    def toggle_rujukan_syarikat(self, checked):
        """Enable/disable rujukan syarikat field"""
        self.entry_rujukan_syarikat.setEnabled(checked)
        if not checked:
            self.entry_rujukan_syarikat.clear()
    
    def show_help_dialog(self):
        """Show comprehensive help dialog"""
        dialog = HelpDialog(self)
        dialog.exec_()
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.document_generator.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':
                os.startfile(temp_path)
            elif os.name == 'posix':
                import subprocess
                subprocess.call(['open', temp_path])
            
            QMessageBox.information(self, "Pratonton", "Dokumen dibuka untuk pratonton.\nSemak temp_preview.docx")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        # Validate fields
        validation_errors = self.validation_manager.validate_before_save()
        if validation_errors:
            QMessageBox.warning(self, "⚠️ Field Tidak Lengkap",
                "Sila lengkapkan maklumat berikut:\n\n" + "\n".join(validation_errors))
            return
        
        # Confirmation dialog
        reply = QMessageBox.question(self, "✅ Sahkan Simpan Dokumen",
            f"Adakah anda pasti untuk simpan dokumen?\n\n"
            f"📋 Rujukan: KE.JB(90)650/05-02/{self.entry_rujukan.text()}\n"
            f"🏢 Syarikat: {self.entry_nama.text()}\n"
            f"📅 Tarikh: {self.entry_tarikh.text()}\n"
            f"📁 Kategori: {self.category} - {self.sub_option}",
            QMessageBox.Yes | QMessageBox.No)
        
        if reply != QMessageBox.Yes:
            return
        
        doc = self.document_generator.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.text().replace('/', '_').replace('\\', '_')
        
        # Add SST/ADAM to filename for penjualan
        if self.sub_option == "penjualan":
            jenis = self.var_sst_adam.upper()
            default_filename = f"{safe_name}_{self.sub_option}_{jenis}.pdf"
        else:
            default_filename = f"{safe_name}_{self.sub_option}.pdf"
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Simpan Dokumen", default_filename,
            "PDF Files (*.pdf);;All Files (*.*)")
        
        if not filename:
            return
        
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            temp_docx = "temp_conversion.docx"
            doc.save(temp_docx)
            
            # Lazy load docx2pdf only when saving
            convert = _lazy_import_docx2pdf()
            if convert:
                convert(temp_docx, filename)
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
            else:
                filename = filename.replace('.pdf', '.docx')
                doc.save(filename)
                QMessageBox.information(self, "Maklumat", 
                    "Library docx2pdf tidak dijumpai. Dokumen disimpan sebagai Word (.docx)")
            
            # Save to database
            try:
                rujukan_kami = f"KE.JB(90)650/05-02/{self.entry_rujukan.text()}"
                rujukan_tuan = ""
                if self.check_rujukan.isChecked():
                    rujukan_tuan = self.entry_rujukan_syarikat.text()
                
                alamat_lines = []
                if self.entry_alamat1.text().strip():
                    alamat_lines.append(self.entry_alamat1.text().strip())
                if self.entry_alamat2.text().strip():
                    alamat_lines.append(self.entry_alamat2.text().strip())
                if self.entry_alamat3.text().strip():
                    alamat_lines.append(self.entry_alamat3.text().strip())
                alamat_combined = "\n".join(alamat_lines)
                
                application_data = {
                    'category': self.category,
                    'sub_option': self.sub_option,
                    'rujukan_kami': rujukan_kami,
                    'rujukan_tuan': rujukan_tuan,
                    'nama_syarikat': self.entry_nama.text(),
                    'alamat': alamat_combined,
                    'tarikh': self.entry_tarikh.text(),
                    'tarikh_islam': self.entry_islam.text(),
                    'nama_pegawai': self.combo_pegawai.currentText(),
                    'status': 'DILULUSKAN',
                    'document_path': filename,
                    'additional_data': {
                        'sst_adam_type': self.var_sst_adam if self.sub_option == "penjualan" else None
                    }
                }
                
                pelupusan_details = {
                    'proses': self.combo_process.currentText() or None,
                    'jenis_barang': self.combo_jenis.currentText() or None,
                    'pengecualian': self.combo_pengecualian.currentText() or None,
                    'amount': self.entry_amount.text() or None,
                    'tarikh_mula': self.entry_pemusnahan_mula.text() if hasattr(self, 'entry_pemusnahan_mula') else None,
                    'tarikh_tamat': self.entry_pemusnahan_tamat.text() if hasattr(self, 'entry_pemusnahan_tamat') else None,
                    'tempoh': self.entry_tempoh.text() if hasattr(self, 'entry_tempoh') else None
                }
                
                app_id = self.db.save_application('pelupusan', application_data, pelupusan_details)
                
                # Clear database cache so history viewer shows new data
                if hasattr(self.db, 'clear_cache'):
                    self.db.clear_cache()
                
                QMessageBox.information(self, "Berjaya", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
            except Exception as db_error:
                QMessageBox.warning(self, "Amaran", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat menyimpan dokumen: {str(e)}")
    
    def center_window(self):
        """Center window on screen"""
        self.setGeometry(
            (QDesktopWidget().screenGeometry().width() - self.width()) // 2,
            (QDesktopWidget().screenGeometry().height() - self.height()) // 2,
            self.width(), self.height()
        )
    
    def on_back_click(self):
        """Go back to Form1 or main menu"""
        if self.parent_window:
            self.parent_window.show()
        self.reject()

