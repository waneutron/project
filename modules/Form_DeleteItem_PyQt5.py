"""
Form_DeleteItem_PyQt5.py - Delete Item Form (PyQt5)
Optimized OOP structure with lazy loading to reduce app load
"""
import os
import sys
from datetime import datetime

# PyQt5 imports
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QColor, QPixmap
from PyQt5.QtWidgets import (QAbstractItemView, QComboBox, QDesktopWidget, QDialog,
                                QFileDialog, QFrame, QGridLayout, QHBoxLayout, QHeaderView,
                                QLabel, QLineEdit, QMessageBox, QPushButton, QScrollArea,
                                QTableWidget, QTableWidgetItem, QTextEdit, QVBoxLayout, QWidget)

# ============================================
# CUSTOM WIDGETS
# ============================================

class SafeLineEdit(QLineEdit):
    """QLineEdit that doesn't close dialog on Enter/Return key"""
    def keyPressEvent(self, event):
        # Allow Tab for moving to next field
        if event.key() == Qt.Key_Tab:
            self.focusNextChild()
            return
        # Allow Shift+Tab for moving to previous field
        elif event.key() == Qt.Key_Backtab:
            self.focusPreviousChild()
            return
        # Ignore Return/Enter - don't let it close the dialog
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            event.ignore()
            return
        # Handle everything else normally
        super().keyPressEvent(event)

# Lazy imports - only import when needed
# Heavy imports moved to methods that use them
_hijri_converter = None
_docx = None
_PIL = None
_docx2pdf = None

def _lazy_import_hijri():
    """Lazy import hijri_converter"""
    global _hijri_converter
    if _hijri_converter is None:
        from hijri_converter import Hijri, Gregorian
        _hijri_converter = (Hijri, Gregorian)
    return _hijri_converter

def _lazy_import_docx():
    """Lazy import docx"""
    global _docx
    if _docx is None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        _docx = (Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement)
    return _docx

def _lazy_import_pil():
    """Lazy import PIL"""
    global _PIL
    if _PIL is None:
        from PIL import Image
        from io import BytesIO
        _PIL = (Image, BytesIO)
    return _PIL

def _lazy_import_docx2pdf():
    """Lazy import docx2pdf"""
    global _docx2pdf
    if _docx2pdf is None:
        try:
            from docx2pdf import convert
            _docx2pdf = convert
        except ImportError:
            _docx2pdf = False
    return _docx2pdf

# Database and helpers (lightweight, can be imported)
from helpers.unified_database import UnifiedDatabase
from helpers.resource_path import get_logo_path, get_template_path


# ============================================
# DIALOG CLASSES (Lightweight, loaded on demand)
# ============================================

class ItemDialog(QDialog):
    """Dialog for adding/editing items - Optimized with __slots__"""
    __slots__ = ('result', 'entry_kod', 'entry_desk', 'entry_tarikh')
    
    def __init__(self, parent, title, initial_values=None):
        super().__init__(parent)
        self.result = None
        self.setWindowTitle(title)
        self.setModal(True)
        self.setFixedSize(600, 300)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(40, 30, 40, 30)
        
        # Title
        title_label = QLabel(title)
        title_label.setStyleSheet("color: #003366; font-size: 18px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Form
        form_widget = QWidget()
        form_layout = QGridLayout(form_widget)
        form_layout.setSpacing(12)
        
        # Kod Tarif
        form_layout.addWidget(QLabel("Kod Tarif:"), 0, 0)
        self.entry_kod = SafeLineEdit()
        self.entry_kod.setMinimumWidth(400)
        form_layout.addWidget(self.entry_kod, 0, 1)
        
        # Deskripsi
        form_layout.addWidget(QLabel("Deskripsi:"), 1, 0)
        self.entry_desk = SafeLineEdit()
        form_layout.addWidget(self.entry_desk, 1, 1)
        
        # Tarikh Kuatkuasa
        form_layout.addWidget(QLabel("Tarikh Kuatkuasa:"), 2, 0)
        self.entry_tarikh = QLineEdit()
        self.entry_tarikh.setText(datetime.now().strftime("%d %B %Y").upper())
        form_layout.addWidget(self.entry_tarikh, 2, 1)
        
        # Pre-fill if editing
        if initial_values:
            self.entry_kod.setText(initial_values[0])
            self.entry_desk.setText(initial_values[1])
            self.entry_tarikh.setText(initial_values[2])
        
        layout.addWidget(form_widget)
        layout.addStretch()
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Simpan")
        btn_save.setStyleSheet("""
            QPushButton {
                background-color: #2E7D32;
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #1B5E20;
            }
        """)
        btn_save.clicked.connect(self.on_ok)
        btn_layout.addWidget(btn_save)
        
        btn_cancel = QPushButton("Batal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #666666;
                color: white;
                font-size: 12px;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #555555;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        layout.addLayout(btn_layout)
        self.entry_kod.setFocus()
    
    def on_ok(self):
        """Save and close"""
        kod = self.entry_kod.text().strip()
        desk = self.entry_desk.text().strip()
        tarikh = self.entry_tarikh.text().strip()
        
        if not kod or not desk:
            QMessageBox.warning(self, "Amaran", "Sila isi Kod Tarif dan Deskripsi")
            return
        
        self.result = (kod, desk, tarikh)
        self.accept()


class AMESListDialog(QDialog):
    """Dialog to show all AMES applications - Optimized"""
    __slots__ = ('db', 'callback', 'all_apps', 'table', 'search_entry')
    
    def __init__(self, parent, database, callback):
        super().__init__(parent)
        self.db = database
        self.callback = callback
        self.all_apps = []
        self.setWindowTitle("Senarai AMES - Pilih Permohonan")
        self.setModal(True)
        self.setGeometry(100, 100, 1200, 600)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Header
        header = QFrame()
        header.setStyleSheet("background-color: #003366;")
        header.setFixedHeight(50)
        header_layout = QVBoxLayout(header)
        title = QLabel("SENARAI PERMOHONAN AMES")
        title.setStyleSheet("color: white; font-size: 13px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(title)
        layout.addWidget(header)
        
        # Search frame
        search_frame = QFrame()
        search_frame.setStyleSheet("background-color: white;")
        search_layout = QHBoxLayout(search_frame)
        search_layout.setContentsMargins(20, 10, 20, 10)
        
        search_layout.addWidget(QLabel("Cari:"))
        self.search_entry = SafeLineEdit()
        self.search_entry.setMinimumWidth(400)
        self.search_entry.textChanged.connect(self.filter_list)
        search_layout.addWidget(self.search_entry)
        
        btn_search = QPushButton("🔍")
        btn_search.clicked.connect(self.filter_list)
        search_layout.addWidget(btn_search)
        
        layout.addWidget(search_frame)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(['ID', 'NO. KELULUSAN', 'NAMA SYARIKAT', 
                                             'KATEGORI', 'TARIKH', 'JML ITEM'])
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.doubleClicked.connect(self.on_select)
        layout.addWidget(self.table)
        
        # Buttons
        btn_frame = QFrame()
        btn_layout = QHBoxLayout(btn_frame)
        btn_layout.setContentsMargins(20, 10, 20, 10)
        
        btn_select = QPushButton("Pilih")
        btn_select.setStyleSheet("background-color: #2E7D32; color: white; padding: 10px 30px;")
        btn_select.clicked.connect(self.on_select)
        btn_layout.addWidget(btn_select)
        
        btn_cancel = QPushButton("Batal")
        btn_cancel.setStyleSheet("background-color: #666666; color: white; padding: 10px 30px;")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        layout.addWidget(btn_frame)
        
        self.load_ames_list()
        self.center_dialog()
    
    def center_dialog(self):
        """Center dialog on screen"""
        self.setGeometry(
            (QDesktopWidget().screenGeometry().width() - self.width()) // 2,
            (QDesktopWidget().screenGeometry().height() - self.height()) // 2,
            self.width(), self.height()
        )
    
    def load_ames_list(self):
        """Load all AMES applications"""
        try:
            apps = self.db.get_all_applications(form_type='ames', limit=500)
            self.all_apps = apps
            
            self.table.setRowCount(len(apps))
            for row, app in enumerate(apps):
                full_app = self.db.get_application_by_id(app['id'])
                ames_details = full_app.get('ames_details', {}) if full_app else {}
                items = full_app.get('items', []) if full_app else []
                
                self.table.setItem(row, 0, QTableWidgetItem(str(app['id'])))
                self.table.setItem(row, 1, QTableWidgetItem(ames_details.get('no_kelulusan', '-')))
                self.table.setItem(row, 2, QTableWidgetItem(app['nama_syarikat']))
                self.table.setItem(row, 3, QTableWidgetItem(ames_details.get('kategori', '-')))
                self.table.setItem(row, 4, QTableWidgetItem(app.get('tarikh', '-')))
                self.table.setItem(row, 5, QTableWidgetItem(str(len(items))))
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat memuatkan senarai: {str(e)}")
    
    def filter_list(self):
        """Filter list based on search"""
        search_text = self.search_entry.text().strip().lower()
        
        if not search_text:
            self.load_ames_list()
            return
        
        filtered = []
        for app in self.all_apps:
            if (search_text in app.get('nama_syarikat', '').lower() or
                search_text in app.get('rujukan_kami', '').lower() or
                search_text in str(app.get('id', '')).lower()):
                filtered.append(app)
        
        self.table.setRowCount(len(filtered))
        for row, app in enumerate(filtered):
            full_app = self.db.get_application_by_id(app['id'])
            ames_details = full_app.get('ames_details', {}) if full_app else {}
            items = full_app.get('items', []) if full_app else []
            
            self.table.setItem(row, 0, QTableWidgetItem(str(app['id'])))
            self.table.setItem(row, 1, QTableWidgetItem(ames_details.get('no_kelulusan', '-')))
            self.table.setItem(row, 2, QTableWidgetItem(app['nama_syarikat']))
            self.table.setItem(row, 3, QTableWidgetItem(ames_details.get('kategori', '-')))
            self.table.setItem(row, 4, QTableWidgetItem(app.get('tarikh', '-')))
            self.table.setItem(row, 5, QTableWidgetItem(str(len(items))))
    
    def on_select(self):
        """Handle selection"""
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Amaran", "Sila pilih permohonan AMES")
            return
        
        app_id = int(self.table.item(selected, 0).text())
        self.callback(app_id)
        self.accept()


class SearchResultsDialog(QDialog):
    """Dialog to show search results - Optimized"""
    __slots__ = ('results', 'callback', 'table')
    
    def __init__(self, parent, results, callback):
        super().__init__(parent)
        self.results = results
        self.callback = callback
        self.setWindowTitle("Keputusan Carian - Pilih Permohonan")
        self.setModal(True)
        self.setGeometry(100, 100, 1000, 400)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Header
        header = QFrame()
        header.setStyleSheet("background-color: #003366;")
        header.setFixedHeight(50)
        header_layout = QVBoxLayout(header)
        title = QLabel(f"KEPUTUSAN CARIAN ({len(results)} rekod)")
        title.setStyleSheet("color: white; font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        header_layout.addWidget(title)
        layout.addWidget(header)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['ID', 'RUJUKAN', 'NAMA SYARIKAT', 'TARIKH'])
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.doubleClicked.connect(self.on_select)
        
        self.table.setRowCount(len(results))
        for row, result in enumerate(results):
            self.table.setItem(row, 0, QTableWidgetItem(str(result['id'])))
            self.table.setItem(row, 1, QTableWidgetItem(result.get('rujukan_kami', '-')))
            self.table.setItem(row, 2, QTableWidgetItem(result.get('nama_syarikat', '-')))
            self.table.setItem(row, 3, QTableWidgetItem(result.get('tarikh', '-')))
        
        layout.addWidget(self.table)
        
        # Buttons
        btn_frame = QFrame()
        btn_layout = QHBoxLayout(btn_frame)
        btn_layout.setContentsMargins(20, 10, 20, 10)
        
        btn_select = QPushButton("Pilih")
        btn_select.setStyleSheet("background-color: #2E7D32; color: white; padding: 10px 30px;")
        btn_select.clicked.connect(self.on_select)
        btn_layout.addWidget(btn_select)
        
        btn_cancel = QPushButton("Batal")
        btn_cancel.setStyleSheet("background-color: #666666; color: white; padding: 10px 30px;")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        layout.addWidget(btn_frame)
        self.center_dialog()
    
    def center_dialog(self):
        """Center dialog on screen"""
        self.setGeometry(
            (QDesktopWidget().screenGeometry().width() - self.width()) // 2,
            (QDesktopWidget().screenGeometry().height() - self.height()) // 2,
            self.width(), self.height()
        )
    
    def on_select(self):
        """Handle selection"""
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Amaran", "Sila pilih permohonan")
            return
        
        app_id = int(self.table.item(selected, 0).text())
        self.callback(app_id)
        self.accept()


# ============================================
# HELPER CLASSES (Loaded only when needed)
# ============================================

class DocumentGenerator:
    """Handles document generation logic - Lazy loaded"""
    __slots__ = ('parent', '_doc_cache')
    
    def __init__(self, parent_form):
        self.parent = parent_form
        self._doc_cache = None  # Cache for document template
    
    def generate_document(self):
        """Generate Word document - Lazy loads docx"""
        Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _lazy_import_docx()
        
        template_file = "delete_item_ames.docx"
        template_path = get_template_path(template_file)
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
            else:
                QMessageBox.warning(self.parent, "Amaran", 
                    f"Templat tidak dijumpai: {template_file}\n\nMenggunakan dokumen kosong...")
                doc = Document()
                doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
        except Exception as e:
            QMessageBox.warning(self.parent, "Amaran", f"Ralat memuat templat: {str(e)}")
            doc = Document()
            doc.add_heading('PERMOHONAN PEMADAMAN/PENAMBAHAN ITEM AMES', 0)
        
        # Use PlaceholderBuilder for standardized placeholders
        from helpers.placeholder_registry import PlaceholderBuilder
        
        # Build replacements using PlaceholderBuilder
        builder = PlaceholderBuilder(self.parent)
        builder.add_rujukan_ames(self.parent.entry_rujukan.text())
        builder.add_nama_syarikat(self.parent.entry_nama.text())
        builder.add_alamat([
            self.parent.entry_alamat1.text().strip(),
            self.parent.entry_alamat2.text().strip(),
            self.parent.entry_alamat3.text().strip()
        ])
        builder.add_tarikh(self.parent.entry_tarikh.text())
        builder.add_tarikh_malay(self.parent.format_tarikh_malay())
        builder.add_tarikh_islam(self.parent.entry_islam.text())
        builder.add_standard('NO_KELULUSAN', self.parent.entry_kelulusan.text())
        builder.add_standard('KATEGORI', self.parent.entry_kategori.text() or "N/A")
        builder.add_nama_pegawai(self.parent.combo_pegawai.currentText())
        
        # Get final replacements dict
        replacements = builder.build()
        
        # Replace placeholders
        from helpers.docx_helper import replace_in_document
        replace_in_document(doc, replacements)
        
        # Handle <<table>> placeholders - replace with actual data tables
        self.replace_table_placeholders(doc)
        
        return doc
    
    def replace_table_placeholders(self, doc):
        """Replace <<table>>, <<table_add>>, <<table_delete>> placeholders with actual data tables"""
        import re
        Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _lazy_import_docx()
        
        # Get items from table
        added_items = []
        deleted_items = []
        
        for row in range(self.parent.table_items.rowCount()):
            status_item = self.parent.table_items.item(row, 1)
            if status_item:
                status = status_item.text()
                if status == "TAMBAH":
                    added_items.append(row)
                elif status == "PADAM":
                    deleted_items.append(row)
        
        # Find all placeholders and process in reverse order
        placeholders_found = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Search for various table placeholders
            if re.search(r'<<\s*table\s*>>', paragraph.text, re.IGNORECASE):
                # <<table>> = both tables (add + delete)
                placeholders_found.append((i, 'both'))
            elif re.search(r'<<\s*table_add\s*>>', paragraph.text, re.IGNORECASE):
                # <<table_add>> = only added items
                placeholders_found.append((i, 'add'))
            elif re.search(r'<<\s*table_delete\s*>>', paragraph.text, re.IGNORECASE):
                # <<table_delete>> = only deleted items
                placeholders_found.append((i, 'delete'))
        
        # Process in reverse order to preserve indices
        for para_idx, table_type in reversed(placeholders_found):
            paragraph = doc.paragraphs[para_idx]
            p_element = paragraph._element
            parent = p_element.getparent()
            
            if table_type == 'both':
                # Insert both tables (delete in reverse)
                if deleted_items:
                    self._insert_delete_table(parent, p_element, deleted_items, RGBColor, qn, OxmlElement)
                if added_items:
                    self._insert_add_table(parent, p_element, added_items, RGBColor, qn, OxmlElement)
            
            elif table_type == 'add':
                if added_items:
                    self._insert_add_table(parent, p_element, added_items, RGBColor, qn, OxmlElement)
            
            elif table_type == 'delete':
                if deleted_items:
                    self._insert_delete_table(parent, p_element, deleted_items, RGBColor, qn, OxmlElement)
            
            # Remove the placeholder paragraph
            parent.remove(p_element)
    
    def _insert_add_table(self, parent, ref_element, added_items, RGBColor, qn, OxmlElement):
        """Insert table for added items with green color"""
        Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _lazy_import_docx()
        
        # Insert title
        title_para = OxmlElement('w:p')
        parent.insert(parent.index(ref_element) + 1, title_para)
        
        run = OxmlElement('w:r')
        title_para.append(run)
        
        rPr = OxmlElement('w:rPr')
        run.append(rPr)
        rPr.append(OxmlElement('w:b'))  # Bold
        
        text = OxmlElement('w:t')
        text.text = 'ITEM YANG DITAMBAH'
        run.append(text)
        
        # Insert table
        table = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        table.append(tblPr)
        
        # Table grid
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(4):
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        table.append(tblGrid)
        
        # Header row
        header_row = OxmlElement('w:tr')
        table.append(header_row)
        
        headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
        for header in headers:
            cell = OxmlElement('w:tc')
            header_row.append(cell)
            
            tcPr = OxmlElement('w:tcPr')
            cell.append(tcPr)
            
            # Green background for added items
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'C8E6C9')
            tcPr.append(shading)
            
            para = OxmlElement('w:p')
            cell.append(para)
            
            pPr = OxmlElement('w:pPr')
            para.append(pPr)
            
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            
            run = OxmlElement('w:r')
            para.append(run)
            
            rPr = OxmlElement('w:rPr')
            rPr.append(OxmlElement('w:b'))
            run.append(rPr)
            
            text = OxmlElement('w:t')
            text.text = header
            run.append(text)
        
        # Data rows
        for idx, row_num in enumerate(added_items, start=1):
            data_row = OxmlElement('w:tr')
            table.append(data_row)
            
            row_data = [
                str(idx),
                self.parent.table_items.item(row_num, 2).text(),
                self.parent.table_items.item(row_num, 3).text(),
                self.parent.table_items.item(row_num, 4).text()
            ]
            
            for cell_value in row_data:
                cell = OxmlElement('w:tc')
                data_row.append(cell)
                
                tcPr = OxmlElement('w:tcPr')
                cell.append(tcPr)
                
                # Green background
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'C8E6C9')
                tcPr.append(shading)
                
                para = OxmlElement('w:p')
                cell.append(para)
                
                run = OxmlElement('w:r')
                para.append(run)
                
                rPr = OxmlElement('w:rPr')
                run.append(rPr)
                
                text = OxmlElement('w:t')
                text.text = str(cell_value) if cell_value else ""
                run.append(text)
        
        parent.insert(parent.index(ref_element) + 1, table)
    
    def _insert_delete_table(self, parent, ref_element, deleted_items, RGBColor, qn, OxmlElement):
        """Insert table for deleted items with red color"""
        Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _lazy_import_docx()
        
        # Insert title
        title_para = OxmlElement('w:p')
        parent.insert(parent.index(ref_element) + 1, title_para)
        
        run = OxmlElement('w:r')
        title_para.append(run)
        
        rPr = OxmlElement('w:rPr')
        run.append(rPr)
        rPr.append(OxmlElement('w:b'))  # Bold
        
        text = OxmlElement('w:t')
        text.text = 'ITEM YANG DIPADAM'
        run.append(text)
        
        # Insert table
        table = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        table.append(tblPr)
        
        # Table grid
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(4):
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        table.append(tblGrid)
        
        # Header row
        header_row = OxmlElement('w:tr')
        table.append(header_row)
        
        headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
        for header in headers:
            cell = OxmlElement('w:tc')
            header_row.append(cell)
            
            tcPr = OxmlElement('w:tcPr')
            cell.append(tcPr)
            
            # Red background for deleted items
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'FFCDD2')
            tcPr.append(shading)
            
            para = OxmlElement('w:p')
            cell.append(para)
            
            pPr = OxmlElement('w:pPr')
            para.append(pPr)
            
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            
            run = OxmlElement('w:r')
            para.append(run)
            
            rPr = OxmlElement('w:rPr')
            rPr.append(OxmlElement('w:b'))
            run.append(rPr)
            
            text = OxmlElement('w:t')
            text.text = header
            run.append(text)
        
        # Data rows
        for idx, row_num in enumerate(deleted_items, start=1):
            data_row = OxmlElement('w:tr')
            table.append(data_row)
            
            row_data = [
                str(idx),
                self.parent.table_items.item(row_num, 2).text(),
                self.parent.table_items.item(row_num, 3).text(),
                self.parent.table_items.item(row_num, 4).text()
            ]
            
            for cell_value in row_data:
                cell = OxmlElement('w:tc')
                data_row.append(cell)
                
                tcPr = OxmlElement('w:tcPr')
                cell.append(tcPr)
                
                # Red background
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FFCDD2')
                tcPr.append(shading)
                
                para = OxmlElement('w:p')
                cell.append(para)
                
                run = OxmlElement('w:r')
                para.append(run)
                
                rPr = OxmlElement('w:rPr')
                run.append(rPr)
                
                text = OxmlElement('w:t')
                text.text = str(cell_value) if cell_value else ""
                run.append(text)
        
        parent.insert(parent.index(ref_element) + 1, table)
    
    def add_items_table_to_document(self, doc):
        """Add items table to document with color coding"""
        Document, Pt, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _lazy_import_docx()
        
        doc.add_page_break()
        doc.add_heading('LAMPIRAN - SENARAI ITEM', level=1)
        
        # Get items from table
        added_items = []
        deleted_items = []
        
        for row in range(self.parent.table_items.rowCount()):
            status_item = self.parent.table_items.item(row, 1)
            if status_item:
                status = status_item.text()
                if status == "TAMBAH":
                    added_items.append(row)
                elif status == "PADAM":
                    deleted_items.append(row)
        
        # Added items section
        if added_items:
            doc.add_heading('ITEM YANG DITAMBAH', level=2)
            para = doc.add_paragraph()
            run = para.add_run('(Ditandakan dengan warna hijau)')
            run.font.color.rgb = RGBColor(76, 175, 80)
            run.bold = True
            
            table = doc.add_table(rows=1 + len(added_items), cols=4)
            headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            for idx, row_num in enumerate(added_items, start=1):
                row = table.rows[idx]
                row.cells[0].text = str(idx)
                row.cells[1].text = self.parent.table_items.item(row_num, 2).text()
                row.cells[2].text = self.parent.table_items.item(row_num, 3).text()
                row.cells[3].text = self.parent.table_items.item(row_num, 4).text()
                
                for cell in row.cells:
                    self.set_cell_background(cell, "C8E6C9")
        
        # Deleted items section
        if deleted_items:
            doc.add_paragraph()
            doc.add_heading('ITEM YANG DIPADAM', level=2)
            para = doc.add_paragraph()
            run = para.add_run('(Ditandakan dengan warna merah)')
            run.font.color.rgb = RGBColor(244, 67, 54)
            run.bold = True
            
            table = doc.add_table(rows=1 + len(deleted_items), cols=4)
            headers = ['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(header)
                run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for idx, row_num in enumerate(deleted_items, start=1):
                row = table.rows[idx]
                row.cells[0].text = str(idx)
                row.cells[1].text = self.parent.table_items.item(row_num, 2).text()
                row.cells[2].text = self.parent.table_items.item(row_num, 3).text()
                row.cells[3].text = self.parent.table_items.item(row_num, 4).text()
                
                for cell in row.cells:
                    self.set_cell_background(cell, "FFCDD2")
    
    def set_cell_background(self, cell, color_hex):
        """Set cell background color"""
        qn, OxmlElement = _lazy_import_docx()[4:6]
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            pass


class TableManager:
    """Manages table operations and color coding - Lightweight"""
    __slots__ = ('parent',)
    
    def __init__(self, parent_form):
        self.parent = parent_form
    
    def add_item(self, status, kod, desk, tarikh):
        """Add item to table with color coding"""
        row = self.parent.table_items.rowCount()
        self.parent.table_items.insertRow(row)
        
        bil = row + 1
        self.parent.table_items.setItem(row, 0, QTableWidgetItem(str(bil)))
        self.parent.table_items.setItem(row, 1, QTableWidgetItem(status))
        self.parent.table_items.setItem(row, 2, QTableWidgetItem(kod))
        self.parent.table_items.setItem(row, 3, QTableWidgetItem(desk))
        self.parent.table_items.setItem(row, 4, QTableWidgetItem(tarikh))
        
        # Set background color based on status
        if status == "TAMBAH":
            color = QColor(200, 230, 201)  # Light green
        elif status == "PADAM":
            color = QColor(255, 205, 210)  # Light red
        else:
            color = QColor(245, 245, 245)  # Light gray
        
        for col in range(5):
            item = self.parent.table_items.item(row, col)
            if item:
                item.setBackground(color)
    
    def mark_as_padam(self, row):
        """Mark existing item as PADAM (red)"""
        # Update status
        status_item = self.parent.table_items.item(row, 1)
        if status_item:
            status_item.setText("PADAM")
        
        # Update background color to red
        red_color = QColor(255, 205, 210)  # Light red
        for col in range(5):
            item = self.parent.table_items.item(row, col)
            if item:
                item.setBackground(red_color)
    
    def renumber_items(self):
        """Renumber items"""
        for row in range(self.parent.table_items.rowCount()):
            item = self.parent.table_items.item(row, 0)
            if item:
                item.setText(str(row + 1))


# ============================================
# MAIN FORM CLASS (Optimized with lazy loading)
# ============================================

class FormDeleteItem(QDialog):
    """Main Delete Item Form - Optimized OOP structure with lazy loading"""
    
    def __init__(self, parent_window=None):
        super().__init__()
        self.parent_window = parent_window
        self.setWindowTitle("Sistem Pemadaman Item - Permohonan Pemadaman Item AMES")
        self.setGeometry(100, 100, 1600, 1000)
        self.setMinimumSize(1200, 800)  # Prevent window from shrinking too much

        # Add borders to text boxes (QTextEdit) only
        self.setStyleSheet("""
            QFrame {
                border: none;
            }
            QLineEdit, QComboBox, QTextEdit {
                border: 1px solid #CCCCCC;
            }
        """)

        self.db = UnifiedDatabase()
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': '#FFFFFF',
            'text_dark': '#1a1a1a',
            'text_light': '#FFFFFF',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_danger': '#C62828',
            'button_info': '#1976D2',
            'item_added': '#4CAF50',
            'item_deleted': '#F44336',
            'button_hover': '#004d99'
        }
        
        # Initialize helper classes (lightweight, no heavy imports)
        self.document_generator = DocumentGenerator(self)
        self.table_manager = TableManager(self)
        
        # Create UI
        self.create_ui()
        self.center_window()
        QTimer.singleShot(100, self.update_tarikh_islam)
    
    def _get_combined_alamat(self):
        """Combine the 3 alamat entries into a single string"""
        alamat_lines = []
        if self.entry_alamat1.text().strip():
            alamat_lines.append(self.entry_alamat1.text().strip())
        if self.entry_alamat2.text().strip():
            alamat_lines.append(self.entry_alamat2.text().strip())
        if self.entry_alamat3.text().strip():
            alamat_lines.append(self.entry_alamat3.text().strip())
        return "\n".join(alamat_lines)
    
    def _get_combined_alamat(self):
        """Combine the 3 alamat entries into a single string"""
        alamat_lines = []
        if self.entry_alamat1.text().strip():
            alamat_lines.append(self.entry_alamat1.text().strip())
        if self.entry_alamat2.text().strip():
            alamat_lines.append(self.entry_alamat2.text().strip())
        if self.entry_alamat3.text().strip():
            alamat_lines.append(self.entry_alamat3.text().strip())
        return "\n".join(alamat_lines)
    
    def create_ui(self):
        """Create the main UI"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Header
        self.create_header(main_layout)
        
        # Scrollable content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet(f"background-color: {self.colors['bg_main']}; border: none;")
        
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(30, 20, 30, 20)
        scroll_layout.setSpacing(20)
        
        # Create sections
        self.create_form_section(scroll_layout)
        self.create_legend_section(scroll_layout)
        self.create_table_section(scroll_layout)
        self.create_button_section(scroll_layout)
        
        scroll_area.setWidget(scroll_widget)
        main_layout.addWidget(scroll_area, stretch=1)
    
    def create_header(self, parent_layout):
        """Create government header - Lazy loads PIL only when needed"""
        header_frame = QFrame()
        header_frame.setFixedHeight(100)
        header_frame.setStyleSheet(f"background-color: {self.colors['primary']};")
        
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(20, 10, 20, 10)
        
        # Logo - Lazy load PIL
        try:
            Image, BytesIO = _lazy_import_pil()
            logo_image = Image.open(get_logo_path())
            if logo_image.mode != 'RGBA':
                logo_image = logo_image.convert('RGBA')
            logo_image = logo_image.resize((60, 60), Image.Resampling.LANCZOS)
            
            bg_rgb = tuple(int(self.colors['primary'][i:i+2], 16) for i in (1, 3, 5))
            background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
            if logo_image.mode == 'RGBA':
                logo_image = Image.alpha_composite(background, logo_image)
                logo_image = logo_image.convert('RGB')
            
            buf = BytesIO()
            logo_image.save(buf, format='PNG')
            pixmap = QPixmap()
            pixmap.loadFromData(buf.getvalue())
            
            logo_label = QLabel()
            logo_label.setPixmap(pixmap.scaled(60, 60, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            logo_label.setStyleSheet(f"background-color: {self.colors['primary']};")
            header_layout.addWidget(logo_label)
        except:
            pass
        
        # Title
        title_widget = QWidget()
        title_layout = QVBoxLayout(title_widget)
        title_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("AMES - PEMADAMAN / PENAMBAHAN ITEM")
        title.setStyleSheet(f"background-color: {self.colors['primary']}; color: white; font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(title)
        
        subtitle = QLabel("Permohonan Pemadaman atau Penambahan Item dari Senarai AMES")
        subtitle.setStyleSheet(f"background-color: {self.colors['primary']}; color: #E0E0E0; font-size: 12px; font-style: italic;")
        subtitle.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(subtitle)
        
        header_layout.addWidget(title_widget, stretch=1)
        
        # Back button
        btn_back = QPushButton("← KEMBALI")
        btn_back.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['secondary']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #004d99;
            }}
        """)
        btn_back.clicked.connect(self.on_back_click)
        header_layout.addWidget(btn_back)
        
        parent_layout.addWidget(header_frame)
        
        # Separator
        separator = QFrame()
        separator.setFixedHeight(3)
        separator.setStyleSheet(f"background-color: {self.colors['accent']};")
        parent_layout.addWidget(separator)
    
    def create_form_section(self, parent_layout):
        """Create form input section"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(25, 20, 25, 20)
        card_layout.setSpacing(15)
        
        section_title = QLabel("Maklumat Permohonan")
        section_title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 18px; font-weight: bold;")
        card_layout.addWidget(section_title)
        
        form_grid = QGridLayout()
        form_grid.setSpacing(10)
        
        row = 0
        
        # Rujukan
        form_grid.addWidget(QLabel("RUJUKAN KAMI:"), row, 0)
        rujukan_widget = QWidget()
        rujukan_layout = QHBoxLayout(rujukan_widget)
        rujukan_layout.setContentsMargins(0, 0, 0, 0)
        rujukan_layout.addWidget(QLabel("KE.JB(90)650/14/AMES/"))
        self.entry_rujukan = SafeLineEdit()
        self.entry_rujukan.setMinimumWidth(200)
        rujukan_layout.addWidget(self.entry_rujukan)
        form_grid.addWidget(rujukan_widget, row, 1)
        
        form_grid.addWidget(QLabel("ALAMAT:"), row, 2)
        self.entry_alamat1 = SafeLineEdit()
        self.entry_alamat1.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat1, row, 3)
        row += 1
        
        # Nama Syarikat
        form_grid.addWidget(QLabel("NAMA SYARIKAT:"), row, 0)
        self.entry_nama = SafeLineEdit()
        form_grid.addWidget(self.entry_nama, row, 1)
        
        self.entry_alamat2 = SafeLineEdit()
        self.entry_alamat2.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat2, row, 3)
        row += 1
        
        # Tarikh
        form_grid.addWidget(QLabel("TARIKH:"), row, 0)
        self.entry_tarikh = QLineEdit()
        self.entry_tarikh.setText(datetime.now().strftime("%d/%m/%Y"))
        
        self.entry_alamat3 = SafeLineEdit()
        self.entry_alamat3.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat3, row, 3)
        self.entry_tarikh.textChanged.connect(self.update_tarikh_islam)
        form_grid.addWidget(self.entry_tarikh, row, 1)
        row += 1
        
        # Tarikh Islam & Nama Pegawai
        form_grid.addWidget(QLabel("TARIKH ISLAM:"), row, 0)
        self.entry_islam = SafeLineEdit()
        self.entry_islam.setReadOnly(True)
        self.entry_islam.setStyleSheet("background-color: #F5F5F5;")
        form_grid.addWidget(self.entry_islam, row, 1)
        
        form_grid.addWidget(QLabel("NAMA PEGAWAI:"), row, 2)
        self.combo_pegawai = QComboBox()
        self.combo_pegawai.setEditable(True)  # Allow typing
        self.combo_pegawai.addItems(["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                    "KHAIRONE DZARWANY BINTI MOHD KAIRI"])
        form_grid.addWidget(self.combo_pegawai, row, 3)
        row += 1
        
        # Divider
        divider = QFrame()
        divider.setFrameShape(QFrame.HLine)
        divider.setStyleSheet(f"color: {self.colors['border']};")
        form_grid.addWidget(divider, row, 0, 1, 4)
        row += 1
        
        # No. Kelulusan AMES
        form_grid.addWidget(QLabel("NO. KELULUSAN AMES:"), row, 0)
        kelulusan_widget = QWidget()
        kelulusan_layout = QHBoxLayout(kelulusan_widget)
        kelulusan_layout.setContentsMargins(0, 0, 0, 0)
        self.entry_kelulusan = SafeLineEdit()
        self.entry_kelulusan.setMinimumWidth(300)
        kelulusan_layout.addWidget(self.entry_kelulusan)
        
        btn_search = QPushButton("🔍 Cari AMES")
        btn_search.setStyleSheet(f"background-color: {self.colors['button_info']}; color: white; padding: 5px 10px;")
        btn_search.clicked.connect(self.search_ames)
        kelulusan_layout.addWidget(btn_search)
        
        btn_list = QPushButton("📋 Pilih dari Senarai")
        btn_list.setStyleSheet(f"background-color: {self.colors['button_secondary']}; color: white; padding: 5px 10px;")
        btn_list.clicked.connect(self.show_ames_list)
        kelulusan_layout.addWidget(btn_list)
        
        form_grid.addWidget(kelulusan_widget, row, 1, 1, 3)
        row += 1
        
        # Kategori
        form_grid.addWidget(QLabel("KATEGORI:"), row, 0)
        self.entry_kategori = SafeLineEdit()
        self.entry_kategori.setReadOnly(True)
        self.entry_kategori.setStyleSheet("background-color: #F5F5F5;")
        form_grid.addWidget(self.entry_kategori, row, 1)
        
        card_layout.addLayout(form_grid)
        parent_layout.addWidget(card)
    
    def create_legend_section(self, parent_layout):
        """Create color legend"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QHBoxLayout(card)
        card_layout.setContentsMargins(25, 15, 25, 15)
        
        card_layout.addWidget(QLabel("LEGENDA WARNA:"))
        
        # Green
        green_frame = QFrame()
        green_frame.setFixedSize(20, 20)
        green_frame.setStyleSheet(f"background-color: {self.colors['item_added']};")
        card_layout.addWidget(green_frame)
        card_layout.addWidget(QLabel("= Item Ditambah (Baru)"))
        
        # Red
        red_frame = QFrame()
        red_frame.setFixedSize(20, 20)
        red_frame.setStyleSheet(f"background-color: {self.colors['item_deleted']};")
        card_layout.addWidget(red_frame)
        card_layout.addWidget(QLabel("= Item Dipadam"))
        
        card_layout.addStretch()
        parent_layout.addWidget(card)
    
    def create_table_section(self, parent_layout):
        """Create items table section"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(25, 20, 25, 20)
        
        title = QLabel("PENGURUSAN ITEM - TAMBAH & PADAM")
        title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 14px; font-weight: bold;")
        card_layout.addWidget(title)
        
        # Instructions
        instructions = QLabel("Pilih item untuk ditandakan sebagai PADAM (merah) atau tambah item baru")
        instructions.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 11px; font-style: italic;")
        card_layout.addWidget(instructions)
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setAlignment(Qt.AlignRight)
        
        btn_add = QPushButton("+ Tambah Item (HIJAU)")
        btn_add.setStyleSheet(f"background-color: {self.colors['item_added']}; color: white; padding: 5px 10px;")
        btn_add.clicked.connect(self.add_item)
        btn_layout.addWidget(btn_add)
        
        btn_delete = QPushButton("- Padam Item (MERAH)")
        btn_delete.setStyleSheet(f"background-color: {self.colors['item_deleted']}; color: white; padding: 5px 10px;")
        btn_delete.clicked.connect(self.delete_item)
        btn_layout.addWidget(btn_delete)
        
        btn_edit = QPushButton("✎ Edit")
        btn_edit.setStyleSheet(f"background-color: {self.colors['button_primary']}; color: white; padding: 5px 10px;")
        btn_edit.clicked.connect(self.edit_item)
        btn_layout.addWidget(btn_edit)
        
        card_layout.addLayout(btn_layout)
        
        # Table
        self.table_items = QTableWidget()
        self.table_items.setColumnCount(5)
        self.table_items.setHorizontalHeaderLabels(['BIL', 'STATUS', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA'])
        self.table_items.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_items.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_items.horizontalHeader().setStretchLastSection(True)
        self.table_items.setAlternatingRowColors(False)
        # Ensure all rows have white background (no stripes)
        self.table_items.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
            }
            QTableWidget::item {
                background-color: white;
            }
        """)
        card_layout.addWidget(self.table_items)
        
        parent_layout.addWidget(card)
    
    def create_button_section(self, parent_layout):
        """Create bottom buttons"""
        button_widget = QWidget()
        button_layout = QHBoxLayout(button_widget)
        button_layout.setAlignment(Qt.AlignCenter)
        button_layout.setSpacing(10)
        
        btn_preview = QPushButton("👁 PRATONTON")
        btn_preview.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_secondary']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
        """)
        btn_preview.clicked.connect(self.on_preview_click)
        button_layout.addWidget(btn_preview)
        
        btn_save = QPushButton("💾 SIMPAN DOKUMEN")
        btn_save.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
        """)
        btn_save.clicked.connect(self.on_save_click)
        button_layout.addWidget(btn_save)
        
        parent_layout.addWidget(button_widget)
    
    def update_tarikh_islam(self):
        """Convert Gregorian to Hijri - Lazy loads hijri_converter"""
        try:
            Hijri, Gregorian = _lazy_import_hijri()
            tarikh_str = self.entry_tarikh.text()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    self.entry_islam.setText(hijri_text)
        except:
            pass
    
    def format_tarikh_malay(self):
        """Format date in Malay"""
        try:
            tarikh_str = self.entry_tarikh.text()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                bulan_melayu = ["Januari", "Februari", "Mac", "April", "Mei", "Jun",
                               "Julai", "Ogos", "September", "Oktober", "November", "Disember"]
                return f"{day:02d} {bulan_melayu[month-1]} {year}"
        except:
            return self.entry_tarikh.text()
    
    def search_ames(self):
        """Search for AMES application"""
        ames_no = self.entry_kelulusan.text().strip()
        if not ames_no:
            QMessageBox.warning(self, "Amaran", "Sila masukkan No. Kelulusan AMES")
            return
        
        try:
            results = self.db.search_applications(ames_no, form_type='ames')
            if not results:
                QMessageBox.information(self, "Tiada Rekod", 
                    f"Tiada rekod AMES dijumpai untuk:\n{ames_no}")
                return
            
            if len(results) > 1:
                dialog = SearchResultsDialog(self, results, self.load_ames_data)
                dialog.exec_()
            else:
                self.load_ames_data(results[0]['id'])
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat mencari AMES: {str(e)}")
    
    def show_ames_list(self):
        """Show list of all AMES applications"""
        dialog = AMESListDialog(self, self.db, self.load_ames_data)
        dialog.exec_()
    
    def load_ames_data(self, application_id):
        """Load AMES data from database"""
        try:
            app = self.db.get_application_by_id(application_id)
            if not app:
                QMessageBox.critical(self, "Ralat", "Rekod tidak dijumpai")
                return
            
            # Fill form fields
            if app.get('rujukan_kami'):
                parts = app['rujukan_kami'].split('/')
                if len(parts) > 0:
                    self.entry_rujukan.setText(parts[-1])
            
            self.entry_nama.setText(app.get('nama_syarikat', ''))
            # Split alamat into 3 fields
            alamat_text = app.get('alamat', '')
            alamat_lines = alamat_text.split('\n') if alamat_text else []
            self.entry_alamat1.setText(alamat_lines[0] if len(alamat_lines) > 0 else '')
            self.entry_alamat2.setText(alamat_lines[1] if len(alamat_lines) > 1 else '')
            self.entry_alamat3.setText(alamat_lines[2] if len(alamat_lines) > 2 else '')
            
            pegawai = app.get('nama_pegawai', '')
            if pegawai:
                # Use setCurrentText for editable combo box (supports custom text)
                self.combo_pegawai.setCurrentText(pegawai)
            
            ames_details = app.get('ames_details', {})
            if ames_details:
                self.entry_kelulusan.setText(ames_details.get('no_kelulusan', ''))
                self.entry_kategori.setText(ames_details.get('kategori', ''))
            
            # Load items into table
            self.table_items.setRowCount(0)
            items = app.get('items', [])
            for item_data in items:
                self.table_manager.add_item("ASAL", 
                    item_data.get('kod_tarif', ''),
                    item_data.get('deskripsi', ''),
                    item_data.get('tarikh_kuatkuasa', ''))
            
            QMessageBox.information(self, "Berjaya", 
                f"Data AMES berjaya dimuatkan!\n\n"
                f"No. Kelulusan: {ames_details.get('no_kelulusan', '')}\n"
                f"Syarikat: {app.get('nama_syarikat', '')}\n"
                f"Jumlah Item Asal: {len(items)}")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat memuatkan data: {str(e)}")
    
    def add_item(self):
        """Add new item (GREEN)"""
        dialog = ItemDialog(self, "Tambah Item (HIJAU - Baru)")
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh = dialog.result
            self.table_manager.add_item("TAMBAH", kod, desk, tarikh)
    
    def delete_item(self):
        """Mark item for deletion (RED) - either mark selected item or add new"""
        selected = self.table_items.currentRow()
        
        if selected >= 0:
            # Mark selected item as PADAM (red)
            current_status = self.table_items.item(selected, 1).text()
            if current_status == "PADAM":
                QMessageBox.information(self, "Info", "Item ini sudah ditandakan sebagai PADAM.")
                return
            
            # Confirm marking as PADAM
            reply = QMessageBox.question(self, "Pengesahan", 
                "Tandakan item ini sebagai PADAM (akan ditandakan dengan warna merah)?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                self.table_manager.mark_as_padam(selected)
        else:
            # No item selected, add new PADAM item
            dialog = ItemDialog(self, "Padam Item (MERAH - Dipadam)")
            if dialog.exec_() == QDialog.Accepted and dialog.result:
                kod, desk, tarikh = dialog.result
                self.table_manager.add_item("PADAM", kod, desk, tarikh)
    
    def edit_item(self):
        """Edit selected item"""
        selected = self.table_items.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Amaran", "Sila pilih item untuk diedit")
            return
        
        kod = self.table_items.item(selected, 2).text()
        desk = self.table_items.item(selected, 3).text()
        tarikh = self.table_items.item(selected, 4).text()
        
        dialog = ItemDialog(self, "Edit Item", (kod, desk, tarikh))
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            new_kod, new_desk, new_tarikh = dialog.result
            status = self.table_items.item(selected, 1).text()
            
            self.table_items.setItem(selected, 2, QTableWidgetItem(new_kod))
            self.table_items.setItem(selected, 3, QTableWidgetItem(new_desk))
            self.table_items.setItem(selected, 4, QTableWidgetItem(new_tarikh))
    
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.document_generator.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview_delete_item.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':
                os.startfile(temp_path)
            elif os.name == 'posix':
                import subprocess
                subprocess.call(['open', temp_path])
            
            QMessageBox.information(self, "Pratonton", "Dokumen dibuka untuk pratonton")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF - Lazy loads docx2pdf"""
        if self.table_items.rowCount() == 0:
            QMessageBox.warning(self, "Amaran", "Sila tambah sekurang-kurangnya satu item")
            return
        
        doc = self.document_generator.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.text().replace('/', '_').replace('\\', '_')
        filename, _ = QFileDialog.getSaveFileName(
            self, "Simpan Dokumen", f"DeleteItem_AMES_{safe_name}.pdf",
            "PDF Files (*.pdf);;All Files (*.*)")
        
        if not filename:
            return
        
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            temp_docx = "temp_delete_item_conversion.docx"
            doc.save(temp_docx)
            
            # Lazy load docx2pdf only when saving
            convert = _lazy_import_docx2pdf()
            if convert:
                convert(temp_docx, filename)
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
            else:
                filename = filename.replace('.pdf', '.docx')
                doc.save(filename)
                QMessageBox.information(self, "Maklumat", 
                    "Library docx2pdf tidak dijumpai. Dokumen disimpan sebagai Word (.docx)")
            
            # Save to database
            self.save_to_database(filename)
            
            QMessageBox.information(self, "Berjaya", 
                f"Dokumen berjaya disimpan:\n{filename}\n\nRekod telah disimpan ke pangkalan data.")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat menyimpan dokumen: {str(e)}")
    
    def save_to_database(self, document_path):
        """Save application to unified database"""
        try:
            application_data = {
                'category': 'AMES',
                'sub_option': 'Delete Item',
                'rujukan_kami': f"KE.JB(90)650/14/AMES/{self.entry_rujukan.text()}",
                'nama_syarikat': self.entry_nama.text(),
                'alamat': self._get_combined_alamat(),
                'tarikh': self.entry_tarikh.text(),
                'tarikh_islam': self.entry_islam.text(),
                'nama_pegawai': self.combo_pegawai.currentText(),
                'status': 'DILULUSKAN',
                'document_path': document_path
            }
            
            specific_details = {
                'no_kelulusan': self.entry_kelulusan.text(),
                'kategori': self.entry_kategori.text(),
                'items': []
            }
            
            # Get all items
            for row in range(self.table_items.rowCount()):
                status = self.table_items.item(row, 1).text()
                specific_details['items'].append({
                    'item_type': status.lower(),
                    'bil': row + 1,
                    'kod_tarif': self.table_items.item(row, 2).text(),
                    'deskripsi': self.table_items.item(row, 3).text(),
                    'tarikh_kuatkuasa': self.table_items.item(row, 4).text()
                })
            
            app_id = self.db.save_application('ames', application_data, specific_details)
            
            # Clear database cache so history viewer shows new data
            if hasattr(self.db, 'clear_cache'):
                self.db.clear_cache()
            
            return app_id
        except Exception as e:
            QMessageBox.critical(self, "Ralat Database", f"Gagal menyimpan ke database: {str(e)}")
            return None
    
    def center_window(self):
        """Center window on screen"""
        self.setGeometry(
            (QDesktopWidget().screenGeometry().width() - self.width()) // 2,
            (QDesktopWidget().screenGeometry().height() - self.height()) // 2,
            self.width(), self.height()
        )
    
    def on_back_click(self):
        """Go back to main menu"""
        if self.parent_window:
            self.parent_window.show()
        self.reject()

