"""
Form3_Government.py - DYNAMIC Government-Styled AMES Form
✨ Enhanced with real-time validation and smart field management
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from hijri_converter import Hijri, Gregorian
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk
import os
import json
import subprocess
from docx2pdf import convert
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from helpers.unified_database import UnifiedDatabase
from helpers.resource_path import get_logo_path, get_template_path

try:
    from helpers.docx_helper import replace_in_document
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value) if value else "")
        return {}

# Import UI components for enhanced UX
try:
    from helpers.ui_components import FieldValidator, Tooltip, NotificationBar
except ImportError:
    FieldValidator = None
    Tooltip = None
    NotificationBar = None


class Form3:
    """Government-styled AMES Form with professional design"""
    
    def __init__(self, root, parent_window=None):
        self.root = root
        self.parent_window = parent_window
        self.root.title("Sistem Pengurusan AMES - Permohonan Kelulusan")
        self.root.geometry("1600x1000")
        self.root.configure(bg='#F5F5F5')
        
        # Initialize database
        self.db = UnifiedDatabase()
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': 'white',
            'text_dark': '#1a1a1a',
            'text_light': 'white',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_danger': '#C62828',
            'button_hover': '#004d99',
            'success': '#4CAF50',
            'warning': '#FF9800',
            'error': '#F44336',
            'info': '#2196F3'
        }
        
        # ✨ DYNAMIC FEATURES
        self.field_validators = {}  # Store field validators
        self.field_states = {}      # Track field states (valid/invalid)
        self.notification_bar = None  # Notification display
        
        # Center window
        self.center_window()
        
        # Make window appear on top
        self.root.attributes('-topmost', True)
        self.root.after_idle(lambda: self.root.attributes('-topmost', False))
        self.root.lift()
        self.root.focus_force()
        
        # Main container
        self.main_container = tk.Frame(root, bg=self.colors['bg_main'])
        self.main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create header (fixed at top)
        self.create_header()
        
        # Create scrollable content area
        self.create_scrollable_content()
        
        # Auto-update tarikh islam
        self.update_tarikh_islam()
        
        # Handle window close
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
    
    def create_scrollable_content(self):
        """Create scrollable canvas for all content"""
        # Container for canvas and scrollbar
        scroll_container = tk.Frame(self.main_container, bg=self.colors['bg_main'])
        scroll_container.pack(fill=tk.BOTH, expand=True)
        
        # Create canvas
        self.canvas = tk.Canvas(scroll_container, bg=self.colors['bg_main'], highlightthickness=0)
        
        # Create scrollbars
        scrollbar_y = ttk.Scrollbar(scroll_container, orient='vertical', command=self.canvas.yview)
        scrollbar_x = ttk.Scrollbar(scroll_container, orient='horizontal', command=self.canvas.xview)
        
        # Create frame inside canvas for all content
        self.scrollable_frame = tk.Frame(self.canvas, bg=self.colors['bg_main'])
        
        # Create window in canvas
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        
        # Configure canvas scrolling
        self.canvas.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
        
        # Pack scrollbars and canvas
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Bind configuration events
        self.scrollable_frame.bind('<Configure>', self.on_frame_configure)
        self.canvas.bind('<Configure>', self.on_canvas_configure)
        
        # Bind mouse wheel scrolling
        self.canvas.bind_all('<MouseWheel>', self.on_mousewheel)
        self.canvas.bind_all('<Button-4>', self.on_mousewheel)  # Linux scroll up
        self.canvas.bind_all('<Button-5>', self.on_mousewheel)  # Linux scroll down
        
        # Now create all form sections inside scrollable_frame
        self.create_form_section()
        self.create_table_section()
        self.create_button_section()
    
    def on_frame_configure(self, event=None):
        """Reset scroll region to encompass the inner frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))
    
    def on_canvas_configure(self, event):
        """Resize the inner frame to match canvas width"""
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas_frame, width=canvas_width)
    
    def on_mousewheel(self, event):
        """Handle mouse wheel scrolling"""
        if event.num == 5 or event.delta < 0:
            self.canvas.yview_scroll(1, 'units')
        elif event.num == 4 or event.delta > 0:
            self.canvas.yview_scroll(-1, 'units')
    
    def create_header(self):
        """Create government header (fixed at top)"""
        # Top banner
        header_frame = tk.Frame(self.main_container, bg=self.colors['primary'], height=100)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Logo section
        logo_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        logo_frame.pack(side=tk.LEFT, padx=20)

        try:
            logo_image = Image.open(get_logo_path())
            
            # Convert to RGBA if not already
            if logo_image.mode != 'RGBA':
                logo_image = logo_image.convert('RGBA')
            
            logo_image = logo_image.resize((60, 60), Image.Resampling.LANCZOS)
            
            # Create a background image with the same color as the frame
            background_color = self.colors['primary']  # '#003366'
            bg_rgb = tuple(int(background_color[i:i+2], 16) for i in (1, 3, 5))
            
            # Create a new RGBA image with the background color
            background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
            
            # Composite the logo onto the background using alpha compositing
            if logo_image.mode == 'RGBA':
                logo_image = Image.alpha_composite(background, logo_image)
                # Convert back to RGB for PhotoImage (Tkinter doesn't support RGBA well)
                logo_image = logo_image.convert('RGB')
            
            logo_photo = ImageTk.PhotoImage(logo_image)
            logo_label = tk.Label(logo_frame, image=logo_photo, bg=self.colors['primary'])
            logo_label.image = logo_photo
            logo_label.pack()
        except Exception as e:
            pass
        
        # Title section
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.pack(side=tk.LEFT, expand=True)
        
        title = tk.Label(title_frame,
                        text="AMES - PERMOHONAN SKIM PENGEKSPORT UTAMA",
                        font=('Arial', 18, 'bold'),
                        bg=self.colors['primary'],
                        fg='white')
        title.pack(pady=(25, 5))
        
        subtitle = tk.Label(title_frame,
                           text="Approved Major Exporter Scheme",
                           font=('Arial', 12, 'italic'),
                           bg=self.colors['primary'],
                           fg='#E0E0E0')
        subtitle.pack()
        
        # ✨ Help Button
        help_btn = tk.Button(header_frame,
                            text="❓ HELP",
                            font=('Arial', 12, 'bold'),
                            bg=self.colors['accent'],
                            fg='white',
                            relief=tk.RAISED,
                            bd=2,
                            padx=20,
                            pady=8,
                            cursor='hand2',
                            command=self.show_help_dialog)
        help_btn.pack(side=tk.RIGHT, padx=10)
        
        # Back button (right)
        self.btn_back = tk.Button(header_frame,
                                  text="← KEMBALI",
                                  font=('Arial', 12, 'bold'),
                                  bg=self.colors['secondary'],
                                  fg='white',
                                  relief=tk.FLAT,
                                  cursor='hand2',
                                  width=18,
                                  height=2,
                                  command=self.on_back_click)
        self.btn_back.pack(side=tk.RIGHT, padx=20)
        
        # Separator
        separator = tk.Frame(self.main_container, bg=self.colors['accent'], height=2)
        separator.pack(fill=tk.X)
    
    def create_form_section(self):
        """Create form input section"""
        # Section container (inside scrollable_frame now)
        section_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        section_container.pack(fill=tk.X, padx=30, pady=20)
        
        # White card
        card = tk.Frame(section_container, bg='white')
        card.pack(fill=tk.X)
        
        # Border
        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, padx=25, pady=20)
        
        # Section title
        section_title = tk.Label(content,
                                text="Maklumat Permohonan",
                                font=('Arial', 16, 'bold'),
                                bg='white',
                                fg=self.colors['primary'])
        section_title.grid(row=0, column=0, columnspan=4, pady=(0, 15), sticky='w')
        
        # Create form fields
        row = 1
        
        def create_label(text, row_num, col_num):
            lbl = tk.Label(content,
                          text=text,
                          font=('Arial', 11, 'bold'),
                          bg='white',
                          fg=self.colors['text_dark'])
            lbl.grid(row=row_num, column=col_num, sticky='e', padx=(5, 10), pady=8)
            return lbl
        
        # Row 1: Rujukan & Alamat
        create_label("RUJUKAN KAMI:", row, 0)
        rujukan_frame = tk.Frame(content, bg='white')
        rujukan_frame.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        tk.Label(rujukan_frame, text="KE.JB(90)650/14/AMES/", bg='white', font=('Arial', 11)).pack(side=tk.LEFT)
        self.entry_rujukan = tk.Entry(rujukan_frame, width=22, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_rujukan.pack(side=tk.LEFT)
        
        create_label("ALAMAT:", row, 2)
        self.text_alamat = tk.Text(content, width=45, height=4, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.text_alamat.grid(row=row, column=3, rowspan=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Row 2: Nama Syarikat
        create_label("NAMA SYARIKAT:", row, 0)
        self.entry_nama = tk.Entry(content, width=42, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_nama.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        row += 1
        
        # Row 3: Tarikh
        create_label("TARIKH:", row, 0)
        self.entry_tarikh = tk.Entry(content, width=42, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_tarikh.insert(0, datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        self.entry_tarikh.bind('<KeyRelease>', lambda e: self.update_tarikh_islam())
        row += 1
        
        # Row 4: Tarikh Islam & Nama Pegawai
        create_label("TARIKH ISLAM:", row, 0)
        self.entry_islam = tk.Entry(content, width=42, font=('Arial', 11), state='readonly', relief=tk.SOLID, bd=1)
        self.entry_islam.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        create_label("NAMA PEGAWAI:", row, 2)
        self.combo_pegawai = ttk.Combobox(content, width=42, font=('Arial', 11))
        self.combo_pegawai['values'] = ["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                        "KHAIRONE DZARWANY BINTI MOHD KAIRI"]
        self.combo_pegawai.current(0)
        self.combo_pegawai.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Divider
        divider = tk.Frame(content, bg=self.colors['border'], height=1)
        divider.grid(row=row, column=0, columnspan=4, sticky='ew', pady=10)
        row += 1
        
        # Row 5: No. Kelulusan & Kategori
        create_label("NO. KELULUSAN AMES:", row, 0)
        self.entry_kelulusan = tk.Entry(content, width=42, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_kelulusan.insert(0, "")
        self.entry_kelulusan.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        create_label("KATEGORI:", row, 2)
        self.combo_kategori = ttk.Combobox(content, width=42, state='readonly', font=('Arial', 11))
        self.combo_kategori['values'] = ["Pedagang", "Pengilang"]
        self.combo_kategori.current(0)
        self.combo_kategori.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Row 6: Tempoh Kelulusan
        create_label("TEMPOH KELULUSAN:", row, 0)
        tempoh_frame = tk.Frame(content, bg='white')
        tempoh_frame.grid(row=row, column=1, columnspan=3, sticky='w', padx=5, pady=8)
        
        self.entry_tarikh_mula = tk.Entry(tempoh_frame, width=18, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_tarikh_mula.insert(0, "")
        self.entry_tarikh_mula.pack(side=tk.LEFT, padx=2)
        
        tk.Label(tempoh_frame, text="hingga", bg='white', font=('Arial', 11, 'bold')).pack(side=tk.LEFT, padx=8)
        
        self.entry_tarikh_tamat = tk.Entry(tempoh_frame, width=18, font=('Arial', 11), relief=tk.SOLID, bd=1)
        self.entry_tarikh_tamat.insert(0, "")
        self.entry_tarikh_tamat.pack(side=tk.LEFT, padx=2)
    
    def create_table_section(self):
        """Create table section based on selected category"""
        self.section_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        self.section_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 20))

        # Default table layout
        self.create_pedagang_table()

        # Rebuild if user switches category
        self.combo_kategori.bind("<<ComboboxSelected>>", lambda e: self.refresh_table_layout())

    def refresh_table_layout(self):
        """Rebuild layout when category changes"""
        for widget in self.section_container.winfo_children():
            widget.destroy()

        kategori = self.combo_kategori.get().lower()
        if kategori == "pengilang":
            self.create_pengilang_tables()
        else:
            self.create_pedagang_table()
        
        # Update scroll region after table change
        self.scrollable_frame.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def create_table_controls(self, parent, add_cmd, edit_cmd, delete_cmd):
        """Reusable button bar for any table"""
        btn_frame = tk.Frame(parent, bg='white')
        btn_frame.pack(anchor='e', pady=(0, 5))

        tk.Button(btn_frame, text="+ Tambah", font=('Arial', 9, 'bold'),
                bg=self.colors['button_success'], fg='white',
                relief=tk.FLAT, cursor='hand2', width=12,
                command=add_cmd).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="✎ Edit", font=('Arial', 9, 'bold'),
                bg=self.colors['button_primary'], fg='white',
                relief=tk.FLAT, cursor='hand2', width=12,
                command=edit_cmd).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="✕ Hapus", font=('Arial', 9, 'bold'),
                bg=self.colors['button_danger'], fg='white',
                relief=tk.FLAT, cursor='hand2', width=12,
                command=delete_cmd).pack(side=tk.LEFT, padx=3)

    def create_treeview(self, parent, jenis):
        """Generic treeview for all table types"""
        columns = {
            "pedagang": ('BIL', 'KOD_TARIF', 'DESKRIPSI', 'TARIKH_KUATKUASA'),
            "bahan": ('BIL', 'KOD_TARIF', 'DESKRIPSI', 'NISBAH', 'TARIKH_KUATKUASA'),
            "barang": ('BIL', 'KOD_TARIF', 'DESKRIPSI', 'TARIKH_KUATKUASA')
        }[jenis]

        frame = tk.Frame(parent, bg='white')
        frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        scroll_y = ttk.Scrollbar(frame, orient='vertical')
        scroll_x = ttk.Scrollbar(frame, orient='horizontal')

        tree = ttk.Treeview(frame, columns=columns, show='headings',
                            yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set, height=6)

        scroll_y.config(command=tree.yview)
        scroll_x.config(command=tree.xview)

        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        tree.pack(fill=tk.BOTH, expand=True)

        # Column headers
        for col in columns:
            tree.heading(col, text=col.replace("_", " "))
            width = 80 if col == "BIL" else 150 if col != "DESKRIPSI" else 350
            tree.column(col, width=width, anchor='center' if col in ("BIL", "TARIKH_KUATKUASA") else 'w')

        return tree

    # =============== PE D A G A N G  ==================

    def create_pedagang_table(self):
        """Create single editable table for Pedagang"""
        card = tk.Frame(self.section_container, bg='white')
        card.pack(fill=tk.BOTH, expand=True)

        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, expand=True, padx=25, pady=20)

        tk.Label(content, text="SENARAI BARANG-BARANG YANG DILULUSKAN (Pedagang)",
                font=('Arial', 14, 'bold'), bg='white', fg=self.colors['primary']).pack(anchor='w', pady=(0, 10))

        # Button controls
        self.create_table_controls(content, self.add_row_pedagang,
                                self.edit_row_pedagang, self.delete_row_pedagang)

        self.tree_pedagang = self.create_treeview(content, "pedagang")

    # =============== P E N G I L A N G  ==================

    def create_pengilang_tables(self):
        """Create dual editable tables for Pengilang"""
        card = tk.Frame(self.section_container, bg='white')
        card.pack(fill=tk.BOTH, expand=True)

        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, expand=True, padx=25, pady=20)

        # Section A
        tk.Label(content, text="A. BAHAN MENTAH, KOMPONEN, BAHAN BUNGKUSAN DAN PEMBUNGKUSAN",
                font=('Arial', 13, 'bold'), bg='white', fg=self.colors['primary']).pack(anchor='w', pady=(0, 5))
        self.create_table_controls(content, self.add_row_bahan, self.edit_row_bahan, self.delete_row_bahan)
        self.tree_bahan = self.create_treeview(content, "bahan")

        # Section B
        tk.Label(content, text="\nB. BARANG SIAP YANG DIKILANGKAN",
                font=('Arial', 13, 'bold'), bg='white', fg=self.colors['primary']).pack(anchor='w', pady=(10, 5))
        self.create_table_controls(content, self.add_row_barang, self.edit_row_barang, self.delete_row_barang)
        self.tree_barang = self.create_treeview(content, "barang")

    # ---------------------- PEDAGANG CRUD ----------------------

    def add_row_pedagang(self):
        dialog = RowDialog(self.root, "Tambah Barangan (Pedagang)")
        if dialog.result:
            kod, desk, tarikh = dialog.result
            bil = len(self.tree_pedagang.get_children()) + 1
            self.tree_pedagang.insert('', 'end', values=(bil, kod, desk, tarikh))

    def edit_row_pedagang(self):
        self._edit_tree_row(self.tree_pedagang, "Edit Barangan (Pedagang)")

    def delete_row_pedagang(self):
        self._delete_tree_row(self.tree_pedagang)

    # ---------------------- PENGILANG CRUD ----------------------

    def add_row_bahan(self):
        dialog = RowDialog(self.root, "Tambah Bahan Mentah", initial_values=("", "", "", "1:1"), with_nisbah=True)
        if dialog.result:
            kod, desk, tarikh, nisbah = dialog.result
            bil = len(self.tree_bahan.get_children()) + 1
            self.tree_bahan.insert('', 'end', values=(bil, kod, desk, nisbah, tarikh))

    def edit_row_bahan(self):
        self._edit_tree_row(self.tree_bahan, "Edit Bahan Mentah", with_nisbah=True)

    def delete_row_bahan(self):
        self._delete_tree_row(self.tree_bahan)

    def add_row_barang(self):
        dialog = RowDialog(self.root, "Tambah Barang Siap")
        if dialog.result:
            kod, desk, tarikh = dialog.result
            bil = len(self.tree_barang.get_children()) + 1
            self.tree_barang.insert('', 'end', values=(bil, kod, desk, tarikh))

    def edit_row_barang(self):
        self._edit_tree_row(self.tree_barang, "Edit Barang Siap")

    def delete_row_barang(self):
        self._delete_tree_row(self.tree_barang)

    # ---------------------- SHARED HELPER FUNCTIONS ----------------------

    def _edit_tree_row(self, tree, title, with_nisbah=False):
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih baris untuk diedit")
            return
        item = selected[0]
        values = tree.item(item)['values']

        # Prepare initial values for dialog
        if with_nisbah:
            # Format: (bil, kod, desk, nisbah, tarikh)
            initial_vals = (values[1], values[2], values[-1], values[3]) if len(values) >= 5 else (values[1], values[2], values[-1], "1:1")
        else:
            # Format: (bil, kod, desk, tarikh)
            initial_vals = (values[1], values[2], values[-1])
        
        # Open dialog
        dialog = RowDialog(self.root, title, initial_values=initial_vals, with_nisbah=with_nisbah)
        if dialog.result:
            if with_nisbah:
                kod, desk, tarikh, nisbah = dialog.result
                tree.item(item, values=(values[0], kod, desk, nisbah, tarikh))
            else:
                kod, desk, tarikh = dialog.result
                tree.item(item, values=(values[0], kod, desk, tarikh))

    def _delete_tree_row(self, tree):
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Amaran", "Sila pilih baris untuk dihapus")
            return

        if messagebox.askyesno("Pengesahan", "Adakah anda pasti ingin menghapus baris ini?"):
            for item in selected:
                tree.delete(item)

            # Renumber rows
            for idx, item in enumerate(tree.get_children(), start=1):
                values = list(tree.item(item)['values'])
                values[0] = idx
                tree.item(item, values=values)

    def create_button_section(self):
        """Create bottom button section"""
        # Button container (inside scrollable_frame now)
        button_container = tk.Frame(self.scrollable_frame, bg=self.colors['bg_main'])
        button_container.pack(fill=tk.X, padx=30, pady=(0, 25))
        
        # Center buttons
        button_frame = tk.Frame(button_container, bg=self.colors['bg_main'])
        button_frame.pack()
        
        # Preview button
        btn_preview = tk.Button(button_frame,
                               text="👁 PRATONTON",
                               font=('Arial', 10, 'bold'),
                               bg=self.colors['button_secondary'],
                               fg='white',
                               relief=tk.FLAT,
                               cursor='hand2',
                               width=18,
                               height=2,
                               command=self.on_preview_click)
        btn_preview.pack(side=tk.LEFT, padx=10)
        
        # Save button
        btn_save = tk.Button(button_frame,
                            text="💾 SIMPAN DOKUMEN",
                            font=('Arial', 10, 'bold'),
                            bg=self.colors['button_success'],
                            fg='white',
                            relief=tk.FLAT,
                            cursor='hand2',
                            width=22,
                            height=2,
                            command=self.on_save_click)
        btn_save.pack(side=tk.LEFT, padx=10)
    
    def update_tarikh_islam(self):
        """Convert Gregorian to Hijri"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    
                    self.entry_islam.config(state='normal')
                    self.entry_islam.delete(0, tk.END)
                    self.entry_islam.insert(0, hijri_text)
                    self.entry_islam.config(state='readonly')
        except:
            pass
    
    def format_tarikh_malay(self):
        """Format date in Malay"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                
                bulan_melayu = ["Januari", "Februari", "Mac", "April", "Mei", "Jun",
                               "Julai", "Ogos", "September", "Oktober", "November", "Disember"]
                
                return f"{day:02d} {bulan_melayu[month-1]} {year}"
        except:
            return self.entry_tarikh.get()
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview_ames.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':  # Windows
                os.startfile(temp_path)
            elif os.name == 'posix':  # macOS/Linux
                subprocess.call(['open', temp_path])
            
            messagebox.showinfo("Pratonton", "Dokumen dibuka untuk pratonton")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        # Validate required fields
        if not self.entry_nama.get().strip():
            messagebox.showerror("Ralat", "Sila isi Nama Syarikat")
            return
        
        if not self.entry_rujukan.get().strip():
            messagebox.showerror("Ralat", "Sila isi Rujukan")
            return
        
        doc = self.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.get().replace('/', '_').replace('\\', '_')
        kategori = self.combo_kategori.get()
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("Fail PDF", "*.pdf"), ("Semua fail", "*.*")],
            initialfile=f"AMES_{safe_name}_{kategori}.pdf"
        )
        
        if not filename:
            return
        
        # Ensure the filename ends with .pdf
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            # Save as temporary Word file first
            temp_docx = "temp_ames_conversion.docx"
            doc.save(temp_docx)
            
            # Convert Word to PDF
            convert(temp_docx, filename)
            
            # Clean up temporary file
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            
            # Save to database
            try:
                # Prepare application data
                rujukan_kami = f"KE.JB(90)650/14/AMES/{self.entry_rujukan.get()}"
                alamat_text = self.text_alamat.get("1.0", tk.END).strip()
                
                application_data = {
                    'category': kategori,
                    'sub_option': None,
                    'rujukan_kami': rujukan_kami,
                    'rujukan_tuan': '',
                    'nama_syarikat': self.entry_nama.get(),
                    'alamat': alamat_text,
                    'tarikh': self.entry_tarikh.get(),
                    'tarikh_islam': self.entry_islam.get(),
                    'nama_pegawai': self.combo_pegawai.get(),
                    'status': 'DILULUSKAN',
                    'document_path': filename,
                    'additional_data': {}
                }
                
                # Prepare AMES-specific details
                ames_details = {
                    'no_kelulusan': self.entry_kelulusan.get(),
                    'kategori': kategori,
                    'tarikh_mula': self.entry_tarikh_mula.get(),
                    'tarikh_tamat': self.entry_tarikh_tamat.get(),
                    'tempoh_kelulusan': f"{self.entry_tarikh_mula.get()} hingga {self.entry_tarikh_tamat.get()}"
                }
                
                # Collect items from tables
                items = []
                
                # Get items from appropriate table(s) based on category
                if kategori.lower() == "pedagang":
                    # Get items from tree_pedagang
                    if hasattr(self, 'tree_pedagang'):
                        for item_id in self.tree_pedagang.get_children():
                            values = self.tree_pedagang.item(item_id)['values']
                            if len(values) >= 4:
                                items.append({
                                    'item_type': 'pedagang',
                                    'bil': values[0],
                                    'kod_tarif': values[1],
                                    'deskripsi': values[2],
                                    'nisbah': None,
                                    'tarikh_kuatkuasa': values[3] if len(values) > 3 else ''
                                })
                else:  # Pengilang
                    # Get items from tree_bahan
                    if hasattr(self, 'tree_bahan'):
                        for item_id in self.tree_bahan.get_children():
                            values = self.tree_bahan.item(item_id)['values']
                            if len(values) >= 5:
                                items.append({
                                    'item_type': 'bahan',
                                    'bil': values[0],
                                    'kod_tarif': values[1],
                                    'deskripsi': values[2],
                                    'nisbah': values[3] if len(values) > 3 else None,
                                    'tarikh_kuatkuasa': values[4] if len(values) > 4 else ''
                                })
                    
                    # Get items from tree_barang
                    if hasattr(self, 'tree_barang'):
                        for item_id in self.tree_barang.get_children():
                            values = self.tree_barang.item(item_id)['values']
                            if len(values) >= 4:
                                items.append({
                                    'item_type': 'barang',
                                    'bil': values[0],
                                    'kod_tarif': values[1],
                                    'deskripsi': values[2],
                                    'nisbah': None,
                                    'tarikh_kuatkuasa': values[3] if len(values) > 3 else ''
                                })
                
                ames_details['items'] = items
                
                # Save to database
                app_id = self.db.save_application('ames', application_data, ames_details)
                
                messagebox.showinfo("Berjaya", 
                    f"Dokumen berjaya disimpan sebagai PDF:\n{filename}\n\n"
                    f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
            except Exception as db_error:
                # Document saved but database save failed
                messagebox.showwarning("Amaran", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
            
        except ImportError:
            messagebox.showerror("Ralat", 
                "Library docx2pdf tidak dijumpai.\n\n"
                "Sila install dengan command:\n"
                "pip install docx2pdf")
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menyimpan dokumen: {str(e)}")
            # Clean up if error occurs
            if os.path.exists("temp_ames_conversion.docx"):
                os.remove("temp_ames_conversion.docx")

    def generate_document(self):
        """Generate AMES document"""
        kategori = self.combo_kategori.get()
        if kategori == "Pedagang":
            template_file = "ames_pedagang.docx"
            # Also try alternative template name
            alt_template = "SURAT KELULUSAN AMES PEDAGANG.docx"
        else:
            template_file = "ames_pengilang.docx"
            alt_template = None
        
        doc = None
        
        # Try to use embedded template storage first
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(template_file)
            if doc is None and alt_template:
                # Try alternative template name
                doc = get_template_document(alt_template)
        except ImportError:
            pass
        
        # Fallback to file system
        if doc is None:
            # Try Templates folder first
            template_path = get_template_path(template_file)
            if os.path.exists(template_path):
                doc = Document(template_path)
            elif alt_template:
                # Try alternative template in Templates folder
                alt_path = get_template_path(alt_template)
                if os.path.exists(alt_path):
                    doc = Document(alt_path)
                # Try root directory
                elif os.path.exists(alt_template):
                    doc = Document(alt_template)
            
            # If still not found, create empty document
            if doc is None:
                messagebox.showwarning("Amaran",
                    f"Templat tidak dijumpai: {template_file}\n\nMenggunakan dokumen kosong...")
                doc = Document()
                doc.add_heading('PERMOHONAN SKIM PENGEKSPORT UTAMA DILULUSKAN (AMES)', 0)
        
        # Prepare replacements
        replacements = {
            '<<RUJUKAN_TUAN>>': '',  # Usually empty, can be added if field exists
            '<<RUJUKAN_KAMI>>': f"KE.JB(90)650/14/AMES/{self.entry_rujukan.get()}",
            '<<NAMA_SYARIKAT>>': self.entry_nama.get().upper(),
            '<<ALAMAT>>': self.text_alamat.get("1.0", tk.END).strip(),
            '<<TARIKH>>': self.entry_tarikh.get(),
            '<<TARIKH_MALAY>>': self.format_tarikh_malay(),
            '<<TARIKH_ISLAM>>': self.entry_islam.get(),
            '<<NO_KELULUSAN>>': self.entry_kelulusan.get(),
            '<<KATEGORI>>': kategori,
            '<<TEMPOH_KELULUSAN>>': f"{self.entry_tarikh_mula.get()} hingga {self.entry_tarikh_tamat.get()}",
            '<<NAMA_PEGAWAI>>': self.combo_pegawai.get().upper(),
        }
        
        # Replace placeholders
        replace_in_document(doc, replacements)
        
        # Add product table
        self.add_table_to_document(doc)
        
        return doc
    
    def add_table_to_document(self, doc):
        """Add Lampiran table(s) to the Word document"""
        kategori = self.combo_kategori.get().lower()

        if kategori == "pedagang":
            self.add_lampiran_a_pedagang(doc)
        else:
            self.add_lampiran_a1_pengilang(doc)

    def add_lampiran_a_pedagang(self, doc):
        """
        Add Lampiran A table for Pedagang to document.
        
        Structure should be:
        - Paragraph 3 (ends with "3.5. Stesen Mengawal...")
        - "SENARAI BARANG-BARANG YANG DILULUSKAN" (heading)
        - "LAMPIRAN A" (heading)  
        - "BAGI PEDAGANG BARANG BERCUKAI" (text)
        - Table with items
        - Paragraph 4 ("Sebagai orang yang telah diluluskan...")
        - Paragraph 5
        - Lampiran B
        
        Note: Due to python-docx limitations, table is added at end of document.
        Template should have placeholder <<LAMPIRAN_A_TABLE>> or structure already in place.
        """
        # First, check if there's a placeholder for the table
        placeholder_found = False
        placeholder_para = None
        
        for para in doc.paragraphs:
            if '<<LAMPIRAN_A_TABLE>>' in para.text or '<<LAMPIRAN_A>>' in para.text:
                placeholder_found = True
                placeholder_para = para
                break
        
        # Find where to insert - look for paragraph 4 or "BAGI PEDAGANG" or Lampiran B
        target_paragraph = None
        if not placeholder_found:
            for para in doc.paragraphs:
                para_text = para.text.strip()
                # Look for "BAGI PEDAGANG BARANG BERCUKAI" - insert table after this
                if 'BAGI PEDAGANG BARANG BERCUKAI' in para_text.upper():
                    target_paragraph = para
                    break
                # Or look for paragraph 4 (starts with "4.")
                elif para_text.startswith('4.') and ('diluluskan' in para_text.lower() or 'dikehendaki' in para_text.lower()):
                    target_paragraph = para
                    break
                # Or Lampiran B as fallback
                elif 'LAMPIRAN B' in para_text.upper():
                    target_paragraph = para
                    break
        
        # Get data from treeview
        items = []
        if hasattr(self, 'tree_pedagang') and self.tree_pedagang:
            for item in self.tree_pedagang.get_children():
                values = self.tree_pedagang.item(item)['values']
                if values and len(values) >= 3:  # Ensure we have at least bil, kod_tarif, deskripsi
                    items.append(values)
        
        # If placeholder found, we'll add table at end but document the position
        # The template should have the structure already, we just need to add the table
        if placeholder_found:
            # Remove placeholder text
            placeholder_para.text = placeholder_para.text.replace('<<LAMPIRAN_A_TABLE>>', '').replace('<<LAMPIRAN_A>>', '')
            if not items:
                placeholder_para.text = 'Tiada item dalam senarai.'
                return
        
        # Insert heading and text before target paragraph if not using placeholder
        if not placeholder_found:
            if target_paragraph:
                # Check if headings already exist
                existing_texts = [p.text.upper() for p in doc.paragraphs]
                if 'BAGI PEDAGANG BARANG BERCUKAI' not in existing_texts:
                    target_paragraph.insert_paragraph_before('BAGI PEDAGANG BARANG BERCUKAI')
                if 'SENARAI BARANG-BARANG YANG DILULUSKAN' not in existing_texts:
                    target_paragraph.insert_paragraph_before('SENARAI BARANG-BARANG YANG DILULUSKAN')
                if 'LAMPIRAN A' not in existing_texts:
                    target_paragraph.insert_paragraph_before('')
                    target_paragraph.insert_paragraph_before('LAMPIRAN A')
                
                if not items:
                    target_paragraph.insert_paragraph_before('Tiada item dalam senarai.')
                    return
            else:
                # Add at end if target not found
                doc.add_heading('LAMPIRAN A', level=1)
                doc.add_paragraph('SENARAI BARANG-BARANG YANG DILULUSKAN')
                doc.add_paragraph('BAGI PEDAGANG BARANG BERCUKAI')
                doc.add_paragraph('')
                if not items:
                    doc.add_paragraph('Tiada item dalam senarai.')
                    return
        
        # Create table (python-docx adds tables at end, but headings are in correct position)
        table = doc.add_table(rows=1 + len(items), cols=4)
        
        # Try to set table style, fallback if style doesn't exist
        try:
            table.style = 'Light Grid Accent 1'
        except (KeyError, ValueError):
            try:
                table.style = 'Table Grid'
            except (KeyError, ValueError):
                try:
                    table.style = 'Light List Accent 1'
                except (KeyError, ValueError):
                    pass
        
        # Headers
        headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            para = cell.paragraphs[0]
            para.clear()  # Clear any existing content
            run = para.add_run(header)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for idx, item_values in enumerate(items, start=1):
            row = table.rows[idx]
            # Safely get values with defaults
            row.cells[0].text = str(item_values[0]) if len(item_values) > 0 and item_values[0] else str(idx)
            row.cells[1].text = str(item_values[1]) if len(item_values) > 1 and item_values[1] else ''
            row.cells[2].text = str(item_values[2]) if len(item_values) > 2 and item_values[2] else ''
            
            # Center align BIL column
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set font for all cells
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
        
        # Merge TARIKH KUATKUASA column for all data rows
        if len(items) > 0:
            try:
                first_tarikh_cell = table.rows[1].cells[3]
                first_tarikh_para = first_tarikh_cell.paragraphs[0]
                first_tarikh_para.clear()
                run = first_tarikh_para.add_run('<<TARIKH_KUATKUASA>>')
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                first_tarikh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if len(items) > 1:
                    for row_idx in range(2, len(items) + 1):
                        cell_to_merge = table.rows[row_idx].cells[3]
                        first_tarikh_cell._tc.merge(cell_to_merge._tc)
            except Exception as e:
                print(f"Warning: Could not merge TARIKH KUATKUASA cells: {e}")
                # Fill each cell individually if merge fails
                for row_idx in range(1, len(items) + 1):
                    cell = table.rows[row_idx].cells[3]
                    para = cell.paragraphs[0]
                    para.clear()
                    run = para.add_run('<<TARIKH_KUATKUASA>>')
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_lampiran_a1_pengilang(self, doc):
        """Add Lampiran A1 tables for Pengilang to document - placed before Lampiran B"""
        # Find where to insert (before Lampiran B if exists)
        insert_position = None
        for i, para in enumerate(doc.paragraphs):
            if 'LAMPIRAN B' in para.text.upper() or 'LAMPIRAN B.' in para.text.upper():
                insert_position = i
                break
        
        # Add heading
        if insert_position is not None:
            para = doc.paragraphs[insert_position]
            para.insert_paragraph_before('LAMPIRAN A1')
            para.insert_paragraph_before('')
        else:
            doc.add_heading('LAMPIRAN A1', level=1)
        
        # Section A: Bahan Mentah
        if insert_position is not None:
            para = doc.paragraphs[insert_position]
            para.insert_paragraph_before('A. Bahan Mentah, Komponen, Bahan Bungkusan dan Pembungkusan')
        else:
            doc.add_paragraph('A. Bahan Mentah, Komponen, Bahan Bungkusan dan Pembungkusan')
        
        # Get data from tree_bahan
        bahan_items = []
        if hasattr(self, 'tree_bahan') and self.tree_bahan:
            for item in self.tree_bahan.get_children():
                values = self.tree_bahan.item(item)['values']
                if values and len(values) >= 4:  # Ensure we have at least bil, kod_tarif, deskripsi, nisbah
                    bahan_items.append(values)
        
        if bahan_items:
            # Get tarikh from first item
            tarikh_kuatkuasa = str(bahan_items[0][4]) if len(bahan_items[0]) > 4 and bahan_items[0][4] else '<<TARIKH_KUATKUASA>>'
            
            # Create table for bahan
            table_bahan = doc.add_table(rows=1 + len(bahan_items), cols=5)
            # Try to set table style, fallback if style doesn't exist
            try:
                table_bahan.style = 'Light Grid Accent 1'
            except (KeyError, ValueError):
                try:
                    table_bahan.style = 'Table Grid'
                except (KeyError, ValueError):
                    try:
                        table_bahan.style = 'Light List Accent 1'
                    except (KeyError, ValueError):
                        pass
            
            # Headers
            headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'NISBAH', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table_bahan.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()  # Clear any existing content
                run = para.add_run(header)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for idx, item_values in enumerate(bahan_items, start=1):
                row = table_bahan.rows[idx]
                # item_values format: (bil, kod_tarif, deskripsi, nisbah, tarikh)
                # Safely get values with defaults
                row.cells[0].text = str(item_values[0]) if len(item_values) > 0 and item_values[0] else str(idx)
                row.cells[1].text = str(item_values[1]) if len(item_values) > 1 and item_values[1] else ''
                row.cells[2].text = str(item_values[2]) if len(item_values) > 2 and item_values[2] else ''
                row.cells[3].text = str(item_values[3]) if len(item_values) > 3 and item_values[3] else ''
                # Don't fill tarikh yet - will merge later
                
                # Center align BIL and NISBAH columns
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set font for all cells
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
            
            # Merge TARIKH KUATKUASA column for all data rows
            if len(bahan_items) > 0:
                try:
                    first_tarikh_cell = table_bahan.rows[1].cells[4]
                    first_tarikh_para = first_tarikh_cell.paragraphs[0]
                    first_tarikh_para.clear()
                    run = first_tarikh_para.add_run('<<TARIKH_KUATKUASA>>')
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    first_tarikh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if len(bahan_items) > 1:
                        for row_idx in range(2, len(bahan_items) + 1):
                            cell_to_merge = table_bahan.rows[row_idx].cells[4]
                            first_tarikh_cell._tc.merge(cell_to_merge._tc)
                except Exception as e:
                    print(f"Warning: Could not merge TARIKH KUATKUASA cells: {e}")
                    # Fill each cell individually if merge fails
                    for row_idx in range(1, len(bahan_items) + 1):
                        cell = table_bahan.rows[row_idx].cells[4]
                        para = cell.paragraphs[0]
                        para.clear()
                        run = para.add_run('<<TARIKH_KUATKUASA>>')
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if insert_position is not None:
                para = doc.paragraphs[insert_position]
                para.insert_paragraph_before('Tiada item dalam senarai.')
            else:
                doc.add_paragraph('Tiada item dalam senarai.')
        
        # Section B: Barang Siap
        if insert_position is not None:
            para = doc.paragraphs[insert_position]
            para.insert_paragraph_before('')
            para.insert_paragraph_before('B. Barang Siap Yang Dikilangkan')
        else:
            doc.add_paragraph()
            doc.add_paragraph('B. Barang Siap Yang Dikilangkan')
        
        # Get data from tree_barang
        barang_items = []
        if hasattr(self, 'tree_barang') and self.tree_barang:
            for item in self.tree_barang.get_children():
                values = self.tree_barang.item(item)['values']
                if values and len(values) >= 3:  # Ensure we have at least bil, kod_tarif, deskripsi
                    barang_items.append(values)
        
        if barang_items:
            # Get tarikh from first item
            tarikh_kuatkuasa = str(barang_items[0][3]) if len(barang_items[0]) > 3 and barang_items[0][3] else '<<TARIKH_KUATKUASA>>'
            
            # Create table for barang
            table_barang = doc.add_table(rows=1 + len(barang_items), cols=4)
            # Try to set table style, fallback if style doesn't exist
            try:
                table_barang.style = 'Light Grid Accent 1'
            except (KeyError, ValueError):
                try:
                    table_barang.style = 'Table Grid'
                except (KeyError, ValueError):
                    try:
                        table_barang.style = 'Light List Accent 1'
                    except (KeyError, ValueError):
                        pass
            
            # Headers
            headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            for idx, header in enumerate(headers):
                cell = table_barang.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()  # Clear any existing content
                run = para.add_run(header)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for idx, item_values in enumerate(barang_items, start=1):
                row = table_barang.rows[idx]
                # item_values format: (bil, kod_tarif, deskripsi, tarikh)
                # Safely get values with defaults
                row.cells[0].text = str(item_values[0]) if len(item_values) > 0 and item_values[0] else str(idx)
                row.cells[1].text = str(item_values[1]) if len(item_values) > 1 and item_values[1] else ''
                row.cells[2].text = str(item_values[2]) if len(item_values) > 2 and item_values[2] else ''
                # Don't fill tarikh yet - will merge later
                
                # Center align BIL column
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set font for all cells
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
            
            # Merge TARIKH KUATKUASA column for all data rows
            if len(barang_items) > 0:
                try:
                    first_tarikh_cell = table_barang.rows[1].cells[3]
                    first_tarikh_para = first_tarikh_cell.paragraphs[0]
                    first_tarikh_para.clear()
                    run = first_tarikh_para.add_run('<<TARIKH_KUATKUASA>>')
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    first_tarikh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if len(barang_items) > 1:
                        for row_idx in range(2, len(barang_items) + 1):
                            cell_to_merge = table_barang.rows[row_idx].cells[3]
                            first_tarikh_cell._tc.merge(cell_to_merge._tc)
                except Exception as e:
                    print(f"Warning: Could not merge TARIKH KUATKUASA cells: {e}")
                    # Fill each cell individually if merge fails
                    for row_idx in range(1, len(barang_items) + 1):
                        cell = table_barang.rows[row_idx].cells[3]
                        para = cell.paragraphs[0]
                        para.clear()
                        run = para.add_run('<<TARIKH_KUATKUASA>>')
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if insert_position is not None:
                para = doc.paragraphs[insert_position]
                para.insert_paragraph_before('Tiada item dalam senarai.')
            else:
                doc.add_paragraph('Tiada item dalam senarai.')
    
    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = 1600
        height = 1000
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def on_back_click(self):
        """Go back to main menu"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.root.destroy()
    
    # ========================================
    # ✨ DYNAMIC FEATURES - Validation  
    # ========================================
    
    def validate_field_realtime(self, field_name, value, field_widget):
        """Real-time validation with visual feedback"""
        is_valid = True
        
        if field_name == "rujukan" and value:
            if len(value) < 2:
                is_valid = False
        elif field_name == "nama_syarikat" and value:
            if len(value) < 3:
                is_valid = False
        elif field_name == "tarikh" and value:
            try:
                datetime.strptime(value, "%d/%m/%Y")
            except:
                is_valid = False
        
        if is_valid:
            field_widget.configure(bg='white')
            self.field_states[field_name] = True
        else:
            field_widget.configure(bg='#FFEBEE')
            self.field_states[field_name] = False
        
        return is_valid
    
    def get_form_completion_percentage(self):
        """Calculate form completion percentage"""
        total_fields = 0
        filled_fields = 0
        
        # Check required fields
        required_fields = [
            ('rujukan', self.entry_rujukan.get().strip()),
            ('nama_syarikat', self.entry_nama.get().strip()),
            ('tarikh', self.entry_tarikh.get().strip()),
            ('alamat', self.text_alamat.get("1.0", tk.END).strip()),
            ('pegawai', self.combo_pegawai.get().strip()),
            ('no_kelulusan', self.entry_kelulusan.get().strip()),
            ('tarikh_mula', self.entry_tarikh_mula.get().strip()),
            ('tarikh_tamat', self.entry_tarikh_tamat.get().strip())
        ]
        
        for field_name, field_value in required_fields:
            total_fields += 1
            if field_value:
                filled_fields += 1
        
        # Check if table has items
        total_fields += 1
        has_items = False
        kategori = self.combo_kategori.get().lower()
        if kategori == "pedagang" and hasattr(self, 'tree_pedagang'):
            has_items = len(self.tree_pedagang.get_children()) > 0
        elif kategori == "pengilang":
            if hasattr(self, 'tree_bahan') and len(self.tree_bahan.get_children()) > 0:
                has_items = True
            if hasattr(self, 'tree_barang') and len(self.tree_barang.get_children()) > 0:
                has_items = True
        
        if has_items:
            filled_fields += 1
        
        if total_fields == 0:
            return 0
        
        return int((filled_fields / total_fields) * 100)
    
    def update_completion_indicator(self):
        """Update form completion progress indicator"""
        percentage = self.get_form_completion_percentage()
        self.root.title(f"Sistem Pengurusan AMES - Permohonan Kelulusan [{percentage}% complete]")
    
    def on_field_change(self, field_name, field_widget):
        """Handle field value change"""
        value = field_widget.get()
        self.validate_field_realtime(field_name, value, field_widget)
        self.update_completion_indicator()
    
    def show_help_dialog(self):
        """Show comprehensive help dialog"""
        help_window = tk.Toplevel(self.root)
        help_window.title("📚 Panduan Form AMES")
        help_window.geometry("700x600")
        help_window.configure(bg='white')
        
        header = tk.Frame(help_window, bg=self.colors['primary'], height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header, text="📚 Panduan Pengisian Borang AMES",
                font=('Arial', 14, 'bold'), bg=self.colors['primary'],
                fg='white').pack(pady=15)
        
        canvas = tk.Canvas(help_window, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(help_window, orient='vertical', command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')
        
        scrollable_frame.bind("<Configure>",
                             lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        help_content = tk.Frame(scrollable_frame, bg='white')
        help_content.pack(padx=30, pady=20, fill=tk.BOTH, expand=True)
        
        help_sections = [
            ("📝 RUJUKAN", "Format: KE.JB(90)650/05-02/XXXX"),
            ("🏢 NAMA SYARIKAT", "Masukkan nama penuh syarikat."),
            ("📅 TARIKH", "Format: DD/MM/YYYY"),
            ("📋 JADUAL AMES", "Tambah/edit/delete item."),
            ("⚡ FEATURES", "Validation, tooltips"),
        ]
        
        for title, content in help_sections:
            section_frame = tk.Frame(help_content, bg='white')
            section_frame.pack(fill=tk.X, pady=10)
            
            tk.Label(section_frame, text=title, font=('Arial', 11, 'bold'),
                    bg='white', fg=self.colors['primary'], anchor='w').pack(fill=tk.X)
            
            tk.Label(section_frame, text=content, font=('Arial', 10),
                    bg='white', fg='#333333', justify=tk.LEFT, anchor='w').pack(fill=tk.X, padx=20, pady=5)
            
            tk.Frame(section_frame, bg='#EEEEEE', height=1).pack(fill=tk.X, pady=5)
        
        btn_frame = tk.Frame(help_window, bg='white')
        btn_frame.pack(fill=tk.X, padx=30, pady=20)
        
        tk.Button(btn_frame, text="Tutup", font=('Arial', 10, 'bold'),
                 bg=self.colors['primary'], fg='white', padx=30, pady=10,
                 cursor='hand2', command=help_window.destroy).pack()
        
        help_window.transient(self.root)
        help_window.grab_set()
        
        help_window.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - help_window.winfo_width()) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - help_window.winfo_height()) // 2
        help_window.geometry(f"+{x}+{y}")
    
    def on_close(self):
        """Handle window close"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            pass
        self.root.destroy()


class RowDialog:
    """Dialog for adding/editing table rows"""
    def __init__(self, parent, title, initial_values=None, with_nisbah=False):
        self.result = None
        self.with_nisbah = with_nisbah
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        height = 380 if with_nisbah else 300
        self.dialog.geometry(f"600x{height}")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.configure(bg='white')
        
        # Center dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - 300
        y = (self.dialog.winfo_screenheight() // 2) - (height // 2)
        self.dialog.geometry(f'600x{height}+{x}+{y}')
        
        # Make window appear on top
        self.dialog.attributes('-topmost', True)
        self.dialog.after_idle(lambda: self.dialog.attributes('-topmost', False))
        self.dialog.lift()
        self.dialog.focus_force()
        
        # Government colors
        colors = {
            'primary': '#003366',
            'success': '#2E7D32',
            'border': '#CCCCCC'
        }
        
        # Title
        title_label = tk.Label(self.dialog,
                            text=title,
                            font=('Arial', 12, 'bold'),
                            bg='white',
                            fg=colors['primary'])
        title_label.pack(pady=(20, 20))
        
        # Form frame
        form_frame = tk.Frame(self.dialog, bg='white')
        form_frame.pack(padx=40, pady=10)
        
        row = 0
        
        # Kod Tarif
        tk.Label(form_frame, text="Kod Tarif:", font=('Arial', 10, 'bold'),
                bg='white').grid(row=row, column=0, padx=10, pady=12, sticky='e')
        self.entry_kod = tk.Entry(form_frame, width=45, font=('Arial', 10), relief=tk.SOLID, bd=1)
        self.entry_kod.grid(row=row, column=1, padx=10, pady=12)
        row += 1
        
        # Deskripsi
        tk.Label(form_frame, text="Deskripsi:", font=('Arial', 10, 'bold'),
                bg='white').grid(row=row, column=0, padx=10, pady=12, sticky='e')
        self.entry_desk = tk.Entry(form_frame, width=45, font=('Arial', 10), relief=tk.SOLID, bd=1)
        self.entry_desk.grid(row=row, column=1, padx=10, pady=12)
        row += 1
        
        # NISBAH (only for bahan)
        if with_nisbah:
            tk.Label(form_frame, text="Nisbah:", font=('Arial', 10, 'bold'),
                    bg='white').grid(row=row, column=0, padx=10, pady=12, sticky='e')
            self.entry_nisbah = tk.Entry(form_frame, width=45, font=('Arial', 10), relief=tk.SOLID, bd=1)
            self.entry_nisbah.insert(0, "1:1")
            self.entry_nisbah.grid(row=row, column=1, padx=10, pady=12)
            row += 1
        
        # Tarikh Kuatkuasa
        tk.Label(form_frame, text="Tarikh Kuatkuasa:", font=('Arial', 10, 'bold'),
                bg='white').grid(row=row, column=0, padx=10, pady=12, sticky='e')
        self.entry_tarikh = tk.Entry(form_frame, width=45, font=('Arial', 10), relief=tk.SOLID, bd=1)
        self.entry_tarikh.insert(0, "01 APRIL 2025")
        self.entry_tarikh.grid(row=row, column=1, padx=10, pady=12)
        
        # Pre-fill if editing
        if initial_values:
            if len(initial_values) >= 3:
                self.entry_kod.delete(0, tk.END)
                self.entry_kod.insert(0, initial_values[0])
                self.entry_desk.delete(0, tk.END)
                self.entry_desk.insert(0, initial_values[1])
                if with_nisbah and len(initial_values) >= 4:
                    self.entry_nisbah.delete(0, tk.END)
                    self.entry_nisbah.insert(0, initial_values[3])
                self.entry_tarikh.delete(0, tk.END)
                self.entry_tarikh.insert(0, initial_values[2] if not with_nisbah else initial_values[-1])
        
        # Buttons
        btn_frame = tk.Frame(self.dialog, bg='white')
        btn_frame.pack(pady=25)
        
        tk.Button(btn_frame, text="Simpan", width=18, height=2, font=('Arial', 12, 'bold'),
                bg=colors['success'], fg='white', relief=tk.FLAT, cursor='hand2',
                command=self.on_ok).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Batal", width=18, height=2, font=('Arial', 12),
                bg='#666666', fg='white', relief=tk.FLAT, cursor='hand2',
                command=self.on_cancel).pack(side=tk.LEFT, padx=8)
        
        # Focus on first field
        self.entry_kod.focus_set()
        
        # Wait for dialog
        self.dialog.wait_window()
    
    def on_ok(self):
        """Save and close"""
        kod = self.entry_kod.get().strip()
        desk = self.entry_desk.get().strip()
        tarikh = self.entry_tarikh.get().strip()
        
        if not kod or not desk:
            messagebox.showwarning("Amaran", "Sila isi Kod Tarif dan Deskripsi")
            return
        
        if self.with_nisbah:
            nisbah = self.entry_nisbah.get().strip() or "1:1"
            self.result = (kod, desk, tarikh, nisbah)
        else:
            self.result = (kod, desk, tarikh)
        
        self.dialog.destroy()
    
    def on_cancel(self):
        """Cancel and close"""
        self.dialog.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = Form3(root)
    root.mainloop()