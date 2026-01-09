"""
Form3_Government_PySide2.py - Form3 AMES (PyQt5)
Optimized OOP structure with lazy loading to reduce app load
"""
import json
import os
import sys
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from hijri_converter import Gregorian, Hijri
from PIL import Image

from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import (QAbstractItemView, QComboBox, QDesktopWidget, QDialog,
                                QFileDialog, QFrame, QGridLayout, QHBoxLayout, QHeaderView,
                                QLabel, QLineEdit, QMessageBox, QPushButton, QScrollArea,
                                QTableWidget, QTableWidgetItem, QTextEdit, QVBoxLayout, QWidget)

from helpers.resource_path import get_logo_path, get_template_path
from helpers.unified_database import UnifiedDatabase

# ============================================
# CUSTOM WIDGETS
# ============================================

class SafeLineEdit(QLineEdit):
    """QLineEdit that doesn't close dialog on Enter/Return key"""
    def keyPressEvent(self, event):
        # Allow Tab for moving to next field
        if event.key() == Qt.Key_Tab:
            self.focusNextChild()
            return
        # Allow Shift+Tab for moving to previous field
        elif event.key() == Qt.Key_Backtab:
            self.focusPreviousChild()
            return
        # Ignore Return/Enter - don't let it close the dialog
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            event.ignore()
            return
        # Handle everything else normally
        super().keyPressEvent(event)

try:
    from helpers.docx_helper import replace_in_document
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value) if value else "")
        return {}

try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class RowDialog(QDialog):
    """Dialog for adding/editing table rows"""
    def __init__(self, parent, title, initial_values=None, with_nisbah=False):
        super().__init__(parent)
        self.result = None
        self.with_nisbah = with_nisbah
        self.setWindowTitle(title)
        self.setModal(True)
        
        height = 380 if with_nisbah else 300
        self.setFixedSize(600, height)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(40, 30, 40, 30)
        
        # Title
        title_label = QLabel(title)
        title_label.setStyleSheet("color: #003366; font-size: 12px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Form
        form_widget = QWidget()
        form_layout = QGridLayout(form_widget)
        form_layout.setSpacing(12)
        
        row = 0
        
        # Kod Tarif
        form_layout.addWidget(QLabel("Kod Tarif:"), row, 0)
        self.entry_kod = SafeLineEdit()
        self.entry_kod.setMinimumWidth(400)
        form_layout.addWidget(self.entry_kod, row, 1)
        row += 1
        
        # Deskripsi
        form_layout.addWidget(QLabel("Deskripsi:"), row, 0)
        self.entry_desk = SafeLineEdit()
        form_layout.addWidget(self.entry_desk, row, 1)
        row += 1
        
        # NISBAH (only for bahan)
        if with_nisbah:
            form_layout.addWidget(QLabel("Nisbah:"), row, 0)
            self.entry_nisbah = SafeLineEdit()
            self.entry_nisbah.setText("1:1")
            form_layout.addWidget(self.entry_nisbah, row, 1)
            row += 1
        
        # Tarikh Kuatkuasa
        form_layout.addWidget(QLabel("Tarikh Kuatkuasa:"), row, 0)
        self.entry_tarikh = SafeLineEdit()
        self.entry_tarikh.setText("01 APRIL 2025")
        form_layout.addWidget(self.entry_tarikh, row, 1)
        
        # Pre-fill if editing
        if initial_values:
            if len(initial_values) >= 3:
                self.entry_kod.setText(initial_values[0])
                self.entry_desk.setText(initial_values[1])
                if with_nisbah and len(initial_values) >= 4:
                    self.entry_nisbah.setText(initial_values[3])
                self.entry_tarikh.setText(initial_values[2] if not with_nisbah else initial_values[-1])
        
        layout.addWidget(form_widget)
        layout.addStretch()
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)
        
        btn_save = QPushButton("Simpan")
        btn_save.setStyleSheet("""
            QPushButton {
                background-color: #2E7D32;
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #1B5E20;
            }
        """)
        btn_save.clicked.connect(self.on_ok)
        btn_layout.addWidget(btn_save)
        
        btn_cancel = QPushButton("Batal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #666666;
                color: white;
                font-size: 12px;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #555555;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        layout.addLayout(btn_layout)
        
        self.entry_kod.setFocus()
    
    def on_ok(self):
        """Save and close"""
        kod = self.entry_kod.text().strip()
        desk = self.entry_desk.text().strip()
        tarikh = self.entry_tarikh.text().strip()
        
        if not kod or not desk:
            QMessageBox.warning(self, "Amaran", "Sila isi Kod Tarif dan Deskripsi")
            return
        
        if self.with_nisbah:
            nisbah = self.entry_nisbah.text().strip() or "1:1"
            self.result = (kod, desk, tarikh, nisbah)
        else:
            self.result = (kod, desk, tarikh)
        
        self.accept()


class VehicleDialog(QDialog):
    """Dialog for adding/editing vehicle rows (Butiran 5D)"""
    def __init__(self, parent, title, initial_values=None):
        super().__init__(parent)
        self.result = None
        self.setWindowTitle(title)
        self.setModal(True)
        self.setFixedSize(600, 350)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(40, 30, 40, 30)
        
        # Title
        title_label = QLabel(title)
        title_label.setStyleSheet("color: #003366; font-size: 12px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Form
        form_widget = QWidget()
        form_layout = QGridLayout(form_widget)
        form_layout.setSpacing(12)
        
        row = 0
        
        # No Chasis (only 2 fields as per document example)
        form_layout.addWidget(QLabel("No Chasis:"), row, 0)
        self.entry_chasis = SafeLineEdit()
        self.entry_chasis.setMinimumWidth(400)
        form_layout.addWidget(self.entry_chasis, row, 1)
        row += 1
        
        # No Enjin
        form_layout.addWidget(QLabel("No Enjin:"), row, 0)
        self.entry_enjin = SafeLineEdit()
        form_layout.addWidget(self.entry_enjin, row, 1)
        
        # Pre-fill if editing
        if initial_values:
            if len(initial_values) >= 2:
                self.entry_chasis.setText(initial_values[0])
                self.entry_enjin.setText(initial_values[1])
        
        layout.addWidget(form_widget)
        layout.addStretch()
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)
        
        btn_save = QPushButton("Simpan")
        btn_save.setStyleSheet("""
            QPushButton {
                background-color: #2E7D32;
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #1B5E20;
            }
        """)
        btn_save.clicked.connect(self.on_ok)
        btn_layout.addWidget(btn_save)
        
        btn_cancel = QPushButton("Batal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #666666;
                color: white;
                font-size: 12px;
                padding: 10px 30px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #555555;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        
        layout.addLayout(btn_layout)
        
        self.entry_chasis.setFocus()
    
    def on_ok(self):
        """Save and close"""
        chasis = self.entry_chasis.text().strip()
        enjin = self.entry_enjin.text().strip()
        
        if not chasis or not enjin:
            QMessageBox.warning(self, "Amaran", "Sila isi No. Chasis dan No. Enjin")
            return
        
        self.result = (chasis, enjin)
        self.accept()


class Form3(QDialog):
    """Government-styled AMES Form with PySide2"""
    
    def __init__(self, parent_window=None):
        super().__init__()
        self.parent_window = parent_window
        self.setWindowTitle("Sistem Pengurusan AMES - Permohonan Kelulusan")
        self.setGeometry(100, 100, 1600, 1000)
        self.setMinimumSize(1200, 800)  # Prevent window from shrinking too much
        
        # Add borders to text boxes (QTextEdit) only
        self.setStyleSheet("""
            QFrame {
                border: none;
            }
            QLineEdit, QComboBox, QTextEdit {
                background-color: #F5F5F5;
                border: 1px solid #CCCCCC;
                font-size: 14px;
            }
        """)
        
        self.db = UnifiedDatabase()
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': '#FFFFFF',
            'text_dark': '#1a1a1a',
            'text_light': '#FFFFFF',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_danger': '#C62828',
            'button_hover': '#004d99',
        }
        
        # Table data storage
        self.table_pedagang_data = []
        self.table_bahan_data = []
        self.table_barang_data = []
        self.table_vehicles_data = []  # Butiran 5D vehicle data
        
        self.create_ui()
        self.center_window()
        self.update_tarikh_islam()
    
    def create_ui(self):
        """Create the main UI"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Header
        self.create_header(main_layout)
        
        # Scrollable content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet(f"background-color: {self.colors['bg_main']}; border: none;")
        
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(30, 20, 30, 20)
        scroll_layout.setSpacing(20)
        
        # Form section
        self.create_form_section(scroll_layout)
        
        # Table section
        self.create_table_section(scroll_layout)
        
        # Button section
        self.create_button_section(scroll_layout)
        
        scroll_area.setWidget(scroll_widget)
        main_layout.addWidget(scroll_area, stretch=1)
    
    def create_header(self, parent_layout):
        """Create government header"""
        header_frame = QFrame()
        header_frame.setFixedHeight(100)
        header_frame.setStyleSheet(f"background-color: {self.colors['primary']};")
        
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(20, 10, 20, 10)
        header_layout.setSpacing(0)
        
        # Logo
        try:
            logo_image = Image.open(get_logo_path())
            if logo_image.mode != 'RGBA':
                logo_image = logo_image.convert('RGBA')
            logo_image = logo_image.resize((60, 60), Image.Resampling.LANCZOS)
            
            bg_rgb = tuple(int(self.colors['primary'][i:i+2], 16) for i in (1, 3, 5))
            background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
            if logo_image.mode == 'RGBA':
                logo_image = Image.alpha_composite(background, logo_image)
                logo_image = logo_image.convert('RGB')
            
            buf = BytesIO()
            logo_image.save(buf, format='PNG')
            pixmap = QPixmap()
            pixmap.loadFromData(buf.getvalue())
            
            logo_label = QLabel()
            logo_label.setPixmap(pixmap.scaled(60, 60, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            logo_label.setStyleSheet(f"background-color: {self.colors['primary']};")
            header_layout.addWidget(logo_label)
        except:
            pass
        
        # Title
        title_widget = QWidget()
        title_layout = QVBoxLayout(title_widget)
        title_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("AMES - PERMOHONAN SKIM PENGEKSPORT UTAMA")
        title.setStyleSheet(f"background-color: {self.colors['primary']}; color: white; font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(title)
        
        subtitle = QLabel("Approved Major Exporter Scheme")
        subtitle.setStyleSheet(f"background-color: {self.colors['primary']}; color: #E0E0E0; font-size: 12px; font-style: italic;")
        subtitle.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(subtitle)
        
        header_layout.addWidget(title_widget, stretch=1)
        
        # Buttons
        btn_help = QPushButton("❓ HELP")
        btn_help.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['accent']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 8px 15px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: {self.colors['button_hover']};
            }}
        """)
        btn_help.clicked.connect(self.show_help_dialog)
        header_layout.addWidget(btn_help)
        
        btn_back = QPushButton("← KEMBALI")
        btn_back.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['secondary']};
                color: white;
                font-size: 12px;
                font-weight: bold;
                padding: 8px 20px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: {self.colors['button_hover']};
            }}
        """)
        btn_back.clicked.connect(self.on_back_click)
        header_layout.addWidget(btn_back)
        
        parent_layout.addWidget(header_frame)
        
        # Separator
        separator = QFrame()
        separator.setFixedHeight(2)
        separator.setStyleSheet(f"background-color: {self.colors['accent']};")
        parent_layout.addWidget(separator)
    
    def create_form_section(self, parent_layout):
        """Create form input section"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(25, 20, 25, 20)
        card_layout.setSpacing(15)
        
        # Section title
        section_title = QLabel("Maklumat Permohonan")
        section_title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 16px; font-weight: bold;")
        card_layout.addWidget(section_title)
        
        # Form grid
        form_grid = QGridLayout()
        form_grid.setSpacing(10)
        
        row = 0
        
        # Rujukan
        form_grid.addWidget(QLabel("RUJUKAN KAMI:"), row, 0)
        rujukan_widget = QWidget()
        rujukan_layout = QHBoxLayout(rujukan_widget)
        rujukan_layout.setContentsMargins(0, 0, 0, 0)
        rujukan_layout.addWidget(QLabel("KE.JB(90)650/14/AMES/"))
        self.entry_rujukan = SafeLineEdit()
        self.entry_rujukan.setMaximumWidth(200)
        rujukan_layout.addWidget(self.entry_rujukan)
        form_grid.addWidget(rujukan_widget, row, 1)
        
        # Alamat
        form_grid.addWidget(QLabel("ALAMAT:"), row, 2)
        self.entry_alamat1 = SafeLineEdit()
        self.entry_alamat1.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat1, row, 3)
        row += 1
        
        # Nama Syarikat
        form_grid.addWidget(QLabel("NAMA SYARIKAT:"), row, 0)
        self.entry_nama = SafeLineEdit()
        form_grid.addWidget(self.entry_nama, row, 1)
        
        self.entry_alamat2 = SafeLineEdit()
        self.entry_alamat2.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat2, row, 3)
        row += 1
        
        # Tarikh
        form_grid.addWidget(QLabel("TARIKH:"), row, 0)
        self.entry_tarikh = QLineEdit()
        self.entry_tarikh.setText(datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.textChanged.connect(self.update_tarikh_islam)
        form_grid.addWidget(self.entry_tarikh, row, 1)
        
        self.entry_alamat3 = SafeLineEdit()
        self.entry_alamat3.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_alamat3, row, 3)
        row += 1
        
        # Tarikh Islam & Nama Pegawai
        form_grid.addWidget(QLabel("TARIKH ISLAM:"), row, 0)
        self.entry_islam = SafeLineEdit()
        self.entry_islam.setReadOnly(True)
        form_grid.addWidget(self.entry_islam, row, 1)
        
        form_grid.addWidget(QLabel("NAMA PEGAWAI:"), row, 2)
        self.combo_pegawai = QComboBox()
        self.combo_pegawai.setEditable(True)  # Allow typing
        self.combo_pegawai.addItems(["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                    "KHAIRONE DZARWANY BINTI MOHD KAIRI"])
        form_grid.addWidget(self.combo_pegawai, row, 3)
        row += 1
        
        # Divider
        divider = QFrame()
        divider.setFrameShape(QFrame.HLine)
        divider.setStyleSheet(f"color: {self.colors['border']};")
        form_grid.addWidget(divider, row, 0, 1, 4)
        row += 1
        
        # No. Kelulusan & Kategori
        form_grid.addWidget(QLabel("NO. KELULUSAN AMES:"), row, 0)
        self.entry_kelulusan = SafeLineEdit()
        form_grid.addWidget(self.entry_kelulusan, row, 1)
        
        form_grid.addWidget(QLabel("KATEGORI:"), row, 2)
        self.combo_kategori = QComboBox()
        self.combo_kategori.addItems(["Pedagang", "Pengilang", "Butiran 5D"])
        self.combo_kategori.currentTextChanged.connect(self.refresh_table_layout)
        form_grid.addWidget(self.combo_kategori, row, 3)
        row += 1
        
        # Butiran 5D specific field: Jenama / Model (hidden by default)
        self.label_jenama = QLabel("JENAMA / MODEL:")
        self.label_jenama.hide()
        form_grid.addWidget(self.label_jenama, row, 0)
        self.entry_jenama_model = SafeLineEdit()
        self.entry_jenama_model.hide()
        self.entry_jenama_model.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 14px;")
        form_grid.addWidget(self.entry_jenama_model, row, 1, 1, 3)
        row += 1
        
        # Tempoh Kelulusan
        form_grid.addWidget(QLabel("TEMPOH KELULUSAN:"), row, 0)
        tempoh_widget = QWidget()
        tempoh_layout = QHBoxLayout(tempoh_widget)
        tempoh_layout.setContentsMargins(0, 0, 0, 0)
        self.entry_tarikh_mula = SafeLineEdit()
        tempoh_layout.addWidget(self.entry_tarikh_mula)
        tempoh_layout.addWidget(QLabel("hingga"))
        self.entry_tarikh_tamat = SafeLineEdit()
        tempoh_layout.addWidget(self.entry_tarikh_tamat)
        form_grid.addWidget(tempoh_widget, row, 1, 1, 3)
        
        card_layout.addLayout(form_grid)
        parent_layout.addWidget(card)
    
    def create_table_section(self, parent_layout):
        """Create table section"""
        self.table_container = QWidget()
        self.table_layout = QVBoxLayout(self.table_container)
        self.table_layout.setContentsMargins(0, 0, 0, 0)
        
        # Default: Pedagang table
        self.create_pedagang_table()
        
        parent_layout.addWidget(self.table_container)
    
    def refresh_table_layout(self):
        """Rebuild layout when category changes"""
        # Clear existing tables
        for i in reversed(range(self.table_layout.count())):
            self.table_layout.itemAt(i).widget().setParent(None)
        
        kategori = self.combo_kategori.currentText().lower()
        
        # Show/hide Butiran 5D specific fields
        if kategori == "butiran 5d":
            self.label_jenama.show()
            self.entry_jenama_model.show()
        else:
            self.label_jenama.hide()
            self.entry_jenama_model.hide()
        
        if kategori == "pengilang":
            self.create_pengilang_tables()
        elif kategori == "butiran 5d":
            self.create_butiran5d_table()
        else:
            self.create_pedagang_table()
    
    def create_pedagang_table(self):
        """Create single editable table for Pedagang"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(25, 20, 25, 20)
        
        title = QLabel("SENARAI BARANG-BARANG YANG DILULUSKAN (Pedagang)")
        title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 14px; font-weight: bold;")
        card_layout.addWidget(title)
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setAlignment(Qt.AlignRight)
        
        btn_add = QPushButton("+ Tambah")
        btn_add.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_add.clicked.connect(self.add_row_pedagang)
        btn_layout.addWidget(btn_add)
        
        btn_edit = QPushButton("✎ Edit")
        btn_edit.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_primary']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_edit.clicked.connect(self.edit_row_pedagang)
        btn_layout.addWidget(btn_edit)
        
        btn_delete = QPushButton("✕ Hapus")
        btn_delete.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_danger']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_delete.clicked.connect(self.delete_row_pedagang)
        btn_layout.addWidget(btn_delete)
        
        card_layout.addLayout(btn_layout)
        
        # Table
        self.table_pedagang = QTableWidget()
        self.table_pedagang.setColumnCount(4)
        self.table_pedagang.setHorizontalHeaderLabels(['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA'])
        self.table_pedagang.horizontalHeader().setStretchLastSection(True)
        self.table_pedagang.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_pedagang.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_pedagang.setAlternatingRowColors(False)
        # Ensure all rows have white background (no stripes)
        self.table_pedagang.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
            }
            QTableWidget::item {
                background-color: white;
            }
        """)
        card_layout.addWidget(self.table_pedagang)
        
        self.table_layout.addWidget(card)
    
    def create_pengilang_tables(self):
        """Create dual editable tables for Pengilang"""
        # Section A
        card_a = QFrame()
        card_a.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_a_layout = QVBoxLayout(card_a)
        card_a_layout.setContentsMargins(25, 20, 25, 20)
        
        title_a = QLabel("A. BAHAN MENTAH, KOMPONEN, BAHAN BUNGKUSAN DAN PEMBUNGKUSAN")
        title_a.setStyleSheet(f"color: {self.colors['primary']}; font-size: 13px; font-weight: bold;")
        card_a_layout.addWidget(title_a)
        
        # Buttons for A
        btn_layout_a = QHBoxLayout()
        btn_layout_a.setAlignment(Qt.AlignRight)
        
        for text, cmd in [("+ Tambah", self.add_row_bahan), 
                         ("✎ Edit", self.edit_row_bahan),
                         ("✕ Hapus", self.delete_row_bahan)]:
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {self.colors['button_success'] if '+ Tambah' in text 
                                      else self.colors['button_primary'] if 'Edit' in text 
                                      else self.colors['button_danger']};
                    color: white;
                    font-size: 9px;
                    font-weight: bold;
                    padding: 5px 10px;
                    border-radius: 3px;
                }}
            """)
            btn.clicked.connect(cmd)
            btn_layout_a.addWidget(btn)
        
        card_a_layout.addLayout(btn_layout_a)
        
        self.table_bahan = QTableWidget()
        self.table_bahan.setColumnCount(5)
        self.table_bahan.setHorizontalHeaderLabels(['BIL', 'KOD TARIF', 'DESKRIPSI', 'NISBAH', 'TARIKH KUATKUASA'])
        self.table_bahan.horizontalHeader().setStretchLastSection(True)
        self.table_bahan.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_bahan.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_bahan.setAlternatingRowColors(False)
        # Ensure all rows have white background (no stripes)
        self.table_bahan.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
            }
            QTableWidget::item {
                background-color: white;
            }
        """)
        card_a_layout.addWidget(self.table_bahan)
        
        self.table_layout.addWidget(card_a)
        
        # Section B
        card_b = QFrame()
        card_b.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_b_layout = QVBoxLayout(card_b)
        card_b_layout.setContentsMargins(25, 20, 25, 20)
        
        title_b = QLabel("B. BARANG SIAP YANG DIKILANGKAN")
        title_b.setStyleSheet(f"color: {self.colors['primary']}; font-size: 13px; font-weight: bold;")
        card_b_layout.addWidget(title_b)
        
        # Buttons for B
        btn_layout_b = QHBoxLayout()
        btn_layout_b.setAlignment(Qt.AlignRight)
        
        for text, cmd in [("+ Tambah", self.add_row_barang), 
                         ("✎ Edit", self.edit_row_barang),
                         ("✕ Hapus", self.delete_row_barang)]:
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {self.colors['button_success'] if '+ Tambah' in text 
                                      else self.colors['button_primary'] if 'Edit' in text 
                                      else self.colors['button_danger']};
                    color: white;
                    font-size: 9px;
                    font-weight: bold;
                    padding: 5px 10px;
                    border-radius: 3px;
                }}
            """)
            btn.clicked.connect(cmd)
            btn_layout_b.addWidget(btn)
        
        card_b_layout.addLayout(btn_layout_b)
        
        self.table_barang = QTableWidget()
        self.table_barang.setColumnCount(4)
        self.table_barang.setHorizontalHeaderLabels(['BIL', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA'])
        self.table_barang.horizontalHeader().setStretchLastSection(True)
        self.table_barang.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_barang.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_barang.setAlternatingRowColors(False)
        # Ensure all rows have white background (no stripes)
        self.table_barang.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
            }
            QTableWidget::item {
                background-color: white;
            }
        """)
        card_b_layout.addWidget(self.table_barang)
        
        self.table_layout.addWidget(card_b)
    
    def create_butiran5d_table(self):
        """Create table for Butiran 5D vehicles"""
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: white;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(25, 20, 25, 20)
        
        title = QLabel("SENARAI KENDERAAN (Butiran 5D)")
        title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 14px; font-weight: bold;")
        card_layout.addWidget(title)
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setAlignment(Qt.AlignRight)
        
        btn_add = QPushButton("+ Tambah")
        btn_add.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_add.clicked.connect(self.add_row_vehicle)
        btn_layout.addWidget(btn_add)
        
        btn_edit = QPushButton("✎ Edit")
        btn_edit.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_primary']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_edit.clicked.connect(self.edit_row_vehicle)
        btn_layout.addWidget(btn_edit)
        
        btn_delete = QPushButton("✕ Hapus")
        btn_delete.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_danger']};
                color: white;
                font-size: 9px;
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 3px;
            }}
        """)
        btn_delete.clicked.connect(self.delete_row_vehicle)
        btn_layout.addWidget(btn_delete)
        
        card_layout.addLayout(btn_layout)
        
        # Table (only 2 columns: No. Chasis and No. Enjin as per document example)
        self.table_vehicles = QTableWidget()
        self.table_vehicles.setColumnCount(2)
        self.table_vehicles.setHorizontalHeaderLabels(['NO CHASIS', 'NO ENJIN'])
        self.table_vehicles.horizontalHeader().setStretchLastSection(True)
        self.table_vehicles.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table_vehicles.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table_vehicles.setAlternatingRowColors(False)
        self.table_vehicles.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
            }
            QTableWidget::item {
                background-color: white;
            }
        """)
        card_layout.addWidget(self.table_vehicles)
        
        self.table_layout.addWidget(card)
    
    def add_row_pedagang(self):
        """Add row to pedagang table"""
        dialog = RowDialog(self, "Tambah Barangan (Pedagang)")
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh = dialog.result
            row = self.table_pedagang.rowCount()
            self.table_pedagang.insertRow(row)
            self.table_pedagang.setItem(row, 0, QTableWidgetItem(str(row + 1)))
            self.table_pedagang.setItem(row, 1, QTableWidgetItem(kod))
            self.table_pedagang.setItem(row, 2, QTableWidgetItem(desk))
            self.table_pedagang.setItem(row, 3, QTableWidgetItem(tarikh))
            self.table_pedagang_data.append((kod, desk, tarikh))
    
    def edit_row_pedagang(self):
        """Edit selected row in pedagang table"""
        selected = self.table_pedagang.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk diedit")
            return
        
        row = selected[0].row()
        kod = self.table_pedagang.item(row, 1).text()
        desk = self.table_pedagang.item(row, 2).text()
        tarikh = self.table_pedagang.item(row, 3).text()
        
        dialog = RowDialog(self, "Edit Barangan (Pedagang)", (kod, desk, tarikh))
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh = dialog.result
            self.table_pedagang.setItem(row, 1, QTableWidgetItem(kod))
            self.table_pedagang.setItem(row, 2, QTableWidgetItem(desk))
            self.table_pedagang.setItem(row, 3, QTableWidgetItem(tarikh))
            if row < len(self.table_pedagang_data):
                self.table_pedagang_data[row] = (kod, desk, tarikh)
    
    def delete_row_pedagang(self):
        """Delete selected row from pedagang table"""
        selected = self.table_pedagang.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk dihapus")
            return
        
        if QMessageBox.question(self, "Pengesahan", "Adakah anda pasti ingin menghapus baris ini?") == QMessageBox.Yes:
            row = selected[0].row()
            self.table_pedagang.removeRow(row)
            if row < len(self.table_pedagang_data):
                self.table_pedagang_data.pop(row)
            # Renumber
            for i in range(self.table_pedagang.rowCount()):
                self.table_pedagang.setItem(i, 0, QTableWidgetItem(str(i + 1)))
    
    def add_row_bahan(self):
        """Add row to bahan table"""
        dialog = RowDialog(self, "Tambah Bahan Mentah", with_nisbah=True)
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh, nisbah = dialog.result
            row = self.table_bahan.rowCount()
            self.table_bahan.insertRow(row)
            self.table_bahan.setItem(row, 0, QTableWidgetItem(str(row + 1)))
            self.table_bahan.setItem(row, 1, QTableWidgetItem(kod))
            self.table_bahan.setItem(row, 2, QTableWidgetItem(desk))
            self.table_bahan.setItem(row, 3, QTableWidgetItem(nisbah))
            self.table_bahan.setItem(row, 4, QTableWidgetItem(tarikh))
            self.table_bahan_data.append((kod, desk, nisbah, tarikh))
    
    def edit_row_bahan(self):
        """Edit selected row in bahan table"""
        selected = self.table_bahan.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk diedit")
            return
        
        row = selected[0].row()
        kod = self.table_bahan.item(row, 1).text()
        desk = self.table_bahan.item(row, 2).text()
        nisbah = self.table_bahan.item(row, 3).text()
        tarikh = self.table_bahan.item(row, 4).text()
        
        dialog = RowDialog(self, "Edit Bahan Mentah", (kod, desk, tarikh, nisbah), with_nisbah=True)
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh, nisbah = dialog.result
            self.table_bahan.setItem(row, 1, QTableWidgetItem(kod))
            self.table_bahan.setItem(row, 2, QTableWidgetItem(desk))
            self.table_bahan.setItem(row, 3, QTableWidgetItem(nisbah))
            self.table_bahan.setItem(row, 4, QTableWidgetItem(tarikh))
            if row < len(self.table_bahan_data):
                self.table_bahan_data[row] = (kod, desk, nisbah, tarikh)
    
    def delete_row_bahan(self):
        """Delete selected row from bahan table"""
        selected = self.table_bahan.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk dihapus")
            return
        
        if QMessageBox.question(self, "Pengesahan", "Adakah anda pasti ingin menghapus baris ini?") == QMessageBox.Yes:
            row = selected[0].row()
            self.table_bahan.removeRow(row)
            if row < len(self.table_bahan_data):
                self.table_bahan_data.pop(row)
            # Renumber
            for i in range(self.table_bahan.rowCount()):
                self.table_bahan.setItem(i, 0, QTableWidgetItem(str(i + 1)))
    
    def add_row_barang(self):
        """Add row to barang table"""
        dialog = RowDialog(self, "Tambah Barang Siap")
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh = dialog.result
            row = self.table_barang.rowCount()
            self.table_barang.insertRow(row)
            self.table_barang.setItem(row, 0, QTableWidgetItem(str(row + 1)))
            self.table_barang.setItem(row, 1, QTableWidgetItem(kod))
            self.table_barang.setItem(row, 2, QTableWidgetItem(desk))
            self.table_barang.setItem(row, 3, QTableWidgetItem(tarikh))
            self.table_barang_data.append((kod, desk, tarikh))
    
    def edit_row_barang(self):
        """Edit selected row in barang table"""
        selected = self.table_barang.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk diedit")
            return
        
        row = selected[0].row()
        kod = self.table_barang.item(row, 1).text()
        desk = self.table_barang.item(row, 2).text()
        tarikh = self.table_barang.item(row, 3).text()
        
        dialog = RowDialog(self, "Edit Barang Siap", (kod, desk, tarikh))
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            kod, desk, tarikh = dialog.result
            self.table_barang.setItem(row, 1, QTableWidgetItem(kod))
            self.table_barang.setItem(row, 2, QTableWidgetItem(desk))
            self.table_barang.setItem(row, 3, QTableWidgetItem(tarikh))
            if row < len(self.table_barang_data):
                self.table_barang_data[row] = (kod, desk, tarikh)
    
    def delete_row_barang(self):
        """Delete selected row from barang table"""
        selected = self.table_barang.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk dihapus")
            return
        
        if QMessageBox.question(self, "Pengesahan", "Adakah anda pasti ingin menghapus baris ini?") == QMessageBox.Yes:
            row = selected[0].row()
            self.table_barang.removeRow(row)
            if row < len(self.table_barang_data):
                self.table_barang_data.pop(row)
            # Renumber
            for i in range(self.table_barang.rowCount()):
                self.table_barang.setItem(i, 0, QTableWidgetItem(str(i + 1)))
    
    def add_row_vehicle(self):
        """Add row to vehicle table (Butiran 5D) - only 2 columns: No. Chasis and No. Enjin"""
        dialog = VehicleDialog(self, "Tambah Kenderaan (Butiran 5D)")
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            chasis, enjin = dialog.result
            row = self.table_vehicles.rowCount()
            self.table_vehicles.insertRow(row)
            self.table_vehicles.setItem(row, 0, QTableWidgetItem(chasis))
            self.table_vehicles.setItem(row, 1, QTableWidgetItem(enjin))
            self.table_vehicles_data.append((chasis, enjin))
    
    def edit_row_vehicle(self):
        """Edit selected row in vehicle table"""
        selected = self.table_vehicles.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk diedit")
            return
        
        row = selected[0].row()
        chasis = self.table_vehicles.item(row, 0).text()
        enjin = self.table_vehicles.item(row, 1).text()
        
        dialog = VehicleDialog(self, "Edit Kenderaan (Butiran 5D)", (chasis, enjin))
        if dialog.exec_() == QDialog.Accepted and dialog.result:
            chasis, enjin = dialog.result
            self.table_vehicles.setItem(row, 0, QTableWidgetItem(chasis))
            self.table_vehicles.setItem(row, 1, QTableWidgetItem(enjin))
            if row < len(self.table_vehicles_data):
                self.table_vehicles_data[row] = (chasis, enjin)
    
    def delete_row_vehicle(self):
        """Delete selected row from vehicle table"""
        selected = self.table_vehicles.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Amaran", "Sila pilih baris untuk dihapus")
            return
        
        if QMessageBox.question(self, "Pengesahan", "Adakah anda pasti ingin menghapus baris ini?") == QMessageBox.Yes:
            row = selected[0].row()
            self.table_vehicles.removeRow(row)
            if row < len(self.table_vehicles_data):
                self.table_vehicles_data.pop(row)

    def create_button_section(self, parent_layout):
        """Create bottom button section"""
        button_widget = QWidget()
        button_layout = QHBoxLayout(button_widget)
        button_layout.setAlignment(Qt.AlignCenter)
        button_layout.setSpacing(10)
        
        btn_preview = QPushButton("👁 PRATONTON")
        btn_preview.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_secondary']};
                color: white;
                font-size: 10px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #555555;
            }}
        """)
        btn_preview.clicked.connect(self.on_preview_click)
        button_layout.addWidget(btn_preview)
        
        btn_save = QPushButton("💾 SIMPAN DOKUMEN")
        btn_save.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 10px;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #1B5E20;
            }}
        """)
        btn_save.clicked.connect(self.on_save_click)
        button_layout.addWidget(btn_save)
        
        parent_layout.addWidget(button_widget)
    
    def update_tarikh_islam(self):
        """Convert Gregorian to Hijri"""
        try:
            tarikh_str = self.entry_tarikh.text()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    self.entry_islam.setText(hijri_text)
        except:
            pass
    
    def format_tarikh_malay(self):
        """Format date in Malay"""
        try:
            tarikh_str = self.entry_tarikh.text()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                bulan_melayu = ["Januari", "Februari", "Mac", "April", "Mei", "Jun",
                               "Julai", "Ogos", "September", "Oktober", "November", "Disember"]
                return f"{day:02d} {bulan_melayu[month-1]} {year}"
        except:
            return self.entry_tarikh.text()
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview_ames.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':
                os.startfile(temp_path)
            elif os.name == 'posix':
                import subprocess
                subprocess.call(['open', temp_path])
            
            QMessageBox.information(self, "Pratonton", "Dokumen dibuka untuk pratonton")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        if not self.entry_nama.text().strip():
            QMessageBox.critical(self, "Ralat", "Sila isi Nama Syarikat")
            return
        
        if not self.entry_rujukan.text().strip():
            QMessageBox.critical(self, "Ralat", "Sila isi Rujukan")
            return
        
        doc = self.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.text().replace('/', '_').replace('\\', '_')
        kategori = self.combo_kategori.currentText()
        
        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Simpan Dokumen",
            f"AMES_{safe_name}_{kategori}.pdf",
            "PDF Files (*.pdf);;All Files (*.*)"
        )
        
        if not filename:
            return
        
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            temp_docx = "temp_ames_conversion.docx"
            doc.save(temp_docx)
            
            if PDF_AVAILABLE:
                convert(temp_docx, filename)
                os.remove(temp_docx)
            else:
                QMessageBox.warning(self, "Amaran",
                    "docx2pdf tidak tersedia. Dokumen disimpan sebagai .docx.\n\n"
                    "Sila install dengan command:\npip install docx2pdf")
                filename = filename.replace('.pdf', '.docx')
                doc.save(filename)
            
            # Save to database
            try:
                rujukan_kami = f"KE.JB(90)650/14/AMES/{self.entry_rujukan.text()}"
                # Combine the 3 alamat entries
                alamat_lines = []
                if self.entry_alamat1.text().strip():
                    alamat_lines.append(self.entry_alamat1.text().strip())
                if self.entry_alamat2.text().strip():
                    alamat_lines.append(self.entry_alamat2.text().strip())
                if self.entry_alamat3.text().strip():
                    alamat_lines.append(self.entry_alamat3.text().strip())
                alamat_text = "\n".join(alamat_lines)
                
                application_data = {
                    'category': kategori,
                    'sub_option': None,
                    'rujukan_kami': rujukan_kami,
                    'rujukan_tuan': '',
                    'nama_syarikat': self.entry_nama.text(),
                    'alamat': alamat_text,
                    'tarikh': self.entry_tarikh.text(),
                    'tarikh_islam': self.entry_islam.text(),
                    'nama_pegawai': self.combo_pegawai.currentText(),
                    'status': 'DILULUSKAN',
                    'document_path': filename,
                    'additional_data': {}
                }
                
                # Prepare details based on category
                if kategori == "Butiran 5D":
                    # Butiran 5D-specific details
                    butiran5d_details = {
                        'no_sijil': self.entry_kelulusan.text(),
                        'tarikh_kuatkuasa': self.entry_tarikh_mula.text(),
                        'sebab_tolak': ''  # Can be added if needed
                    }
                    
                    # Collect vehicles (only 2 columns: No. Chasis and No. Enjin)
                    vehicles = []
                    for row in range(self.table_vehicles.rowCount()):
                        vehicles.append({
                            'bil': row + 1,
                            'no_chasis': self.table_vehicles.item(row, 0).text(),
                            'no_enjin': self.table_vehicles.item(row, 1).text()
                        })
                    
                    butiran5d_details['vehicles'] = vehicles
                    
                    app_id = self.db.save_application('butiran5d', application_data, butiran5d_details)
                    
                    # Clear database cache so history viewer shows new data
                    if hasattr(self.db, 'clear_cache'):
                        self.db.clear_cache()
                    
                    QMessageBox.information(self, "Berjaya", 
                        f"Dokumen disimpan: {filename}\n\n"
                        f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
                else:
                    # AMES-specific details
                    ames_details = {
                        'no_kelulusan': self.entry_kelulusan.text(),
                        'kategori': kategori,
                        'tarikh_mula': self.entry_tarikh_mula.text(),
                        'tarikh_tamat': self.entry_tarikh_tamat.text(),
                        'tempoh_kelulusan': f"{self.entry_tarikh_mula.text()} hingga {self.entry_tarikh_tamat.text()}"
                    }
                    
                    # Collect items from tables
                    items = []
                    
                    # Get items from appropriate table(s) based on category
                    if kategori == "Pedagang":
                        for row in range(self.table_pedagang.rowCount()):
                            items.append({
                                'bil': row + 1,
                                'kod_tarif': self.table_pedagang.item(row, 1).text(),
                                'deskripsi': self.table_pedagang.item(row, 2).text(),
                                'tarikh_kuatkuasa': self.table_pedagang.item(row, 3).text()
                            })
                    else:  # Pengilang
                        for row in range(self.table_barang.rowCount()):
                            items.append({
                                'bil': row + 1,
                                'kod_tarif': self.table_barang.item(row, 1).text(),
                                'deskripsi': self.table_barang.item(row, 2).text(),
                                'tarikh_kuatkuasa': self.table_barang.item(row, 3).text(),
                                'nisbah': self.table_barang.item(row, 4).text() if self.table_barang.columnCount() > 4 else None
                            })
                        
                        for row in range(self.table_bahan.rowCount()):
                            items.append({
                                'bil': len(items) + 1,
                                'kod_tarif': self.table_bahan.item(row, 1).text(),
                                'deskripsi': self.table_bahan.item(row, 2).text(),
                                'tarikh_kuatkuasa': self.table_bahan.item(row, 4).text(),
                                'nisbah': self.table_bahan.item(row, 3).text()
                            })
                    
                    ames_details['items'] = items
                
                app_id = self.db.save_application('ames', application_data, ames_details)
                
                # Clear database cache so history viewer shows new data
                if hasattr(self.db, 'clear_cache'):
                    self.db.clear_cache()
                
                QMessageBox.information(self, "Berjaya", 
                    f"Dokumen disimpan: {filename}\n\n"
                    f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
            except Exception as db_error:
                QMessageBox.warning(self, "Amaran", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
            
            QMessageBox.information(self, "Berjaya", f"Dokumen disimpan: {filename}")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat menyimpan dokumen: {str(e)}")
            if os.path.exists("temp_ames_conversion.docx"):
                os.remove("temp_ames_conversion.docx")
    
    def generate_document(self):
        """Generate AMES or Butiran 5D document"""
        kategori = self.combo_kategori.currentText()
        if kategori == "Pedagang":
            template_file = "ames_pedagang.docx"
        elif kategori == "Butiran 5D":
            # Determine which Butiran 5D template based on some criteria
            # For now, default to "Lulus" - can be enhanced later
            template_file = "surat kelulusan butiran 5D (Lulus).docx"
        else:
            template_file = "ames_pengilang.docx"
        
        doc = None
        
        # Try to load template
        template_path = get_template_path(template_file)
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            QMessageBox.warning(self, "Amaran",
                f"Templat tidak dijumpai: {template_file}\n\nMenggunakan dokumen kosong...")
            doc = Document()
            doc.add_heading('PERMOHONAN SKIM PENGEKSPORT UTAMA DILULUSKAN (AMES)', 0)
        
        # Use PlaceholderBuilder for standardized placeholders
        from helpers.placeholder_registry import PlaceholderBuilder
        
        # Build replacements using PlaceholderBuilder
        builder = PlaceholderBuilder(self)
        builder.add_rujukan_ames(self.entry_rujukan.text())
        builder.add_nama_syarikat(self.entry_nama.text())
        builder.add_alamat([
            self.entry_alamat1.text().strip(),
            self.entry_alamat2.text().strip(),
            self.entry_alamat3.text().strip()
        ])
        builder.add_tarikh(self.entry_tarikh.text())
        builder.add_tarikh_malay(self.format_tarikh_malay())
        builder.add_tarikh_islam(self.entry_islam.text())
        builder.add_standard('NO_KELULUSAN', self.entry_kelulusan.text())
        builder.add_standard('KATEGORI', kategori)
        builder.add_standard('TEMPOH_KELULUSAN', f"{self.entry_tarikh_mula.text()} hingga {self.entry_tarikh_tamat.text()}")
        builder.add_nama_pegawai(self.combo_pegawai.currentText())
        
        # Butiran 5D specific fields
        if kategori == "Butiran 5D":
            builder.add_standard('NO_SIJIL_PENGECUALIAN', self.entry_kelulusan.text())
            # Format tarikh kuatkuasa (convert to Malay format if needed)
            tarikh_kuatkuasa = self.entry_tarikh_mula.text()
            try:
                # Try to format if it's in DD/MM/YYYY format
                if '/' in tarikh_kuatkuasa:
                    day, month, year = map(int, tarikh_kuatkuasa.split('/'))
                    bulan_melayu = ["JANUARI", "FEBRUARI", "MAC", "APRIL", "MEI", "JUN",
                                   "JULAI", "OGOS", "SEPTEMBER", "OKTOBER", "NOVEMBER", "DISEMBER"]
                    tarikh_kuatkuasa = f"{day:02d} {bulan_melayu[month-1]} {year}"
            except:
                pass
            builder.add_standard('TARIKH_KUATKUASA_SIJIL', tarikh_kuatkuasa)
            builder.add_standard('JENAMA_MODEL', self.entry_jenama_model.text() if hasattr(self, 'entry_jenama_model') else '')
        
        # Get final replacements dict
        replacements = builder.build()
        
        replace_in_document(doc, replacements)
        
        # Handle <<table>> placeholders - replace with actual data tables
        self.replace_table_placeholders(doc, kategori)
        
        return doc
    
    def add_table_to_document_pedagang(self, doc):
        """Add table for pedagang"""
        if self.table_pedagang.rowCount() == 0:
            return
        
        table = doc.add_table(rows=1 + self.table_pedagang.rowCount(), cols=4)
        headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
        
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(header)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for row_idx in range(self.table_pedagang.rowCount()):
            for col_idx in range(4):
                cell = table.rows[row_idx + 1].cells[col_idx]
                item = self.table_pedagang.item(row_idx, col_idx)
                cell.text = item.text() if item else ""
    
    def add_table_to_document_pengilang(self, doc):
        """Add tables for pengilang"""
        # Add bahan table
        if self.table_bahan.rowCount() > 0:
            doc.add_paragraph("A. BAHAN MENTAH, KOMPONEN, BAHAN BUNGKUSAN DAN PEMBUNGKUSAN")
            table_bahan = doc.add_table(rows=1 + self.table_bahan.rowCount(), cols=5)
            headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'NISBAH', 'TARIKH KUATKUASA']
            
            for idx, header in enumerate(headers):
                cell = table_bahan.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(header)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            for row_idx in range(self.table_bahan.rowCount()):
                for col_idx in range(5):
                    cell = table_bahan.rows[row_idx + 1].cells[col_idx]
                    item = self.table_bahan.item(row_idx, col_idx)
                    cell.text = item.text() if item else ""
        
        # Add barang table
        if self.table_barang.rowCount() > 0:
            doc.add_paragraph("B. BARANG SIAP YANG DIKILANGKAN")
            table_barang = doc.add_table(rows=1 + self.table_barang.rowCount(), cols=4)
            headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
            
            for idx, header in enumerate(headers):
                cell = table_barang.rows[0].cells[idx]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(header)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            for row_idx in range(self.table_barang.rowCount()):
                for col_idx in range(4):
                    cell = table_barang.rows[row_idx + 1].cells[col_idx]
                    item = self.table_barang.item(row_idx, col_idx)
                    cell.text = item.text() if item else ""
    
    def replace_table_placeholders(self, doc, kategori):
        """Replace <<table>> placeholders with actual data tables"""
        import re
        
        # Find all <<table>> placeholders and their parent elements
        placeholders_found = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Search for <<table>> or <<TABLE>> placeholders
            if re.search(r'<<\s*table\s*>>', paragraph.text, re.IGNORECASE):
                placeholders_found.append(i)
        
        # Process in reverse order so indices don't shift when we modify doc
        for para_idx in reversed(placeholders_found):
            paragraph = doc.paragraphs[para_idx]
            p_element = paragraph._element
            parent = p_element.getparent()
            
            # Get the table to insert based on category
            if kategori == "Pedagang":
                table_data = self.get_pedagang_table_data()
                headers = ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA']
                self._insert_table_after_element(parent, p_element, None, table_data, headers)
            
            elif kategori == "Butiran 5D":
                table_data = self.get_vehicles_table_data()
                headers = ['NO CHASIS', 'NO ENJIN']
                self._insert_table_after_element(parent, p_element, None, table_data, headers)
            
            else:  # Pengilang - insert both tables
                bahan_data = self.get_bahan_table_data()
                barang_data = self.get_barang_table_data()
                
                # Insert Barang section first (in reverse order)
                self._insert_table_after_element(parent, p_element,
                    "B. BARANG SIAP YANG DIKILANGKAN",
                    barang_data,
                    ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'TARIKH KUATKUASA'])
                
                # Insert Bahan section
                self._insert_table_after_element(parent, p_element,
                    "A. BAHAN MENTAH, KOMPONEN, BAHAN BUNGKUSAN DAN PEMBUNGKUSAN",
                    bahan_data,
                    ['BIL.', 'KOD TARIF', 'DESKRIPSI', 'NISBAH', 'TARIKH KUATKUASA'])
            
            # Remove the placeholder paragraph
            parent.remove(p_element)
    
    def _insert_table_after_element(self, parent, ref_element, title, table_data, headers):
        """Insert table into document after reference element"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Only insert if there's data
        if not table_data:
            return
        
        # Add title paragraph if provided
        if title:
            title_para = OxmlElement('w:p')
            parent.insert(parent.index(ref_element) + 1, title_para)
            
            # Create title paragraph element
            pPr = OxmlElement('w:pPr')
            title_para.append(pPr)
            
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            title_para.append(run)
            run.append(rPr)
            
            text = OxmlElement('w:t')
            text.text = title
            run.append(text)
        
        # Create table
        tbl = OxmlElement('w:tbl')
        
        # Add table properties
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
        
        # Create table grid (column widths)
        tblGrid = OxmlElement('w:tblGrid')
        for _ in headers:
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        tbl.append(tblGrid)
        
        # Add header row
        header_row = OxmlElement('w:tr')
        tbl.append(header_row)
        
        for header in headers:
            cell = OxmlElement('w:tc')
            header_row.append(cell)
            
            tcPr = OxmlElement('w:tcPr')
            cell.append(tcPr)
            
            para = OxmlElement('w:p')
            cell.append(para)
            
            pPr = OxmlElement('w:pPr')
            para.append(pPr)
            
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            
            run = OxmlElement('w:r')
            para.append(run)
            
            rPr = OxmlElement('w:rPr')
            rPr.append(OxmlElement('w:b'))  # Bold
            run.append(rPr)
            
            text = OxmlElement('w:t')
            text.text = header
            run.append(text)
        
        # Add data rows
        for row_data in table_data:
            data_row = OxmlElement('w:tr')
            tbl.append(data_row)
            
            for cell_value in row_data:
                cell = OxmlElement('w:tc')
                data_row.append(cell)
                
                tcPr = OxmlElement('w:tcPr')
                cell.append(tcPr)
                
                para = OxmlElement('w:p')
                cell.append(para)
                
                run = OxmlElement('w:r')
                para.append(run)
                
                rPr = OxmlElement('w:rPr')
                run.append(rPr)
                
                text = OxmlElement('w:t')
                text.text = str(cell_value) if cell_value else ""
                run.append(text)
        
        # Insert table after reference element
        parent.insert(parent.index(ref_element) + 1, tbl)
    
    def get_pedagang_table_data(self):
        """Extract pedagang table data as list of lists"""
        data = []
        for row in range(self.table_pedagang.rowCount()):
            row_data = []
            for col in range(4):
                item = self.table_pedagang.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        return data
    
    def get_bahan_table_data(self):
        """Extract bahan table data as list of lists"""
        data = []
        for row in range(self.table_bahan.rowCount()):
            row_data = []
            for col in range(5):
                item = self.table_bahan.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        return data
    
    def get_barang_table_data(self):
        """Extract barang table data as list of lists"""
        data = []
        for row in range(self.table_barang.rowCount()):
            row_data = []
            for col in range(4):
                item = self.table_barang.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        return data
    
    def get_vehicles_table_data(self):
        """Extract vehicles table data as list of lists"""
        data = []
        for row in range(self.table_vehicles.rowCount()):
            row_data = []
            for col in range(2):
                item = self.table_vehicles.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        return data
    
    def show_help_dialog(self):
        """Show help dialog"""
        QMessageBox.information(self, "Bantuan", 
            "Bantuan untuk menggunakan sistem AMES:\n\n"
            "1. Isi semua maklumat permohonan\n"
            "2. Pilih kategori (Pedagang atau Pengilang)\n"
            "3. Tambah item dalam jadual\n"
            "4. Klik Pratonton untuk melihat dokumen\n"
            "5. Klik Simpan untuk menyimpan sebagai PDF")
    
    def center_window(self):
        """Center window on screen"""
        frame_geometry = self.frameGeometry()
        try:
            desktop = QDesktopWidget()
            screen_geometry = desktop.availableGeometry(desktop.primaryScreen())
        except:
            from PyQt5.QtWidgets import QApplication
            screen_geometry = QApplication.primaryScreen().availableGeometry()
        
        center_point = screen_geometry.center()
        frame_geometry.moveCenter(center_point)
        self.move(frame_geometry.topLeft())
        self.raise_()
        self.activateWindow()
    
    def on_back_click(self):
        """Go back to main menu"""
        # Just close this dialog, don't close parent
        self.reject()  # Use reject instead of accept to just close dialog
    
    def _get_combined_alamat(self):
        """Combine the 3 alamat entries into a single string"""
        alamat_lines = []
        if self.entry_alamat1.text().strip():
            alamat_lines.append(self.entry_alamat1.text().strip())
        if self.entry_alamat2.text().strip():
            alamat_lines.append(self.entry_alamat2.text().strip())
        if self.entry_alamat3.text().strip():
            alamat_lines.append(self.entry_alamat3.text().strip())
        return "\n".join(alamat_lines)


if __name__ == "__main__":
    from PyQt5.QtWidgets import QApplication
    app = QApplication(sys.argv)
    window = Form3()
    window.show()
    sys.exit(app.exec_())

