"""
Form_SignUp.py - Government-Styled Sign Up B Form
Professional government interface for Sign Up B registration
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from hijri_converter import Hijri, Gregorian
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk
import os
from docx2pdf import convert
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from helpers.unified_database import UnifiedDatabase
from helpers.resource_path import get_logo_path, get_template_path

try:
    from helpers.docx_helper import replace_in_document
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value) if value else "")
        return {}


class FormSignUp:
    """Government-styled Sign Up B Form with professional design"""
    
    def __init__(self, root, parent_window=None):
        self.root = root
        self.parent_window = parent_window
        self.root.title("Sistem Pendaftaran Sign Up B - Borang Pendaftaran")
        self.root.geometry("1600x1000")
        self.root.configure(bg='#F5F5F5')
        self.db = UnifiedDatabase()
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': 'white',
            'text_dark': '#1a1a1a',
            'text_light': 'white',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_hover': '#004d99'
        }
        
        self.create_widgets()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        
        # Center window
        self.center_window()
        
        # Make window appear on top
        self.root.attributes('-topmost', True)
        self.root.after_idle(lambda: self.root.attributes('-topmost', False))
        self.root.lift()
        self.root.focus_force()
    
    def create_widgets(self):
        """Create government-styled interface"""
        # Header
        self.create_header()
        
        # Main content area
        main_container = tk.Frame(self.root, bg=self.colors['bg_main'])
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Canvas with scrollbar for form
        canvas = tk.Canvas(main_container, bg=self.colors['bg_main'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_container, orient='vertical', command=canvas.yview)
        
        scrollable_frame = tk.Frame(canvas, bg=self.colors['bg_main'])
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Create form in white card
        self.create_form(scrollable_frame)
        
        # Bind mouse wheel
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
    
    def create_header(self):
        """Create government header"""
        header_frame = tk.Frame(self.root, bg=self.colors['primary'], height=100)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Logo (left side)
        logo_frame = tk.Frame(header_frame, bg=self.colors['primary'], width=100)
        logo_frame.pack(side=tk.LEFT, padx=20, pady=10)
        
        try:
            logo_image = Image.open(get_logo_path())
            
            # Convert to RGBA if not already
            if logo_image.mode != 'RGBA':
                logo_image = logo_image.convert('RGBA')
            
            logo_image = logo_image.resize((80, 80), Image.Resampling.LANCZOS)
            
            # Create a background image with the same color as the frame
            background_color = self.colors['primary']  # '#003366'
            bg_rgb = tuple(int(background_color[i:i+2], 16) for i in (1, 3, 5))
            
            # Create a new RGBA image with the background color
            background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
            
            # Composite the logo onto the background using alpha compositing
            if logo_image.mode == 'RGBA':
                logo_image = Image.alpha_composite(background, logo_image)
                # Convert back to RGB for PhotoImage (Tkinter doesn't support RGBA well)
                logo_image = logo_image.convert('RGB')
            
            logo_photo = ImageTk.PhotoImage(logo_image)
            logo_label = tk.Label(logo_frame, image=logo_photo, bg=self.colors['primary'])
            logo_label.image = logo_photo
            logo_label.pack()
        except Exception as e:
            pass
        
        # Title section (center)
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=20)
        
        title = tk.Label(title_frame,
                        text="PENDAFTARAN SIGN UP B",
                        font=('Arial', 22, 'bold'),
                        bg=self.colors['primary'],
                        fg='white')
        title.pack(pady=(20, 5))
        
        subtitle = tk.Label(title_frame,
                           text="Sistem Pendaftaran Jadual B",
                           font=('Arial', 14),
                           bg=self.colors['primary'],
                           fg='#E0E0E0')
        subtitle.pack()
        
        # Back button (right side)
        if self.parent_window:
            btn_back = tk.Button(header_frame,
                                text="← KEMBALI",
                                font=('Arial', 14, 'bold'),
                                bg=self.colors['secondary'],
                                fg='white',
                                relief=tk.FLAT,
                                cursor='hand2',
                                width=15,
                                height=2,
                                command=self.on_back_click)
            btn_back.pack(side=tk.RIGHT, padx=20, pady=25)
        
        # Separator
        separator = tk.Frame(self.root, bg=self.colors['accent'], height=2)
        separator.pack(fill=tk.X)
    
    def create_form(self, parent):
        """Create form fields"""
        # White card container
        card = tk.Frame(parent, bg='white', relief=tk.FLAT, bd=0)
        card.pack(fill=tk.BOTH, expand=True, padx=40, pady=30)
        
        # Add border effect
        border_frame = tk.Frame(card, bg=self.colors['border'], bd=1)
        border_frame.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        form_frame = tk.Frame(border_frame, bg='white')
        form_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Create inner frame for grid layout
        grid_frame = tk.Frame(form_frame, bg='white')
        grid_frame.pack(fill=tk.BOTH, expand=True)
        
        # Form title
        form_title = tk.Label(grid_frame,
                             text="Maklumat Pendaftaran",
                             font=('Arial', 20, 'bold'),
                             bg='white',
                             fg=self.colors['primary'])
        form_title.grid(row=0, column=0, columnspan=2, sticky='w', pady=(0, 20))
        
        # Form fields grid
        row = 1
        
        # Rujukan Tuan
        tk.Label(grid_frame,
                text="Rujukan Tuan:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_rujukan_tuan = tk.Entry(grid_frame,
                                          font=('Arial', 13),
                                          width=50,
                                          relief=tk.SOLID,
                                          bd=1)
        self.entry_rujukan_tuan.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Rujukan Kami
        tk.Label(grid_frame,
                text="Rujukan Kami:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_rujukan_kami = tk.Entry(grid_frame,
                                          font=('Arial', 13),
                                          width=50,
                                          relief=tk.SOLID,
                                          bd=1)
        self.entry_rujukan_kami.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Tarikh
        tk.Label(grid_frame,
                text="Tarikh:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_tarikh = tk.Entry(grid_frame,
                                    font=('Arial', 13),
                                    width=20,
                                    relief=tk.SOLID,
                                    bd=1)
        self.entry_tarikh.insert(0, datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        self.entry_tarikh.bind('<KeyRelease>', lambda e: self.update_tarikh_islam())
        row += 1
        
        # Tarikh Islam
        tk.Label(grid_frame,
                text="Tarikh (Hijri):",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_islam = tk.Entry(grid_frame,
                                   font=('Arial', 13),
                                   width=30,
                                   relief=tk.SOLID,
                                   bd=1,
                                   state='readonly')
        self.entry_islam.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        self.update_tarikh_islam()
        row += 1
        
        # Separator
        separator = tk.Frame(grid_frame, bg=self.colors['border'], height=1)
        separator.grid(row=row, column=0, columnspan=2, sticky='ew', padx=10, pady=20)
        row += 1
        
        # Business Name (Nama Perniagaan)
        tk.Label(grid_frame,
                text="Nama Perniagaan: *",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_business_name = tk.Entry(grid_frame,
                                          font=('Arial', 13),
                                          width=50,
                                          relief=tk.SOLID,
                                          bd=1)
        self.entry_business_name.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Business Address (Alamat Perniagaan)
        tk.Label(grid_frame,
                text="Alamat Perniagaan: *",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='nw', padx=10, pady=10)
        self.text_business_address = tk.Text(grid_frame,
                                            font=('Arial', 13),
                                            width=50,
                                            height=4,
                                            relief=tk.SOLID,
                                            bd=1,
                                            wrap=tk.WORD)
        self.text_business_address.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Schedule Type (Jenis Jadual)
        tk.Label(grid_frame,
                text="Jenis Jadual: *",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.combo_schedule_type = ttk.Combobox(grid_frame,
                                               font=('Arial', 13),
                                               width=47,
                                               state='readonly',
                                               values=['Jadual B', 'Jadual A', 'Jadual C'])
        self.combo_schedule_type.set('Jadual B')
        self.combo_schedule_type.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Registration Number (Nombor Pendaftaran)
        tk.Label(grid_frame,
                text="Nombor Pendaftaran:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_registration_number = tk.Entry(grid_frame,
                                                 font=('Arial', 13),
                                                 width=50,
                                                 relief=tk.SOLID,
                                                 bd=1)
        self.entry_registration_number.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Salutation (Sapaan)
        tk.Label(grid_frame,
                text="Sapaan:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.combo_salutation = ttk.Combobox(grid_frame,
                                            font=('Arial', 13),
                                            width=47,
                                            state='readonly',
                                            values=['Tuan', 'Puan', 'Encik', 'Cik'])
        self.combo_salutation.set('Tuan')
        self.combo_salutation.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Email
        tk.Label(grid_frame,
                text="E-mel:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_email = tk.Entry(grid_frame,
                                    font=('Arial', 13),
                                    width=50,
                                    relief=tk.SOLID,
                                    bd=1)
        self.entry_email.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Talian (Phone)
        tk.Label(grid_frame,
                text="Talian:",
                font=('Arial', 13, 'bold'),
                bg='white',
                fg=self.colors['text_dark']).grid(row=row, column=0, sticky='w', padx=10, pady=10)
        self.entry_talian = tk.Entry(grid_frame,
                                    font=('Arial', 13),
                                    width=50,
                                    relief=tk.SOLID,
                                    bd=1)
        self.entry_talian.grid(row=row, column=1, sticky='w', padx=10, pady=10)
        row += 1
        
        # Separator before checklist
        separator2 = tk.Frame(grid_frame, bg=self.colors['border'], height=1)
        separator2.grid(row=row, column=0, columnspan=2, sticky='ew', padx=10, pady=20)
        row += 1
        
        # Senarai Semak (Checklist) Section
        checklist_label = tk.Label(grid_frame,
                                  text="Senarai Semak: *",
                                  font=('Arial', 14, 'bold'),
                                  bg='white',
                                  fg=self.colors['primary'])
        checklist_label.grid(row=row, column=0, columnspan=2, sticky='w', padx=10, pady=(10, 5))
        row += 1
        
        # Checklist items frame
        checklist_frame = tk.Frame(grid_frame, bg='white', relief=tk.SOLID, bd=1)
        checklist_frame.grid(row=row, column=0, columnspan=2, sticky='ew', padx=10, pady=10)
        
        # Default checklist items for Sign Up B
        self.checklist_items = [
            "Dokumen Pendaftaran Perniagaan",
            "Sijil Pendaftaran Syarikat (SSM)",
            "Dokumen Identiti Pemohon",
            "Alamat Perniagaan yang Sah",
            "Maklumat Bank Perniagaan",
            "Lesen Perniagaan (jika berkenaan)",
            "Dokumen Cukai (jika berkenaan)",
            "Lain-lain Dokumen Sokongan"
        ]
        
        # Store checkboxes
        self.checklist_vars = {}
        checklist_inner = tk.Frame(checklist_frame, bg='white')
        checklist_inner.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Create checkboxes in 2 columns
        for idx, item in enumerate(self.checklist_items):
            var = tk.BooleanVar()
            self.checklist_vars[item] = var
            
            checkbox = tk.Checkbutton(checklist_inner,
                                     text=item,
                                     variable=var,
                                     font=('Arial', 14),
                                     bg='white',
                                     fg=self.colors['text_dark'],
                                     activebackground='white',
                                     activeforeground=self.colors['text_dark'],
                                     anchor='w')
            
            # Arrange in 2 columns
            col = idx % 2
            checkbox_row = idx // 2
            checkbox.grid(row=checkbox_row, column=col, sticky='w', padx=10, pady=5)
        
        row += 1
        
        # Buttons
        button_frame = tk.Frame(grid_frame, bg='white')
        button_frame.grid(row=row, column=0, columnspan=2, pady=30)
        
        btn_save = tk.Button(button_frame,
                           text="💾 SIMPAN",
                           font=('Arial', 14, 'bold'),
                           bg=self.colors['button_success'],
                           fg='white',
                           relief=tk.FLAT,
                           cursor='hand2',
                           width=20,
                           height=2,
                           command=self.on_save_click)
        btn_save.pack(side=tk.LEFT, padx=10)
        btn_save.bind('<Enter>', lambda e: btn_save.config(bg='#1B5E20'))
        btn_save.bind('<Leave>', lambda e: btn_save.config(bg=self.colors['button_success']))
    
    def update_tarikh_islam(self):
        """Update Hijri date based on Gregorian date"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    
                    self.entry_islam.config(state='normal')
                    self.entry_islam.delete(0, tk.END)
                    self.entry_islam.insert(0, hijri_text)
                    self.entry_islam.config(state='readonly')
        except:
            pass
    
    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = 1600
        height = 1000
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def generate_document(self):
        """Generate Word document with replacements"""
        template_file = "signUpB.docx"
        
        # Try to use embedded template storage first
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(template_file)
            if doc is None:
                # Fallback to file system
                template_path = get_template_path(template_file)
                if not os.path.exists(template_path):
                    messagebox.showerror("Ralat", f"Templat tidak dijumpai: {template_file}")
                    return None
                doc = Document(template_path)
        except ImportError:
            # Fallback to file system if template_storage not available
            template_path = get_template_path(template_file)
            if not os.path.exists(template_path):
                messagebox.showerror("Ralat", f"Templat tidak dijumpai: {template_path}")
                return None
            doc = Document(template_path)
        
        # Get selected checklist items
        selected_checklist = self.get_selected_checklist_items()
        checklist_text = "\n".join([f"✓ {item}" for item in selected_checklist])
        
        # Prepare replacements
        replacements = {
            '<<RUJUKAN_TUAN>>': self.entry_rujukan_tuan.get(),
            '<<RUJUKAN_KAMI>>': self.entry_rujukan_kami.get(),
            '<<TARIKH>>': self.entry_tarikh.get(),
            '<<TARIKH_ISLAM>>': self.entry_islam.get(),
            '<<BUSINESS_NAME>>': self.entry_business_name.get().upper(),
            '<<BUSINESS_ADDRESS>>': self.text_business_address.get("1.0", tk.END).strip(),
            '<<SCHEDULE_TYPE>>': self.combo_schedule_type.get(),
            '<<REGISTRATION_NUMBER>>': self.entry_registration_number.get(),
            '<<SALUTATION>>': self.combo_salutation.get(),
            '<<SENARAI_SEMAK>>': checklist_text,
            '<<CHECKLIST>>': checklist_text,
            # Also handle MERGEFIELD format
            'MERGEFIELD Business_name': self.entry_business_name.get().upper(),
            'MERGEFIELD ALAMAT': self.text_business_address.get("1.0", tk.END).strip(),
            '< Nama pemohon >': self.entry_business_name.get().upper(),
            '< Alamat pemohon >': self.text_business_address.get("1.0", tk.END).strip(),
            'Ruj. Tuan': self.entry_rujukan_tuan.get(),
            'Ruj. Kami': self.entry_rujukan_kami.get(),
            'Tarikh': self.entry_tarikh.get()
        }
        
        # Replace placeholders
        replace_in_document(doc, replacements)
        
        return doc
    
    def get_selected_checklist_items(self):
        """Get list of selected checklist items"""
        selected = []
        for item, var in self.checklist_vars.items():
            if var.get():
                selected.append(item)
        return selected
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        # Validate required fields
        if not self.entry_business_name.get().strip():
            messagebox.showwarning("Amaran", "Sila isi Nama Perniagaan.")
            self.entry_business_name.focus()
            return
        
        if not self.text_business_address.get("1.0", tk.END).strip():
            messagebox.showwarning("Amaran", "Sila isi Alamat Perniagaan.")
            self.text_business_address.focus()
            return
        
        if not self.combo_schedule_type.get():
            messagebox.showwarning("Amaran", "Sila pilih Jenis Jadual.")
            return
        
        # Validate checklist - at least one item must be selected
        selected_items = self.get_selected_checklist_items()
        if not selected_items:
            messagebox.showwarning("Amaran", "Sila pilih sekurang-kurangnya satu item dalam Senarai Semak.")
            return
        
        try:
            # Generate document
            doc = self.generate_document()
            if doc is None:
                return
            
            # Save as DOCX first
            default_filename = f"SignUpB_{self.entry_business_name.get().strip().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Simpan Dokumen Sign Up B"
            )
            
            if filename:
                doc.save(filename)
                
                # Try to convert to PDF
                try:
                    pdf_filename = filename.replace('.docx', '.pdf')
                    convert(filename, pdf_filename)
                    messagebox.showinfo("Berjaya", 
                        f"Dokumen berjaya disimpan:\n{filename}\n\n"
                        f"PDF juga telah dijana:\n{pdf_filename}")
                except Exception as pdf_error:
                    messagebox.showinfo("Berjaya", 
                        f"Dokumen berjaya disimpan:\n{filename}\n\n"
                        f"Nota: Penukaran ke PDF gagal: {str(pdf_error)}")
                
                # Save to database
                try:
                    application_data = {
                        'form_type': 'signupb',
                        'rujukan_tuan': self.entry_rujukan_tuan.get(),
                        'rujukan_kami': self.entry_rujukan_kami.get(),
                        'tarikh': self.entry_tarikh.get(),
                        'tarikh_islam': self.entry_islam.get(),
                        'document_path': filename
                    }
                    
                    signupb_details = {
                        'email': self.entry_email.get(),
                        'talian': self.entry_talian.get()
                    }
                    
                    app_id = self.db.save_application('signupb', application_data, signupb_details)
                    
                    messagebox.showinfo("Berjaya", 
                        f"Dokumen berjaya disimpan:\n{filename}\n\n"
                        f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
                except Exception as db_error:
                    messagebox.showwarning("Amaran", 
                        f"Dokumen berjaya disimpan:\n{filename}\n\n"
                        f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
            
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menyimpan dokumen: {str(e)}")
    
    def on_back_click(self):
        """Go back to main menu"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.root.destroy()
    
    def on_close(self):
        """Handle window close"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = FormSignUp(root)
    root.mainloop()

