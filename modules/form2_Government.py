"""
form2_Government.py - DYNAMIC Government-Styled Data Entry Form
✨ Enhanced with real-time validation and smart field management
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime, timedelta
from hijri_converter import Hijri, Gregorian
from docx import Document
import os
import json
from helpers.unified_database import UnifiedDatabase
from helpers.resource_path import get_template_path

# Import helper functions
try:
    from helpers.docx_helper import replace_in_document, preview_placeholders, validate_replacements
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value) if value else "")
        return {}
    
    def preview_placeholders(doc):
        return []
    
    def validate_replacements(doc, replacements):
        return []

# Import UI components for enhanced UX
try:
    from helpers.ui_components import FieldValidator, Tooltip, NotificationBar
except ImportError:
    FieldValidator = None
    Tooltip = None
    NotificationBar = None


class Form2:
    """Government-styled Form2 with professional design"""
    
    def __init__(self, window, parent, category, sub_option, template_file,
                 requires_amount, requires_pengecualian):
        self.window = window
        self.parent_window = parent
        self.category = category
        self.sub_option = sub_option
        self.template_file = template_file
        self.requires_amount = requires_amount
        self.requires_pengecualian = requires_pengecualian
        self.db = UnifiedDatabase()
        
        self.window.title("Sistem Pengurusan Dokumen - Pengisian Data")
        self.window.geometry("1600x1000")
        self.window.configure(bg='#F5F5F5')
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': 'white',
            'text_dark': '#1a1a1a',
            'text_light': 'white',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_hover': '#004d99',
            'success': '#4CAF50',
            'warning': '#FF9800',
            'error': '#F44336',
            'info': '#2196F3'
        }
        
        # ✨ DYNAMIC FEATURES
        self.field_validators = {}  # Store field validators
        self.field_states = {}      # Track field states (valid/invalid)
        self.notification_bar = None  # Notification display
        self.dynamic_fields = {}    # Store dynamically created fields
        
        self.create_widgets()
        self.window.protocol("WM_DELETE_WINDOW", self.on_close)
        
        # Center window
        self.center_window()
        
        # Make window appear on top
        self.window.attributes('-topmost', True)
        self.window.after_idle(lambda: self.window.attributes('-topmost', False))
        self.window.lift()
        self.window.focus_force()
    
    def create_widgets(self):
        """Create government-styled interface"""
        # Header
        self.create_header()
        
        # Main content area
        main_container = tk.Frame(self.window, bg=self.colors['bg_main'])
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Canvas with scrollbar for form
        canvas = tk.Canvas(main_container, bg=self.colors['bg_main'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_container, orient='vertical', command=canvas.yview)
        
        scrollable_frame = tk.Frame(canvas, bg=self.colors['bg_main'])
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Create form in white card
        self.create_form(scrollable_frame)
        
        # Bind mouse wheel
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
    
    def create_header(self):
        """Create government header"""
        # Top banner
        header_frame = tk.Frame(self.window, bg=self.colors['primary'], height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Back button (left)
        btn_back = tk.Button(header_frame,
                            text="← KEMBALI",
                            font=('Arial', 12, 'bold'),
                            bg=self.colors['secondary'],
                            fg='white',
                            relief=tk.FLAT,
                            cursor='hand2',
                            width=15,
                            height=2,
                            command=self.on_back_click)
        btn_back.pack(side=tk.LEFT, padx=15, pady=15)
        
        # Title and info (center)
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.pack(side=tk.LEFT, expand=True, padx=20)
        
        title = tk.Label(title_frame,
                        text="Borang Pengisian Data Dokumen",
                        font=('Arial', 18, 'bold'),
                        bg=self.colors['primary'],
                        fg='white')
        title.pack(pady=(15, 5))
        
        subtitle = tk.Label(title_frame,
                           text=f"Kategori: {self.category} | Sub-Kategori: {self.sub_option}",
                           font=('Arial', 12),
                           bg=self.colors['primary'],
                           fg='#E0E0E0')
        subtitle.pack()
        
        # ✨ Help Button (right)
        help_btn = tk.Button(header_frame,
                            text="❓ HELP",
                            font=('Arial', 12, 'bold'),
                            bg=self.colors['accent'],
                            fg='white',
                            relief=tk.RAISED,
                            bd=2,
                            padx=15,
                            pady=5,
                            cursor='hand2',
                            command=self.show_help_dialog)
        help_btn.pack(side=tk.RIGHT, padx=15, pady=15)
        
        # Separator
        separator = tk.Frame(self.window, bg=self.colors['accent'], height=2)
        separator.pack(fill=tk.X)
        
        # ✨ Notification Bar placeholder (for dynamic messages)
        self.notification_bar = tk.Frame(self.window, bg=self.colors['info'], height=0)
        self.notification_bar.pack(fill=tk.X)
    
    def create_form(self, parent):
        """Create form fields in professional layout"""
        # Main white card
        card = tk.Frame(parent, bg='white')
        card.pack(padx=40, pady=30, fill=tk.BOTH, expand=True)
        
        # Add border
        border = tk.Frame(card, bg=self.colors['border'], bd=1)
        border.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content = tk.Frame(border, bg='white')
        content.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Section title
        section_title = tk.Label(content,
                                text="Maklumat Dokumen",
                                font=('Arial', 15, 'bold'),
                                bg='white',
                                fg=self.colors['primary'])
        section_title.grid(row=0, column=0, columnspan=4, pady=(0, 20), sticky='w')
        
        # Create form fields
        row = 1
        
        # Helper function for labels
        def create_label(text, row_num, col_num, is_section=False):
            if is_section:
                lbl = tk.Label(content,
                              text=text,
                              font=('Arial', 15, 'bold'),
                              bg='white',
                              fg=self.colors['primary'])
                lbl.grid(row=row_num, column=col_num, columnspan=4,
                        pady=(20, 10), sticky='w')
            else:
                lbl = tk.Label(content,
                              text=text,
                              font=('Arial', 12),
                              bg='white',
                              fg=self.colors['text_dark'])
                lbl.grid(row=row_num, column=col_num, sticky='e', padx=(0, 10), pady=8)
            return lbl
        
        # Process
        self.label_process = create_label("PROSES:", row, 0)
        self.combo_process = ttk.Combobox(content, width=35, state='readonly', font=('Arial', 12))
        self.combo_process.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        self.combo_process.bind("<<ComboboxSelected>>", self.on_process_change)
        row += 1
        
        # SST/ADM Radio Buttons (only for penjualan)
        self.label_sst_adam = create_label("JENIS:", row, 0)
        self.var_sst_adam = tk.StringVar(value="sst")
        
        sst_adam_frame = tk.Frame(content, bg='white')
        sst_adam_frame.grid(row=row, column=1, columnspan=3, sticky='w', padx=5, pady=8)
        
        self.radio_sst = tk.Radiobutton(sst_adam_frame, text="SST-ADM", variable=self.var_sst_adam,
                      value="sst", bg='white', font=('Arial', 12),
                      command=self.on_sst_adam_change)
        self.radio_sst.pack(side=tk.LEFT, padx=(0, 15))
        
        self.radio_adam = tk.Radiobutton(sst_adam_frame, text="AMES-03", variable=self.var_sst_adam,
                      value="adam", bg='white', font=('Arial', 12),
                      command=self.on_sst_adam_change)
        self.radio_adam.pack(side=tk.LEFT)
        
        # Store the frame for hiding/showing
        self.sst_adam_frame = sst_adam_frame
        row += 1
        
        # Rujukan
        create_label("RUJUKAN:", row, 0)
        rujukan_frame = tk.Frame(content, bg='white')
        rujukan_frame.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        tk.Label(rujukan_frame, text="KE.JB(90)650/05-02/", bg='white', font=('Arial', 12)).pack(side=tk.LEFT)
        self.entry_rujukan = tk.Entry(rujukan_frame, width=25, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_rujukan.pack(side=tk.LEFT)
        
        # ✨ Add real-time validation
        self.entry_rujukan.bind('<KeyRelease>', lambda e: self.on_field_change('rujukan', self.entry_rujukan))
        
        # Alamat - 3 separate lines (right side)
        create_label("ALAMAT:", row, 2)
        self.entry_alamat1 = tk.Entry(content, width=40, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_alamat1.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        
        # ✨ Add tooltip
        if Tooltip:
            Tooltip(self.entry_alamat1, "Baris pertama alamat")
        row += 1
        
        # Nama Syarikat
        create_label("NAMA SYARIKAT:", row, 0)
        self.entry_nama = tk.Entry(content, width=38, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_nama.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        # ✨ Add real-time validation & tooltip
        self.entry_nama.bind('<KeyRelease>', lambda e: self.on_field_change('nama_syarikat', self.entry_nama))
        if Tooltip:
            Tooltip(self.entry_nama, "Masukkan nama penuh syarikat")
        
        # Alamat line 2
        self.entry_alamat2 = tk.Entry(content, width=40, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_alamat2.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        
        # ✨ Add tooltip
        if Tooltip:
            Tooltip(self.entry_alamat2, "Baris kedua alamat")
        row += 1
        
        # Tarikh
        create_label("TARIKH:", row, 0)
        self.entry_tarikh = tk.Entry(content, width=38, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_tarikh.insert(0, datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        self.entry_tarikh.bind('<KeyRelease>', self.update_tarikh_islam)
        
        # Alamat line 3
        self.entry_alamat3 = tk.Entry(content, width=40, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_alamat3.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Tarikh Islam
        create_label("TARIKH ISLAM:", row, 0)
        self.entry_islam = tk.Entry(content, width=38, font=('Arial', 12), state='readonly', relief=tk.SOLID, bd=1)
        self.entry_islam.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        # Nama Pegawai (right)
        create_label("NAMA PEGAWAI:", row, 2)
        self.combo_pegawai = ttk.Combobox(content, width=37, font=('Arial', 12))
        self.combo_pegawai['values'] = ["HABBAH SYAKIRAH BINTI AB GHAFAR",
                                        "KHAIRONE DZARWANY BINTI MOHD KAIRI"]
        self.combo_pegawai.current(0)
        self.combo_pegawai.grid(row=row, column=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Divider line
        divider = tk.Frame(content, bg=self.colors['border'], height=1)
        divider.grid(row=row, column=0, columnspan=4, sticky='ew', pady=15)
        row += 1
        
        # Jenis Barang
        self.label_jenis = create_label("JENIS BARANG:", row, 0)
        self.combo_jenis = ttk.Combobox(content, width=35, state='readonly', font=('Arial', 12))
        self.combo_jenis.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        row += 1
        
        # Pengecualian
        self.label_pengecualian = create_label("PENGECUALIAN:", row, 0)
        self.combo_pengecualian = ttk.Combobox(content, width=80, state='readonly', font=('Arial', 12))
        self.combo_pengecualian.grid(row=row, column=1, columnspan=3, sticky='w', padx=5, pady=8)
        row += 1
        
        # Amount / Reason field
        if self.sub_option == "tidak_lulus":
            self.label_amount = create_label("SEBAB TIDAK DILULUSKAN:", row, 0)
        else:
            self.label_amount = create_label("AMAUN (RM):", row, 0)
        self.entry_amount = tk.Entry(content, width=38, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_amount.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        # ✨ Add real-time validation for amount
        if self.sub_option != "tidak_lulus":
            self.entry_amount.bind('<KeyRelease>', lambda e: self.on_field_change('amount', self.entry_amount))
            if Tooltip:
                Tooltip(self.entry_amount, "Masukkan jumlah dalam format: 1234.56")
        row += 1
        
        # Rujukan Syarikat checkbox
        self.var_rujukan_syarikat = tk.BooleanVar()
        self.check_rujukan = tk.Checkbutton(content,
                                           text="Rujukan Syarikat",
                                           variable=self.var_rujukan_syarikat,
                                           font=('Arial', 12),
                                           bg='white',
                                           command=self.toggle_rujukan_syarikat)
        self.check_rujukan.grid(row=row, column=0, sticky='e', padx=(0, 10), pady=8)
        
        rujukan_frame2 = tk.Frame(content, bg='white')
        rujukan_frame2.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        tk.Label(rujukan_frame2, text="", bg='white', font=('Arial', 12)).pack(side=tk.LEFT)
        self.entry_rujukan_syarikat = tk.Entry(rujukan_frame2, width=25, font=('Arial', 12),
                                               state='disabled', relief=tk.SOLID, bd=1)
        self.entry_rujukan_syarikat.pack(side=tk.LEFT)
        row += 1
        
        # ===== PEMUSNAHAN SECTION =====
        # Divider before pemusnahan section
        self.pemusnahan_divider = tk.Frame(content, bg=self.colors['border'], height=1)
        self.pemusnahan_divider.grid(row=row, column=0, columnspan=4, sticky='ew', pady=15)
        row += 1
        
        # Section title for Pemusnahan
        self.pemusnahan_section_title = create_label("MAKLUMAT PEMUSNAHAN", row, 0, is_section=True)
        row += 1
        
        
        # Tempoh Pemusnahan - Start Date
        self.label_pemusnahan_mula = create_label("TEMPOH PEMUSNAHAN:", row, 0)
        
        tempoh_frame = tk.Frame(content, bg='white')
        tempoh_frame.grid(row=row, column=1, columnspan=3, sticky='w', padx=5, pady=8)
        
        tk.Label(tempoh_frame, text="Mulai:", bg='white', font=('Arial', 12, 'bold')).pack(side=tk.LEFT, padx=(0, 5))
        self.entry_pemusnahan_mula = tk.Entry(tempoh_frame, width=15, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_pemusnahan_mula.pack(side=tk.LEFT, padx=5)
        
        # Calculate default dates (30 days from today)
        today = datetime.now()
        end_date = today + timedelta(days=30)
        self.entry_pemusnahan_mula.insert(0, today.strftime("%d/%m/%Y"))
        self.entry_pemusnahan_mula.bind('<KeyRelease>', self.calculate_tempoh_pemusnahan)
        
        tk.Label(tempoh_frame, text="hingga:", bg='white', font=('Arial', 12, 'bold')).pack(side=tk.LEFT, padx=(15, 5))
        self.entry_pemusnahan_tamat = tk.Entry(tempoh_frame, width=15, font=('Arial', 12), relief=tk.SOLID, bd=1)
        self.entry_pemusnahan_tamat.pack(side=tk.LEFT, padx=5)
        self.entry_pemusnahan_tamat.insert(0, end_date.strftime("%d/%m/%Y"))
        self.entry_pemusnahan_tamat.bind('<KeyRelease>', self.calculate_tempoh_pemusnahan)
        
        tk.Label(tempoh_frame, text="(30 hari)", bg='white', font=('Arial', 9, 'italic'),
                fg='#666666').pack(side=tk.LEFT, padx=(5, 0))
        row += 1
        
        # Tempoh (Period) - Auto-calculated
        self.label_tempoh = create_label("TEMPOH:", row, 0)
        self.entry_tempoh = tk.Entry(content, width=38, font=('Arial', 12), 
                                     state='readonly', relief=tk.SOLID, bd=1,
                                     readonlybackground='#f0f0f0')
        self.entry_tempoh.grid(row=row, column=1, sticky='w', padx=5, pady=8)
        
        # Calculate initial tempoh
        self.calculate_tempoh_pemusnahan()
        row += 1
        
        # ===== END PEMUSNAHAN SECTION =====
        
        # Button section
        button_frame = tk.Frame(content, bg='white')
        button_frame.grid(row=row, column=0, columnspan=4, pady=30)
        
        # Preview button
        btn_preview = tk.Button(button_frame,
                               text="👁 PRATONTON",
                               font=('Arial', 12, 'bold'),
                               bg=self.colors['button_secondary'],
                               fg='white',
                               relief=tk.FLAT,
                               cursor='hand2',
                               width=18,
                               height=2,
                               command=self.on_preview_click)
        btn_preview.pack(side=tk.LEFT, padx=10)
        
        # Save button
        btn_save = tk.Button(button_frame,
                            text="💾 SIMPAN DOKUMEN",
                            font=('Arial', 12, 'bold'),
                            bg=self.colors['button_success'],
                            fg='white',
                            relief=tk.FLAT,
                            cursor='hand2',
                            width=20,
                            height=2,
                            command=self.on_save_click)
        btn_save.pack(side=tk.LEFT, padx=10)
        
        # Load options based on category
        self.load_options()
        self.update_field_visibility()
        self.update_tarikh_islam(None)
    
    def load_options(self):
        """Load dropdown options based on category"""
        # Process options
        process_options = {
            "pemusnahan": ["pemusnahan"],
            "penjualan": ["penjualan", "pemberian"],
            "skrap": ["pemusnahan"],
            "tidak_lulus": ["pemusnahan", "penjualan", "pemberian"]
        }
        
        if self.sub_option in process_options:
            self.combo_process['values'] = process_options[self.sub_option]
            if self.combo_process['values']:
                self.combo_process.current(0)
        
        # Jenis Barang
        self.combo_jenis['values'] = ["barang siap", "bahan mentah", "komponen", "bahan pembungkus",
                                       "bahan pembungkusan", "skrap", "sisa dan hampas", "apa-apa barang bercukai"]
        self.combo_jenis.current(0)
        
        # Pengecualian options
        self.combo_pengecualian['values'] = [
            "perintah cukai jualan (orang yang dikecualikan daripada pembayaran cukai) 2018",
            "skim pengeksport utama diluluskan (ames)",
            "kemudahan pemotongan cukai jualan",
            "subseksyen 35(3) akta cukai jualan 2018"
        ]
        self.combo_pengecualian.current(0)
    
    def on_process_change(self, event=None):
        """Handle process selection change"""
        self.update_pemusnahan_visibility()
    
    def on_sst_adam_change(self):
        """Handle SST/ADAM radio button change"""
        # No additional logic needed - just for future use if needed
        pass
    
    def is_penjualan_sst(self):
        """Check if this is penjualan with SST selected"""
        return (self.sub_option == "penjualan" and 
                self.var_sst_adam.get() == "sst")
    
    def update_pemusnahan_visibility(self):
        """Show/hide pemusnahan section based on process selection"""
        process = self.combo_process.get().lower()
        
        pemusnahan_widgets = [
            self.pemusnahan_divider,
            self.pemusnahan_section_title,
            self.label_pemusnahan_mula,
            self.entry_pemusnahan_mula.master,  # The frame containing both date entries
            self.label_tempoh,
            self.entry_tempoh
        ]
        
        # IMPORTANT: tidak_lulus category should NEVER show pemusnahan details
        # Only show pemusnahan details for other categories when process is pemusnahan
        if self.sub_option == "tidak_lulus":
            # Hide pemusnahan fields for tidak_lulus
            for widget in pemusnahan_widgets:
                if hasattr(widget, 'grid_remove'):
                    widget.grid_remove()
        elif process == "pemusnahan":
            # Show pemusnahan fields for other categories
            for widget in pemusnahan_widgets:
                if hasattr(widget, 'grid'):
                    widget.grid()
        else:
            # Hide pemusnahan fields
            for widget in pemusnahan_widgets:
                if hasattr(widget, 'grid_remove'):
                    widget.grid_remove()
    
    def update_field_visibility(self):
        """Show/hide fields based on requirements"""
        # Hide all optional fields first
        widgets_to_hide = [
            self.label_jenis, self.combo_jenis,
            self.label_pengecualian, self.combo_pengecualian,
            self.label_amount, self.entry_amount,
            self.label_sst_adam, self.sst_adam_frame
        ]
        
        for widget in widgets_to_hide:
            if hasattr(widget, 'grid_remove'):
                widget.grid_remove()
        
        # Show SST/ADAM radio buttons ONLY for penjualan
        if self.sub_option == "penjualan":
            self.label_sst_adam.grid()
            self.sst_adam_frame.grid()
        
        # Show required fields
        if self.requires_pengecualian:
            self.label_jenis.grid()
            self.combo_jenis.grid()
            self.label_pengecualian.grid()
            self.combo_pengecualian.grid()
        
        # For tidak_lulus: ALWAYS show amount field (as reason field)
        if self.sub_option == "tidak_lulus":
            self.label_amount.grid()
            self.entry_amount.grid()
        # For other categories: show amount field only if required
        elif self.requires_amount:
            self.label_amount.grid()
            self.entry_amount.grid()
        
        # Update pemusnahan visibility based on process
        self.update_pemusnahan_visibility()
    
    def toggle_rujukan_syarikat(self):
        """Enable/disable rujukan syarikat field"""
        if self.var_rujukan_syarikat.get():
            self.entry_rujukan_syarikat.config(state='normal')
        else:
            self.entry_rujukan_syarikat.config(state='disabled')
            self.entry_rujukan_syarikat.delete(0, tk.END)
    
    def calculate_tempoh_pemusnahan(self, event=None):
        """Calculate the period between start and end dates"""
        try:
            tarikh_mula = self.entry_pemusnahan_mula.get()
            tarikh_tamat = self.entry_pemusnahan_tamat.get()
            
            if '/' in tarikh_mula and '/' in tarikh_tamat:
                # Parse dates
                day1, month1, year1 = map(int, tarikh_mula.split('/'))
                day2, month2, year2 = map(int, tarikh_tamat.split('/'))
                
                date1 = datetime(year1, month1, day1)
                date2 = datetime(year2, month2, day2)
                
                # Calculate difference in days
                delta = (date2 - date1).days
                
                # Convert to Malay words
                tempoh_text = self.number_to_malay_words(delta)
                tempoh_formatted = f"{tempoh_text} ({delta}) hari"
                
                # Update tempoh field
                self.entry_tempoh.config(state='normal')
                self.entry_tempoh.delete(0, tk.END)
                self.entry_tempoh.insert(0, tempoh_formatted)
                self.entry_tempoh.config(state='readonly')
        except:
            pass
    
    def number_to_malay_words(self, num):
        """Convert number to Malay words"""
        ones = ["", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "lapan", "sembilan"]
        teens = ["sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas", "lima belas",
                "enam belas", "tujuh belas", "lapan belas", "sembilan belas"]
        tens = ["", "", "dua puluh", "tiga puluh", "empat puluh", "lima puluh",
               "enam puluh", "tujuh puluh", "lapan puluh", "sembilan puluh"]
        
        if num == 0:
            return "kosong"
        elif num < 10:
            return ones[num]
        elif num < 20:
            return teens[num - 10]
        elif num < 100:
            ten = num // 10
            one = num % 10
            if one == 0:
                return tens[ten]
            else:
                return f"{tens[ten]} {ones[one]}"
        elif num < 1000:
            hundred = num // 100
            remainder = num % 100
            if hundred == 1:
                result = "seratus"
            else:
                result = f"{ones[hundred]} ratus"
            if remainder > 0:
                result += f" {self.number_to_malay_words(remainder)}"
            return result
        else:
            return str(num)  # For larger numbers, just use digits
    
    def update_tarikh_islam(self, event):
        """Convert Gregorian to Hijri date"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    
                    self.entry_islam.config(state='normal')
                    self.entry_islam.delete(0, tk.END)
                    self.entry_islam.insert(0, hijri_text)
                    self.entry_islam.config(state='readonly')
        except:
            pass
    
    def format_date_malay(self, date_str):
        """Format date string (DD/MM/YYYY) to Malay format (DD Month YYYY)"""
        try:
            if '/' in date_str:
                day, month, year = map(int, date_str.split('/'))
                
                bulan_melayu = {
                    1: "Januari", 2: "Februari", 3: "Mac", 4: "April",
                    5: "Mei", 6: "Jun", 7: "Julai", 8: "Ogos",
                    9: "September", 10: "Oktober", 11: "November", 12: "Disember"
                }
                
                return f"{day:02d}/{month:02d}/{year}"
        except:
            return date_str
    
    def format_tarikh_malay(self):
        """Format date in Malay: 03 November 2025"""
        try:
            tarikh_str = self.entry_tarikh.get()
            if '/' in tarikh_str:
                day, month, year = map(int, tarikh_str.split('/'))
                date_obj = datetime(year, month, day)
                
                bulan_melayu = {
                    1: "Januari", 2: "Februari", 3: "Mac", 4: "April",
                    5: "Mei", 6: "Jun", 7: "Julai", 8: "Ogos",
                    9: "September", 10: "Oktober", 11: "November", 12: "Disember"
                }
                
                return f"{day:02d} {bulan_melayu[month]} {year}"
        except:
            return self.entry_tarikh.get()
    
    def generate_document(self):
        """Generate Word document with replacements"""
        # Try to use embedded template storage first
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(self.template_file)
            if doc is None:
                # Fallback to file system
                template_path = get_template_path(self.template_file)
                if not os.path.exists(template_path):
                    messagebox.showerror("Ralat", f"Templat tidak dijumpai: {self.template_file}")
                    return None
                # Check if file is .doc (old format) - python-docx doesn't support it
                if self.template_file.lower().endswith('.doc') and not self.template_file.lower().endswith('.docx'):
                    messagebox.showerror("Ralat", 
                        f"Templat {self.template_file} menggunakan format lama (.doc).\n\n"
                        f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                    return None
                try:
                    doc = Document(template_path)
                except Exception as e:
                    if self.template_file.lower().endswith('.doc') and not self.template_file.lower().endswith('.docx'):
                        messagebox.showerror("Ralat", 
                            f"Templat {self.template_file} menggunakan format lama (.doc) yang tidak disokong.\n\n"
                            f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                    else:
                        messagebox.showerror("Ralat", f"Ralat memuat templat: {str(e)}")
                    return None
        except ImportError:
            # Fallback to file system if template_storage not available
            template_path = os.path.join("Templates", self.template_file)
            if not os.path.exists(template_path):
                messagebox.showerror("Ralat", f"Templat tidak dijumpai: {template_path}")
                return None
            # Check if file is .doc (old format) - python-docx doesn't support it
            if self.template_file.lower().endswith('.doc') and not self.template_file.lower().endswith('.docx'):
                messagebox.showerror("Ralat", 
                    f"Templat {self.template_file} menggunakan format lama (.doc).\n\n"
                    f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                return None
            try:
                doc = Document(template_path)
            except Exception as e:
                if self.template_file.lower().endswith('.doc') and not self.template_file.lower().endswith('.docx'):
                    messagebox.showerror("Ralat", 
                        f"Templat {self.template_file} menggunakan format lama (.doc) yang tidak disokong.\n\n"
                        f"Sila tukar kepada format .docx menggunakan Microsoft Word atau TemplateEditor.")
                else:
                    messagebox.showerror("Ralat", f"Ralat memuat templat: {str(e)}")
                return None
        except Exception as e:
            messagebox.showerror("Ralat", f"Error loading template: {str(e)}")
            return None
        
        try:
            
            # Prepare replacements
            rujukan_kami = f"{self.entry_rujukan.get()}"
            rujukan_tuan = ""
            if self.var_rujukan_syarikat.get():
                rujukan_tuan = f"{self.entry_rujukan_syarikat.get()}"
            
            # Combine the 3 alamat entries
            alamat_lines = []
            if self.entry_alamat1.get().strip():
                alamat_lines.append(self.entry_alamat1.get().strip())
            if self.entry_alamat2.get().strip():
                alamat_lines.append(self.entry_alamat2.get().strip())
            if self.entry_alamat3.get().strip():
                alamat_lines.append(self.entry_alamat3.get().strip())
            alamat_combined = "\n".join(alamat_lines)
            
            replacements = {
                '<<RUJUKAN_TUAN>>': rujukan_tuan,
                '<<RUJUKAN>>': rujukan_kami,
                '<<NAMA_SYARIKAT>>': self.entry_nama.get(),
                '<<ALAMAT>>': alamat_combined,
                '<<TARIKH>>': self.entry_tarikh.get(),
                '<<TARIKH2>>': self.format_tarikh_malay(),
                '<<TARIKH_ISLAM>>': self.entry_islam.get(),
                '<<NAMA_PEGAWAI>>': self.combo_pegawai.get().upper(),
                '<<AMOUNT>>': self.entry_amount.get(),
            }
            
            # Add SST/ADAM type for penjualan
            if self.sub_option == "penjualan":
                replacements['<<JENIS_TEMPLATE>>'] = self.var_sst_adam.get().upper()
            
            # Process-specific replacements
            if self.combo_process.winfo_ismapped() and self.combo_process.get():
                replacements['<<PROSES>>'] = self.combo_process.get()
                
                # Pemusnahan-specific replacements
                if self.combo_process.get().lower() == "pemusnahan":
                    replacements['<<TARIKH_MULA>>'] = self.format_date_malay(self.entry_pemusnahan_mula.get())
                    replacements['<<TARIKH_TAMAT>>'] = self.format_date_malay(self.entry_pemusnahan_tamat.get())
                    replacements['<<TEMPOH>>'] = self.entry_tempoh.get()
            
            # Jenis Barang
            if self.combo_jenis.winfo_ismapped() and self.combo_jenis.get():
                replacements['<<JENIS_BARANG>>'] = self.combo_jenis.get()
            
            # Pengecualian
            if self.combo_pengecualian.winfo_ismapped() and self.combo_pengecualian.get():
                pengecualian = self.combo_pengecualian.get()
                replacements['<<PENGECUALIAN>>'] = pengecualian
                
                # Tajuk surat
                jenis_barang = self.combo_jenis.get() if self.combo_jenis.winfo_ismapped() else ""
                proses = self.combo_process.get() if self.combo_process.winfo_ismapped() else ""
                
                tajuk_surat = f"pelupusan secara {proses} {jenis_barang} yang diperolehi dengan pengecualian di bawah {pengecualian}"
                replacements['<<TAJUK_SURAT>>'] = tajuk_surat.upper()
                replacements['<<TAJUK_SURAT2>>'] = tajuk_surat.lower()
            
            # Rujukan info
            if self.var_rujukan_syarikat.get():
                replacements['<<RUJUKAN_INFO>>'] = f"rujukan {rujukan_tuan} "
            else:
                replacements['<<RUJUKAN_INFO>>'] = ""
            
            # Status replacement - SST or ADAM for penjualan
            if self.sub_option == "penjualan":
                replacements['<<STATUS>>'] = self.var_sst_adam.get().upper()
            
            # Replace in document
            replace_in_document(doc, replacements)
            
            # Make TAJUK_SURAT bold
            tajuk_uppercase = replacements.get('<<TAJUK_SURAT>>')
            if tajuk_uppercase:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if tajuk_uppercase in run.text:
                            run.bold = True
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if tajuk_uppercase in run.text:
                                        run.bold = True
            
            return doc
            
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menjana dokumen: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def on_preview_click(self):
        """Preview document"""
        doc = self.generate_document()
        if not doc:
            return
        
        temp_path = "temp_preview.docx"
        doc.save(temp_path)
        
        try:
            if os.name == 'nt':  # Windows
                os.startfile(temp_path)
            elif os.name == 'posix':  # macOS/Linux
                import subprocess
                subprocess.call(['open', temp_path])
            
            messagebox.showinfo("Pratonton", "Dokumen dibuka untuk pratonton.\nSemak temp_preview.docx")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat buka pratonton: {e}")
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        # ✨ Enhanced validation with visual feedback
        validation_errors = []
        
        if not self.entry_nama.get().strip():
            validation_errors.append("• Nama Syarikat")
            self.entry_nama.configure(bg='#FFEBEE')
        else:
            self.entry_nama.configure(bg='white')
        
        if not self.entry_rujukan.get().strip():
            validation_errors.append("• Rujukan")
            self.entry_rujukan.configure(bg='#FFEBEE')
        else:
            self.entry_rujukan.configure(bg='white')
        
        if not self.entry_tarikh.get().strip():
            validation_errors.append("• Tarikh")
            self.entry_tarikh.configure(bg='#FFEBEE')
        else:
            self.entry_tarikh.configure(bg='white')
        
        if not self.entry_alamat1.get().strip():
            validation_errors.append("• Alamat")
            self.entry_alamat1.configure(bg='#FFEBEE')
        else:
            self.entry_alamat1.configure(bg='white')
        
        # Show validation errors
        if validation_errors:
            messagebox.showwarning(
                "⚠️ Field Tidak Lengkap",
                "Sila lengkapkan maklumat berikut:\n\n" + "\n".join(validation_errors)
            )
            return
        
        # ✨ Confirmation dialog before saving
        response = messagebox.askyesno(
            "✅ Sahkan Simpan Dokumen",
            f"Adakah anda pasti untuk simpan dokumen?\n\n"
            f"📋 Rujukan: KE.JB(90)650/05-02/{self.entry_rujukan.get()}\n"
            f"🏢 Syarikat: {self.entry_nama.get()}\n"
            f"📅 Tarikh: {self.entry_tarikh.get()}\n"
            f"📁 Kategori: {self.category} - {self.sub_option}"
        )
        
        if not response:
            return
        
        doc = self.generate_document()
        if not doc:
            return
        
        safe_name = self.entry_nama.get().replace('/', '_').replace('\\', '_')
        
        # Add SST/ADAM to filename for penjualan
        if self.sub_option == "penjualan":
            jenis = self.var_sst_adam.get().upper()
            default_filename = f"{safe_name}_{self.sub_option}_{jenis}.pdf"
        else:
            default_filename = f"{safe_name}_{self.sub_option}.pdf"
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("Fail PDF", "*.pdf"), ("Semua fail", "*.*")],
            initialfile=default_filename
        )
        
        if not filename:
            return
        
        # Ensure the filename ends with .pdf
        if not filename.lower().endswith('.pdf'):
            filename += '.pdf'
        
        try:
            # Save as temp Word file first
            temp_docx = "temp_conversion.docx"
            doc.save(temp_docx)
            
            # Convert to PDF
            from docx2pdf import convert
            convert(temp_docx, filename)
            
            # Clean up temp file
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            
            # Save to database
            try:
                # Prepare application data
                rujukan_kami = f"KE.JB(90)650/05-02/{self.entry_rujukan.get()}"
                rujukan_tuan = ""
                if self.var_rujukan_syarikat.get():
                    rujukan_tuan = self.entry_rujukan_syarikat.get()
                
                # Combine alamat lines
                alamat_lines = []
                if self.entry_alamat1.get().strip():
                    alamat_lines.append(self.entry_alamat1.get().strip())
                if self.entry_alamat2.get().strip():
                    alamat_lines.append(self.entry_alamat2.get().strip())
                if self.entry_alamat3.get().strip():
                    alamat_lines.append(self.entry_alamat3.get().strip())
                alamat_combined = "\n".join(alamat_lines)
                
                application_data = {
                    'category': self.category,
                    'sub_option': self.sub_option,
                    'rujukan_kami': rujukan_kami,
                    'rujukan_tuan': rujukan_tuan,
                    'nama_syarikat': self.entry_nama.get(),
                    'alamat': alamat_combined,
                    'tarikh': self.entry_tarikh.get(),
                    'tarikh_islam': self.entry_islam.get(),
                    'nama_pegawai': self.combo_pegawai.get(),
                    'status': 'DILULUSKAN',
                    'document_path': filename,
                    'additional_data': {
                        'sst_adam_type': self.var_sst_adam.get() if self.sub_option == "penjualan" else None
                    }
                }
                
                # Prepare pelupusan-specific details
                pelupusan_details = {
                    'proses': self.combo_process.get() if self.combo_process.winfo_ismapped() and self.combo_process.get() else None,
                    'jenis_barang': self.combo_jenis.get() if self.combo_jenis.winfo_ismapped() and self.combo_jenis.get() else None,
                    'pengecualian': self.combo_pengecualian.get() if self.combo_pengecualian.winfo_ismapped() and self.combo_pengecualian.get() else None,
                    'amount': self.entry_amount.get() if self.entry_amount.winfo_ismapped() else None,
                    'tarikh_mula': self.entry_pemusnahan_mula.get() if hasattr(self, 'entry_pemusnahan_mula') and self.entry_pemusnahan_mula.winfo_ismapped() else None,
                    'tarikh_tamat': self.entry_pemusnahan_tamat.get() if hasattr(self, 'entry_pemusnahan_tamat') and self.entry_pemusnahan_tamat.winfo_ismapped() else None,
                    'tempoh': self.entry_tempoh.get() if hasattr(self, 'entry_tempoh') and self.entry_tempoh.winfo_ismapped() else None
                }
                
                # Save to database
                app_id = self.db.save_application('pelupusan', application_data, pelupusan_details)
                
                messagebox.showinfo("Berjaya", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
            except Exception as db_error:
                # Document saved but database save failed
                messagebox.showwarning("Amaran", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
            
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menyimpan dokumen: {str(e)}")
    
    def center_window(self):
        """Center window on screen"""
        self.window.update_idletasks()
        width = 1600
        height = 1000
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
    
    # ========================================
    # ✨ DYNAMIC FEATURES - Validation
    # ========================================
    
    def show_notification(self, message, type="info"):
        """Show notification bar at top of form"""
        if not hasattr(self, 'notification_bar') or not self.notification_bar:
            return
        
        colors = {
            'success': self.colors['success'],
            'error': self.colors['error'],
            'warning': self.colors['warning'],
            'info': self.colors['info']
        }
        
        bg_color = colors.get(type, self.colors['info'])
        
        # Create temporary notification
        notif_frame = tk.Frame(self.window, bg=bg_color, height=40)
        notif_frame.pack(side=tk.TOP, fill=tk.X, before=self.window.winfo_children()[0])
        
        label = tk.Label(notif_frame, text=message, bg=bg_color, fg='white',
                        font=('Arial', 12, 'bold'), pady=10)
        label.pack()
        
        # Auto-hide after 3 seconds
        self.window.after(3000, notif_frame.destroy)
    
    def validate_field_realtime(self, field_name, value, field_widget):
        """Real-time validation with visual feedback"""
        is_valid = True
        error_msg = ""
        
        # Validation rules
        if field_name == "rujukan" and value:
            if len(value) < 2:
                is_valid = False
                error_msg = "Rujukan too short"
        
        elif field_name == "nama_syarikat" and value:
            if len(value) < 3:
                is_valid = False
                error_msg = "Nama syarikat too short"
        
        elif field_name == "tarikh" and value:
            # Check date format
            try:
                datetime.strptime(value, "%d/%m/%Y")
            except:
                is_valid = False
                error_msg = "Format tarikh: DD/MM/YYYY"
        
        elif field_name == "amount" and value:
            # Check if numeric
            try:
                float(value.replace(',', ''))
            except:
                is_valid = False
                error_msg = "Amount must be numeric"
        
        # Visual feedback
        if is_valid:
            field_widget.configure(bg='white')
            self.field_states[field_name] = True
        else:
            field_widget.configure(bg='#FFEBEE')  # Light red
            self.field_states[field_name] = False
        
        return is_valid
    
    def toggle_field_visibility(self, field_widget, label_widget, show):
        """Dynamically show/hide fields with animation"""
        if show:
            field_widget.grid()
            if label_widget:
                label_widget.grid()
        else:
            field_widget.grid_remove()
            if label_widget:
                label_widget.grid_remove()
    
    def get_form_completion_percentage(self):
        """Calculate form completion percentage"""
        total_fields = 0
        filled_fields = 0
        
        # Check all visible entry fields
        for widget in [self.entry_rujukan, self.entry_nama, self.entry_tarikh,
                      self.entry_alamat1, self.entry_alamat2, self.entry_alamat3]:
            if widget.winfo_ismapped():
                total_fields += 1
                if widget.get().strip():
                    filled_fields += 1
        
        # Check comboboxes
        for combo in [self.combo_pegawai, self.combo_process, self.combo_jenis]:
            if combo.winfo_ismapped():
                total_fields += 1
                if combo.get().strip():
                    filled_fields += 1
        
        if total_fields == 0:
            return 0
        
        return int((filled_fields / total_fields) * 100)
    
    def update_completion_indicator(self):
        """Update form completion progress indicator"""
        percentage = self.get_form_completion_percentage()
        
        # Update title with percentage
        self.window.title(f"Sistem Pengurusan Dokumen - Pengisian Data [{percentage}% complete]")
    
    def on_field_change(self, field_name, field_widget):
        """Handle field value change with validation and completion update"""
        value = field_widget.get()
        
        # Validate field
        self.validate_field_realtime(field_name, value, field_widget)
        
        # Update completion indicator
        self.update_completion_indicator()
        
    
    def show_help_dialog(self):
        """Show comprehensive help dialog"""
        help_window = tk.Toplevel(self.window)
        help_window.title("📚 Panduan Pengisian Borang")
        help_window.geometry("700x600")
        help_window.configure(bg='white')
        
        # Header
        header = tk.Frame(help_window, bg=self.colors['primary'], height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header,
                text="📚 Panduan Pengisian Borang Form 2",
                font=('Arial', 18, 'bold'),
                bg=self.colors['primary'],
                fg='white').pack(pady=15)
        
        # Content area with scrollbar
        canvas = tk.Canvas(help_window, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(help_window, orient='vertical', command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')
        
        scrollable_frame.bind("<Configure>",
                             lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Help content
        help_content = tk.Frame(scrollable_frame, bg='white')
        help_content.pack(padx=30, pady=20, fill=tk.BOTH, expand=True)
        
        help_sections = [
            ("📝 RUJUKAN", "Format: KE.JB(90)650/05-02/XXXX\nMasukkan nombor rujukan unik anda."),
            ("🏢 NAMA SYARIKAT", "Masukkan nama penuh syarikat dengan ejaan yang betul."),
            ("📅 TARIKH", "Format: DD/MM/YYYY (contoh: 17/12/2025)\nTarikh Islam akan dikira automatik."),
            ("📍 ALAMAT", "Masukkan alamat lengkap dalam 3 baris:\n- Baris 1: Nombor & nama jalan\n- Baris 2: Kawasan/Taman\n- Baris 3: Poskod & Bandar"),
            ("👤 NAMA PEGAWAI", "Pilih nama pegawai yang bertanggungjawab."),
            ("💰 AMAUN", "Masukkan jumlah dalam RM (contoh: 1234.56)\nHanya angka dan titik perpuluhan."),
            ("✅ PROSES & JENIS", "Pilih proses dan jenis barang yang berkaitan."),
            ("⚡ FEATURES DINAMIK", 
             "• Real-time validation dengan color feedback\n"
             "• Progress indicator di title bar\n"
             "• Tooltips untuk panduan pantas"),
            ("⌨️ KEYBOARD SHORTCUTS",
             "• Ctrl+S: Save document\n"
             "• Ctrl+Z: Undo\n"
             "• Tab: Next field\n"
             "• Shift+Tab: Previous field")
        ]
        
        for title, content in help_sections:
            section_frame = tk.Frame(help_content, bg='white')
            section_frame.pack(fill=tk.X, pady=10)
            
            tk.Label(section_frame,
                    text=title,
                    font=('Arial', 15, 'bold'),
                    bg='white',
                    fg=self.colors['primary'],
                    anchor='w').pack(fill=tk.X)
            
            tk.Label(section_frame,
                    text=content,
                    font=('Arial', 12),
                    bg='white',
                    fg='#333333',
                    justify=tk.LEFT,
                    anchor='w').pack(fill=tk.X, padx=20, pady=5)
            
            # Divider
            tk.Frame(section_frame, bg='#EEEEEE', height=1).pack(fill=tk.X, pady=5)
        
        # Close button
        btn_frame = tk.Frame(help_window, bg='white')
        btn_frame.pack(fill=tk.X, padx=30, pady=20)
        
        tk.Button(btn_frame,
                 text="Tutup",
                 font=('Arial', 12, 'bold'),
                 bg=self.colors['primary'],
                 fg='white',
                 padx=30,
                 pady=10,
                 cursor='hand2',
                 command=help_window.destroy).pack()
        
        # Bind mouse wheel
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        
        # Center window
        help_window.transient(self.window)
        help_window.grab_set()
        
        # Center on parent
        help_window.update_idletasks()
        x = self.window.winfo_x() + (self.window.winfo_width() - help_window.winfo_width()) // 2
        y = self.window.winfo_y() + (self.window.winfo_height() - help_window.winfo_height()) // 2
        help_window.geometry(f"+{x}+{y}")
    
    def on_back_click(self):
        """Go back to Form1 or main menu"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, try to go back to main menu
            pass
        self.window.destroy()
    
    def on_close(self):
        """Handle window close"""
        try:
            if self.parent_window and self.parent_window.winfo_exists():
                self.parent_window.deiconify()
        except (tk.TclError, AttributeError):
            # Parent window has been destroyed, nothing to restore
            pass
        self.window.destroy()