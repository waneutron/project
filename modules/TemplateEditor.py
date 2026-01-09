"""
TemplateEditor.py - Template Management System for Non-IT Users
User-friendly interface to view, edit, and manage document templates
Templates are now embedded within this file - no external files needed!
"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from docx import Document
import os
import shutil
from datetime import datetime
import re
import base64
import io
import tempfile

# Import helper function
try:
    from helpers.docx_helper import preview_placeholders
except ImportError:
    def preview_placeholders(doc):
        """Extract placeholders from document"""
        placeholders = set()
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = re.findall(r'<<[^>]+>>', text)
            placeholders.update(matches)
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'<<[^>]+>>', text)
                        placeholders.update(matches)
        return sorted(list(placeholders))


# ==================== EMBEDDED TEMPLATE STORAGE ====================
# Templates are stored as base64-encoded strings
# To add/update templates: Use the TemplateEditor UI or encode files manually

class EmbeddedTemplateStorage:
    """Storage for embedded templates with version tracking"""
    
    def __init__(self):
        # Dictionary to store templates: {filename: {'content': base64, 'metadata': {...}}}
        self.templates = {}
        self._load_default_templates()
    
    def _load_default_templates(self):
        """Load default empty templates structure"""
        # Template categories for organization (based on template_categories_table.md)
        # Templates can belong to multiple categories (e.g., pelupusan_penjualan is both APPROVAL and DISPOSAL)
        template_categories = {
            'APPROVAL': [
                'ames_pedagang.docx',
                'ames_pengilang.docx',
                'surat kelulusan butiran 5D (Lulus).docx',
                'pelupusan_penjualan.docx',  # Can be approval for sale
                'pelupusan_skrap.docx'  # Can be approval for scrap
            ],
            'REJECTION': [
                'pelupusan_tidak_lulus.docx',
                'surat kelulusan butiran 5D (tidak lulus).docx'
            ],
            'DISPOSAL': [
                'pelupusan_pemusnahan.docx',
                'pelupusan_penjualan.docx',
                'pelupusan_skrap.docx'
            ],
            'REGISTRATION': [
                'signUpB.docx'
            ],
            'Delete Item': [
                'delete_item.docx',
                'delete_item_ames.docx'
            ],
            'Lain-lain': [
                'batal_sijil.docx'
            ]
        }
        
        # Initialize all templates with metadata structure
        for category, templates in template_categories.items():
            for filename in templates:
                if filename not in self.templates:
                    self.templates[filename] = {
                        'content': None,
                        'metadata': {
                            'category': category,
                            'version': '1.0',
                            'created_date': None,
                            'modified_date': None,
                            'is_new': True,
                            'description': self._get_template_description(filename, category)
                        }
                    }
        
        # Try to load from external files if they exist (migration)
        from helpers.resource_path import get_templates_dir
        templates_dir = get_templates_dir()
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit to prevent hanging on huge files
        
        if os.path.exists(templates_dir):
            try:
                template_keys = list(self.templates.keys())
                for filename in template_keys:
                    filepath = os.path.join(templates_dir, filename)
                    if os.path.exists(filepath):
                        try:
                            # Check file size first
                            try:
                                file_size = os.path.getsize(filepath)
                                if file_size > MAX_FILE_SIZE:
                                    print(f"Warning: Template {filename} is too large ({file_size / 1024 / 1024:.1f}MB), skipping")
                                    continue
                            except (OSError, IOError) as size_error:
                                print(f"Warning: Could not get size for {filename}: {size_error}")
                                continue
                            
                            try:
                                with open(filepath, 'rb') as f:
                                    content = f.read()
                            except (IOError, OSError, PermissionError) as read_error:
                                print(f"Warning: Could not read template {filename}: {read_error}")
                                continue
                            
                            try:
                                file_stat = os.stat(filepath)
                            except (OSError, IOError):
                                file_stat = None
                            
                            # Preserve existing metadata if available
                            existing_metadata = self.templates.get(filename, {}).get('metadata', {})
                            
                            try:
                                encoded_content = base64.b64encode(content).decode('utf-8')
                            except Exception as encode_error:
                                print(f"Warning: Could not encode template {filename}: {encode_error}")
                                continue
                            
                            self.templates[filename] = {
                                'content': encoded_content,
                                'metadata': {
                                    'category': existing_metadata.get('category', 'Lain-lain'),
                                    'version': existing_metadata.get('version', '1.0'),
                                    'created_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat() if (file_stat and existing_metadata.get('created_date') is None) else existing_metadata.get('created_date'),
                                    'modified_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat() if file_stat else datetime.now().isoformat(),
                                    'is_new': False,  # Migrated from old system
                                    'description': existing_metadata.get('description', self._get_template_description(filename, existing_metadata.get('category', 'Lain-lain')))
                                }
                            }
                        except Exception as e:
                            print(f"Warning: Could not load template {filename} from file system: {e}")
                            import traceback
                            traceback.print_exc()
                            continue  # Keep default structure
            except Exception as e:
                print(f"Warning: Error processing templates directory: {e}")
                import traceback
                traceback.print_exc()
            
            # Also check for any .docx files in Templates folder that aren't in our list
            try:
                files = os.listdir(templates_dir)
                for file in files:
                    if file.endswith('.docx') and file not in self.templates:
                        # Add new template found in folder
                        filepath = os.path.join(templates_dir, file)
                        try:
                            # Check file size first
                            try:
                                file_size = os.path.getsize(filepath)
                                if file_size > MAX_FILE_SIZE:
                                    print(f"Warning: Template {file} is too large ({file_size / 1024 / 1024:.1f}MB), skipping")
                                    continue
                            except (OSError, IOError) as size_error:
                                print(f"Warning: Could not get size for {file}: {size_error}")
                                continue
                            
                            try:
                                with open(filepath, 'rb') as f:
                                    content = f.read()
                            except (IOError, OSError, PermissionError) as read_error:
                                print(f"Warning: Could not read template {file}: {read_error}")
                                continue
                            
                            try:
                                file_stat = os.stat(filepath)
                            except (OSError, IOError):
                                file_stat = None
                            
                            try:
                                encoded_content = base64.b64encode(content).decode('utf-8')
                            except Exception as encode_error:
                                print(f"Warning: Could not encode template {file}: {encode_error}")
                                continue
                            
                            category = self._detect_category(file)
                            self.templates[file] = {
                                'content': encoded_content,
                                'metadata': {
                                    'category': category,
                                    'version': '1.0',
                                    'created_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat() if file_stat else datetime.now().isoformat(),
                                    'modified_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat() if file_stat else datetime.now().isoformat(),
                                    'is_new': False,
                                    'description': self._get_template_description(file, category)
                                }
                            }
                        except Exception as e:
                            print(f"Warning: Could not load template {file} from file system: {e}")
                            import traceback
                            traceback.print_exc()
            except Exception as e:
                print(f"Warning: Could not scan Templates folder: {e}")
                import traceback
                traceback.print_exc()
    
    def _get_template_description(self, filename, category):
        """Get description for template (based on template_categories_table.md)"""
        descriptions = {
            # APPROVAL templates
            'ames_pedagang.docx': 'Templat Kelulusan AMES untuk Pedagang (APPROVAL)',
            'ames_pengilang.docx': 'Templat Kelulusan AMES untuk Pengilang (APPROVAL)',
            'surat kelulusan butiran 5D (Lulus).docx': 'Templat Kelulusan Butiran 5D (APPROVAL)',
            'pelupusan_penjualan.docx': 'Templat Kelulusan Pelupusan melalui Penjualan (APPROVAL/DISPOSAL)',
            'pelupusan_skrap.docx': 'Templat Kelulusan Pelupusan Skrap (APPROVAL/DISPOSAL)',
            
            # REJECTION templates
            'pelupusan_tidak_lulus.docx': 'Templat Penolakan Pelupusan (REJECTION)',
            'surat kelulusan butiran 5D (tidak lulus).docx': 'Templat Penolakan Butiran 5D (REJECTION)',
            
            # DISPOSAL templates
            'pelupusan_pemusnahan.docx': 'Templat Pelupusan melalui Pemusnahan (DISPOSAL)',
            
            # REGISTRATION templates
            'signUpB.docx': 'Templat Pendaftaran Sign Up B (REGISTRATION)',
            
            # Delete Item templates
            'delete_item.doc': 'Templat untuk pemadaman item (format lama)',
            'delete_item_ames.docx': 'Templat untuk pemadaman item AMES (format baru)',
            
            # Other templates
            'batal_sijil.doc': 'Templat untuk pembatalan sijil'
        }
        return descriptions.get(filename, f'Templat {category}')
    
    def get_template(self, filename):
        """Get template content as Document object"""
        if filename not in self.templates:
            return None
        
        # Check if file is .doc (old format) - python-docx doesn't support it
        if filename.lower().endswith('.doc') and not filename.lower().endswith('.docx'):
            print(f"Warning: Template {filename} is in old .doc format. python-docx only supports .docx files.")
            print(f"Please convert {filename} to .docx format or use TemplateEditor to upload a .docx version.")
            return None
        
        template_data = self.templates[filename]
        if template_data is None or template_data.get('content') is None:
            return None
        
        try:
            # Decode base64 to bytes
            content = base64.b64decode(template_data['content'])
            # Create Document from bytes
            doc_stream = io.BytesIO(content)
            return Document(doc_stream)
        except Exception as e:
            # Check if it's a .doc file causing the error
            if filename.lower().endswith('.doc') and not filename.lower().endswith('.docx'):
                print(f"Error loading template {filename}: Old .doc format not supported. Please convert to .docx")
            else:
                print(f"Error loading template {filename}: {e}")
            return None
    
    def get_template_metadata(self, filename):
        """Get template metadata"""
        if filename not in self.templates:
            return None
        return self.templates[filename].get('metadata', {})
    
    def save_template(self, filename, doc, is_update=False):
        """Save template from Document object"""
        try:
            # Save Document to bytes
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            content = doc_stream.read()
            # Encode to base64
            
            # Get or create metadata
            if filename not in self.templates:
                category = self._detect_category(filename)
                self.templates[filename] = {
                    'content': None,
                    'metadata': {
                        'category': category,
                        'version': '1.0',
                        'created_date': None,
                        'modified_date': None,
                        'is_new': True,
                        'description': self._get_template_description(filename, category)
                    }
                }
            
            metadata = self.templates[filename].get('metadata', {})
            now = datetime.now().isoformat()
            
            # Update metadata
            if metadata.get('created_date') is None:
                metadata['created_date'] = now
            metadata['modified_date'] = now
            metadata['is_new'] = False  # No longer new after first save
            if is_update:
                # Increment version on update
                try:
                    version = float(metadata.get('version', '1.0'))
                    metadata['version'] = f"{version + 0.1:.1f}"
                except:
                    metadata['version'] = '1.1'
            
            self.templates[filename] = {
                'content': base64.b64encode(content).decode('utf-8'),
                'metadata': metadata
            }
            return True
        except Exception as e:
            print(f"Error saving template {filename}: {e}")
            return False
    
    def _detect_category(self, filename):
        """Detect template category from filename"""
        filename_lower = filename.lower()
        if 'pelupusan' in filename_lower:
            return 'Pelupusan'
        elif 'ames' in filename_lower:
            return 'AMES'
        elif 'delete' in filename_lower:
            return 'Delete Item'
        else:
            return 'Lain-lain'
    
    def save_template_from_file(self, filename, filepath, is_new=True):
        """Save template from file path"""
        try:
            with open(filepath, 'rb') as f:
                content = f.read()
            
            file_stat = os.stat(filepath)
            category = self._detect_category(filename)
            now = datetime.now().isoformat()
            
            self.templates[filename] = {
                'content': base64.b64encode(content).decode('utf-8'),
                'metadata': {
                    'category': category,
                    'version': '1.0' if is_new else '1.1',
                    'created_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat() if is_new else now,
                    'modified_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                    'is_new': is_new,
                    'description': self._get_template_description(filename, category)
                }
            }
            return True
        except Exception as e:
            print(f"Error saving template from file {filename}: {e}")
            return False
    
    def list_templates(self, category=None):
        """List all available templates, optionally filtered by category"""
        templates = []
        for name, data in self.templates.items():
            if data and data.get('content') is not None:
                if category is None or data.get('metadata', {}).get('category') == category:
                    templates.append(name)
        return templates
    
    def has_template(self, filename):
        """Check if template exists"""
        return (filename in self.templates and 
                self.templates[filename] is not None and
                self.templates[filename].get('content') is not None)
    
    def is_new_template(self, filename):
        """Check if template is marked as new"""
        if filename not in self.templates:
            return False
        return self.templates[filename].get('metadata', {}).get('is_new', False)
    
    def export_template_to_file(self, filename, filepath):
        """Export template to file"""
        if not self.has_template(filename):
            return False
        
        try:
            content = base64.b64decode(self.templates[filename]['content'])
            with open(filepath, 'wb') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Error exporting template {filename}: {e}")
            return False
    
    def delete_template(self, filename):
        """Delete template"""
        if filename in self.templates:
            # Keep metadata but clear content
            metadata = self.templates[filename].get('metadata', {})
            self.templates[filename] = {
                'content': None,
                'metadata': metadata
            }
            return True
        return False
    
    def get_templates_by_category(self):
        """Get templates organized by category"""
        categories = {}
        for filename, data in self.templates.items():
            if data and data.get('content') is not None:
                category = data.get('metadata', {}).get('category', 'Lain-lain')
                if category not in categories:
                    categories[category] = []
                categories[category].append(filename)
        return categories
    
    def add_template_from_file(self, filepath, new_filename=None):
        """Add template from external file"""
        if not os.path.exists(filepath):
            return False
        
        filename = new_filename or os.path.basename(filepath)
        return self.save_template_from_file(filename, filepath)


class TemplateEditor:
    """Template Editor for managing document templates"""
    
    def __init__(self, root, parent_window=None):
        self.root = root
        self.parent_window = parent_window
        self.root.title("Pengurusan Templat Dokumen")
        self.root.geometry("1600x1000")
        self.root.configure(bg='#F5F5F5')
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#F5F5F5',
            'bg_white': 'white',
            'text_dark': '#1a1a1a',
            'text_light': 'white',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_success': '#2E7D32',
            'button_danger': '#C62828',
            'button_warning': '#F57C00',
            'button_hover': '#004d99'
        }
        
        # Temporary directory for editing
        self.temp_dir = tempfile.mkdtemp(prefix="template_editor_")
        
        # Backup storage (in memory)
        self.backups = {}
        
        # Initialize templates list
        self.templates = []
        
        # Center window
        self.center_window()
        
        # Make window appear on top
        self.root.attributes('-topmost', True)
        self.root.after_idle(lambda: self.root.attributes('-topmost', False))
        self.root.lift()
        self.root.focus_force()
        
        # Create interface first (show window immediately)
        self.create_header()
        self.create_main_content()
        
        # Update window to ensure it's displayed
        self.root.update()
        
        # Load templates asynchronously after window is shown (with delay to ensure UI is ready)
        self.root.after(200, self.load_templates_async)
        
        # Handle window close
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
    
    def _is_template_storage_ready(self):
        """Check if template storage is initialized"""
        return hasattr(self, 'template_storage') and self.template_storage is not None
    
    def load_templates_async(self):
        """Load templates asynchronously to avoid blocking UI"""
        try:
            # Show loading message
            if hasattr(self, 'template_listbox') and self.template_listbox:
                try:
                    self.template_listbox.delete(0, tk.END)
                    self.template_listbox.insert(0, "Loading templates...")
                    self.root.update_idletasks()
                except:
                    pass
            
            # Initialize template storage (this may take time)
            try:
                self.template_storage = EmbeddedTemplateStorage()
            except Exception as storage_error:
                print(f"Error initializing template storage: {storage_error}")
                import traceback
                traceback.print_exc()
                # Create minimal storage on error
                from types import SimpleNamespace
                self.template_storage = SimpleNamespace()
                self.template_storage.templates = {}
                self.template_storage.list_templates = lambda category=None: []
                self.template_storage.has_template = lambda filename: False
                self.template_storage.is_new_template = lambda filename: False
                self.template_storage.get_template_metadata = lambda filename: {}
                self.template_storage.get_template = lambda filename: None
                self.template_storage.save_template = lambda filename, doc, is_update=False: False
                self.template_storage.save_template_from_file = lambda filename, filepath, is_new=True: False
                self.template_storage.export_template_to_file = lambda filename, filepath: False
                self.template_storage.delete_template = lambda filename: False
                self.template_storage.add_template_from_file = lambda filepath, new_filename=None: False
                self.template_storage.get_templates_by_category = lambda: {}
                if hasattr(self, 'template_listbox') and self.template_listbox:
                    try:
                        self.template_listbox.delete(0, tk.END)
                        self.template_listbox.insert(0, f"Warning: Template storage initialization failed. Using minimal storage.")
                    except:
                        pass
                return
            
            # Reload templates in UI
            if hasattr(self, 'reload_templates'):
                try:
                    self.reload_templates()
                except Exception as reload_error:
                    print(f"Error reloading templates: {reload_error}")
                    import traceback
                    traceback.print_exc()
                    if hasattr(self, 'template_listbox') and self.template_listbox:
                        try:
                            self.template_listbox.delete(0, tk.END)
                            self.template_listbox.insert(0, f"Error loading template list: {reload_error}")
                        except:
                            pass
        except Exception as e:
            import traceback
            traceback.print_exc()
            # Show error but don't block
            if hasattr(self, 'template_listbox') and self.template_listbox:
                try:
                    self.template_listbox.delete(0, tk.END)
                    self.template_listbox.insert(0, f"Error loading templates: {e}")
                except:
                    pass
            # Ensure we have a storage object even if it's minimal
            if not hasattr(self, 'template_storage'):
                from types import SimpleNamespace
                self.template_storage = SimpleNamespace()
                self.template_storage.templates = {}
                self.template_storage.list_templates = lambda category=None: []
                self.template_storage.has_template = lambda filename: False
                self.template_storage.is_new_template = lambda filename: False
                self.template_storage.get_template_metadata = lambda filename: {}
                self.template_storage.get_template = lambda filename: None
                self.template_storage.save_template = lambda filename, doc, is_update=False: False
                self.template_storage.save_template_from_file = lambda filename, filepath, is_new=True: False
                self.template_storage.export_template_to_file = lambda filename, filepath: False
                self.template_storage.delete_template = lambda filename: False
                self.template_storage.add_template_from_file = lambda filepath, new_filename=None: False
                self.template_storage.get_templates_by_category = lambda: {}
    
    def center_window(self):
        """Center window on screen"""
        self.root.update_idletasks()
        width = 1600
        height = 1000
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def create_header(self):
        """Create header section"""
        header_frame = tk.Frame(self.root, bg=self.colors['primary'], height=100)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Title
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.pack(side=tk.LEFT, expand=True, padx=30)
        
        title = tk.Label(title_frame,
                        text="PENGURUSAN TEMPLAT DOKUMEN",
                        font=('Arial', 16, 'bold'),
                        bg=self.colors['primary'],
                        fg='white')
        title.pack(pady=(25, 5))
        
        subtitle = tk.Label(title_frame,
                           text="Sistem Pengurusan dan Penyuntingan Templat",
                           font=('Arial', 11),
                           bg=self.colors['primary'],
                           fg='#E0E0E0')
        subtitle.pack()
        
        # Form Editor button (prominent)
        btn_form_editor = tk.Button(header_frame,
                                   text="⚙️ FORM EDITOR",
                                   font=('Arial', 10, 'bold'),
                                   bg='#9C27B0',
                                   fg='white',
                                   relief=tk.FLAT,
                                   cursor='hand2',
                                   width=18,
                                   height=2,
                                   command=self.open_form_editor_menu)
        btn_form_editor.pack(side=tk.RIGHT, padx=10, pady=25)
        
        # Help button (always visible)
        btn_help = tk.Button(header_frame,
                            text="❓ BANTUAN",
                            font=('Arial', 10, 'bold'),
                            bg='#1976D2',
                            fg='white',
                            relief=tk.FLAT,
                            cursor='hand2',
                            width=15,
                            height=2,
                            command=self.open_help)
        btn_help.pack(side=tk.RIGHT, padx=10, pady=25)
        
        # Back button
        if self.parent_window:
            btn_back = tk.Button(header_frame,
                                text="← KEMBALI",
                                font=('Arial', 10, 'bold'),
                                bg=self.colors['secondary'],
                                fg='white',
                                relief=tk.FLAT,
                                cursor='hand2',
                                width=15,
                                height=2,
                                command=self.on_back_click)
            btn_back.pack(side=tk.RIGHT, padx=10, pady=25)
        
        # Separator
        separator = tk.Frame(self.root, bg=self.colors['accent'], height=2)
        separator.pack(fill=tk.X)
    
    def create_main_content(self):
        """Create main content area"""
        main_container = tk.Frame(self.root, bg=self.colors['bg_main'])
        main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Left panel - Template list
        left_panel = tk.Frame(main_container, bg='white')
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=False, padx=(0, 15))
        left_panel.config(width=400)
        
        # Border
        border_left = tk.Frame(left_panel, bg=self.colors['border'], bd=1)
        border_left.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content_left = tk.Frame(border_left, bg='white')
        content_left.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Template list title
        tk.Label(content_left,
                text="Senarai Templat",
                font=('Arial', 14, 'bold'),
                bg='white',
                fg=self.colors['primary']).pack(anchor='w', pady=(0, 10))
        
        # Category filter buttons
        filter_frame = tk.Frame(content_left, bg='white')
        filter_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.filter_category = tk.StringVar(value="Semua")
        categories = ["Semua", "APPROVAL", "REJECTION", "DISPOSAL", "REGISTRATION", "Delete Item", "Lain-lain", "🆕 Baru Sahaja"]
        
        for i, cat in enumerate(categories):
            btn = tk.Radiobutton(filter_frame,
                               text=cat,
                               variable=self.filter_category,
                               value=cat,
                               font=('Arial', 9),
                               bg='white',
                               command=self.apply_category_filter)
            btn.grid(row=i//3, column=i%3, sticky='w', padx=5, pady=2)
        
        # Search box
        search_frame = tk.Frame(content_left, bg='white')
        search_frame.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(search_frame,
                text="Cari:",
                font=('Arial', 10),
                bg='white').pack(side=tk.LEFT, padx=(0, 10))
        
        self.search_var = tk.StringVar()
        self.search_var.trace('w', self.filter_templates)
        search_entry = tk.Entry(search_frame,
                               textvariable=self.search_var,
                               font=('Arial', 10),
                               relief=tk.SOLID,
                               bd=1)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Template listbox with scrollbar
        list_frame = tk.Frame(content_left, bg='white')
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.template_listbox = tk.Listbox(list_frame,
                                          font=('Arial', 11),
                                          yscrollcommand=scrollbar.set,
                                          selectmode=tk.SINGLE)
        self.template_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.template_listbox.yview)
        
        self.template_listbox.bind('<<ListboxSelect>>', self.on_template_select)
        
        # Action buttons
        btn_frame = tk.Frame(content_left, bg='white')
        btn_frame.pack(fill=tk.X)
        
        tk.Button(btn_frame,
                 text="📝 Buka dalam Word",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_primary'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.open_in_word).pack(fill=tk.X, pady=5)
        
        tk.Button(btn_frame,
                 text="✏️ Edit Templat (Teks)",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_success'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.edit_template_text).pack(fill=tk.X, pady=5)
        
        tk.Button(btn_frame,
                 text="📋 Salin Templat",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_success'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.copy_template).pack(fill=tk.X, pady=5)
        
        tk.Button(btn_frame,
                 text="💾 Eksport ke Fail",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_warning'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.export_template).pack(fill=tk.X, pady=5)
        
        tk.Button(btn_frame,
                 text="🗑️ Padam Templat",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_danger'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.delete_template).pack(fill=tk.X, pady=5)
        
        # Separator for Form Editor section
        separator_frame = tk.Frame(content_left, bg=self.colors['border'], height=2)
        separator_frame.pack(fill=tk.X, pady=15)
        
        # Form Editor section title
        tk.Label(content_left,
                text="Form Editor",
                font=('Arial', 12, 'bold'),
                bg='white',
                fg=self.colors['primary']).pack(anchor='w', pady=(10, 5))
        
        tk.Label(content_left,
                text="Edit forms, fields, and mappings",
                font=('Arial', 9),
                bg='white',
                fg='#666666').pack(anchor='w', pady=(0, 10))
        
        # Form selection frame
        form_select_frame = tk.Frame(content_left, bg='white')
        form_select_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(form_select_frame,
                text="Select Form:",
                font=('Arial', 9, 'bold'),
                bg='white').pack(side=tk.LEFT, padx=(0, 5))
        
        self.form_selector = ttk.Combobox(form_select_frame,
                                         values=['Form2', 'Form3', 'FormSignUp', 'FormDeleteItem'],
                                         state='readonly',
                                         width=20,
                                         font=('Arial', 9))
        self.form_selector.set('Form2')
        self.form_selector.pack(side=tk.LEFT)
        
        # Form Editor button
        tk.Button(content_left,
                 text="⚙️ Edit Selected Form",
                 font=('Arial', 11, 'bold'),
                 bg='#9C27B0',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.open_form_editor).pack(fill=tk.X, pady=5)
        
        tk.Button(content_left,
                 text="➕ Create New Form",
                 font=('Arial', 10, 'bold'),
                 bg='#1976D2',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.create_new_form).pack(fill=tk.X, pady=5)
        
        tk.Button(btn_frame,
                 text="🔄 Muat Semula",
                 font=('Arial', 10, 'bold'),
                 bg=self.colors['button_primary'],
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=self.reload_templates).pack(fill=tk.X, pady=5)
        
        # Right panel - Template text editor
        right_panel = tk.Frame(main_container, bg='white')
        right_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Border
        border_right = tk.Frame(right_panel, bg=self.colors['border'], bd=1)
        border_right.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        content_right = tk.Frame(border_right, bg='white')
        content_right.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Template info title
        self.info_title = tk.Label(content_right,
                                  text="Pilih templat untuk mengedit",
                                  font=('Arial', 14, 'bold'),
                                  bg='white',
                                  fg=self.colors['primary'])
        self.info_title.pack(anchor='w', pady=(0, 10))
        
        # Template name display
        self.template_name_label = tk.Label(content_right,
                                          text="",
                                          font=('Arial', 11),
                                          bg='white',
                                          fg=self.colors['text_dark'])
        self.template_name_label.pack(anchor='w', pady=(0, 10))
        
        # Template text editor section with helpful info
        editor_header_frame = tk.Frame(content_right, bg='white')
        editor_header_frame.pack(fill=tk.X, pady=(0, 5))
        
        editor_label = tk.Label(editor_header_frame,
                               text="Edit Kandungan Templat:",
                               font=('Arial', 12, 'bold'),
                               bg='white',
                               fg=self.colors['primary'])
        editor_label.pack(side=tk.LEFT)
        
        # Status indicator
        self.editor_status_label = tk.Label(editor_header_frame,
                                           text="",
                                           font=('Arial', 9),
                                           bg='white',
                                           fg='#666666')
        self.editor_status_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # Placeholder helper button
        btn_placeholder_helper = tk.Button(editor_header_frame,
                                          text="💡 Bantuan Placeholder",
                                          font=('Arial', 9),
                                          bg='#E3F2FD',
                                          fg=self.colors['primary'],
                                          relief=tk.FLAT,
                                          cursor='hand2',
                                          command=self.show_placeholder_helper)
        btn_placeholder_helper.pack(side=tk.RIGHT)
        
        # Text editor frame
        editor_frame = tk.Frame(content_right, bg='white')
        editor_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        editor_scroll = ttk.Scrollbar(editor_frame)
        editor_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.template_text_editor = tk.Text(editor_frame,
                                           font=('Courier', 11),
                                           wrap=tk.WORD,
                                           yscrollcommand=editor_scroll.set,
                                           bg='#FAFAFA',
                                           relief=tk.SOLID,
                                           bd=2,
                                           padx=10,
                                           pady=10,
                                           state=tk.DISABLED)
        self.template_text_editor.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        editor_scroll.config(command=self.template_text_editor.yview)
        
        # Bind text change to update status
        self.template_text_editor.bind('<KeyRelease>', self.on_text_change)
        self.template_text_editor.bind('<Button-1>', self.on_text_change)
        
        # Character/word count and placeholder count
        stats_frame = tk.Frame(content_right, bg='white')
        stats_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.stats_label = tk.Label(stats_frame,
                                   text="",
                                   font=('Arial', 9),
                                   bg='white',
                                   fg='#666666')
        self.stats_label.pack(side=tk.LEFT)
        
        # Action buttons frame
        btn_save_frame = tk.Frame(content_right, bg='white')
        btn_save_frame.pack(fill=tk.X, pady=(15, 0))
        
        # Left side buttons
        btn_left_frame = tk.Frame(btn_save_frame, bg='white')
        btn_left_frame.pack(side=tk.LEFT)
        
        self.btn_save_template = tk.Button(btn_left_frame,
                                          text="💾 Simpan Templat",
                                          font=('Arial', 11, 'bold'),
                                          bg=self.colors['button_success'],
                                          fg='white',
                                          relief=tk.FLAT,
                                          cursor='hand2',
                                          command=self.save_template_from_text,
                                          state=tk.DISABLED,
                                          padx=15,
                                          pady=8)
        self.btn_save_template.pack(side=tk.LEFT, padx=(0, 10))
        
        # Add hover effect
        def on_save_enter(e):
            if self.btn_save_template['state'] == 'normal':
                self.btn_save_template.config(bg='#2E8B57')
        def on_save_leave(e):
            if self.btn_save_template['state'] == 'normal':
                self.btn_save_template.config(bg=self.colors['button_success'])
        self.btn_save_template.bind('<Enter>', on_save_enter)
        self.btn_save_template.bind('<Leave>', on_save_leave)
        
        self.btn_new_template = tk.Button(btn_left_frame,
                                         text="➕ Templat Baru",
                                         font=('Arial', 11, 'bold'),
                                         bg=self.colors['button_primary'],
                                         fg='white',
                                         relief=tk.FLAT,
                                         cursor='hand2',
                                         command=self.create_new_template,
                                         padx=15,
                                         pady=8)
        self.btn_new_template.pack(side=tk.LEFT, padx=(0, 10))
        
        # Preview button
        self.btn_preview = tk.Button(btn_left_frame,
                                    text="👁️ Pratonton",
                                    font=('Arial', 11, 'bold'),
                                    bg='#9C27B0',
                                    fg='white',
                                    relief=tk.FLAT,
                                    cursor='hand2',
                                    command=self.preview_template,
                                    state=tk.DISABLED,
                                    padx=15,
                                    pady=8)
        self.btn_preview.pack(side=tk.LEFT)
        
        # Right side - validation status
        self.validation_label = tk.Label(btn_save_frame,
                                         text="",
                                         font=('Arial', 9),
                                         bg='white',
                                         fg='#666666')
        self.validation_label.pack(side=tk.RIGHT, padx=(10, 0))
        
        # Quick tips section (collapsible)
        tips_frame = tk.Frame(content_right, bg='#F5F5F5', relief=tk.SOLID, bd=1)
        tips_frame.pack(fill=tk.X, pady=(20, 0))
        
        tips_header = tk.Frame(tips_frame, bg='#E8E8E8')
        tips_header.pack(fill=tk.X, padx=1, pady=1)
        
        self.tips_expanded = tk.BooleanVar(value=False)
        tips_toggle = tk.Checkbutton(tips_header,
                                    text="💡 Tips Pantas",
                                    font=('Arial', 10, 'bold'),
                                    bg='#E8E8E8',
                                    variable=self.tips_expanded,
                                    command=self.toggle_tips,
                                    cursor='hand2')
        tips_toggle.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.tips_content = tk.Frame(tips_frame, bg='white')
        
        instructions_text = """✓ Taip atau paste kandungan templat di atas
✓ Gunakan placeholder seperti <<NAMA_SYARIKAT>>, <<TARIKH>>, <<ALAMAT>>
✓ Setiap baris baru akan menjadi paragraph baru dalam Word
✓ Klik "💾 Simpan Templat" untuk menyimpan
✓ Gunakan "👁️ Pratonton" untuk melihat hasil sebelum simpan
✓ Klik "💡 Bantuan Placeholder" untuk senarai lengkap placeholder"""
        
        self.tips_label = tk.Label(self.tips_content,
                                   text=instructions_text,
                                   font=('Arial', 10),
                                   bg='white',
                                   fg=self.colors['text_dark'],
                                   justify=tk.LEFT,
                                   anchor='w',
                                   padx=15,
                                   pady=10)
        self.tips_label.pack(anchor='w', fill=tk.X)
        
        # Initialize current editing template
        self.current_editing_template = None
        
        # Templates will be loaded asynchronously in load_templates_async()
        # Show placeholder message
        self.template_listbox.insert(0, "Loading templates...")
    
    def apply_category_filter(self):
        """Apply category filter to template list"""
        if not hasattr(self, 'template_storage') or not self._is_template_storage_ready():
            if hasattr(self, 'template_listbox'):
                self.template_listbox.delete(0, tk.END)
                self.template_listbox.insert(0, "Loading templates...")
            return
        
        if not hasattr(self, 'templates'):
            self.templates = []
        
        category = self.filter_category.get()
        self.template_listbox.delete(0, tk.END)
        
        try:
            if category == "Semua":
                templates_to_show = self.templates if self.templates else []
            elif category == "🆕 Baru Sahaja":
                templates_to_show = [t for t in (self.templates if self.templates else []) if self.template_storage.is_new_template(t)]
            else:
                templates_to_show = self.template_storage.list_templates(category=category)
                # Also include empty templates in this category
                for name, data in self.template_storage.templates.items():
                    if data and data.get('metadata', {}).get('category') == category:
                        if name not in templates_to_show:
                            templates_to_show.append(name)
            
            # Apply search filter if any
            search_term = self.search_var.get().lower()
            
            for template in sorted(templates_to_show):
                if search_term and search_term not in template.lower():
                    continue
                
                has_template = self.template_storage.has_template(template)
                is_new = self.template_storage.is_new_template(template)
                metadata = self.template_storage.get_template_metadata(template)
                
                status = "✓" if has_template else "○"
                new_indicator = " 🆕" if (is_new and has_template) else ""
                version = metadata.get('version', 'N/A') if metadata else 'N/A'
                template_category = metadata.get('category', '') if metadata else ''
                
                display_name = f"{status} {template}{new_indicator} [v{version}]"
                if template_category:
                    display_name += f" ({template_category})"
                
                self.template_listbox.insert(tk.END, display_name)
        except Exception as e:
            self.template_listbox.insert(0, f"Error filtering templates: {e}")
    
    def reload_templates(self):
        """Reload template list from embedded storage"""
        if not hasattr(self, 'template_storage') or not self._is_template_storage_ready():
            if hasattr(self, 'template_listbox'):
                self.template_listbox.delete(0, tk.END)
                self.template_listbox.insert(0, "Loading templates...")
            return
        
        if not hasattr(self, 'template_listbox'):
            return
        
        try:
            self.template_listbox.delete(0, tk.END)
        except:
            pass
        
        if not hasattr(self, 'templates'):
            self.templates = []
        else:
            self.templates = []
        
        try:
            # Get all templates from embedded storage
            if not hasattr(self.template_storage, 'list_templates'):
                self.template_listbox.insert(0, "Template storage not ready")
                return
            
            template_names = self.template_storage.list_templates()
            
            # Also include template names that exist in storage (even if empty)
            if hasattr(self.template_storage, 'templates'):
                all_template_names = list(self.template_storage.templates.keys())
            else:
                all_template_names = []
            
            # Combine and sort
            self.templates = sorted(set(template_names + all_template_names))
            
            for template in self.templates:
                try:
                    has_template = self.template_storage.has_template(template) if hasattr(self.template_storage, 'has_template') else False
                    is_new = self.template_storage.is_new_template(template) if hasattr(self.template_storage, 'is_new_template') else False
                    metadata = self.template_storage.get_template_metadata(template) if hasattr(self.template_storage, 'get_template_metadata') else {}
                    
                    # Build display string with status indicators
                    status = "✓" if has_template else "○"
                    new_indicator = " 🆕" if (is_new and has_template) else ""
                    version = metadata.get('version', 'N/A') if metadata else 'N/A'
                    category = metadata.get('category', '') if metadata else ''
                    
                    # Format: ✓ Template Name 🆕 [v1.0] (Category)
                    display_name = f"{status} {template}{new_indicator} [v{version}]"
                    if category:
                        display_name += f" ({category})"
                    
                    self.template_listbox.insert(tk.END, display_name)
                except Exception as e:
                    # Skip this template if there's an error
                    print(f"Error processing template {template}: {e}")
                    continue
        except Exception as e:
            import traceback
            traceback.print_exc()
            try:
                self.template_listbox.delete(0, tk.END)
                self.template_listbox.insert(0, f"Error loading templates: {e}")
            except:
                pass
    
    def filter_templates(self, *args):
        """Filter templates based on search - applies both search and category filter"""
        self.apply_category_filter()  # This will also apply search filter
    
    def on_template_select(self, event):
        """Handle template selection"""
        selection = self.template_listbox.curselection()
        if not selection:
            return
        
        display_name = self.template_listbox.get(selection[0])
        # Extract template name (remove status, version, category info)
        # Format: "✓ Template Name 🆕 [v1.0] (Category)"
        template_name = display_name
        # Remove status indicators
        template_name = template_name.replace("✓ ", "").replace("○ ", "")
        # Remove new indicator
        template_name = template_name.replace(" 🆕", "")
        # Remove version info [vX.X]
        import re
        template_name = re.sub(r'\s*\[v[\d.]+\]', '', template_name)
        # Remove category (Category)
        template_name = re.sub(r'\s*\([^)]+\)$', '', template_name).strip()
        
        self.load_template_info(template_name)
    
    def load_template_info(self, template_name):
        """Load and display template information - now shows in text editor"""
        if not hasattr(self, 'template_storage'):
            self.info_title.config(text="Templat: Loading...")
            self.template_name_label.config(text="Nama: Loading...")
            return
        
        # Just update the title, text editor will be populated when user clicks edit
        try:
            metadata = self.template_storage.get_template_metadata(template_name)
            
            if self.template_storage.has_template(template_name):
                version = metadata.get('version', 'N/A') if metadata else 'N/A'
                category = metadata.get('category', '') if metadata else ''
                is_new = metadata.get('is_new', False) if metadata else False
                new_label = " 🆕 BARU" if is_new else ""
                
                title_text = f"Templat: {template_name} [v{version}]{new_label}"
                if category:
                    title_text += f" - {category}"
                self.info_title.config(text=title_text)
                self.template_name_label.config(text=f"Nama: {template_name}")
            else:
                self.info_title.config(text=f"Templat: {template_name} (Kosong)")
                self.template_name_label.config(text=f"Nama: {template_name}")
        except Exception as e:
            self.info_title.config(text=f"Templat: {template_name} (Error)")
            self.template_name_label.config(text=f"Error: {e}")
        
        # Disable editor until user clicks edit
        self.template_text_editor.config(state=tk.DISABLED)
        self.template_text_editor.delete('1.0', tk.END)
        self.template_text_editor.insert('1.0', "Klik '✏️ Edit Templat (Teks)' untuk mengedit kandungan templat ini.")
        self.btn_save_template.config(state=tk.DISABLED)
    
    def open_in_word(self):
        """Open selected template in Microsoft Word"""
        if not self._is_template_storage_ready():
            messagebox.showwarning("Amaran", "Templates are still loading. Please wait...")
            return
        
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Amaran", "Sila pilih templat terlebih dahulu")
            return
        
        display_name = self.template_listbox.get(selection[0])
        # Extract template name (remove status, version, category info)
        # Format: "✓ Template Name 🆕 [v1.0] (Category)"
        template_name = display_name
        # Remove status indicators
        template_name = template_name.replace("✓ ", "").replace("○ ", "")
        # Remove new indicator
        template_name = template_name.replace(" 🆕", "")
        # Remove version info [vX.X]
        import re
        template_name = re.sub(r'\s*\[v[\d.]+\]', '', template_name)
        # Remove category (Category)
        template_name = re.sub(r'\s*\([^)]+\)$', '', template_name).strip()
        
        # Check if template has content
        if not self.template_storage.has_template(template_name):
            # Try to auto-load from Templates folder
            from helpers.resource_path import get_template_path
            template_path = get_template_path(template_name)
            if os.path.exists(template_path):
                try:
                    if hasattr(self.template_storage, 'add_template_from_file'):
                        if self.template_storage.add_template_from_file(template_path, template_name):
                            messagebox.showinfo("Berjaya", f"Templat {template_name} telah dimuatkan. Cuba buka semula.")
                            self.reload_templates()
                            return
                    else:
                        # Try using save_template_from_file directly
                        if hasattr(self.template_storage, 'save_template_from_file'):
                            if self.template_storage.save_template_from_file(template_name, template_path, is_new=False):
                                messagebox.showinfo("Berjaya", f"Templat {template_name} telah dimuatkan. Cuba buka semula.")
                                self.reload_templates()
                                return
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    messagebox.showerror("Ralat", f"Tidak dapat memuatkan templat: {str(e)}")
                    return
            else:
                messagebox.showwarning("Amaran", "Templat ini kosong. Sila import templat terlebih dahulu.")
                return
        
        try:
            # Ensure temp directory exists
            if not os.path.exists(self.temp_dir):
                os.makedirs(self.temp_dir)
            
            # Export template to temporary file
            temp_path = os.path.join(self.temp_dir, template_name)
            if not self.template_storage.export_template_to_file(template_name, temp_path):
                messagebox.showerror("Ralat", "Tidak dapat mengeksport templat")
                return
            
            # Verify file was created
            if not os.path.exists(temp_path):
                messagebox.showerror("Ralat", f"Fail tidak dapat dicipta: {temp_path}")
                return
            
            # Open file with default application (Word)
            try:
                if os.name == 'nt':  # Windows
                    os.startfile(temp_path)
                elif os.name == 'posix':  # macOS/Linux
                    import subprocess
                    subprocess.call(['open', temp_path])
                else:  # Linux
                    import subprocess
                    subprocess.call(['xdg-open', temp_path])
                
                messagebox.showinfo("Berjaya", 
                                  f"Templat dibuka dalam Microsoft Word.\n\n"
                                  f"Selepas membuat perubahan:\n"
                                  f"1. Simpan fail dalam Word (Ctrl+S)\n"
                                  f"2. Tutup Word\n"
                                  f"3. Klik '➕ Import dari Fail' dan pilih fail yang disimpan\n"
                                  f"4. Perubahan akan disimpan dalam sistem\n\n"
                                  f"Lokasi fail sementara: {temp_path}")
            except Exception as open_error:
                # Try alternative method for Windows
                if os.name == 'nt':
                    try:
                        import subprocess
                        subprocess.Popen([temp_path], shell=True)
                        messagebox.showinfo("Berjaya", 
                                          f"Templat dibuka dalam Microsoft Word.\n\n"
                                          f"Selepas membuat perubahan:\n"
                                          f"1. Simpan fail dalam Word (Ctrl+S)\n"
                                          f"2. Tutup Word\n"
                                          f"3. Klik '➕ Import dari Fail' dan pilih fail yang disimpan\n"
                                          f"4. Perubahan akan disimpan dalam sistem\n\n"
                                          f"Lokasi fail sementara: {temp_path}")
                    except Exception as e2:
                        messagebox.showerror("Ralat", 
                                           f"Tidak dapat membuka fail dalam Word.\n\n"
                                           f"Ralat: {str(e2)}\n\n"
                                           f"Fail telah dieksport ke: {temp_path}\n\n"
                                           f"Sila buka fail ini secara manual dalam Microsoft Word.")
                else:
                    messagebox.showerror("Ralat", f"Tidak dapat membuka templat: {str(open_error)}")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat membuka templat: {str(e)}\n\nLokasi fail sementara: {self.temp_dir}")
    
    def import_template(self):
        """Import template from external file"""
        filepath = filedialog.askopenfilename(
            title="Pilih Fail Templat",
            filetypes=[
                ("Word Documents", "*.docx *.doc"),
                ("All Files", "*.*")
            ]
        )
        
        if not filepath:
            return
        
        # Ask for template name
        default_name = os.path.basename(filepath)
        template_name = simpledialog.askstring(
            "Import Templat",
            f"Masukkan nama untuk templat:\n\n"
            f"Fail: {default_name}",
            initialvalue=default_name
        )
        
        if not template_name:
            return
        
        # Ensure .docx extension
        if not template_name.endswith(('.docx', '.doc')):
            template_name += '.docx'
        
        # Check if template exists
        if self.template_storage.has_template(template_name):
            if not messagebox.askyesno("Amaran", f"Templat {template_name} sudah wujud. Gantikan?"):
                return
        
        try:
            # Check if this is a new template or updating existing
            is_new = not self.template_storage.has_template(template_name)
            
            # Use save_template_from_file directly since it supports is_new parameter
            if hasattr(self.template_storage, 'save_template_from_file'):
                success = self.template_storage.save_template_from_file(template_name, filepath, is_new=is_new)
            elif hasattr(self.template_storage, 'add_template_from_file'):
                success = self.template_storage.add_template_from_file(filepath, template_name)
            else:
                success = False
            
            if success:
                status = "diimport" if is_new else "dikemaskini"
                messagebox.showinfo("Berjaya", f"Templat berjaya {status}: {template_name}")
                self.reload_templates()
                # Select the newly imported template
                for i in range(self.template_listbox.size()):
                    item = self.template_listbox.get(i)
                    if template_name in item:
                        self.template_listbox.selection_set(i)
                        self.template_listbox.see(i)
                        self.on_template_select(None)
                        break
            else:
                messagebox.showerror("Ralat", "Tidak dapat mengimport templat")
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat mengimport templat: {str(e)}")
    
    def copy_template(self):
        """Create a copy of the selected template"""
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Amaran", "Sila pilih templat terlebih dahulu")
            return
        
        display_name = self.template_listbox.get(selection[0])
        template_name = display_name.replace("✓ ", "").replace("○ ", "")
        
        if not self.template_storage.has_template(template_name):
            messagebox.showwarning("Amaran", "Templat ini kosong. Tidak dapat disalin.")
            return
        
        # Ask for new name
        new_name = simpledialog.askstring("Salin Templat",
                                         f"Masukkan nama baru untuk templat:\n\n"
                                         f"Templat asal: {template_name}",
                                         initialvalue=f"copy_of_{template_name}")
        
        if not new_name:
            return
        
        # Ensure .docx extension
        if not new_name.endswith(('.docx', '.doc')):
            new_name += '.docx'
        
        # Check if template exists
        if self.template_storage.has_template(new_name):
            if not messagebox.askyesno("Amaran", f"Templat {new_name} sudah wujud. Gantikan?"):
                return
        
        try:
            # Get template document
            doc = self.template_storage.get_template(template_name)
            if doc and self.template_storage.save_template(new_name, doc):
                messagebox.showinfo("Berjaya", f"Templat berjaya disalin sebagai: {new_name}")
                self.reload_templates()
            else:
                messagebox.showerror("Ralat", "Tidak dapat menyalin templat")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat menyalin templat: {str(e)}")
    
    def export_template(self):
        """Export template to external file"""
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Amaran", "Sila pilih templat terlebih dahulu")
            return
        
        display_name = self.template_listbox.get(selection[0])
        template_name = display_name.replace("✓ ", "").replace("○ ", "")
        
        if not self.template_storage.has_template(template_name):
            messagebox.showwarning("Amaran", "Templat ini kosong. Tidak dapat dieksport.")
            return
        
        filepath = filedialog.asksaveasfilename(
            title="Simpan Templat",
            defaultextension=".docx",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Word 97-2003", "*.doc"),
                ("All Files", "*.*")
            ],
            initialfile=template_name
        )
        
        if not filepath:
            return
        
        try:
            if self.template_storage.export_template_to_file(template_name, filepath):
                messagebox.showinfo("Berjaya", f"Templat berjaya dieksport ke:\n{filepath}")
            else:
                messagebox.showerror("Ralat", "Tidak dapat mengeksport templat")
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat mengeksport templat: {str(e)}")
    
    def delete_template(self):
        """Delete template from storage"""
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Amaran", "Sila pilih templat terlebih dahulu")
            return
        
        display_name = self.template_listbox.get(selection[0])
        template_name = display_name.replace("✓ ", "").replace("○ ", "")
        
        if not self.template_storage.has_template(template_name):
            messagebox.showwarning("Amaran", "Templat ini sudah kosong.")
            return
        
        if messagebox.askyesno("Pengesahan", 
                              f"Adakah anda pasti ingin memadam templat:\n{template_name}?\n\n"
                              f"Tindakan ini tidak boleh dibatalkan."):
            try:
                if self.template_storage.delete_template(template_name):
                    messagebox.showinfo("Berjaya", f"Templat {template_name} berjaya dipadam")
                    self.reload_templates()
                else:
                    messagebox.showerror("Ralat", "Tidak dapat memadam templat")
            except Exception as e:
                messagebox.showerror("Ralat", f"Tidak dapat memadam templat: {str(e)}")
    
    def edit_template_text(self):
        """Open text editor for selected template"""
        # Check if template storage is ready
        if not self._is_template_storage_ready():
            messagebox.showwarning("Amaran", "Templates masih sedang dimuatkan. Sila tunggu sebentar dan cuba lagi.")
            return
        
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Amaran", "Sila pilih templat terlebih dahulu")
            return
        
        try:
            display_name = self.template_listbox.get(selection[0])
        except Exception as e:
            messagebox.showerror("Ralat", f"Tidak dapat membaca pilihan templat: {e}")
            return
        
        # Extract template name
        template_name = display_name
        template_name = template_name.replace("✓ ", "").replace("○ ", "")
        template_name = template_name.replace(" 🆕", "")
        import re
        template_name = re.sub(r'\s*\[v[\d.]+\]', '', template_name)
        template_name = re.sub(r'\s*\([^)]+\)$', '', template_name).strip()
        
        if not template_name:
            messagebox.showwarning("Amaran", "Nama templat tidak sah")
            return
        
        self.current_editing_template = template_name
        
        # Load existing template content or start fresh
        template_text = ""
        try:
            if not hasattr(self.template_storage, 'get_template'):
                template_text = ""
            else:
                try:
                    doc = self.template_storage.get_template(template_name)
                    if doc is None:
                        template_text = ""
                    else:
                        # Convert Word document to text
                        text_content = []
                        try:
                            # Safely access paragraphs
                            if hasattr(doc, 'paragraphs'):
                                for para in doc.paragraphs:
                                    try:
                                        if para and hasattr(para, 'text') and para.text.strip():
                                            text_content.append(para.text)
                                    except Exception as para_error:
                                        print(f"Error reading paragraph: {para_error}")
                                        continue
                            
                            # Safely access tables
                            if hasattr(doc, 'tables'):
                                for table in doc.tables:
                                    try:
                                        if table and hasattr(table, 'rows'):
                                            for row in table.rows:
                                                try:
                                                    if row and hasattr(row, 'cells'):
                                                        row_text = []
                                                        for cell in row.cells:
                                                            try:
                                                                if cell and hasattr(cell, 'paragraphs'):
                                                                    cell_text = ' '.join([p.text for p in cell.paragraphs if p and hasattr(p, 'text') and p.text.strip()])
                                                                    if cell_text:
                                                                        row_text.append(cell_text)
                                                            except Exception as cell_error:
                                                                print(f"Error reading cell: {cell_error}")
                                                                continue
                                                        if row_text:
                                                            text_content.append(' | '.join(row_text))
                                                except Exception as row_error:
                                                    print(f"Error reading row: {row_error}")
                                                    continue
                                    except Exception as table_error:
                                        print(f"Error reading table: {table_error}")
                                        continue
                            
                            template_text = '\n'.join(text_content)
                        except Exception as convert_error:
                            import traceback
                            traceback.print_exc()
                            print(f"Error converting template to text: {convert_error}")
                            template_text = ""
                except Exception as get_error:
                    import traceback
                    traceback.print_exc()
                    print(f"Error getting template: {get_error}")
                    template_text = ""
        except Exception as e:
            template_text = ""
            import traceback
            traceback.print_exc()
            print(f"Error loading template: {e}")
            # Don't show messagebox here, just use empty text
        
        # Update UI - with comprehensive error handling
        try:
            if hasattr(self, 'info_title') and self.info_title:
                self.info_title.config(text=f"Edit Templat: {template_name}")
        except Exception as e:
            print(f"Error updating info_title: {e}")
        
        try:
            if hasattr(self, 'template_name_label') and self.template_name_label:
                self.template_name_label.config(text=f"Nama: {template_name}")
        except Exception as e:
            print(f"Error updating template_name_label: {e}")
        
        try:
            if hasattr(self, 'template_text_editor') and self.template_text_editor:
                self.template_text_editor.config(state=tk.NORMAL)
                self.template_text_editor.delete('1.0', tk.END)
                self.template_text_editor.insert('1.0', template_text)
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Error updating template_text_editor: {e}")
            messagebox.showerror("Ralat", f"Tidak dapat membuka editor teks: {e}")
            return
        
        try:
            if hasattr(self, 'btn_save_template') and self.btn_save_template:
                self.btn_save_template.config(state=tk.NORMAL)
        except Exception as e:
            print(f"Error updating btn_save_template: {e}")
        
        try:
            if hasattr(self, 'btn_preview') and self.btn_preview:
                self.btn_preview.config(state=tk.NORMAL)
        except Exception as e:
            print(f"Error updating btn_preview: {e}")
        
        try:
            if hasattr(self, 'editor_status_label') and self.editor_status_label:
                self.editor_status_label.config(text="✏️ Sedang mengedit...", fg='#1976D2')
        except Exception as e:
            print(f"Error updating editor_status_label: {e}")
        
        # Update stats
        try:
            if hasattr(self, 'update_stats'):
                self.update_stats()
        except Exception as e:
            print(f"Error updating stats: {e}")
        
        try:
            if hasattr(self, 'validate_template'):
                self.validate_template()
        except Exception as e:
            print(f"Error validating template: {e}")
    
    def create_new_template(self):
        """Create a new template"""
        template_name = simpledialog.askstring(
            "Templat Baru",
            "Masukkan nama templat baru:",
            initialvalue="template_baru.docx"
        )
        
        if not template_name:
            return
        
        # Ensure .docx extension
        if not template_name.endswith(('.docx', '.doc')):
            template_name += '.docx'
        
        # Check if template exists
        if self.template_storage.has_template(template_name):
            if not messagebox.askyesno("Amaran", f"Templat {template_name} sudah wujud. Gantikan?"):
                return
        
        self.current_editing_template = template_name
        
        # Update UI
        self.info_title.config(text=f"Templat Baru: {template_name}")
        self.template_name_label.config(text=f"Nama: {template_name}")
        self.template_text_editor.config(state=tk.NORMAL)
        self.template_text_editor.delete('1.0', tk.END)
        self.template_text_editor.insert('1.0', "Masukkan kandungan templat di sini...\n\nGunakan placeholder seperti:\n<<NAMA_SYARIKAT>>\n<<TARIKH>>\n<<ALAMAT>>")
        self.btn_save_template.config(state=tk.NORMAL)
        
        # Add to list if not exists
        if template_name not in self.template_storage.templates:
            category = self._detect_category(template_name)
            self.template_storage.templates[template_name] = {
                'content': None,
                'metadata': {
                    'category': category,
                    'version': '1.0',
                    'created_date': datetime.now().isoformat(),
                    'modified_date': datetime.now().isoformat(),
                    'is_new': True,
                    'description': f'Templat baru: {template_name}'
                }
            }
        
        self.reload_templates()
    
    def _detect_category(self, filename):
        """Detect template category from filename"""
        filename_lower = filename.lower()
        if 'pelupusan' in filename_lower:
            return 'DISPOSAL'
        elif 'ames' in filename_lower:
            return 'APPROVAL'
        elif 'delete' in filename_lower:
            return 'Delete Item'
        elif 'signup' in filename_lower or 'sign_up' in filename_lower:
            return 'REGISTRATION'
        else:
            return 'Lain-lain'
    
    def save_template_from_text(self):
        """Save template from text editor to Word document"""
        if not hasattr(self, 'current_editing_template') or not self.current_editing_template:
            messagebox.showwarning("Amaran", "Tiada templat dipilih untuk disimpan")
            return
        
        template_name = self.current_editing_template
        text_content = self.template_text_editor.get('1.0', tk.END).strip()
        
        if not text_content:
            messagebox.showwarning("Amaran", "Kandungan templat tidak boleh kosong")
            return
        
        try:
            # Convert text to Word document
            doc = Document()
            
            # Split by lines and create paragraphs
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():  # Only add non-empty lines
                    para = doc.add_paragraph(line.strip())
                    # Apply formatting
                    from docx.shared import Pt
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
            
            # Save template
            is_new = not self.template_storage.has_template(template_name)
            
            if self.template_storage.save_template(template_name, doc, is_update=not is_new):
                status = "dicipta" if is_new else "dikemaskini"
                messagebox.showinfo("✅ Berjaya!", 
                    f"Templat berjaya {status}!\n\n"
                    f"Nama: {template_name}\n"
                    f"Status: Siap digunakan dalam sistem")
                self.reload_templates()
                self.editor_status_label.config(text="✅ Disimpan!", fg='#2E7D32')
                
                # Select the saved template
                for i in range(self.template_listbox.size()):
                    item = self.template_listbox.get(i)
                    if template_name in item:
                        self.template_listbox.selection_set(i)
                        self.template_listbox.see(i)
                        break
            else:
                messagebox.showerror("❌ Ralat", 
                    "Tidak dapat menyimpan templat.\n\n"
                    "Sila cuba lagi atau hubungi pentadbir sistem.")
        except Exception as e:
            messagebox.showerror("Ralat", f"Ralat menyimpan templat: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def open_form_editor_menu(self):
        """Open form editor menu dialog"""
        menu_window = tk.Toplevel(self.root)
        menu_window.title("Form Editor - Select Form")
        menu_window.geometry("500x400")
        menu_window.configure(bg='#F5F5F5')
        menu_window.transient(self.root)
        menu_window.grab_set()
        
        # Center window
        menu_window.update_idletasks()
        x = (menu_window.winfo_screenwidth() // 2) - 250
        y = (menu_window.winfo_screenheight() // 2) - 200
        menu_window.geometry(f'500x400+{x}+{y}')
        # Make window appear on top
        menu_window.attributes('-topmost', True)
        menu_window.after_idle(lambda: menu_window.attributes('-topmost', False))
        menu_window.lift()
        menu_window.focus_force()
        
        # Header
        header = tk.Frame(menu_window, bg='#003366', height=80)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header,
                text="⚙️ Form Editor",
                font=('Arial', 16, 'bold'),
                bg='#003366',
                fg='white').pack(pady=25)
        
        # Content
        content = tk.Frame(menu_window, bg='white', padx=30, pady=30)
        content.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(content,
                text="Select a form to edit:",
                font=('Arial', 12, 'bold'),
                bg='white',
                fg='#003366').pack(anchor='w', pady=(0, 15))
        
        # Form buttons
        forms = [
            ('Form2', 'Pelupusan & Lain-lain Forms'),
            ('Form3', 'AMES Forms'),
            ('FormSignUp', 'Sign Up Registration'),
            ('FormDeleteItem', 'Delete Item AMES')
        ]
        
        for form_name, description in forms:
            btn_frame = tk.Frame(content, bg='white')
            btn_frame.pack(fill=tk.X, pady=8)
            
            btn = tk.Button(btn_frame,
                          text=f"📝 {form_name}",
                          font=('Arial', 11, 'bold'),
                          bg='#9C27B0',
                          fg='white',
                          relief=tk.FLAT,
                          cursor='hand2',
                          width=25,
                          command=lambda fn=form_name: self.open_form_editor(fn, menu_window))
            btn.pack(side=tk.LEFT, padx=(0, 10))
            
            tk.Label(btn_frame,
                    text=description,
                    font=('Arial', 9),
                    bg='white',
                    fg='#666666').pack(side=tk.LEFT)
        
        # Create new form button
        tk.Label(content,
                text="Or create a new form:",
                font=('Arial', 10, 'bold'),
                bg='white',
                fg='#003366').pack(anchor='w', pady=(20, 10))
        
        tk.Button(content,
                 text="➕ Create New Form",
                 font=('Arial', 11, 'bold'),
                 bg='#1976D2',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 width=25,
                 command=lambda: self.create_new_form(menu_window)).pack(pady=10)
        
        # Close button
        tk.Button(content,
                 text="Tutup",
                 font=('Arial', 10),
                 bg='#666666',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 width=15,
                 command=menu_window.destroy).pack(pady=(20, 0))
    
    def open_form_editor(self, form_name=None, parent_window=None):
        """Open form editor for specified form"""
        try:
            # Check if root window exists and is valid
            if not hasattr(self, 'root') or not self.root:
                messagebox.showerror("Error", "Parent window is not available")
                return
            
            try:
                from helpers.dynamic_form_editor import DynamicFormEditor
            except ImportError as import_error:
                import traceback
                traceback.print_exc()
                messagebox.showerror("Error", 
                    f"Cannot import form editor module: {import_error}\n\n"
                    "Please ensure helpers/dynamic_form_editor.py exists.")
                return
            
            # If form_name is not provided, try to get it from form_selector
            if form_name is None:
                if hasattr(self, 'form_selector') and self.form_selector:
                    try:
                        form_name = self.form_selector.get()
                        if not form_name or form_name.strip() == "":
                            form_name = None
                    except Exception as e:
                        print(f"Error getting form from selector: {e}")
                        import traceback
                        traceback.print_exc()
                        form_name = None
                else:
                    form_name = None
            
            if not form_name:
                messagebox.showwarning("Warning", "Please select a form to edit from the dropdown")
                return
            
            # Close parent window if provided
            if parent_window:
                try:
                    if parent_window.winfo_exists():
                        parent_window.destroy()
                except Exception as e:
                    print(f"Error closing parent window: {e}")
            
            # Open form editor with error handling
            try:
                # Ensure root window is valid
                if not self.root.winfo_exists():
                    messagebox.showerror("Error", "Parent window has been closed")
                    return
                
                DynamicFormEditor(self.root, form_name)
            except tk.TclError as tcl_error:
                import traceback
                traceback.print_exc()
                messagebox.showerror("Error", 
                    f"Window error: {tcl_error}\n\n"
                    "The parent window may have been closed.")
            except Exception as editor_error:
                import traceback
                traceback.print_exc()
                error_msg = str(editor_error)
                messagebox.showerror("Error", 
                    f"Error opening form editor: {error_msg}\n\n"
                    f"Form: {form_name}\n\n"
                    "Please check the console for more details.")
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Error", f"Unexpected error: {e}")
    
    def create_new_form(self, parent_window=None):
        """Create a new form"""
        try:
            from tkinter import simpledialog
            
            # Close parent window if provided
            if parent_window:
                try:
                    parent_window.destroy()
                except Exception as e:
                    print(f"Error closing parent window: {e}")
            
            form_name = simpledialog.askstring("New Form", "Enter form name:")
            if form_name:
                try:
                    from helpers.dynamic_form_editor import DynamicFormEditor
                    DynamicFormEditor(self.root, form_name)
                except ImportError as e:
                    import traceback
                    traceback.print_exc()
                    messagebox.showerror("Error", f"Cannot import form editor module: {e}\n\nPlease ensure helpers/dynamic_form_editor.py exists.")
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    messagebox.showerror("Error", f"Error creating form: {e}")
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Error", f"Error in create_new_form: {e}")
    
    def on_back_click(self):
        """Go back to main menu"""
        if self.parent_window:
            self.parent_window.deiconify()
        self.root.destroy()
    
    def open_help(self):
        """Open help/documentation window"""
        help_window = tk.Toplevel(self.root)
        help_window.title("Bantuan - Panduan Pengurusan Templat")
        help_window.geometry("900x700")
        help_window.configure(bg='#F5F5F5')
        
        # Center window
        help_window.update_idletasks()
        width = 900
        height = 700
        x = (help_window.winfo_screenwidth() // 2) - (width // 2)
        y = (help_window.winfo_screenheight() // 2) - (height // 2)
        help_window.geometry(f'{width}x{height}+{x}+{y}')
        # Make window appear on top
        help_window.attributes('-topmost', True)
        help_window.after_idle(lambda: help_window.attributes('-topmost', False))
        help_window.lift()
        help_window.focus_force()
        
        # Header
        header = tk.Frame(help_window, bg='#003366', height=60)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header,
                text="📚 Panduan Pengurusan Templat",
                font=('Arial', 16, 'bold'),
                bg='#003366',
                fg='white').pack(pady=15)
        
        # Content with scrollbar
        content_frame = tk.Frame(help_window, bg='#F5F5F5')
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        canvas = tk.Canvas(content_frame, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(content_frame, orient='vertical', command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Help content
        help_content = tk.Text(scrollable_frame,
                              font=('Arial', 11),
                              wrap=tk.WORD,
                              bg='white',
                              padx=30,
                              pady=30,
                              relief=tk.FLAT)
        help_content.pack(fill=tk.BOTH, expand=True)
        
        help_text = """
═══════════════════════════════════════════════════════════════
                    PANDUAN PENGURUSAN TEMPLAT
═══════════════════════════════════════════════════════════════

1. CARA MENAMBAH TEMPLAT BARU
───────────────────────────────────────────────────────────────

   Langkah 1: Sediakan fail Word (.docx) dengan format yang dikehendaki
   Langkah 2: Klik butang "➕ Import dari Fail" di panel kiri
   Langkah 3: Pilih fail Word anda
   Langkah 4: Masukkan nama untuk templat (contoh: "templat_baru.docx")
   Langkah 5: Templat akan disimpan dalam sistem

   Nota: Nama templat mestilah unik dan berakhir dengan .docx atau .doc


2. CARA MENYUNTING TEMPLAT
───────────────────────────────────────────────────────────────

   Langkah 1: Pilih templat dari senarai di sebelah kiri
   Langkah 2: Klik butang "📝 Buka dalam Word"
   Langkah 3: Fail akan dibuka dalam Microsoft Word
   Langkah 4: Buat perubahan yang diperlukan
   Langkah 5: Simpan fail dalam Word (Ctrl+S)
   Langkah 6: Tutup Word
   Langkah 7: Klik "➕ Import dari Fail" semula dan pilih fail yang sama
   Langkah 8: Pilih "Gantikan" apabila diminta

   Nota: Pastikan anda menggunakan pemegang tempat (placeholders) yang betul


3. PEMEGANG TEMPAT (PLACEHOLDERS)
───────────────────────────────────────────────────────────────

   Pemegang tempat adalah kod khas yang akan digantikan dengan data sebenar
   semasa dokumen dijana. Format: <<NAMA_PLACEHOLDER>>

   Contoh Pemegang Tempat Biasa:
   
   • <<NAMA_SYARIKAT>>     - Nama syarikat (akan ditukar kepada huruf besar)
   • <<ALAMAT>>            - Alamat syarikat (boleh berbilang baris)
   • <<TARIKH>>            - Tarikh (format: DD/MM/YYYY)
   • <<TARIKH2>>           - Tarikh (format: DD Month YYYY, contoh: 10 September 2025)
   • <<TARIKH_ISLAM>>      - Tarikh dalam kalendar Islam (contoh: 17 Rabiulawal 1447H)
   • <<RUJUKAN>>           - Nombor rujukan
   • <<NAMA_PEGAWAI>>      - Nama pegawai (huruf besar)
   • <<AMOUNT>>            - Jumlah dalam Ringgit Malaysia (RM)
   • <<PROSES>>            - Proses (pemusnahan/penjualan)
   • <<JENIS_BARANG>>      - Jenis barang
   • <<PENGECUALIAN>>      - Pengecualian
   • <<TAJUK_SURAT>>       - Tajuk surat (huruf besar dan tebal)
   • <<TAJUK_SURAT2>>      - Tajuk surat (huruf kecil untuk badan surat)


4. PEMEGANG TEMPAT KHAS MENGIKUT JENIS TEMPLAT
───────────────────────────────────────────────────────────────

   A. TEMPLAT PELUPUSAN:
      • <<RUJUKAN_TUAN>>   - Rujukan syarikat (jika ada)
      • <<RUJUKAN>>        - Rujukan kami
      • <<TARIKH_MULA>>    - Tarikh mula pemusnahan
      • <<TARIKH_TAMAT>>   - Tarikh tamat pemusnahan
      • <<TEMPOH>>         - Tempoh pemusnahan (dalam perkataan Melayu)
      • <<STATUS>>         - Status (SST-ADM atau AMES-03)

   B. TEMPLAT AMES:
      • <<RUJUKAN_KAMI>>   - Rujukan kami (format: KE.JB(90)650/14/AMES/XXX)
      • <<NO_KELULUSAN>>   - Nombor kelulusan AMES
      • <<KATEGORI>>       - Kategori (Pedagang atau Pengilang)
      • <<TEMPOH_KELULUSAN>> - Tempoh kelulusan (contoh: 01 April 2025 hingga 31 Mac 2027)
      • <<TARIKH_MALAY>>   - Tarikh dalam format Melayu

   C. TEMPLAT DELETE ITEM:
      • <<NO_KELULUSAN>>   - Nombor kelulusan AMES
      • Semua pemegang tempat AMES juga boleh digunakan


5. CARA MENGGUNAKAN PEMEGANG TEMPAT
───────────────────────────────────────────────────────────────

   Langkah 1: Buka templat dalam Word
   Langkah 2: Cari teks yang ingin digantikan dengan data dinamik
   Langkah 3: Gantikan teks tersebut dengan pemegang tempat
   
   Contoh:
   
   Sebelum: "Kepada: ABC SDN BHD"
   Selepas:  "Kepada: <<NAMA_SYARIKAT>>"
   
   Sebelum: "Tarikh: 10 September 2025"
   Selepas:  "Tarikh: <<TARIKH2>>"


6. FORMATTING DALAM TEMPLAT
───────────────────────────────────────────────────────────────

   • Teks dalam <<NAMA_SYARIKAT>> akan secara automatik menjadi HURUF BESAR
   • Teks dalam <<TAJUK_SURAT>> akan menjadi HURUF BESAR dan TEBAL
   • Anda boleh menggunakan format Word biasa (bold, italic, dll) pada teks lain
   • Jangan format pemegang tempat itu sendiri - biarkan sebagai teks biasa


7. CARA MEMBUAT SALINAN TEMPLAT
───────────────────────────────────────────────────────────────

   Langkah 1: Pilih templat yang ingin disalin
   Langkah 2: Klik butang "📋 Salin Templat"
   Langkah 3: Masukkan nama baru untuk salinan
   Langkah 4: Salinan akan dibuat dengan nama baru

   Kegunaan: Membuat variasi templat tanpa mengubah yang asal


8. CARA MENGEKSPORT TEMPLAT
───────────────────────────────────────────────────────────────

   Langkah 1: Pilih templat yang ingin dieksport
   Langkah 2: Klik butang "💾 Eksport ke Fail"
   Langkah 3: Pilih lokasi untuk menyimpan fail
   Langkah 4: Fail Word akan disimpan di lokasi yang dipilih

   Kegunaan: Membuat sandaran atau berkongsi templat dengan pengguna lain


9. CARA MEMADAM TEMPLAT
───────────────────────────────────────────────────────────────

   Langkah 1: Pilih templat yang ingin dipadam
   Langkah 2: Klik butang "🗑️ Padam Templat"
   Langkah 3: Sahkan pemadaman
   Langkah 4: Templat akan dipadam dari sistem

   AMARAN: Tindakan ini tidak boleh dibatalkan!


10. TIPS DAN AMALAN TERBAIK
───────────────────────────────────────────────────────────────

   ✓ Sentiasa buat salinan sebelum mengubah templat penting
   ✓ Gunakan nama templat yang jelas dan deskriptif
   ✓ Pastikan semua pemegang tempat ditulis dengan betul (case-sensitive)
   ✓ Uji templat dengan data sampel sebelum digunakan
   ✓ Simpan templat Word anda sebagai sandaran sebelum import
   ✓ Gunakan format standard JKDM untuk konsistensi
   ✓ Semak senarai pemegang tempat yang tersedia sebelum membuat templat


11. MENYELESAIKAN MASALAH
───────────────────────────────────────────────────────────────

   Masalah: Pemegang tempat tidak digantikan
   Penyelesaian: 
   - Pastikan ejaan pemegang tempat betul (case-sensitive)
   - Pastikan menggunakan format <<NAMA>> dengan tanda kurung berganda
   - Semak senarai pemegang tempat yang tersedia

   Masalah: Templat tidak boleh dibuka dalam Word
   Penyelesaian:
   - Pastikan Microsoft Word dipasang
   - Pastikan fail tidak rosak
   - Cuba eksport templat dan buka secara manual

   Masalah: Perubahan tidak tersimpan
   Penyelesaian:
   - Pastikan anda menyimpan fail dalam Word sebelum import semula
   - Pastikan anda memilih "Gantikan" apabila diminta


═══════════════════════════════════════════════════════════════
                    SELAMAT MENGGUNAKAN!
═══════════════════════════════════════════════════════════════
        """
        
        help_content.insert('1.0', help_text)
        help_content.config(state=tk.DISABLED)
        
        # Close button
        btn_close = tk.Button(help_window,
                             text="Tutup",
                             font=('Arial', 10, 'bold'),
                             bg='#003366',
                             fg='white',
                             relief=tk.FLAT,
                             cursor='hand2',
                             width=15,
                             height=2,
                             command=help_window.destroy)
        btn_close.pack(pady=10)
        
        # Bind mouse wheel to canvas only (not bind_all to avoid errors when window is closed)
        def on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except tk.TclError:
                pass  # Canvas already destroyed
        
        canvas.bind("<MouseWheel>", on_mousewheel)
        # Also bind to scrollable_frame for better scrolling
        scrollable_frame.bind("<MouseWheel>", on_mousewheel)
    
    def on_text_change(self, event=None):
        """Handle text changes in editor"""
        self.update_stats()
        self.validate_template()
    
    def update_stats(self):
        """Update character/word/placeholder statistics"""
        if not hasattr(self, 'template_text_editor'):
            return
        
        try:
            text = self.template_text_editor.get('1.0', tk.END)
            chars = len(text) - 1  # Exclude newline at end
            words = len(text.split())
            placeholders = len(re.findall(r'<<[^>]+>>', text))
            
            stats_text = f"📊 {chars} aksara | {words} perkataan | {placeholders} placeholder"
            self.stats_label.config(text=stats_text)
        except:
            pass
    
    def validate_template(self):
        """Validate template and show warnings"""
        if not hasattr(self, 'template_text_editor'):
            return
        
        try:
            text = self.template_text_editor.get('1.0', tk.END).strip()
            
            if not text or len(text) < 10:
                self.validation_label.config(text="⚠️ Kandungan terlalu pendek", fg='#F57C00')
                return
            
            # Check for placeholders
            placeholders = re.findall(r'<<[^>]+>>', text)
            if not placeholders:
                self.validation_label.config(text="💡 Tiada placeholder dijumpai", fg='#9C27B0')
            else:
                self.validation_label.config(text=f"✅ {len(placeholders)} placeholder dijumpai", fg='#2E7D32')
        except:
            pass
    
    def toggle_tips(self):
        """Toggle tips section visibility"""
        if self.tips_expanded.get():
            self.tips_content.pack(fill=tk.X, padx=1, pady=(0, 1))
        else:
            self.tips_content.pack_forget()
    
    def show_placeholder_helper(self):
        """Show placeholder helper window"""
        helper_window = tk.Toplevel(self.root)
        helper_window.title("Bantuan Placeholder")
        helper_window.geometry("600x500")
        helper_window.configure(bg='white')
        
        # Center window
        helper_window.update_idletasks()
        x = (helper_window.winfo_screenwidth() // 2) - (600 // 2)
        y = (helper_window.winfo_screenheight() // 2) - (500 // 2)
        helper_window.geometry(f'600x500+{x}+{y}')
        # Make window appear on top
        helper_window.attributes('-topmost', True)
        helper_window.lift()
        helper_window.focus_force()
        
        # Header
        header = tk.Frame(helper_window, bg='#003366', height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header,
                text="💡 Senarai Placeholder yang Tersedia",
                font=('Arial', 14, 'bold'),
                bg='#003366',
                fg='white').pack(pady=12)
        
        # Content
        content_frame = tk.Frame(helper_window, bg='white')
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        scrollbar = ttk.Scrollbar(content_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        text_widget = tk.Text(content_frame,
                             font=('Arial', 10),
                             wrap=tk.WORD,
                             yscrollcommand=scrollbar.set,
                             bg='white',
                             padx=15,
                             pady=15)
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=text_widget.yview)
        
        helper_text = """PLACEHOLDER UMUM:

<<NAMA_SYARIKAT>>     - Nama syarikat (huruf besar)
<<ALAMAT>>            - Alamat syarikat
<<TARIKH>>            - Tarikh (DD/MM/YYYY)
<<TARIKH2>>           - Tarikh (DD Month YYYY)
<<TARIKH_ISLAM>>      - Tarikh Islam
<<RUJUKAN>>           - Nombor rujukan
<<RUJUKAN_KAMI>>      - Rujukan kami
<<RUJUKAN_TUAN>>      - Rujukan tuan
<<NAMA_PEGAWAI>>      - Nama pegawai (huruf besar)
<<AMOUNT>>            - Jumlah (RM)

PLACEHOLDER PELUPUSAN:

<<PROSES>>            - Proses (pemusnahan/penjualan)
<<JENIS_BARANG>>      - Jenis barang
<<PENGECUALIAN>>      - Pengecualian
<<TARIKH_MULA>>       - Tarikh mula
<<TARIKH_TAMAT>>      - Tarikh tamat
<<TEMPOH>>            - Tempoh
<<TAJUK_SURAT>>       - Tajuk surat (huruf besar, tebal)
<<TAJUK_SURAT2>>      - Tajuk surat (huruf kecil)
<<STATUS>>            - Status

PLACEHOLDER AMES:

<<NO_KELULUSAN>>      - Nombor kelulusan AMES
<<KATEGORI>>          - Kategori (Pedagang/Pengilang)
<<TEMPOH_KELULUSAN>>  - Tempoh kelulusan
<<TARIKH_MALAY>>      - Tarikh format Melayu

CARA PENGGUNAAN:
1. Taip placeholder dalam format <<NAMA_PLACEHOLDER>>
2. Pastikan ejaan betul (case-sensitive)
3. Sistem akan gantikan dengan data sebenar semasa jana dokumen"""
        
        text_widget.insert('1.0', helper_text)
        text_widget.config(state=tk.DISABLED)
        
        # Close button
        tk.Button(helper_window,
                 text="Tutup",
                 font=('Arial', 10, 'bold'),
                 bg='#003366',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=helper_window.destroy,
                 padx=20,
                 pady=5).pack(pady=10)
    
    def preview_template(self):
        """Preview template with sample data"""
        if not hasattr(self, 'current_editing_template') or not self.current_editing_template:
            messagebox.showwarning("Amaran", "Tiada templat untuk dipratonton")
            return
        
        text_content = self.template_text_editor.get('1.0', tk.END).strip()
        if not text_content:
            messagebox.showwarning("Amaran", "Kandungan templat kosong")
            return
        
        # Create preview window
        preview_window = tk.Toplevel(self.root)
        preview_window.title("Pratonton Templat")
        preview_window.geometry("800x600")
        preview_window.configure(bg='white')
        
        # Center window
        preview_window.update_idletasks()
        x = (preview_window.winfo_screenwidth() // 2) - (800 // 2)
        y = (preview_window.winfo_screenheight() // 2) - (600 // 2)
        preview_window.geometry(f'800x600+{x}+{y}')
        # Make window appear on top
        preview_window.attributes('-topmost', True)
        preview_window.after_idle(lambda: preview_window.attributes('-topmost', False))
        preview_window.lift()
        preview_window.focus_force()
        
        # Header
        header = tk.Frame(preview_window, bg='#003366', height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        tk.Label(header,
                text="👁️ Pratonton Templat (dengan Data Sampel)",
                font=('Arial', 14, 'bold'),
                bg='#003366',
                fg='white').pack(pady=12)
        
        # Preview content
        content_frame = tk.Frame(preview_window, bg='white')
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        scrollbar = ttk.Scrollbar(content_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        preview_text = tk.Text(content_frame,
                              font=('Arial', 11),
                              wrap=tk.WORD,
                              yscrollcommand=scrollbar.set,
                              bg='#FAFAFA',
                              padx=15,
                              pady=15)
        preview_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=preview_text.yview)
        
        # Replace placeholders with sample data
        sample_data = {
            '<<NAMA_SYARIKAT>>': 'ABC SDN BHD',
            '<<ALAMAT>>': '123 Jalan ABC\n12345 Kuala Lumpur',
            '<<TARIKH>>': '15/01/2025',
            '<<TARIKH2>>': '15 Januari 2025',
            '<<TARIKH_ISLAM>>': '15 Jamadilawal 1446H',
            '<<RUJUKAN>>': 'KE.JB(90)650/14/001',
            '<<RUJUKAN_KAMI>>': 'KE.JB(90)650/14/001',
            '<<RUJUKAN_TUAN>>': 'REF/2025/001',
            '<<NAMA_PEGAWAI>>': 'AHMAD BIN ALI',
            '<<AMOUNT>>': 'RM 1,000.00',
            '<<PROSES>>': 'Pemusnahan',
            '<<JENIS_BARANG>>': 'Barang Terkawal',
            '<<PENGECUALIAN>>': 'Item 5D',
            '<<TARIKH_MULA>>': '01 Januari 2025',
            '<<TARIKH_TAMAT>>': '31 Disember 2025',
            '<<TEMPOH>>': 'Dua belas bulan',
            '<<TAJUK_SURAT>>': 'KELULUSAN PELUPUSAN',
            '<<TAJUK_SURAT2>>': 'kelulusan pelupusan',
            '<<STATUS>>': 'SST',
            '<<NO_KELULUSAN>>': 'AMES/2025/001',
            '<<KATEGORI>>': 'Pedagang',
            '<<TEMPOH_KELULUSAN>>': '01 April 2025 hingga 31 Mac 2027',
            '<<TARIKH_MALAY>>': '15 Januari 2025'
        }
        
        preview_content = text_content
        for placeholder, value in sample_data.items():
            preview_content = preview_content.replace(placeholder, value)
        
        preview_text.insert('1.0', preview_content)
        preview_text.config(state=tk.DISABLED)
        
        # Close button
        tk.Button(preview_window,
                 text="Tutup",
                 font=('Arial', 10, 'bold'),
                 bg='#003366',
                 fg='white',
                 relief=tk.FLAT,
                 cursor='hand2',
                 command=preview_window.destroy,
                 padx=20,
                 pady=5).pack(pady=10)
    
    def on_close(self):
        """Handle window close"""
        # Clean up temporary directory
        try:
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception:
            pass
        
        if self.parent_window:
            self.parent_window.deiconify()
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = TemplateEditor(root)
    root.mainloop()

