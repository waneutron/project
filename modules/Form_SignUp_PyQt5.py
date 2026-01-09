"""
Form_SignUp_PyQt5.py - Sign Up Form (PyQt5)
Optimized OOP structure with lazy loading to reduce app load
"""
import os
import sys
from datetime import datetime

from docx import Document
from hijri_converter import Gregorian, Hijri
from PIL import Image

from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import (QAbstractItemView, QCheckBox, QComboBox, QDesktopWidget, QDialog,
                                QFileDialog, QFrame, QGridLayout, QHBoxLayout, QLabel, QLineEdit,
                                QMessageBox, QPushButton, QScrollArea, QTableWidget,
                                QTableWidgetItem, QTextEdit, QVBoxLayout, QWidget)

from helpers.resource_path import get_logo_path, get_template_path
from helpers.unified_database import UnifiedDatabase

# ============================================
# CUSTOM WIDGETS
# ============================================

class SafeLineEdit(QLineEdit):
    """QLineEdit that doesn't close dialog on Enter/Return key"""
    def keyPressEvent(self, event):
        # Allow Tab for moving to next field
        if event.key() == Qt.Key_Tab:
            self.focusNextChild()
            return
        # Allow Shift+Tab for moving to previous field
        elif event.key() == Qt.Key_Backtab:
            self.focusPreviousChild()
            return
        # Ignore Return/Enter - don't let it close the dialog
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            event.ignore()
            return
        # Handle everything else normally
        super().keyPressEvent(event)

try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

_docx2pdf = None

def _lazy_import_docx2pdf():
    """Lazy import docx2pdf"""
    global _docx2pdf
    if _docx2pdf is None:
        try:
            from docx2pdf import convert
            _docx2pdf = convert
        except ImportError:
            _docx2pdf = False
    return _docx2pdf

try:
    from helpers.docx_helper import replace_in_document
except ImportError:
    def replace_in_document(doc, replacements):
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value) if value else "")
        return {}


class FormSignUp(QDialog):
    """Government-styled Sign Up B Form with professional design"""
    
    def __init__(self, parent_window=None):
        super().__init__()
        self.parent_window = parent_window
        self.setWindowTitle("Sistem Pendaftaran Sign Up B - Borang Pendaftaran")
        self.setGeometry(100, 100, 1600, 1000)
        self.setMinimumSize(1200, 800)  # Prevent window from shrinking too much
        
        # Government colors
        self.colors = {
            'primary': '#003366',
            'secondary': '#004080',
            'accent': '#006699',
            'bg_main': '#FFFFFF',
            'bg_white': '#FFFFFF',
            'text_dark': '#1a1a1a',
            'text_light': '#FFFFFF',
            'border': '#CCCCCC',
            'button_primary': '#003366',
            'button_secondary': '#666666',
            'button_success': '#2E7D32',
            'button_hover': '#004d99',
            'success': '#4CAF50',
            'warning': '#FF9800',
            'error': '#F44336',
            'info': '#2196F3'
        }
        
        self.db = UnifiedDatabase()
        
        # Store checkboxes
        self.checklist_vars = {}
        
        # Create UI
        self.create_ui()
        
        # Apply global styles ONCE after UI creation
        self.apply_global_styles()
        
        self.center_window()
        self.update_tarikh_islam()
    
    def apply_global_styles(self):
        """Apply global styles once to reduce redraws"""
        style_sheet = """
            QFrame {
                border: none;
                background-color: #FFFFFF;
            }
            QWidget {
                background-color: #FFFFFF;
            }
            QLineEdit, QComboBox, QTextEdit {
                background-color: #F5F5F5;
                border: 1px solid #CCCCCC;
                font-size: 14px;
                padding: 8px;
                color: #1a1a1a;
            }
            QLineEdit:focus, QComboBox:focus, QTextEdit:focus {
                border: 2px solid #006699;
                background-color: #F5F5F5;
            }
            QPushButton {
                font-weight: bold;
                border-radius: 5px;
                padding: 8px;
            }
            QLabel {
                background-color: #FFFFFF;
                color: #1a1a1a;
            }
            QCheckBox {
                background-color: #FFFFFF;
                color: #1a1a1a;
            }
            QScrollArea {
                background-color: #FFFFFF;
            }
        """
        self.setStyleSheet(style_sheet)
    
    def create_ui(self):
        """Create government-styled interface"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Header
        self.create_header(main_layout)
        
        # Scrollable content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("background-color: #FFFFFF; border: none;")
        
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(40, 30, 40, 30)
        scroll_layout.setSpacing(20)
        
        # Create form
        self.create_form(scroll_layout)
        
        scroll_area.setWidget(scroll_widget)
        main_layout.addWidget(scroll_area, stretch=1)
    
    def create_header(self, parent_layout):
        """Create government header"""
        header_frame = QFrame()
        header_frame.setFixedHeight(100)
        header_frame.setStyleSheet(f"background-color: {self.colors['primary']};")
        
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(20, 10, 20, 10)
        header_layout.setSpacing(20)
        
        # Logo (left side)
        logo_widget = QWidget()
        logo_widget.setStyleSheet(f"background-color: {self.colors['primary']};")
        logo_layout = QHBoxLayout(logo_widget)
        logo_layout.setContentsMargins(0, 0, 0, 0)
        
        try:
            logo_image = Image.open(get_logo_path())
            if logo_image.mode != 'RGBA':
                logo_image = logo_image.convert('RGBA')
            logo_image = logo_image.resize((80, 80), Image.Resampling.LANCZOS)
            
            background_color = self.colors['primary']
            bg_rgb = tuple(int(background_color[i:i+2], 16) for i in (1, 3, 5))
            background = Image.new('RGBA', logo_image.size, bg_rgb + (255,))
            
            if logo_image.mode == 'RGBA':
                logo_image = Image.alpha_composite(background, logo_image)
                logo_image = logo_image.convert('RGB')
            
            from io import BytesIO
            buf = BytesIO()
            logo_image.save(buf, format='PNG')
            pixmap = QPixmap()
            pixmap.loadFromData(buf.getvalue())
            
            logo_label = QLabel()
            logo_label.setPixmap(pixmap.scaled(80, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            logo_label.setStyleSheet(f"background-color: {self.colors['primary']};")
            logo_layout.addWidget(logo_label)
        except Exception:
            pass
        
        header_layout.addWidget(logo_widget)
        
        # Title section (center)
        title_widget = QWidget()
        title_widget.setStyleSheet(f"background-color: {self.colors['primary']};")
        title_layout = QVBoxLayout(title_widget)
        title_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("PENDAFTARAN SIGN UP B")
        title.setStyleSheet(f"background-color: {self.colors['primary']}; color: white; font-size: 22px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(title)
        
        subtitle = QLabel("Sistem Pendaftaran Jadual B")
        subtitle.setStyleSheet(f"background-color: {self.colors['primary']}; color: #E0E0E0; font-size: 14px;")
        subtitle.setAlignment(Qt.AlignCenter)
        title_layout.addWidget(subtitle)
        
        header_layout.addWidget(title_widget, stretch=1)
        
        # Back button (right side)
        if self.parent_window:
            btn_back = QPushButton("← KEMBALI")
            btn_back.setStyleSheet(f"""
                QPushButton {{
                    background-color: {self.colors['secondary']};
                    color: white;
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px 20px;
                    border-radius: 5px;
                    min-width: 150px;
                    min-height: 40px;
                }}
                QPushButton:hover {{
                    background-color: {self.colors['button_hover']};
                }}
            """)
            btn_back.clicked.connect(self.on_back_click)
            header_layout.addWidget(btn_back)
        
        parent_layout.addWidget(header_frame)
        
        # Separator
        separator = QFrame()
        separator.setFixedHeight(2)
        separator.setStyleSheet(f"background-color: {self.colors['accent']};")
        parent_layout.addWidget(separator)
    
    def create_form(self, parent_layout):
        """Create form fields"""
        # Helper function to create styled labels like Form2
        def create_label(text):
            label = QLabel(text)
            label.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 14px; font-weight: bold; padding: 0px; margin: 0px;")
            label.setContentsMargins(0, 0, 0, 0)  # Remove any margins
            return label
        
        # White card container
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color: #FFFFFF;
                border: none;
                border-radius: 5px;
            }
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(30, 30, 30, 30)
        card_layout.setSpacing(20)
        
        # Form title
        form_title = QLabel("📋 BORANG PENDAFTARAN SIGN UP B")
        form_title.setStyleSheet(f"color: {self.colors['primary']}; font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        card_layout.addWidget(form_title)
        
        # Description
        desc_label = QLabel("Sila isi maklumat di bawah dengan lengkap. Medan bertanda * adalah wajib.")
        desc_label.setStyleSheet(f"color: #666666; font-size: 12px; margin-bottom: 15px;")
        desc_label.setWordWrap(True)
        card_layout.addWidget(desc_label)
        
        # Form fields grid
        grid_widget = QWidget()
        grid_layout = QGridLayout(grid_widget)
        grid_layout.setSpacing(12)
        grid_layout.setContentsMargins(0, 0, 0, 0)
        
        row = 0
        
        # Rujukan Tuan
        grid_layout.addWidget(create_label("Rujukan Tuan:"), row, 0)
        
        self.entry_rujukan_tuan = SafeLineEdit()
        self.entry_rujukan_tuan.setMinimumWidth(500)
        self.entry_rujukan_tuan.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_rujukan_tuan, row, 1)
        row += 1
        
        # Rujukan Kami
        grid_layout.addWidget(create_label("Rujukan Kami:"), row, 0)
        
        self.entry_rujukan_kami = SafeLineEdit()
        self.entry_rujukan_kami.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_rujukan_kami, row, 1)
        row += 1
        
        # Tarikh
        grid_layout.addWidget(create_label("Tarikh:"), row, 0)
        
        self.entry_tarikh = SafeLineEdit()
        self.entry_tarikh.setText(datetime.now().strftime("%d/%m/%Y"))
        self.entry_tarikh.setStyleSheet("background-color: #F5F5F5; font-size: 13px; padding: 5px;")
        self.entry_tarikh.setMaximumWidth(200)
        self.entry_tarikh.textChanged.connect(self.update_tarikh_islam)
        grid_layout.addWidget(self.entry_tarikh, row, 1)
        row += 1
        
        # Tarikh Islam
        grid_layout.addWidget(create_label("Tarikh (Hijri):"), row, 0)
        
        self.entry_islam = SafeLineEdit()
        self.entry_islam.setReadOnly(True)
        self.entry_islam.setMaximumWidth(300)
        grid_layout.addWidget(self.entry_islam, row, 1)
        row += 1
        
        # Separator
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setStyleSheet(f"color: {self.colors['border']};")
        grid_layout.addWidget(separator1, row, 0, 1, 2)
        row += 1
        
        # Business Name
        grid_layout.addWidget(create_label("Nama Perniagaan: *"), row, 0)
        
        self.entry_business_name = SafeLineEdit()
        self.entry_business_name.setStyleSheet("background-color: #F5F5F5; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_business_name, row, 1)
        row += 1
        
        # Business Address
        label_address = QLabel("Alamat Perniagaan: *")
        label_address.setStyleSheet(f"color: {self.colors['text_dark']}; font-size: 13px; font-weight: bold;")
        grid_layout.addWidget(label_address, row, 0, Qt.AlignTop)
        
        self.text_business_address = QTextEdit()
        self.text_business_address.setMaximumHeight(100)
        self.text_business_address.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 8px;")
        grid_layout.addWidget(self.text_business_address, row, 1)
        row += 1
        
        # Schedule Type
        grid_layout.addWidget(create_label("Jenis Jadual: *"), row, 0)
        
        self.combo_schedule_type = QComboBox()
        self.combo_schedule_type.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.combo_schedule_type, row, 1)
        row += 1
        
        # Registration Number
        grid_layout.addWidget(create_label("Nombor Pendaftaran:"), row, 0)
        
        self.entry_registration_number = SafeLineEdit()
        self.entry_registration_number.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_registration_number, row, 1)
        row += 1
        
        # Salutation
        grid_layout.addWidget(create_label("Sapaan:"), row, 0)
        
        self.combo_salutation = QComboBox()
        self.combo_salutation.addItems(['Tuan', 'Puan', 'Encik', 'Cik'])
        self.combo_salutation.setCurrentText('Tuan')
        self.combo_salutation.setEditable(False)
        self.combo_salutation.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.combo_salutation, row, 1)
        row += 1
        
        # Email
        grid_layout.addWidget(create_label("E-mel:"), row, 0)
        
        self.entry_email = SafeLineEdit()
        self.entry_email.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_email, row, 1)
        row += 1
        
        # Talian
        grid_layout.addWidget(create_label("Talian:"), row, 0)
        
        self.entry_talian = SafeLineEdit()
        self.entry_talian.setStyleSheet("background-color: #F5F5F5; border: 1px solid #CCCCCC; font-size: 13px; padding: 5px;")
        grid_layout.addWidget(self.entry_talian, row, 1)
        row += 1
        
        # Separator before checklist
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setStyleSheet(f"color: {self.colors['border']};")
        grid_layout.addWidget(separator2, row, 0, 1, 2)
        row += 1
        
        # Checklist Section
        checklist_label = QLabel("Senarai Semak: *")
        checklist_label.setStyleSheet(f"color: {self.colors['primary']}; font-size: 14px; font-weight: bold;")
        grid_layout.addWidget(checklist_label, row, 0, 1, 2)
        row += 1
        
        # Checklist items frame
        checklist_frame = QFrame()
        checklist_frame.setStyleSheet("background-color: white; border: 1px solid #CCCCCC; border-radius: 5px;")
        checklist_layout = QGridLayout(checklist_frame)
        checklist_layout.setContentsMargins(15, 15, 15, 15)
        checklist_layout.setSpacing(10)
        
        checklist_items = [
            "Dokumen Pendaftaran Perniagaan",
            "Sijil Pendaftaran Syarikat (SSM)",
            "Dokumen Identiti Pemohon",
            "Alamat Perniagaan yang Sah",
            "Maklumat Bank Perniagaan",
            "Lesen Perniagaan (jika berkenaan)",
            "Dokumen Cukai (jika berkenaan)",
            "Lain-lain Dokumen Sokongan"
        ]
        
        for idx, item in enumerate(checklist_items):
            checkbox = QCheckBox(item)
            checkbox.setStyleSheet("font-size: 14px;")
            self.checklist_vars[item] = checkbox
            col = idx % 2
            checkbox_row = idx // 2
            checklist_layout.addWidget(checkbox, checkbox_row, col)
        
        grid_layout.addWidget(checklist_frame, row, 0, 1, 2)
        row += 1
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.setAlignment(Qt.AlignCenter)
        button_layout.setSpacing(15)
        
        btn_edit_template = QPushButton("✏️ EDIT TEMPLATE")
        btn_edit_template.setStyleSheet(f"""
            QPushButton {{
                background-color: #FF9800;
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 15px 30px;
                border-radius: 5px;
                min-width: 200px;
                min-height: 50px;
            }}
            QPushButton:hover {{
                background-color: #F57C00;
            }}
        """)
        btn_edit_template.clicked.connect(self.open_template_editor)
        button_layout.addWidget(btn_edit_template)
        
        btn_preview = QPushButton("👁️ PREVIEW")
        btn_preview.setStyleSheet(f"""
            QPushButton {{
                background-color: #2196F3;
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 15px 30px;
                border-radius: 5px;
                min-width: 200px;
                min-height: 50px;
            }}
            QPushButton:hover {{
                background-color: #1976D2;
            }}
        """)
        btn_preview.clicked.connect(self.on_preview_click)
        button_layout.addWidget(btn_preview)
        
        btn_save = QPushButton("💾 SIMPAN")
        btn_save.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.colors['button_success']};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 15px 30px;
                border-radius: 5px;
                min-width: 200px;
                min-height: 50px;
            }}
            QPushButton:hover {{
                background-color: #1B5E20;
            }}
        """)
        btn_save.clicked.connect(self.on_save_click)
        button_layout.addWidget(btn_save)
        
        grid_layout.addLayout(button_layout, row, 0, 1, 2)
        
        card_layout.addWidget(grid_widget)
        parent_layout.addWidget(card)
    
    def open_template_editor(self):
        """Open template editor for SignUp template"""
        try:
            from modules.TemplateEditor import TemplateEditor
            editor = TemplateEditor(self, template_name='signUpB.docx', form_type='signupb')
            editor.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Tidak dapat membuka template editor: {str(e)}")
    
    def on_preview_click(self):
        """Preview document before saving"""
        # Validate basic fields for preview
        if not self.entry_business_name.text().strip():
            QMessageBox.warning(self, "Amaran", "Sila isi Nama Perniagaan untuk preview.")
            return
        
        try:
            doc = self.generate_document()
            if doc is None:
                return
            
            temp_path = "temp_preview_signup.docx"
            doc.save(temp_path)
            
            if os.name == 'nt':
                os.startfile(temp_path)
            elif os.name == 'posix':
                import subprocess
                subprocess.call(['open', temp_path])
            
            QMessageBox.information(self, "Pratonton", 
                "Dokumen dibuka untuk pratonton.\n\n"
                "Semak temp_preview_signup.docx")
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Tidak dapat buka pratonton: {str(e)}")
    
    def update_tarikh_islam(self):
        """Update Hijri date based on Gregorian date"""
        try:
            tarikh_str = self.entry_tarikh.text()
            if '/' in tarikh_str:
                parts = tarikh_str.split('/')
                if len(parts) == 3:
                    day, month, year = map(int, parts)
                    hijri = Gregorian(year, month, day).to_hijri()
                    
                    bulan_hijri = ["Muharam", "Safar", "Rabiul Awal", "Rabiul Akhir",
                                  "Jamadil Awal", "Jamadil Akhir", "Rejab", "Syaaban",
                                  "Ramadhan", "Syawal", "Zulkaedah", "Zulhijjah"]
                    
                    hijri_text = f"{hijri.day} {bulan_hijri[hijri.month-1]} {hijri.year}H"
                    self.entry_islam.setText(hijri_text)
        except:
            pass
    
    def center_window(self):
        """Center window on screen"""
        frame_geometry = self.frameGeometry()
        try:
            from PyQt5.QtWidgets import QDesktopWidget
            desktop = QDesktopWidget()
            screen_geometry = desktop.availableGeometry(desktop.primaryScreen())
        except:
            screen_geometry = self.screen().availableGeometry()
        
        center_point = screen_geometry.center()
        frame_geometry.moveCenter(center_point)
        self.move(frame_geometry.topLeft())
        self.raise_()
        self.activateWindow()
    
    def generate_document(self):
        """Generate Word document with replacements"""
        template_file = "signUpB.docx"
        
        try:
            from helpers.template_storage import get_template_document
            doc = get_template_document(template_file)
            if doc is None:
                template_path = get_template_path(template_file)
                if not os.path.exists(template_path):
                    QMessageBox.critical(self, "Ralat", f"Templat tidak dijumpai: {template_file}")
                    return None
                doc = Document(template_path)
        except ImportError:
            template_path = get_template_path(template_file)
            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Ralat", f"Templat tidak dijumpai: {template_path}")
                return None
            doc = Document(template_path)
        
        # Get selected checklist items
        selected_checklist = self.get_selected_checklist_items()
        checklist_text = "\n".join([f"✓ {item}" for item in selected_checklist])
        
        # Use PlaceholderBuilder for standardized placeholders
        from helpers.placeholder_registry import PlaceholderBuilder
        
        # Build replacements using PlaceholderBuilder
        builder = PlaceholderBuilder(self)
        builder.add_standard('RUJUKAN_TUAN', self.entry_rujukan_tuan.text())
        builder.add_standard('RUJUKAN_KAMI', self.entry_rujukan_kami.text())
        builder.add_tarikh(self.entry_tarikh.text())
        builder.add_tarikh_islam(self.entry_islam.text())
        # Use alias: BUSINESS_NAME -> NAMA_SYARIKAT
        builder.add_nama_syarikat(self.entry_business_name.text())
        builder.add_standard('BUSINESS_NAME', self.entry_business_name.text().upper())  # Keep for backward compatibility
        builder.add_standard('BUSINESS_ADDRESS', self.text_business_address.toPlainText().strip())
        builder.add_standard('SCHEDULE_TYPE', self.combo_schedule_type.currentText())
        builder.add_standard('REGISTRATION_NUMBER', self.entry_registration_number.text())
        builder.add_standard('SALUTATION', self.combo_salutation.currentText())
        builder.add_standard('SENARAI_SEMAK', checklist_text)
        builder.add_standard('CHECKLIST', checklist_text)
        
        # Get final replacements dict
        replacements = builder.build()
        
        # Replace placeholders
        replace_in_document(doc, replacements)
        
        return doc
    
    def get_selected_checklist_items(self):
        """Get list of selected checklist items"""
        selected = []
        for item, checkbox in self.checklist_vars.items():
            if checkbox.isChecked():
                selected.append(item)
        return selected
    
    def on_save_click(self):
        """Save document as PDF and to database"""
        # Comprehensive validation
        validation_errors = []
        
        if not self.entry_rujukan_tuan.text().strip():
            validation_errors.append("Rujukan Tuan (Penerima) tidak diisi")
        
        if not self.entry_rujukan_kami.text().strip():
            validation_errors.append("Rujukan Kami tidak diisi")
        
        if not self.entry_business_name.text().strip():
            validation_errors.append("Nama Perniagaan tidak diisi")
        
        if not self.text_business_address.toPlainText().strip():
            validation_errors.append("Alamat Perniagaan tidak diisi")
        
        if not self.entry_email.text().strip():
            validation_errors.append("Email tidak diisi")
        
        if not self.entry_talian.text().strip():
            validation_errors.append("Talian tidak diisi")
        
        if not self.combo_schedule_type.currentText():
            validation_errors.append("Jenis Jadual tidak dipilih")
        
        selected_items = self.get_selected_checklist_items()
        if not selected_items:
            validation_errors.append("Sekurang-kurangnya satu item dalam Senarai Semak mesti dipilih")
        
        if validation_errors:
            QMessageBox.warning(self, "⚠️ Maklumat Tidak Lengkap",
                "Sila lengkapkan maklumat berikut:\n\n" + "\n".join([f"• {err}" for err in validation_errors]))
            return
        
        # Confirmation dialog
        reply = QMessageBox.question(self, "✅ Sahkan Simpan Dokumen",
            f"Adakah anda pasti untuk simpan dokumen?\n\n"
            f"📋 Rujukan Tuan: {self.entry_rujukan_tuan.text()}\n"
            f"📋 Rujukan Kami: {self.entry_rujukan_kami.text()}\n"
            f"🏢 Nama Perniagaan: {self.entry_business_name.text()}\n"
            f"📅 Tarikh: {self.entry_tarikh.text()}\n"
            f"📊 Jadual: {self.combo_schedule_type.currentText()}",
            QMessageBox.Yes | QMessageBox.No)
        
        if reply != QMessageBox.Yes:
            return
        
        try:
            # Generate document
            doc = self.generate_document()
            if doc is None:
                return
            
            # Save as PDF directly (like other modules)
            safe_name = self.entry_business_name.text().strip().replace('/', '_').replace('\\', '_')
            default_filename = f"SignUpB_{safe_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
            
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Simpan Dokumen Sign Up B sebagai PDF",
                default_filename,
                "PDF Files (*.pdf);;Word Documents (*.docx);;All Files (*.*)"
            )
            
            if not filename:
                return
            
            # Determine file type from extension
            if filename.lower().endswith('.docx'):
                # Save as DOCX
                doc.save(filename)
                
                # Try to convert to PDF
                if PDF_AVAILABLE:
                    try:
                        convert = _lazy_import_docx2pdf()
                        if convert:
                            pdf_filename = filename.replace('.docx', '.pdf')
                            convert(filename, pdf_filename)
                    except:
                        pass
            else:
                # Save as PDF
                if not filename.lower().endswith('.pdf'):
                    filename += '.pdf'
                
                # Save DOCX first, then convert
                temp_docx = filename.replace('.pdf', '.docx')
                doc.save(temp_docx)
                
                if PDF_AVAILABLE:
                    try:
                        convert = _lazy_import_docx2pdf()
                        if convert:
                            convert(temp_docx, filename)
                            # Remove temp DOCX
                            try:
                                os.remove(temp_docx)
                            except:
                                pass
                    except Exception as pdf_error:
                        QMessageBox.warning(self, "Amaran", 
                            f"Dokumen disimpan sebagai Word:\n{temp_docx}\n\n"
                            f"Penukaran ke PDF gagal: {str(pdf_error)}")
                        filename = temp_docx
            
            # Save to database
            try:
                application_data = {
                    'form_type': 'signupb',
                    'rujukan_tuan': self.entry_rujukan_tuan.text(),
                    'rujukan_kami': self.entry_rujukan_kami.text(),
                    'tarikh': self.entry_tarikh.text(),
                    'tarikh_islam': self.entry_islam.text(),
                    'nama_syarikat': self.entry_business_name.text(),
                    'document_path': filename
                }
                
                signupb_details = {
                    'email': self.entry_email.text(),
                    'talian': self.entry_talian.text()
                }
                
                app_id = self.db.save_application('signupb', application_data, signupb_details)
                
                # Clear database cache so history viewer shows new data
                if hasattr(self.db, 'clear_cache'):
                    self.db.clear_cache()
                
                QMessageBox.information(self, "✅ Berjaya", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Rekod telah disimpan ke pangkalan data (ID: {app_id})")
                
                # Close form after successful save
                self.reject()
            except Exception as db_error:
                QMessageBox.warning(self, "Amaran", 
                    f"Dokumen berjaya disimpan:\n{filename}\n\n"
                    f"Tetapi ralat menyimpan ke pangkalan data:\n{str(db_error)}")
        
        except Exception as e:
            QMessageBox.critical(self, "Ralat", f"Ralat menyimpan dokumen: {str(e)}")
    
    def on_back_click(self):
        """Go back to main menu"""
        self.reject()


if __name__ == "__main__":
    from PyQt5.QtWidgets import QApplication
    app = QApplication(sys.argv)
    window = FormSignUp()
    window.show()
    sys.exit(app.exec_())

