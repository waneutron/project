"""
template_field_validator.py - Validate Template-Field Completeness
Check if semua placeholders dalam template ada input fields

Supports scanning across all forms:
- Form2 (generic dynamic editor)
- Form3 (AMES system)
- Form_DeleteItem_PyQt5 (Delete Item)
- Form_SignUp_PyQt5 (Sign Up/Registration)
"""
import os
import re
from docx import Document
from helpers.resource_path import get_template_path

class TemplateFieldMapper:
    """Helper class untuk scan templates across all forms"""
    
    # Form templates mapping for easy reference
    # All templates are stored in /Templates folder and loaded via get_template_path()
    FORM_TEMPLATES = {
        'Form1': {
            'name': 'Pelupusan',
            'templates': {
                'pelupusan': [
                    'pelupusan_pemusnahan.docx',
                    'pelupusan_penjualan.docx',
                    'pelupusan_skrap.docx',
                    'pelupusan_tidak_lulus.docx'
                ],
                'pembatalan': [
                    'batal_sijil.docx'
                ]
            }
        },
        'Form2': {
            'name': 'Dynamic Form Editor (form2_Government_PyQt5)',
            'templates': {
                'pelupusan': [
                    'pelupusan_pemusnahan.docx',
                    'pelupusan_penjualan.docx',
                    'pelupusan_skrap.docx',
                    'pelupusan_tidak_lulus.docx'
                ],
                'butiran_5d': [
                    'surat kelulusan butiran 5D (Lulus).docx',
                    'surat kelulusan butiran 5D (tidak lulus).docx'
                ]
            }
        },
        'Form3': {
            'name': 'AMES System (Form3_Government_PySide2)',
            'templates': {
                'ames': [
                    'ames_pedagang.docx',
                    'ames_pengilang.docx'
                ],
                'butiran_5d': [
                    'surat kelulusan butiran 5D (Lulus).docx'
                ]
            }
        },
        'FormDeleteItem': {
            'name': 'Delete Item (Form_DeleteItem_PyQt5)',
            'templates': {
                'delete_item': [
                    'delete_item.docx'
                ]
            }
        },
        'FormSignUp': {
            'name': 'Sign Up Registration (Form_SignUp_PyQt5)',
            'templates': {
                'signup': [
                    'signUpB.docx'
                ]
            }
        }
    }
    
    def scan_template_placeholders(self, template_file, form_name=None):
        """Scan template untuk detect semua placeholders (<<PLACEHOLDER>>)
        
        Args:
            template_file: Template file name to scan
            form_name: Optional form name (Form2, Form3, FormDeleteItem, FormSignUp)
        
        Returns:
            sorted list of placeholders found
        """
        placeholders = set()
        
        try:
            # Try embedded storage first
            try:
                from helpers.template_storage import get_template_document
                doc = get_template_document(template_file)
                if doc is None:
                    # Fallback to file system
                    template_path = get_template_path(template_file)
                    if os.path.exists(template_path):
                        doc = Document(template_path)
            except ImportError:
                # Fallback to file system
                template_path = get_template_path(template_file)
                if os.path.exists(template_path):
                    doc = Document(template_path)
                else:
                    return []
            
            if not doc:
                return []
            
            # Scan paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    matches = re.findall(r'<<([^>]+)>>', paragraph.text)
                    placeholders.update(matches)
            
            # Scan tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'<<([^>]+)>>', paragraph.text)
                            placeholders.update(matches)
            
            return sorted(list(placeholders))
            
        except Exception as e:
            print(f"Error scanning template '{template_file}' for form '{form_name}': {e}")
            return []
    
    def scan_all_forms_templates(self):
        """Scan all templates across all forms
        
        Returns:
            dict with structure {form_name: {category: [template_results]}}
        """
        results = {}
        
        for form_name, form_config in self.FORM_TEMPLATES.items():
            results[form_name] = {
                'name': form_config['name'],
                'templates': {}
            }
            
            # Scan templates for this form
            for category, templates in form_config['templates'].items():
                results[form_name]['templates'][category] = []
                
                for template_file in templates:
                    placeholders = self.scan_template_placeholders(template_file, form_name)
                    
                    template_path = get_template_path(template_file)
                    exists = os.path.exists(template_path)
                    
                    results[form_name]['templates'][category].append({
                        'file': template_file,
                        'exists': exists,
                        'placeholder_count': len(placeholders),
                        'placeholders': placeholders
                    })
        
        return results
    
    def scan_form_templates(self, form_name):
        """Scan templates for a specific form
        
        Args:
            form_name: Form name (Form2, Form3, FormDeleteItem, FormSignUp)
        
        Returns:
            dict with template scan results
        """
        if form_name not in self.FORM_TEMPLATES:
            return {'error': f'Form {form_name} not found'}
        
        form_config = self.FORM_TEMPLATES[form_name]
        result = {
            'form_name': form_name,
            'form_display_name': form_config['name'],
            'templates': {}
        }
        
        for category, templates in form_config['templates'].items():
            result['templates'][category] = []
            
            for template_file in templates:
                placeholders = self.scan_template_placeholders(template_file, form_name)
                template_path = get_template_path(template_file)
                exists = os.path.exists(template_path)
                
                result['templates'][category].append({
                    'file': template_file,
                    'path': template_path,
                    'exists': exists,
                    'placeholder_count': len(placeholders),
                    'placeholders': placeholders
                })
        
        return result
    
    def get_form_templates_list(self, form_name):
        """Get all template files for a form
        
        Args:
            form_name: Form name
        
        Returns:
            list of template files for this form
        """
        if form_name not in self.FORM_TEMPLATES:
            return []
        
        form_config = self.FORM_TEMPLATES[form_name]
        templates = []
        
        for category, template_list in form_config['templates'].items():
            templates.extend(template_list)
        
        return templates


class TemplateFieldValidator:
    """Validate template-field completeness"""
    
    def __init__(self):
        self.mapper = TemplateFieldMapper()
    
    def validate_template_completeness(self, form_name, template_file, custom_fields=None, existing_mapping=None):
        """
        Validate jika semua placeholders dalam template ada fields
        
        Args:
            form_name: Form name (e.g., 'Form2')
            template_file: Template file name
            custom_fields: List of custom fields from form
            existing_mapping: Existing field mapping dict
        
        Returns:
            dict dengan validation results
        """
        # Scan template untuk placeholders
        placeholders = self.mapper.scan_template_placeholders(template_file)
        
        if not placeholders:
            return {
                'is_complete': True,
                'message': 'No placeholders found in template',
                'missing_fields': [],
                'unmapped_placeholders': [],
                'warnings': [],
                'completeness_percent': 100,
                'total_placeholders': 0,
                'mapped_count': 0
            }
        
        # Get mapped placeholders
        mapped_placeholders = set()
        if existing_mapping and existing_mapping.get('field_mappings'):
            for field_id, map_info in existing_mapping['field_mappings'].items():
                placeholder = map_info['placeholder'].replace('<<', '').replace('>>', '')
                mapped_placeholders.add(placeholder)
        
        # Also check dari custom_fields jika provided
        if custom_fields:
            for field in custom_fields:
                field_id = field.get('field_id', '')
                # Try to match dengan placeholders
                field_label = field.get('label', '').upper()
                for placeholder in placeholders:
                    if (field_id.upper() in placeholder.upper() or 
                        placeholder.upper() in field_id.upper() or
                        field_label in placeholder.upper() or
                        placeholder.upper() in field_label):
                        mapped_placeholders.add(placeholder)
        
        # Check untuk unmapped placeholders
        unmapped = []
        for placeholder in placeholders:
            if placeholder not in mapped_placeholders:
                unmapped.append(placeholder)
        
        # Calculate completeness
        total_placeholders = len(placeholders)
        mapped_count = len(mapped_placeholders)
        completeness_percent = (mapped_count / total_placeholders * 100) if total_placeholders > 0 else 100
        
        is_complete = len(unmapped) == 0
        
        # Generate suggestions untuk missing fields
        suggestions = []
        if unmapped:
            for placeholder in unmapped:
                # Generate field suggestion dari placeholder name
                field_suggestion = self._generate_field_suggestion(placeholder)
                suggestions.append({
                    'placeholder': placeholder,
                    'suggested_field': field_suggestion
                })
        
        return {
            'is_complete': is_complete,
            'completeness_percent': completeness_percent,
            'total_placeholders': total_placeholders,
            'mapped_count': mapped_count,
            'unmapped_placeholders': unmapped,
            'suggestions': suggestions,
            'all_placeholders': placeholders,
            'mapped_placeholders': list(mapped_placeholders),
            'message': self._generate_message(is_complete, unmapped, completeness_percent)
        }
    
    def _generate_field_suggestion(self, placeholder):
        """Generate field suggestion dari placeholder name"""
        # Remove << and >>
        name = placeholder.replace('<<', '').replace('>>', '')
        
        # Convert to readable label
        label = name.replace('_', ' ').title()
        
        # Generate field_id
        field_id = f"entry_{name.lower()}"
        
        # Determine field type based on name
        field_type = 'text'
        if 'ALAMAT' in name or 'ADDRESS' in name:
            field_type = 'textarea'
        elif 'TARIKH' in name or 'DATE' in name:
            field_type = 'date'
        elif 'NO' in name or 'NUMBER' in name or 'NOMBOR' in name:
            field_type = 'number'
        
        return {
            'field_id': field_id,
            'label': label,
            'type': field_type,
            'placeholder': f"<<{placeholder}>>"
        }
    
    def _generate_message(self, is_complete, unmapped, percent):
        """Generate user-friendly message"""
        if is_complete:
            return f"Semua placeholders ada input fields ({percent:.0f}% complete)"
        else:
            return f"{len(unmapped)} placeholders tiada input fields ({percent:.0f}% complete)"
    
    def validate_before_generate(self, form_name, template_file, form_fields_dict, mapping=None):
        """
        Validate sebelum generate document - check jika semua mapped fields ada value
        
        Args:
            form_name: Form name
            template_file: Template file name
            form_fields_dict: Dictionary dengan field values dari form
            mapping: Field mapping dict (optional)
        
        Returns:
            dict dengan validation results
        """
        if not mapping:
            # Try to load mapping
            try:
                from helpers.template_field_mapper import TemplateFieldMapper
                mapper = TemplateFieldMapper()
                mapping = mapper.load_mapping(form_name, template_file)
            except:
                mapping = None
        
        if not mapping or not mapping.get('field_mappings'):
            return {
                'is_valid': True,
                'missing_values': [],
                'warnings': [],
                'message': 'No field mapping found - using default fields'
            }
        
        missing_values = []
        warnings = []
        
        # Check setiap mapped field
        for field_id, map_info in mapping['field_mappings'].items():
            placeholder = map_info['placeholder']
            field_label = map_info['field_label']
            
            # Get value dari form_fields_dict
            value = form_fields_dict.get(field_id, '').strip() if field_id in form_fields_dict else ''
            
            if not value:
                missing_values.append({
                    'field_id': field_id,
                    'field_label': field_label,
                    'placeholder': placeholder
                })
        
        is_valid = len(missing_values) == 0
        
        return {
            'is_valid': is_valid,
            'missing_values': missing_values,
            'warnings': warnings,
            'message': self._generate_validation_message(is_valid, missing_values)
        }
    
    def _generate_validation_message(self, is_valid, missing_values):
        """Generate validation message"""
        if is_valid:
            return "Semua fields telah diisi"
        else:
            fields_list = "\n".join([f"  • {mv['field_label']} ({mv['placeholder']})" 
                                   for mv in missing_values])
            return f"Sila isi fields berikut:\n\n{fields_list}"

