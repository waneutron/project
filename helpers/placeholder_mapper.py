"""
placeholder_mapper.py - Backend for placeholder mapping system
"""
import json
import os
import re
from pathlib import Path


class PlaceholderMapper:
    """Manages persistent placeholder mappings"""
    
    def __init__(self, config_file='config/placeholder_mappings.json'):
        self.config_file = config_file
        self.mappings = {}
        self.load_mappings()
    
    def load_mappings(self):
        """Load saved mappings"""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    self.mappings = json.load(f)
            except Exception as e:
                print(f"[PlaceholderMapper] Error: {e}")
                self.mappings = {}
        else:
            os.makedirs(os.path.dirname(self.config_file) or '.', exist_ok=True)
            self.mappings = {}
    
    def save_mappings(self):
        """Save mappings to file"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.mappings, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            print(f"[PlaceholderMapper] Error: {e}")
            return False
    
    def get_template_mapping(self, template_file):
        """Get mapping for template"""
        return self.mappings.get(template_file)
    
    def is_template_configured(self, template_file):
        """Check if template configured"""
        return template_file in self.mappings
    
    def set_template_mapping(self, template_file, mapping):
        """Set mapping for template"""
        self.mappings[template_file] = mapping
        self.save_mappings()
    
    def scan_template_placeholders(self, template_path):
        """Scan template for placeholders"""
        try:
            from docx import Document
            doc = Document(template_path)
            placeholders = set()
            
            for para in doc.paragraphs:
                found = re.findall(r'<<[^>]+>>', para.text)
                placeholders.update(found)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            found = re.findall(r'<<[^>]+>>', para.text)
                            placeholders.update(found)
            
            return sorted(list(placeholders))
        except Exception as e:
            print(f"[PlaceholderMapper] Error scanning: {e}")
            return []
    
    def get_field_options(self):
        """Get available form fields"""
        return {
            'entry_nama': 'Nama Syarikat',
            'entry_rujukan': 'Rujukan',
            'entry_rujukan_syarikat': 'Rujukan Syarikat',
            'entry_tarikh': 'Tarikh',
            'entry_islam': 'Tarikh Islam',
            'entry_alamat1': 'Alamat Baris 1',
            'entry_alamat2': 'Alamat Baris 2',
            'entry_alamat3': 'Alamat Baris 3',
            'combo_pegawai': 'Nama Pegawai',
            'combo_process': 'Proses',
            'combo_jenis': 'Jenis Barang',
            'combo_pengecualian': 'Pengecualian',
            'entry_amount': 'Amaun/Sebab',
            'entry_pemusnahan_mula': 'Tarikh Mula',
            'entry_pemusnahan_tamat': 'Tarikh Tamat',
            'entry_tempoh': 'Tempoh',
            'var_sst_adam': 'Jenis SST/ADAM',
            'COMPUTED:alamat_full': 'Alamat Penuh',
            'COMPUTED:rujukan_full': 'Rujukan Penuh (with prefix)',
            'COMPUTED:tarikh_malay': 'Tarikh Format Melayu',
            'CUSTOM': 'Custom Value (manual input)'
        }
    
    def apply_mapping(self, form_instance, template_file):
        """Apply mapping to get replacements"""
        mapping = self.get_template_mapping(template_file)
        if not mapping:
            return {}
        
        replacements = {}
        for placeholder, field_source in mapping.items():
            value = self._get_field_value(form_instance, field_source)
            replacements[placeholder] = value
        
        return replacements
    
    def _get_field_value(self, form, field_source):
        """Get value from form field"""
        if field_source.startswith('COMPUTED:'):
            compute_type = field_source.split(':')[1]
            
            if compute_type == 'alamat_full':
                lines = []
                for attr in ['entry_alamat1', 'entry_alamat2', 'entry_alamat3']:
                    if hasattr(form, attr):
                        val = getattr(form, attr).text().strip()
                        if val:
                            lines.append(val)
                return '\n'.join(lines)
            
            elif compute_type == 'rujukan_full':
                if hasattr(form, 'entry_rujukan'):
                    ruj = form.entry_rujukan.text().strip()
                    return f"KE.JB(90)650/05-02/{ruj}" if ruj else ""
            
            elif compute_type == 'tarikh_malay':
                if hasattr(form, 'date_converter'):
                    return form.date_converter.format_tarikh_malay()
        
        if field_source.startswith('CUSTOM:'):
            return field_source.split(':', 1)[1]
        
        if hasattr(form, field_source):
            field = getattr(form, field_source)
            if hasattr(field, 'text'):
                return field.text()
            elif hasattr(field, 'currentText'):
                return field.currentText()
            elif hasattr(field, 'isChecked'):
                return "Ya" if field.isChecked() else "Tidak"
            else:
                return str(field)
        
        return ""