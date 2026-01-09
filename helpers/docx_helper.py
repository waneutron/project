"""
docx_helper.py - Advanced Word Document Text Replacement
Handles placeholders in paragraphs, tables, headers, footers, and text boxes
"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import re


def replace_text_in_paragraph(paragraph, search_text, replace_text):
    """
    Replace text in a paragraph while preserving formatting
    Handles text split across multiple runs
    """
    if search_text in paragraph.text:
        # Get the full text
        full_text = paragraph.text
        
        # Check if replacement is needed
        if search_text not in full_text:
            return False
        
        # Replace the text
        new_text = full_text.replace(search_text, replace_text)
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ''
        
        # Add new text to first run (preserves first run's formatting)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        
        return True
    return False


def replace_text_in_runs(paragraph, search_text, replace_text):
    """
    More advanced replacement that handles text split across runs
    Preserves formatting of each run and applies Arial 11
    """
    from docx.shared import Pt
    
    # Build complete text from runs
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if search_text not in full_text:
        return False
    
    # Find all occurrences
    replaced = False
    while search_text in full_text:
        # Replace one occurrence
        full_text = full_text.replace(search_text, replace_text, 1)
        replaced = True
    
    if not replaced:
        return False
    
    # Clear all runs except first and apply formatting
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = full_text
            # Apply Arial font size 11
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        else:
            run.text = ''
    
    return True


def replace_in_document(doc, replacements):
    """
    Replace all placeholders in document
    Searches in paragraphs, tables, headers, and footers
    Applies Arial 11 font to all text
    Supports placeholder aliases for backward compatibility
    
    Args:
        doc: Document object
        replacements: dict of {placeholder: value}
    """
    from docx.shared import Pt
    from helpers.placeholder_registry import PlaceholderRegistry
    
    replaced_count = {}
    
    # Expand replacements with aliases for backward compatibility
    expanded_replacements = replacements.copy()
    for placeholder, value in replacements.items():
        standard_name = PlaceholderRegistry.get_standard_name(placeholder)
        standard_placeholder = f"<<{standard_name}>>"
        
        # If placeholder is not already in standard format, add alias
        if placeholder != standard_placeholder:
            expanded_replacements[standard_placeholder] = value
            # Also keep original for templates that might use it
            if placeholder not in expanded_replacements:
                expanded_replacements[placeholder] = value
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for search_text, replace_text in expanded_replacements.items():
            if replace_text is None:
                replace_text = ""
            if replace_text_in_runs(paragraph, search_text, str(replace_text)):
                replaced_count[search_text] = replaced_count.get(search_text, 0) + 1
        
        # Apply Arial 11 to all runs in paragraph
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for search_text, replace_text in expanded_replacements.items():
                        if replace_text is None:
                            replace_text = ""
                        if replace_text_in_runs(paragraph, search_text, str(replace_text)):
                            replaced_count[search_text] = replaced_count.get(search_text, 0) + 1
                    
                    # Apply Arial 11 to all runs in table
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
    
    # Replace in headers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            for search_text, replace_text in expanded_replacements.items():
                if replace_text is None:
                    replace_text = ""
                if replace_text_in_runs(paragraph, search_text, str(replace_text)):
                    replaced_count[search_text] = replaced_count.get(search_text, 0) + 1
            
            # Apply Arial 11 to header
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            for search_text, replace_text in expanded_replacements.items():
                if replace_text is None:
                    replace_text = ""
                if replace_text_in_runs(paragraph, search_text, str(replace_text)):
                    replaced_count[search_text] = replaced_count.get(search_text, 0) + 1
            
            # Apply Arial 11 to footer
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    return replaced_count


def replace_with_formatting(paragraph, search_text, replace_text, bold=False, uppercase=False):
    """
    Replace text with specific formatting
    """
    if uppercase:
        replace_text = replace_text.upper()
    
    if replace_text_in_runs(paragraph, search_text, replace_text):
        if bold:
            for run in paragraph.runs:
                if replace_text in run.text:
                    run.bold = True
        return True
    return False


def find_and_replace_with_format(doc, search_text, replace_text, bold=False, uppercase=False):
    """
    Find and replace with formatting options
    """
    if uppercase:
        replace_text = replace_text.upper()
    
    replaced = False
    
    # In paragraphs
    for paragraph in doc.paragraphs:
        if replace_with_formatting(paragraph, search_text, replace_text, bold, uppercase):
            replaced = True
    
    # In tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_with_formatting(paragraph, search_text, replace_text, bold, uppercase):
                        replaced = True
    
    return replaced


def preview_placeholders(doc):
    """
    Find all placeholders in document (text between << and >>)
    Returns list of unique placeholders
    """
    placeholders = set()
    pattern = r'<<([^>]+)>>'
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        placeholders.update(matches)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.update(matches)
    
    return sorted(list(placeholders))


def validate_replacements(doc, replacements):
    """
    Check which placeholders were not replaced
    Returns list of missing placeholders
    """
    found_placeholders = preview_placeholders(doc)
    provided_keys = [key.strip('<>') for key in replacements.keys()]
    
    missing = [p for p in found_placeholders if p not in provided_keys]
    return missing


# Example usage
if __name__ == "__main__":
    # Test the functions
    doc = Document('template.docx')
    
    # Preview all placeholders
    placeholders = preview_placeholders(doc)
    print("Found placeholders:", placeholders)
    
    # Define replacements
    replacements = {
        '<<RUJUKAN>>': 'REF001',
        '<<NAMA_SYARIKAT>>': 'ABC Sdn Bhd',
        '<<TARIKH>>': '01/01/2024'
    }
    
    # Replace
    counts = replace_in_document(doc, replacements)
    print("Replaced:", counts)
    
    # Check for missing
    missing = validate_replacements(doc, replacements)
    if missing:
        print("Warning: These placeholders were not replaced:", missing)
    
    # Save
    doc.save('output.docx')
    print("Document saved!")


def add_standard_header_footer(doc):
    """
    Add standard government header and footer to all document sections
    Reads configuration from header_footer_config.json
    Only adds if header/footer are empty or if force_replace is enabled
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Load configuration
    try:
        from helpers.header_footer_config import HeaderFooterConfig
        config_manager = HeaderFooterConfig()
        header_config = config_manager.get_header_config()
        footer_config = config_manager.get_footer_config()
    except Exception as e:
        print(f"Error loading header/footer config: {e}")
        return
    
    # Skip if both disabled
    if not header_config.get('enabled', True) and not footer_config.get('enabled', True):
        return
    
    # Check if we should force replace existing headers/footers
    force_replace = header_config.get('force_replace', False) or footer_config.get('force_replace', False)
    
    for section in doc.sections:
        # Set header/footer to appear only on first page
        section.different_first_page_header_footer = True
        
        # Clear regular header/footer so they don't appear on subsequent pages
        # Clear regular header (for pages 2+)
        regular_header = section.header
        for paragraph in regular_header.paragraphs:
            paragraph.clear()
        # Clear regular footer (for pages 2+)
        regular_footer = section.footer
        for paragraph in regular_footer.paragraphs:
            paragraph.clear()
        
        # ========== HEADER (First Page Only) ==========
        if header_config.get('enabled', True):
            # Use first page header (not the default header)
            header = section.first_page_header
            
            # Check if header already has content
            has_existing_content = False
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    has_existing_content = True
                    break
            # Also check tables
            if not has_existing_content:
                for table in header.tables:
                    if table:
                        has_existing_content = True
                        break
            
            # Only replace if empty or force_replace is True
            if not has_existing_content or force_replace:
                # Clear existing header content
                for paragraph in list(header.paragraphs):
                    paragraph.clear()
                
                # Remove all tables from header
                try:
                    from docx.oxml.ns import qn
                    header_xml = header._element
                    # Find and remove all tables
                    tables_to_remove = []
                    for element in header_xml.iter():
                        if element.tag.endswith('}tbl'):
                            tables_to_remove.append(element)
                    for tbl in tables_to_remove:
                        parent = tbl.getparent()
                        if parent is not None:
                            parent.remove(tbl)
                except Exception:
                    pass  # If removal fails, continue anyway
                
                # Create header table for layout
                header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
                header_table.autofit = False
                header_table.columns[0].width = Inches(1.2)
                header_table.columns[1].width = Inches(4.5)
                header_table.columns[2].width = Inches(1.8)
                
                header_row = header_table.rows[0]
                
                # Left cell - Leave empty (no logo on left side)
                logo_cell = header_row.cells[0]
                logo_para = logo_cell.paragraphs[0]
                logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Left cell is empty - logo only on right side
                
                # Middle cell - Department info
                dept_cell = header_row.cells[1]
                dept_para = dept_cell.paragraphs[0]
                dept_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Department name
                dept_name_text = header_config.get('department_name', '')
                if dept_name_text:
                    dept_name = dept_para.add_run(dept_name_text + '\n')
                    dept_name.font.name = 'Arial'
                    dept_name.font.size = Pt(11)
                    dept_name.font.bold = True
                    dept_name.font.color.rgb = RGBColor(0, 0, 0)
                
                # Address lines
                address_lines = header_config.get('address_lines', [])
                for line in address_lines:
                    if line.strip():
                        address_run = dept_para.add_run(line.strip() + '\n')
                        address_run.font.name = 'Arial'
                        address_run.font.size = Pt(9)
                        address_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Right cell - Logo and Contact info
                contact_config = header_config.get('contact', {})
                contact_cell = header_row.cells[2]
                contact_para = contact_cell.paragraphs[0]
                contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Insert logo image at top of right cell
                try:
                    from helpers.resource_path import get_logo_path
                    import os
                    logo_path = get_logo_path()
                    if os.path.exists(logo_path):
                        logo_run = contact_para.add_run()
                        logo_run.add_picture(logo_path, width=Inches(0.6))
                        contact_para.add_run('\n')  # Add space after logo
                    else:
                        # If logo not found, try relative path
                        if os.path.exists('logo.png'):
                            logo_run = contact_para.add_run()
                            logo_run.add_picture('logo.png', width=Inches(0.6))
                            contact_para.add_run('\n')  # Add space after logo
                except Exception as e:
                    print(f"Could not insert logo in right cell: {e}")
                
                # Phone
                phone_label_text = contact_config.get('phone_label', '')
                if phone_label_text:
                    phone_label = contact_para.add_run(phone_label_text + ' ')
                    phone_label.font.name = 'Arial'
                    phone_label.font.size = Pt(9)
                    phone_label.font.bold = True
                    
                    phone_value_text = contact_config.get('phone_value', '') or ''
                    if phone_value_text:
                        phone_value = contact_para.add_run(phone_value_text + '\n')
                        phone_value.font.name = 'Arial'
                        phone_value.font.size = Pt(9)
                    else:
                        contact_para.add_run('\n')
                
                # Email/Web
                email_label_text = contact_config.get('email_label', '')
                if email_label_text:
                    email_label = contact_para.add_run(email_label_text + ' ')
                    email_label.font.name = 'Arial'
                    email_label.font.size = Pt(9)
                    email_label.font.bold = True
                    
                    email_value_text = contact_config.get('email_value', '') or ''
                    if email_value_text:
                        email_value = contact_para.add_run(email_value_text)
                        email_value.font.name = 'Arial'
                        email_value.font.size = Pt(9)
                
                # Border line
                if header_config.get('show_border', True):
                    header_border = header.add_paragraph()
                    header_border_run = header_border.add_run("_" * 100)
                    header_border_run.font.size = Pt(1)
                    header_border_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # ========== FOOTER (First Page Only) ==========
        if footer_config.get('enabled', True):
            # Use first page footer (not the default footer)
            footer = section.first_page_footer
            
            # Check if footer already has content
            has_existing_content = False
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    has_existing_content = True
                    break
            
            # Only replace if empty or force_replace is True
            if not has_existing_content or force_replace:
                # Clear existing footer content
                for paragraph in footer.paragraphs:
                    paragraph.clear()
            
            # Border line
            if footer_config.get('show_border', True):
                footer_border = footer.add_paragraph()
                footer_border_run = footer_border.add_run("_" * 100)
                footer_border_run.font.size = Pt(1)
                footer_border_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Yellow bar
            if footer_config.get('yellow_bar', True):
                yellow_bar = footer.add_paragraph()
                yellow_bar_run = yellow_bar.add_run("█" * 50)
                yellow_bar_run.font.size = Pt(6)
                yellow_bar_run.font.color.rgb = RGBColor(255, 204, 0)
                yellow_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Blue-grey bar
            if footer_config.get('blue_bar', True):
                blue_bar = footer.add_paragraph()
                blue_bar_run = blue_bar.add_run("█" * 50)
                blue_bar_run.font.size = Pt(8)
                blue_bar_run.font.color.rgb = RGBColor(100, 100, 120)
                blue_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Motto text
            motto_text = footer_config.get('motto_text', '')
            if motto_text:
                motto_para = footer.add_paragraph()
                motto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                motto_run = motto_para.add_run(motto_text)
                motto_run.font.name = 'Arial'
                motto_run.font.size = Pt(footer_config.get('motto_font_size', 10))
                motto_run.font.bold = footer_config.get('motto_bold', True)
                motto_run.font.color.rgb = RGBColor(50, 50, 50)
