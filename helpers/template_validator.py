"""
template_validator.py - Validate templates before use
"""
import re
import os
from pathlib import Path
from docx import Document

class TemplateValidator:
    """Validate Word templates"""
    
    def __init__(self):
        # Required placeholders by category
        self.required_placeholders = {
            'common': [
                '<<NAMA_SYARIKAT>>',
                '<<RUJUKAN>>',
                '<<TARIKH>>'
            ],
            'pelupusan': [
                '<<PROSES>>',
                '<<JENIS_BARANG>>',
                '<<PENGECUALIAN>>'
            ],
            'ames': [
                '<<NAMA_PEGAWAI>>',
                '<<TARIKH_ISLAM>>'
            ],
            'signupb': [
                '<<NAMA_PERNIAGAAN>>',
                '<<NO_PENDAFTARAN>>'
            ]
        }
        
        # Module field mappings - what placeholders each module uses
        self.module_field_mappings = {
            'Form2': {
                'module_name': 'Form2_Government_PyQt5',
                'required_fields': ['NAMA_SYARIKAT', 'RUJUKAN', 'TARIKH', 'PROSES', 'JENIS_BARANG'],
                'optional_fields': ['PENGECUALIAN', 'CATATAN']
            },
            'Form3': {
                'module_name': 'Form3_Government_PySide2',
                'required_fields': ['NAMA_PEGAWAI', 'TARIKH_ISLAM', 'TARIKH', 'RUJUKAN', 'BUTIRAN5D'],
                'optional_fields': ['CATATAN_AMES', 'MAKLUMAT_TAMBAHAN']
            },
            'FormDeleteItem': {
                'module_name': 'Form_DeleteItem_PyQt5',
                'required_fields': ['NAMA_ITEM', 'TARIKH', 'SEBAB_PEMADAMAN'],
                'optional_fields': ['KELULUSAN', 'CATATAN']
            },
            'FormSignUp': {
                'module_name': 'Form_SignUp_PyQt5',
                'required_fields': ['NAMA_PERNIAGAAN', 'NO_PENDAFTARAN', 'TARIKH', 'ALAMAT'],
                'optional_fields': ['TELEFON', 'EMAIL', 'CONTACT']
            }
        }
        
        # Validation rules
        self.rules = {
            'placeholder_format': r'^<<[A-Z_0-9]+>>$',
            'no_nested_placeholders': r'<<.*<<.*>>.*>>',
            'min_paragraphs': 3,
            'max_file_size_mb': 5
        }
    
    def validate_template(self, doc_path_or_doc, category='common'):
        """
        Validate a template
        
        Returns:
            dict: {
                'valid': bool,
                'errors': list,
                'warnings': list,
                'info': dict
            }
        """
        result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'info': {}
        }
        
        try:
            # Load document
            if isinstance(doc_path_or_doc, str):
                doc = Document(doc_path_or_doc)
            else:
                doc = doc_path_or_doc
            
            # Extract placeholders
            placeholders = self._extract_all_placeholders(doc)
            result['info']['placeholders_found'] = len(placeholders)
            result['info']['placeholder_list'] = sorted(list(placeholders))
            
            # Check required placeholders
            required = self.required_placeholders.get(category, [])
            missing = []
            for placeholder in required:
                if placeholder not in placeholders:
                    missing.append(placeholder)
            
            if missing:
                result['warnings'].append(f"Missing recommended placeholders: {', '.join(missing)}")
            
            # Check placeholder format
            invalid_formats = []
            for placeholder in placeholders:
                if not re.match(self.rules['placeholder_format'], placeholder):
                    invalid_formats.append(placeholder)
            
            if invalid_formats:
                result['warnings'].append(f"Invalid placeholder format: {', '.join(invalid_formats)}")
            
            # Check for nested placeholders
            for paragraph in doc.paragraphs:
                if re.search(self.rules['no_nested_placeholders'], paragraph.text):
                    result['errors'].append(f"Nested placeholders found")
                    result['valid'] = False
                    break
            
            # Check document structure
            if len(doc.paragraphs) < self.rules['min_paragraphs']:
                result['warnings'].append(f"Document has only {len(doc.paragraphs)} paragraphs")
            
            # Check for empty placeholders
            empty_placeholders = [p for p in placeholders if len(p) <= 4]  # <<>>
            if empty_placeholders:
                result['errors'].append(f"Empty placeholders found: {empty_placeholders}")
                result['valid'] = False
            
            # Document info
            result['info']['total_paragraphs'] = len(doc.paragraphs)
            result['info']['total_tables'] = len(doc.tables)
            result['info']['has_header'] = len(doc.sections[0].header.paragraphs) > 0 if doc.sections else False
            result['info']['has_footer'] = len(doc.sections[0].footer.paragraphs) > 0 if doc.sections else False
            
        except Exception as e:
            result['valid'] = False
            result['errors'].append(f"Validation failed: {str(e)}")
        
        return result
    
    def _extract_all_placeholders(self, doc):
        """Extract all placeholders from document"""
        placeholders = set()
        
        # From paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(r'<<[^>]+>>', paragraph.text)
            placeholders.update(matches)
        
        # From tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(r'<<[^>]+>>', paragraph.text)
                        placeholders.update(matches)
        
        # From headers/footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                matches = re.findall(r'<<[^>]+>>', paragraph.text)
                placeholders.update(matches)
            
            for paragraph in section.footer.paragraphs:
                matches = re.findall(r'<<[^>]+>>', paragraph.text)
                placeholders.update(matches)
        
        return placeholders
    
    def validate_replacements(self, doc, replacements):
        """Validate that replacements match template placeholders"""
        result = {
            'valid': True,
            'errors': [],
            'warnings': []
        }
        
        # Get template placeholders
        template_placeholders = self._extract_all_placeholders(doc)
        replacement_keys = set(replacements.keys())
        
        # Check for missing replacements
        missing = template_placeholders - replacement_keys
        if missing:
            result['warnings'].append(f"Placeholders without replacements: {', '.join(sorted(missing))}")
        
        # Check for unused replacements
        unused = replacement_keys - template_placeholders
        if unused:
            result['warnings'].append(f"Replacements not in template: {', '.join(sorted(unused))}")
        
        # Check for empty replacement values
        empty = [k for k, v in replacements.items() if not v or str(v).strip() == '']
        if empty:
            result['warnings'].append(f"Empty values for: {', '.join(sorted(empty))}")
        
        return result
    
    def generate_validation_report(self, template_path, category='common'):
        """Generate comprehensive validation report"""
        validation = self.validate_template(template_path, category)
        
        report = []
        report.append("=" * 60)
        report.append("TEMPLATE VALIDATION REPORT")
        report.append("=" * 60)
        report.append(f"Template: {template_path}")
        report.append(f"Category: {category}")
        report.append(f"Status: {'VALID' if validation['valid'] else 'INVALID'}")
        report.append("")
        
        if validation['errors']:
            report.append("ERRORS:")
            for error in validation['errors']:
                report.append(f"  - {error}")
            report.append("")
        
        if validation['warnings']:
            report.append("WARNINGS:")
            for warning in validation['warnings']:
                report.append(f"  - {warning}")
            report.append("")
        
        report.append("INFORMATION:")
        for key, value in validation['info'].items():
            if key == 'placeholder_list':
                report.append(f"  Placeholders:")
                for placeholder in value:
                    report.append(f"    - {placeholder}")
            else:
                report.append(f"  {key}: {value}")
        
        report.append("=" * 60)
        
        return "\n".join(report)
    
    def scan_all_templates(self, template_dir='Templates'):
        """
        Scan all templates in directory and extract their placeholders
        
        Returns:
            dict: {
                'template_name': {
                    'path': str,
                    'placeholders': set,
                    'valid': bool,
                    'errors': list
                }
            }
        """
        templates_info = {}
        
        if not os.path.exists(template_dir):
            return templates_info
        
        # Scan all .docx files
        for file in Path(template_dir).glob('*.docx'):
            template_name = file.name
            try:
                doc = Document(str(file))
                placeholders = self._extract_all_placeholders(doc)
                
                # Validate
                validation = self.validate_template(doc, 'common')
                
                templates_info[template_name] = {
                    'path': str(file),
                    'placeholders': sorted(list(placeholders)),
                    'valid': validation['valid'],
                    'errors': validation['errors'],
                    'warnings': validation['warnings']
                }
            except Exception as e:
                templates_info[template_name] = {
                    'path': str(file),
                    'placeholders': [],
                    'valid': False,
                    'errors': [f"Failed to load: {str(e)}"],
                    'warnings': []
                }
        
        return templates_info
    
    def match_template_to_module(self, template_name, module_name):
        """
        Check if template matches a module's requirements
        
        Args:
            template_name: Name of template file
            module_name: Name of module (Form2, Form3, FormDeleteItem, FormSignUp)
        
        Returns:
            dict: {
                'match_score': float (0-100),
                'is_compatible': bool,
                'required_fields_found': list,
                'required_fields_missing': list,
                'extra_placeholders': list,
                'recommendation': str
            }
        """
        result = {
            'match_score': 0,
            'is_compatible': False,
            'required_fields_found': [],
            'required_fields_missing': [],
            'extra_placeholders': [],
            'recommendation': ''
        }
        
        # Get module requirements
        if module_name not in self.module_field_mappings:
            result['recommendation'] = f"Unknown module: {module_name}"
            return result
        
        module_config = self.module_field_mappings[module_name]
        required_fields = module_config['required_fields']
        optional_fields = module_config['optional_fields']
        
        # Get template placeholders
        template_info = self.scan_all_templates()
        if template_name not in template_info:
            result['recommendation'] = f"Template not found: {template_name}"
            return result
        
        template_placeholders = template_info[template_name]['placeholders']
        
        # Extract field names from placeholders (remove << and >>)
        placeholder_fields = set()
        for ph in template_placeholders:
            field_name = ph.replace('<<', '').replace('>>', '')
            placeholder_fields.add(field_name)
        
        # Check required fields
        required_found = 0
        for field in required_fields:
            if field in placeholder_fields:
                result['required_fields_found'].append(field)
                required_found += 1
            else:
                result['required_fields_missing'].append(field)
        
        # Check optional fields
        optional_found = 0
        for field in optional_fields:
            if field in placeholder_fields:
                optional_found += 1
        
        # Check for extra/unknown placeholders
        all_expected = set(required_fields + optional_fields)
        for field in placeholder_fields:
            if field not in all_expected:
                result['extra_placeholders'].append(field)
        
        # Calculate match score
        required_percentage = (required_found / len(required_fields) * 100) if required_fields else 100
        optional_percentage = (optional_found / len(optional_fields) * 100) if optional_fields else 0
        
        result['match_score'] = (required_percentage * 0.7) + (optional_percentage * 0.3)
        
        # Determine compatibility
        result['is_compatible'] = required_found == len(required_fields)
        
        # Generate recommendation
        if result['is_compatible']:
            if result['extra_placeholders']:
                result['recommendation'] = f"✓ Compatible with {module_name}. Has {len(result['extra_placeholders'])} extra placeholders."
            else:
                result['recommendation'] = f"✓ Perfectly compatible with {module_name}!"
        else:
            missing_count = len(result['required_fields_missing'])
            result['recommendation'] = f"✗ Incompatible - missing {missing_count} required fields: {', '.join(result['required_fields_missing'])}"
        
        return result
    
    def select_templates_for_module(self, module_name, template_dir='Templates', min_score=60):
        """
        Scan all templates and return list of compatible ones for a module
        
        Args:
            module_name: Module name (Form2, Form3, FormDeleteItem, FormSignUp)
            template_dir: Directory containing templates
            min_score: Minimum match score (0-100) to include
        
        Returns:
            dict: {
                'module': str,
                'compatible_templates': [
                    {
                        'template_name': str,
                        'match_score': float,
                        'required_found': int,
                        'required_total': int,
                        'recommendation': str,
                        'path': str
                    }
                ],
                'incompatible_templates': [...]
            }
        """
        result = {
            'module': module_name,
            'compatible_templates': [],
            'incompatible_templates': []
        }
        
        # Scan all templates
        templates = self.scan_all_templates(template_dir)
        
        # Match each template to module
        for template_name, template_info in templates.items():
            if not template_info['valid']:
                result['incompatible_templates'].append({
                    'template_name': template_name,
                    'reason': 'Invalid template',
                    'errors': template_info['errors'],
                    'path': template_info['path']
                })
                continue
            
            match_result = self.match_template_to_module(template_name, module_name)
            
            template_entry = {
                'template_name': template_name,
                'match_score': match_result['match_score'],
                'required_found': len(match_result['required_fields_found']),
                'required_total': len(self.module_field_mappings[module_name]['required_fields']),
                'recommendation': match_result['recommendation'],
                'path': template_info['path'],
                'is_compatible': match_result['is_compatible']
            }
            
            if match_result['is_compatible'] and match_result['match_score'] >= min_score:
                result['compatible_templates'].append(template_entry)
            else:
                result['incompatible_templates'].append(template_entry)
        
        # Sort by score
        result['compatible_templates'].sort(key=lambda x: x['match_score'], reverse=True)
        result['incompatible_templates'].sort(key=lambda x: x.get('match_score', 0), reverse=True)
        
        return result
    
    def get_template_field_report(self, template_name, template_dir='Templates'):
        """
        Get detailed report of all fields in a template and which modules can use it
        
        Returns:
            dict: {
                'template': str,
                'placeholders': list,
                'modules_compatibility': {
                    'Form2': {'score': float, 'compatible': bool},
                    'Form3': {...}
                }
            }
        """
        result = {
            'template': template_name,
            'placeholders': [],
            'modules_compatibility': {}
        }
        
        templates = self.scan_all_templates(template_dir)
        if template_name not in templates:
            result['error'] = f"Template not found: {template_name}"
            return result
        
        result['placeholders'] = templates[template_name]['placeholders']
        
        # Check compatibility with all modules
        for module_name in self.module_field_mappings.keys():
            match_result = self.match_template_to_module(template_name, module_name)
            result['modules_compatibility'][module_name] = {
                'score': match_result['match_score'],
                'compatible': match_result['is_compatible'],
                'required_found': len(match_result['required_fields_found']),
                'required_missing': match_result['required_fields_missing']
            }
        
        return result


if __name__ == '__main__':
    validator = TemplateValidator()
    print("Template Validator initialized")
    print("Use: validator.validate_template('path/to/template.docx', 'category')")
    print("Or: validator.select_templates_for_module('Form2')")


